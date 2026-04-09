from pathlib import Path
import re

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


DOWNLOADS = Path(r"C:\Users\tumsu\Downloads")
MANUAL_DIR = Path(__file__).resolve().parent
PIC_DIR = MANUAL_DIR / "pic"
OUTPUT_PATH = MANUAL_DIR / "Purchase_Manual_with_Auto_PO_4_3.docx"

NORMAL_FONT = "Angsana New"
NORMAL_SIZE = 18
CAPTION_SIZE = 16
IMAGE_WIDTH = 4.7
CAPTION_COL_WIDTH = 1.8


FIELD_ROWS = [
    ["Field / ปุ่ม", "ความหมาย", "กดหรือแก้ไขได้หรือไม่", "กดแล้วเกิดอะไรขึ้น"],
    ["New", "สร้างรายการเติมสต็อกใหม่", "ได้", "ระบบเพิ่มแถวใหม่ด้านล่างสุดให้กรอกข้อมูล Replenishment"],
    ["Product", "สินค้าเป้าหมายของการเติมสต็อก", "ได้", "ใช้เป็นฐานข้อมูลของ orderpoint รายการนั้น"],
    ["Location", "คลังหรือ location ปลายทางที่ต้องการให้มีสินค้า", "ได้", "ระบบใช้คำนวณ On Hand และ Forecast ของ location นั้น"],
    ["On Hand", "ยอดคงเหลือจริง ณ location นั้น", "ไม่ได้", "ระบบคำนวณอัตโนมัติจาก stock ปัจจุบัน"],
    ["Forecast", "ยอดคาดการณ์หลังรวมของเข้าและของออก", "ไม่ได้", "ระบบคำนวณอัตโนมัติจาก stock movement และ demand"],
    ["Route", "วิธีเติมสต็อก เช่น Buy, Manufacture, Transfer", "ได้", "บอกว่าระบบจะเติมสต็อกด้วยวิธีใด"],
    ["Trigger", "วิธี trigger การสั่งเติม", "ได้", "กำหนดว่าจะใช้ Manual หรือ Auto"],
    ["Min", "จุดต่ำสุดที่ยอมให้มีสินค้า", "ได้", "ใช้เป็นเกณฑ์การคำนวณเติมสต็อก"],
    ["Max", "เป้าหมายหลังเติมสต็อก", "ได้", "ใช้เป็นเป้าหมายในการคำนวณจำนวนเติม"],
    ["To Order", "จำนวนที่จะสั่งเติม", "ได้บางกรณี", "ถ้าเป็น Manual ผู้ใช้สามารถกำหนดจำนวนได้เอง"],
    ["UoM", "หน่วยนับสินค้า", "ไม่ได้", "ดึงจาก master data ของสินค้า"],
    ["Order", "สั่งเติมทันที", "ได้", "ระบบเรียก Replenish ทันทีตาม route ที่กำหนด"],
    ["Automate", "เปลี่ยนรายการเป็น Auto และสั่งทันที", "ได้", "ระบบเปลี่ยน Trigger เป็น Auto แล้ว Replenish ต่อทันที"],
    ["Snooze", "พักรายการ Manual ชั่วคราว", "ได้บางกรณี", "ซ่อนรายการนี้ชั่วคราวจากมุมมอง Not Snoozed"],
]

ROUTE_ROWS = [
    ["Route", "ใช้เมื่อไร", "ระบบจะสร้างอะไร"],
    ["Buy", "ต้องการให้ระบบไปสร้างเอกสารซื้อ", "สร้าง procurement ฝั่งซื้อ และไปสร้างหรืออัปเดต RFQ / PO"],
    ["Manufacture", "ต้องการเติมสต็อกโดยการผลิต", "ไปสร้าง MO"],
    ["Replenish on Order (MTO)", "ต้องการสร้าง supply ตาม demand จริง", "ไปตาม flow ของ MTO"],
    ["Auto Transfer ...", "ต้องการดึงของจาก location หรือคลังอื่น", "ไปสร้าง internal transfer"],
    ["Resupply Subcontractor on Order", "ใช้กับงาน subcontract", "ไปตาม flow ของ subcontractor"],
]

TRIGGER_ROWS = [
    ["Trigger", "ความหมาย", "ผลต่อการทำงาน"],
    ["Manual", "ผู้ใช้เป็นคนตัดสินใจกด Order หรือ Automate เอง", "เหมาะกับกรณีที่ต้องการ review ก่อนสั่ง"],
    ["Auto", "ใช้เป็น reordering rule อัตโนมัติ", "ระบบใช้ Min, Max และ Forecast คำนวณปริมาณที่จะเติม"],
]

ACTION_ROWS = [
    ["ปุ่ม", "หน้าที่", "กดแล้วเกิดอะไรขึ้น"],
    ["Order", "สั่งเติมทันที", "ระบบเรียก Replenish ทันที ถ้า Route = Buy จะไปสร้างหรืออัปเดต RFQ / PO"],
    ["Automate", "เปลี่ยนรายการนี้เป็น Auto แล้วสั่งทันที", "ระบบเปลี่ยน Trigger เป็น Auto แล้ว Replenish ต่อทันที"],
    ["Snooze", "พักรายการชั่วคราว", "ระบบซ่อนรายการชั่วคราวจากมุมมอง Not Snoozed"],
]

CONDITION_ROWS = [
    ["เงื่อนไข", "คำอธิบาย"],
    ["Product มี Vendor / Supplier Info", "ถ้าไม่มีผู้ขาย ระบบอาจสร้าง RFQ / PO ไม่ได้"],
    ["Route = Buy", "ถ้า route ไม่ใช่ Buy ระบบจะไปสร้างเอกสารชนิดอื่นแทน"],
    ["To Order มากกว่า 0 หรือระบบคำนวณได้", "ถ้าจำนวนเป็น 0 จะไม่มีอะไรให้สั่ง"],
    ["ผู้ใช้มีสิทธิ์ตาม flow", "ผู้ใช้ต้องเข้าถึง Inventory และ Purchase ในส่วนที่เกี่ยวข้องได้"],
]


def purchase_doc_path() -> Path:
    matches = sorted(DOWNLOADS.glob("*Purchase* (1).docx"))
    if not matches:
        raise FileNotFoundError("Purchase manual source file not found")
    return matches[0]


def set_run_font(run, size, bold=None):
    run.font.name = NORMAL_FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), NORMAL_FONT)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def set_paragraph_format(paragraph, align=None, space_before=None, space_after=None):
    paragraph.paragraph_format.line_spacing = 1.0
    if align is not None:
        paragraph.alignment = align
    if space_before is not None:
        paragraph.paragraph_format.space_before = Pt(space_before)
    if space_after is not None:
        paragraph.paragraph_format.space_after = Pt(space_after)


def insert_paragraph_before(doc, anchor, text="", style="normal", align=None, font_size=NORMAL_SIZE, bold=False, space_before=None, space_after=None):
    paragraph = doc.add_paragraph(style=style)
    if text:
        run = paragraph.add_run(text)
        set_run_font(run, font_size, bold=bold)
    set_paragraph_format(paragraph, align=align, space_before=space_before, space_after=space_after)
    anchor._p.addprevious(paragraph._p)
    return paragraph


def insert_picture_before(doc, anchor, image_name, caption_text):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    remove_table_borders(table)

    image_cell = table.cell(0, 0)
    caption_cell = table.cell(0, 1)
    image_cell.width = Inches(IMAGE_WIDTH)
    caption_cell.width = Inches(CAPTION_COL_WIDTH)
    image_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    caption_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    image_para = image_cell.paragraphs[0]
    image_para.style = "normal"
    image_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_run = image_para.add_run()
    image_run.add_picture(str(PIC_DIR / image_name), width=Inches(IMAGE_WIDTH))

    caption_para = caption_cell.paragraphs[0]
    caption_para.style = "normal"
    caption_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    caption_run = caption_para.add_run(caption_text)
    set_run_font(caption_run, CAPTION_SIZE)
    set_paragraph_format(caption_para)

    anchor._p.addprevious(table._tbl)
    insert_paragraph_before(doc, anchor, "", style="normal")


def remove_table_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "nil")


def insert_table_before(doc, anchor, rows):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.text = ""
            para = cell.paragraphs[0]
            run = para.add_run(str(value))
            set_run_font(run, NORMAL_SIZE, bold=(r_idx == 0))
            set_paragraph_format(para)
    anchor._p.addprevious(table._tbl)
    insert_paragraph_before(doc, anchor, "", style="normal")


def replace_paragraph_text(paragraph, new_text):
    if paragraph.runs:
        first = paragraph.runs[0]
        first.text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        run = paragraph.add_run(new_text)
        set_run_font(run, NORMAL_SIZE)


def renumber_chapter4_figures(paragraphs):
    pattern = re.compile(r"รูป 4\.(\d+)")
    for paragraph in paragraphs:
        text = paragraph.text
        if not text:
            continue

        def repl(match):
            num = int(match.group(1))
            if num >= 5:
                return f"รูป 4.{num + 10}"
            return match.group(0)

        new_text = pattern.sub(repl, text)
        if new_text != text:
            replace_paragraph_text(paragraph, new_text)


def build():
    source = purchase_doc_path()
    doc = Document(str(source))

    start_idx = None
    anchor_idx = None
    for i, para in enumerate(doc.paragraphs):
        txt = para.text.strip()
        if txt.startswith("4.3 "):
            start_idx = i
        if txt.startswith("4.4 "):
            anchor_idx = i
            break

    if start_idx is None or anchor_idx is None:
        raise ValueError("Could not find section 4.3/4.4 in source document")

    anchor = doc.paragraphs[anchor_idx]

    renumber_chapter4_figures(doc.paragraphs[anchor_idx:])

    insert_paragraph_before(
        doc,
        anchor,
        "หน้าจอนี้ใช้สำหรับการสร้างรายการเติมสต็อก (Replenishment) เพื่อให้ระบบสร้างเอกสารจัดหาสินค้าอัตโนมัติตามเส้นทางที่กำหนด โดยในกรณีที่เลือก Route เป็น Buy ระบบจะสร้างหรืออัปเดตเอกสาร RFQ / Purchase Order ให้โดยอัตโนมัติ",
        style="normal",
    )
    insert_paragraph_before(doc, anchor, "", style="normal")
    insert_paragraph_before(
        doc,
        anchor,
        "การเข้าถึง (Menu Path): ไปที่แอป Inventory > Operations > Replenishment",
        style="normal",
        bold=True,
    )
    insert_paragraph_before(doc, anchor, "", style="normal")

    insert_picture_before(doc, anchor, "1.png", "รูป 4.5 แอป Inventory")
    insert_paragraph_before(doc, anchor, "คลิกเมนู Operations > Replenishment เพื่อเข้าสู่หน้าจอรายการเติมสต็อก", style="normal")
    insert_paragraph_before(doc, anchor, "", style="normal")

    insert_picture_before(doc, anchor, "2.png", "รูป 4.6 เมนู Operations > Replenishment")
    insert_paragraph_before(doc, anchor, "หน้าจอ Replenishment ใช้สำหรับสร้างรายการเติมสต็อกและกำหนดวิธีจัดหาสินค้า โดยหากกำหนด Route เป็น Buy ระบบจะสร้างหรืออัปเดต RFQ / Purchase Order ให้อัตโนมัติ", style="normal")
    insert_paragraph_before(doc, anchor, "", style="normal")

    insert_picture_before(doc, anchor, "11.png", "รูป 4.7 ภาพรวมหน้าจอ Replenishment")
    insert_paragraph_before(doc, anchor, "ส่วนประกอบภายในหน้าจอ Replenishment", style="normal")
    insert_paragraph_before(doc, anchor, "", style="normal")
    insert_table_before(doc, anchor, FIELD_ROWS)

    insert_paragraph_before(doc, anchor, "ขั้นตอนการทำงาน (Step-by-step):", style="normal", bold=True, space_before=12, space_after=12)

    insert_paragraph_before(doc, anchor, "1.\tกดปุ่ม New เพื่อเพิ่มรายการเติมสต็อกใหม่", style="normal")
    insert_paragraph_before(doc, anchor, "", style="normal")
    insert_picture_before(doc, anchor, "5.png", "รูป 4.8 ปุ่ม New สำหรับสร้างรายการ Replenishment")

    insert_paragraph_before(doc, anchor, "2.\tเลือกสินค้าในช่อง Product เพื่อระบุว่าสินค้าใดต้องการให้ระบบเติมสต็อก", style="normal")
    insert_paragraph_before(doc, anchor, "", style="normal")
    insert_picture_before(doc, anchor, "6.png", "รูป 4.9 การเลือกสินค้าในช่อง Product")

    insert_paragraph_before(doc, anchor, "3.\tเลือก Location ปลายทางที่ต้องการให้มีสินค้าในระบบ", style="normal")
    insert_paragraph_before(doc, anchor, "", style="normal")
    insert_picture_before(doc, anchor, "7.png", "รูป 4.10 การเลือก Location ของรายการเติมสต็อก")

    insert_paragraph_before(doc, anchor, "4.\tเลือก Route เป็น Buy เพื่อให้ระบบจัดซื้อสร้าง RFQ / Purchase Order ให้อัตโนมัติ", style="normal")
    insert_paragraph_before(doc, anchor, "", style="normal")
    insert_picture_before(doc, anchor, "8.png", "รูป 4.11 การเลือก Route สำหรับการเติมสต็อก")

    insert_paragraph_before(doc, anchor, "5.\tเลือก Trigger เป็น Manual หากต้องการให้ผู้ใช้ตัดสินใจกดสั่งเอง หรือเลือก Auto หากต้องการเก็บเป็นกฎการเติมสต็อกอัตโนมัติ", style="normal")
    insert_paragraph_before(doc, anchor, "", style="normal")
    insert_picture_before(doc, anchor, "9.png", "รูป 4.12 การเลือก Trigger ของรายการเติมสต็อก")

    insert_paragraph_before(doc, anchor, "6.\tระบุค่า Min, Max และ/หรือ To Order ตามนโยบายการเติมสต็อก แล้วกด Save เพื่อบันทึกรายการ", style="normal")
    insert_paragraph_before(doc, anchor, "", style="normal")
    insert_picture_before(doc, anchor, "10.png", "รูป 4.13 การกด Save เพื่อบันทึกรายการ")

    insert_paragraph_before(doc, anchor, "7.\tหลังบันทึกแล้ว หากรายการมีจำนวนที่ต้องสั่ง ระบบจะแสดงปุ่ม Order, Automate และ Snooze ให้เลือกใช้งานตามสิทธิ์และเงื่อนไขของรายการ", style="normal")
    insert_paragraph_before(doc, anchor, "", style="normal")
    insert_picture_before(doc, anchor, "3.png", "รูป 4.14 ปุ่ม Order / Automate / Snooze บนหน้าจอ Replenishment")

    insert_paragraph_before(doc, anchor, "ตารางอธิบายค่าใน Route", style="normal", bold=True, space_before=12, space_after=12)
    insert_table_before(doc, anchor, ROUTE_ROWS)

    insert_paragraph_before(doc, anchor, "ตารางอธิบายค่าใน Trigger", style="normal", bold=True, space_before=12, space_after=12)
    insert_table_before(doc, anchor, TRIGGER_ROWS)

    insert_paragraph_before(doc, anchor, "ตารางอธิบายปุ่ม Action", style="normal", bold=True, space_before=12, space_after=12)
    insert_table_before(doc, anchor, ACTION_ROWS)

    insert_paragraph_before(doc, anchor, "เงื่อนไขที่ต้องพร้อมเพื่อให้เกิด Auto PO จริง", style="normal", bold=True, space_before=12, space_after=12)
    insert_table_before(doc, anchor, CONDITION_ROWS)

    insert_paragraph_before(doc, anchor, "สรุป: ผู้ใช้สร้างรายการ Replenishment โดยระบุสินค้า, Location, Trigger และ Route หากเลือก Route = Buy และกด Order หรือ Automate ระบบจะสร้าง procurement ฝั่งซื้อและนำไปสร้างหรืออัปเดต RFQ / Purchase Order ให้อัตโนมัติ", style="normal")
    insert_paragraph_before(doc, anchor, "", style="normal")

    doc.save(str(OUTPUT_PATH))
    print(OUTPUT_PATH)


if __name__ == "__main__":
    build()

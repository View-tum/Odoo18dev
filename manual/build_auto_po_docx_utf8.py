from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


MANUAL_DIR = Path(__file__).resolve().parent
OUTPUT_DOCX = MANUAL_DIR / "Auto_PO_from_Replenishment_with_images_TH.docx"
PIC_DIR = MANUAL_DIR / "pic"


def pick_source_markdown() -> Path:
    candidates = [p for p in MANUAL_DIR.glob("*.md") if p.name != "Auto_PO_from_Replenishment_TH.md"]
    if not candidates:
        raise FileNotFoundError("No source markdown found")
    return sorted(candidates, key=lambda p: p.stat().st_mtime, reverse=True)[0]


STEP_IMAGES = {
    1: ["1.png"],
    2: ["2.png"],
    3: ["3.png"],
    4: ["5.png"],
    5: ["6.png"],
    6: ["7.png"],
    7: ["8.png"],
    8: ["9.png"],
    9: ["10.png"],
    10: ["11.png", "4.png"],
    11: ["3.png"],
}


def set_font(doc: Document) -> None:
    for style_name in ["Normal", "Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[style_name]
        style.font.name = "TH Sarabun New"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "TH Sarabun New")
        style.font.size = Pt(14 if style_name == "Normal" else 18)


def add_table(doc: Document, raw_rows: list[str]) -> None:
    rows = []
    for raw in raw_rows:
        parts = [p.strip() for p in raw.strip().strip("|").split("|")]
        if parts and all(re.fullmatch(r"[:\- ]+", p or "") for p in parts):
            continue
        rows.append(parts)
    if not rows:
        return
    max_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j in range(max_cols):
            text = row[j] if j < len(row) else ""
            para = table.cell(i, j).paragraphs[0]
            run = para.add_run(text)
            if i == 0:
                run.bold = True


def add_step_images(doc: Document, line: str) -> None:
    match = re.match(r"^##\s+Step\s+(\d+):", line)
    if not match:
        return
    step_no = int(match.group(1))
    for image_name in STEP_IMAGES.get(step_no, []):
        image_path = PIC_DIR / image_name
        if image_path.exists():
            doc.add_picture(str(image_path), width=Inches(6.9))
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.add_run(image_name).italic = True


def build() -> Path:
    source_md = pick_source_markdown()
    text = source_md.read_text(encoding="utf-8")
    lines = text.splitlines()

    doc = Document()
    set_font(doc)

    table_buffer: list[str] = []

    def flush_table() -> None:
        nonlocal table_buffer
        if table_buffer:
            add_table(doc, table_buffer)
            table_buffer = []

    for line in lines:
        stripped = line.strip()

        if stripped.startswith("|") and stripped.endswith("|"):
            table_buffer.append(line)
            continue

        flush_table()

        if not stripped:
            doc.add_paragraph("")
            continue

        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
            continue

        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
            add_step_images(doc, line)
            continue

        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
            continue

        if re.match(r"^\d+\.\s+", line):
            para = doc.add_paragraph(style="List Number")
            para.add_run(re.sub(r"^\d+\.\s+", "", line))
            continue

        if line.startswith("- "):
            para = doc.add_paragraph(style="List Bullet")
            para.add_run(line[2:].strip())
            continue

        doc.add_paragraph(line)

    flush_table()
    doc.save(OUTPUT_DOCX)
    return OUTPUT_DOCX


if __name__ == "__main__":
    print(build())

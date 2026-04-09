from pathlib import Path
import re
import os

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from inventory_manual_annotations import ANNOTATIONS

BASE_DIR = Path(__file__).resolve().parent
SOURCE_MD = BASE_DIR / "inventory_manual_source_th.md"
OUTPUT_DOCX = BASE_DIR / "Inventory_Manual_View_IV_01_to_IV_23.docx"
ANNOTATED_DIR = BASE_DIR / "inventory_manual_assets_annotated"
ASSET_DIR = BASE_DIR / "inventory_manual_assets"
ORIGINAL_ASSET_DIR = ANNOTATED_DIR / "inventory_manual_assets"
USE_ANNOTATED = os.getenv("INVENTORY_MANUAL_USE_ANNOTATED") == "1"

FONT_NAME = "Angsana New"
NORMAL_SIZE = 18
H1_SIZE = 22
H2_SIZE = 20
TITLE_SIZE = 26
IMAGE_WIDTH = 7.1


def set_font(run, size=Pt(NORMAL_SIZE), bold=None):
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run.font.size = size
    if bold is not None:
        run.bold = bold


def configure_doc(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.header_distance = Inches(0.5)
    section.footer_distance = Inches(0.5)

    for style_name, size, bold in [
        ("Normal", Pt(NORMAL_SIZE), False),
        ("Heading 1", Pt(H1_SIZE), True),
        ("Heading 2", Pt(H2_SIZE), True),
        ("Heading 3", Pt(NORMAL_SIZE), True),
        ("Title", Pt(TITLE_SIZE), True),
    ]:
        style = doc.styles[style_name]
        style.font.name = FONT_NAME
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
        style.font.size = size
        style.font.bold = bold


def add_paragraph(doc: Document, text: str, style=None, align=None, size=Pt(NORMAL_SIZE), bold=False):
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    set_font(run, size=size, bold=bold)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_after = Pt(0)
    if align is not None:
        p.alignment = align
    return p


def add_table(doc: Document, raw_rows: list[str]):
    rows = []
    for raw in raw_rows:
        parts = [p.strip() for p in raw.strip().strip("|").split("|")]
        if parts and all(re.fullmatch(r"[:\- ]+", p or "") for p in parts):
            continue
        rows.append(parts)
    if not rows:
        return
    max_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j in range(max_cols):
            text = row[j] if j < len(row) else ""
            para = table.cell(i, j).paragraphs[0]
            run = para.add_run(text)
            set_font(run, size=Pt(NORMAL_SIZE), bold=(i == 0))
            para.paragraph_format.line_spacing = 1.0
            para.paragraph_format.space_after = Pt(0)


def add_image(doc: Document, image_name: str, caption: str):
    image_path = ASSET_DIR / image_name
    if not image_path.exists():
        image_path = ORIGINAL_ASSET_DIR / image_name
    if USE_ANNOTATED:
        annotated_path = ANNOTATED_DIR / image_name
        if annotated_path.exists():
            image_path = annotated_path
    if not image_path.exists():
        raise FileNotFoundError(f"Missing image: {image_path}")
    doc.add_picture(str(image_path), width=Inches(IMAGE_WIDTH))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_paragraph(doc, caption, align=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(16))
    if USE_ANNOTATED and image_name in ANNOTATIONS:
        legend = "จุดอ้างอิงในภาพ: " + "  ".join(
            f"[{item['id']}] {item['legend']}" for item in ANNOTATIONS[image_name]
        )
        add_paragraph(doc, legend, size=Pt(16), bold=True)


def build() -> Path:
    text = SOURCE_MD.read_text(encoding="utf-8")
    lines = text.splitlines()

    doc = Document()
    configure_doc(doc)

    table_buffer: list[str] = []

    def flush_table():
        nonlocal table_buffer
        if table_buffer:
            add_table(doc, table_buffer)
            table_buffer = []

    for line in lines:
        stripped = line.strip()

        if stripped.startswith("|") and stripped.endswith("|"):
            table_buffer.append(line)
            continue

        flush_table()

        if not stripped:
            doc.add_paragraph("")
            continue

        img_match = re.match(r"^!\[(.+)\]\((.+)\)$", stripped)
        if img_match:
            add_image(doc, img_match.group(2), img_match.group(1))
            continue

        if line.startswith("# "):
            add_paragraph(doc, line[2:].strip(), style="Title", align=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(TITLE_SIZE), bold=True)
            continue

        if line.startswith("## "):
            text = line[3:].strip()
            if text == "2. Inventory Module":
                doc.add_section(WD_SECTION.NEW_PAGE)
            add_paragraph(doc, text, style="Heading 1", size=Pt(H1_SIZE), bold=True)
            continue

        if line.startswith("### "):
            add_paragraph(doc, line[4:].strip(), style="Heading 2", size=Pt(H2_SIZE), bold=True)
            continue

        if line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(line[2:].strip())
            set_font(run)
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_after = Pt(0)
            continue

        add_paragraph(doc, line)

    flush_table()
    doc.save(OUTPUT_DOCX)
    return OUTPUT_DOCX


if __name__ == "__main__":
    print(build())

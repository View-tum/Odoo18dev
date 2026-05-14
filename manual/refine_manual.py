import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT

# Paths
base_dir = r"c:\365_project\TheCool18e\Dev\manual\Accouting_Manual\generated_20260408\docx"
input_path = os.path.join(base_dir, "6_Fixed_Asset_Improved.docx")
output_path = os.path.join(base_dir, "6_Fixed_Asset_Refined.docx")
new_image_path = r"C:\Users\tumsu\.gemini\antigravity\brain\bd2314b9-b5e5-468c-bf59-e214cdee1d93\odoo18_correct_disposal_je_1776926845191.png"

# Tone & Wording Mapping
wording_replacements = {
    "วัตถุประสงค์": "เป้าหมายและประโยชน์",
    "Menu Path": "เส้นทางเข้าสู่เมนู",
    "รูปพาเข้าเมนู": "ขั้นตอนการเข้าใช้งาน",
    "Account Path": "การกำหนดค่าบัญชี",
    "Default Account": "บัญชีตั้งต้น",
    "หมายเหตุ": "ข้อควรระวังและคำแนะนำเพิ่มเติม",
    "Field สำคัญ": "ข้อมูลที่ต้องระบุ",
    "จุดตรวจ": "สิ่งที่ต้องตรวจสอบ",
}

def set_column_width(column, width):
    for cell in column.cells:
        cell.width = width

def refine_document():
    if not os.path.exists(input_path):
        print(f"Error: File not found at {input_path}")
        return

    doc = Document(input_path)

    # 1. Replace Wording in Paragraphs
    for i, para in enumerate(doc.paragraphs):
        # Special case: Remove the negative warning about 510000 and replace with positive info
        if "510000" in para.text and "ต้นทุนขาย" in para.text:
            if "ภาพที่ 6.5-2" in para.text or i > 280: # Context of disposal
                para.text = "ขาบัญชีที่ถูกต้องสำหรับการตัดจำหน่าย (Disposal) คือการล้างมูลค่าทุนและค่าเสื่อมสะสมออก และรับรู้ผลต่างที่บัญชีกำไร/ขาดทุนจากการจำหน่ายสินทรัพย์"
        
        for old, new in wording_replacements.items():
            if old in para.text:
                para.text = para.text.replace(old, new)

    # 2. Iterate Tables for Layout
    for table in doc.tables:
        if len(table.columns) >= 2:
            first_cell_text = table.cell(0, 0).text.strip().lower()
            if any(key in first_cell_text for key in ["#", "ลำดับ", "no.", "id", "step"]):
                set_column_width(table.columns[0], Inches(0.4))
                if len(table.columns) >= 2:
                    set_column_width(table.columns[1], Inches(4.5))
            elif "field" in first_cell_text or "ข้อมูล" in first_cell_text:
                set_column_width(table.columns[0], Inches(1.2))
                if len(table.columns) >= 2:
                    set_column_width(table.columns[1], Inches(3.7))

    # 3. Replace Image for Disposal (ภาพที่ 6.5-2)
    for i, para in enumerate(doc.paragraphs):
        if "ภาพที่ 6.5-2 Journal Entry จากการตัดจำหน่ายทรัพย์สิน" in para.text:
            # The image is likely in the paragraph BEFORE this caption
            if i > 0:
                prev_para = doc.paragraphs[i-1]
                # Remove existing content in the previous paragraph
                for run in prev_para.runs:
                    run.text = ""
                # Add the new image to the previous paragraph
                run = prev_para.add_run()
                if os.path.exists(new_image_path):
                    run.add_picture(new_image_path, width=Inches(6.0))
                
                # Insert a clear explanation right after the caption
                new_para = doc.paragraphs[i+1].insert_paragraph_before("")
                run = new_para.add_run("📊 การบันทึกบัญชีที่ถูกต้อง (Correct Journal Entry)")
                run.bold = True
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(0, 102, 204) # Professional Blue
                
                # Insert a table showing the correct mapping
                table = doc.add_table(rows=1, cols=3)
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'รายการบัญชี'
                hdr_cells[1].text = 'Dr / Cr'
                hdr_cells[2].text = 'คำอธิบาย'
                
                row_data = [
                    ('ค่าเสื่อมราคาสะสม (Accumulated Depreciation)', 'Dr', 'ล้างยอดสะสมเดิมออก'),
                    ('กำไร/ขาดทุนจากการจำหน่ายสินทรัพย์ (Gain/Loss on Disposal)', 'Dr', 'รับรู้มูลค่าคงเหลือที่ตัดออก'),
                    ('บัญชีสินทรัพย์ถาวร (Fixed Asset Account)', 'Cr', 'ล้างมูลค่าทุนเดิมออก')
                ]
                
                for item, drcr, note in row_data:
                    row_cells = table.add_row().cells
                    row_cells[0].text = item
                    row_cells[1].text = drcr
                    row_cells[2].text = note

    doc.save(output_path)
    print(f"Saved refined document to: {output_path}")

if __name__ == "__main__":
    refine_document()

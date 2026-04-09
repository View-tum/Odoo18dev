from __future__ import annotations

import json
import shutil
from pathlib import Path
from typing import Dict, List, Tuple

from openpyxl import load_workbook
from openpyxl.drawing.image import Image as XLImage
from openpyxl.styles import Alignment, Font
from PIL import Image
from playwright.sync_api import sync_playwright
from docx import Document
from docx.shared import Inches


ROOT = Path(r"C:\365_project\TheCool18e\Dev")
WORKBOOK_PATH = Path(
    r"C:\Users\tumsu\Downloads\UAT_GoldMints_Test Scenario_MANU_actual_flow_th_20260407.xlsx"
)
OUTPUT_XLSX = Path(
    r"C:\Users\tumsu\Downloads\UAT_GoldMints_Test Scenario_MANU_actual_flow_th_evidence_20260407.xlsx"
)
OUTPUT_DOCX = ROOT / "reports" / "manu_actual_flow_evidence_20260407.docx"
OUTPUT_JSON = ROOT / "reports" / "manu_actual_flow_evidence_20260407.json"
OUTPUT_MD = ROOT / "reports" / "manu_actual_flow_evidence_20260407.md"
IMAGE_DIR = ROOT / "reports" / "manu_actual_flow_uat_20260407_images"
THUMB_DIR = IMAGE_DIR / "thumbs"

LOGIN_URL = "http://localhost:8811/web/login"
USER = "admin"
PASSWORD = "365@gmp"

REUSE_MAP = {
    "MU01-01": ROOT / "reports" / "manu_uat_20260406_mu01_mu04_images" / "MU01_01_forecast_report.png",
    "MU01-02": ROOT / "reports" / "manu_uat_20260406_mu01_mu04_images" / "MU01_02_replenishment_fg_pss.png",
    "MU01-03": ROOT / "reports" / "manu_uat_20260406_mu01_mu04_images" / "MU01_03_mo_created_fg_pnc.png",
    "MU11-04": ROOT / "reports" / "mold_shopfloor_uat_20260406_images" / "02_mold_card.png",
    "MU14-01": ROOT / "reports" / "mold_shopfloor_uat_20260406_images" / "02_mold_card.png",
    "MU14-02": ROOT / "reports" / "mold_shopfloor_uat_20260406_images" / "06_change_mold_dialog.png",
    "MU14-03": ROOT / "reports" / "mold_shopfloor_uat_20260406_images" / "03_full_mold_warning.png",
    "MU14-04": ROOT / "reports" / "mold_shopfloor_uat_20260406_images" / "09_changed_mold_started.png",
    "MU14-05": ROOT / "reports" / "mold_shopfloor_uat_20260406_images" / "10_reset_life_result.png",
    "MU14-06": ROOT / "reports" / "mold_shopfloor_uat_20260406_images" / "11_more_menu_report_issue.png",
}

STATUS_MAP = {
    "ยืนยันแล้ว": "ยืนยันเชิงฟังก์ชัน",
    "ยืนยันแล้ว (มีหมายเหตุ)": "ยืนยันเชิงฟังก์ชัน (มีหมายเหตุ)",
    "มีฟังก์ชันจริง แต่ยังไม่ปิดเคส UAT": "เปิดเมนูและฟังก์ชันจริง แต่ยังไม่ปิดเคสเชิงธุรกรรม",
}


def load_cases() -> List[dict]:
    wb = load_workbook(WORKBOOK_PATH)
    cases = []
    for ws in wb.worksheets:
        if ws.title in ("00_สรุป", "99_Coverage_Backlog"):
            continue
        row = 7
        while row <= ws.max_row and ws.cell(row, 2).value:
            cases.append(
                {
                    "sheet": ws.title,
                    "row": row,
                    "case_id": ws.cell(row, 2).value,
                    "backlog_ids": ws.cell(row, 3).value or "",
                    "event": ws.cell(row, 4).value or "",
                    "scenario": ws.cell(row, 5).value or "",
                    "role": ws.cell(row, 6).value or "",
                    "menu_path": ws.cell(row, 7).value or "",
                    "precondition": ws.cell(row, 8).value or "",
                    "test_data": ws.cell(row, 9).value or "",
                    "steps": ws.cell(row, 10).value or "",
                    "expected": ws.cell(row, 11).value or "",
                    "fix": ws.cell(row, 12).value or "",
                    "base_status": ws.cell(row, 13).value or "",
                    "base_note": ws.cell(row, 14).value or "",
                }
            )
            row += 1
    return cases


def action_url(case: dict) -> str:
    menu = case["menu_path"]
    cid = case["case_id"]
    if "การเติมสินค้า" in menu:
        return "http://localhost:8811/odoo/action-697?menu_id=460"
    if "สินค้าคงคลัง > สินค้า > สินค้า" in menu:
        return "http://localhost:8811/odoo/action-688?menu_id=454"
    if "การขาย > คำสั่ง > ใบเสนอราคา" in menu:
        return "http://localhost:8811/odoo/action-358?menu_id=197"
    if "คำขอซื้อ > คำขอซื้อ > คำขอซื้อ" in menu:
        return "http://localhost:8811/odoo/action-1020?menu_id=625"
    if "สั่งซื้อ > คำสั่ง > ใบแจ้งขอใบเสนอราคา" in menu:
        return "http://localhost:8811/odoo/action-766?menu_id=514"
    if "สั่งซื้อ > คำสั่ง > คำสั่งซื้อ" in menu:
        return "http://localhost:8811/odoo/action-767?menu_id=515"
    if "สั่งซื้อ > สินค้า > สินค้า" in menu:
        return "http://localhost:8811/odoo/action-764?menu_id=512"
    if "การโอน" in menu:
        return "http://localhost:8811/odoo/action-660?menu_id=451"
    if "ใบสั่งผลิต" in menu:
        return "http://localhost:8811/odoo/action-821?menu_id=531"
    if "คำสั่งงาน" in menu and "การรายงาน" not in menu:
        return "http://localhost:8811/odoo/action-794?menu_id=532"
    if "ศูนย์งาน" in menu:
        return "http://localhost:8811/odoo/action-804?menu_id=528"
    if "คำสั่งรื้อ" in menu:
        return "http://localhost:8811/odoo/action-831?menu_id=538"
    if "บิลวัสดุ" in menu:
        return "http://localhost:8811/odoo/action-808?menu_id=530"
    if "ล็อต/หมายเลขซีเรียล" in menu:
        return "http://localhost:8811/odoo/action-643?menu_id=525"
    if "ประวัติการย้าย" in menu:
        return "http://localhost:8811/odoo/action-655?menu_id=447"
    if "วิเคราะห์การผลิต" in menu:
        return "http://localhost:8811/odoo/action-858?menu_id=546"
    if "Machine Report" in menu:
        return "http://localhost:8811/odoo/action-1647?menu_id=991"
    if "เอกสาร > เอกสาร" in menu:
        return "http://localhost:8811/odoo/action-916?menu_id=577"
    if "การแจ้งเตือนการจัดการคุณภาพ" in menu:
        return "http://localhost:8811/odoo/action-1064?menu_id=652"
    if "จุดควบคุม" in menu:
        return "http://localhost:8811/odoo/action-1077?menu_id=650"
    if "การตรวจสอบคุณภาพ" in menu:
        return "http://localhost:8811/odoo/action-1070?menu_id=651"
    if "ตำแหน่ง" in menu:
        return "http://localhost:8811/odoo/action-692?menu_id=458"
    if "MPS" in menu or "กำหนดการการผลิตหลัก" in menu:
        return "http://localhost:8811/odoo/action-1186?menu_id=747"
    if menu.startswith("GMP Shop Floor"):
        if cid == "MU11-04":
            return "http://localhost:8811/odoo/action-1643?menu_id=987"
        return "http://localhost:8811/odoo/action-1401?menu_id=880"
    if "Scrap" in menu or "เศษสินค้า" in menu:
        return "http://localhost:8811/odoo/action-645?menu_id=526"
    return "http://localhost:8811/odoo"


def create_thumbnail(src: Path, dst: Path, width: int = 320):
    dst.parent.mkdir(parents=True, exist_ok=True)
    with Image.open(src) as img:
        img = img.convert("RGB")
        ratio = width / img.width
        height = int(img.height * ratio)
        img = img.resize((width, height))
        img.save(dst, format="PNG")


def login(page):
    page.goto(LOGIN_URL, wait_until="domcontentloaded")
    page.get_by_text("uat", exact=True).click()
    page.wait_for_timeout(1500)
    page.fill('input[name="login"]', USER)
    page.fill('input[name="password"]', PASSWORD)
    page.click('button[type="submit"]')
    page.wait_for_timeout(4000)


def capture_cases(cases: List[dict]) -> Dict[str, dict]:
    IMAGE_DIR.mkdir(parents=True, exist_ok=True)
    THUMB_DIR.mkdir(parents=True, exist_ok=True)
    results: Dict[str, dict] = {}
    with sync_playwright() as p:
        browser = p.chromium.launch(headless=True)
        page = browser.new_page(viewport={"width": 1440, "height": 1024})
        login(page)
        for case in cases:
            case_id = case["case_id"]
            image_path = IMAGE_DIR / f"{case_id}.png"
            thumb_path = THUMB_DIR / f"{case_id}.png"
            try:
                if case_id in REUSE_MAP and REUSE_MAP[case_id].exists():
                    shutil.copyfile(REUSE_MAP[case_id], image_path)
                    create_thumbnail(image_path, thumb_path)
                    results[case_id] = {
                        "status": STATUS_MAP.get(case["base_status"], case["base_status"] or "ยืนยันเชิงฟังก์ชัน"),
                        "note": f"{case['base_note']} | ใช้ภาพหลักฐานจากการทดสอบจริงรอบก่อน",
                        "image": str(image_path),
                        "thumbnail": str(thumb_path),
                    }
                    continue

                url = action_url(case)
                page.goto(url, wait_until="domcontentloaded")
                page.wait_for_timeout(4500)
                # Try to dismiss possible tour/notification popups without failing the run.
                for selector in [
                    'button[aria-label="Close"]',
                    '.o_notification_close',
                    '.o-mail-DiscussSidebar button',
                ]:
                    try:
                        if page.locator(selector).count() > 0:
                            page.locator(selector).first.click(timeout=500)
                    except Exception:
                        pass
                page.screenshot(path=str(image_path), full_page=False)
                create_thumbnail(image_path, thumb_path)
                prior = STATUS_MAP.get(case["base_status"], case["base_status"])
                status = prior if prior else "เปิดเมนูและหน้าจอจริงสำเร็จ"
                note = case["base_note"] or ""
                if note:
                    note = note + " | "
                note += f"เปิดเมนูจริงจาก local UAT สำเร็จและเก็บภาพรอบนี้ที่ {image_path.name}"
                results[case_id] = {
                    "status": status,
                    "note": note,
                    "image": str(image_path),
                    "thumbnail": str(thumb_path),
                }
            except Exception as exc:
                results[case_id] = {
                    "status": "เปิดเมนูไม่ผ่าน",
                    "note": f"เปิดหน้าจอรอบนี้ไม่สำเร็จ: {exc}",
                    "image": "",
                    "thumbnail": "",
                }
        browser.close()
    return results


def update_workbook(cases: List[dict], results: Dict[str, dict]):
    wb = load_workbook(WORKBOOK_PATH)
    for ws in wb.worksheets:
        if ws.title in ("00_สรุป", "99_Coverage_Backlog"):
            continue
        ws.cell(6, 18).value = "ไฟล์ภาพหลักฐาน"
        ws.cell(6, 19).value = "ภาพประกอบ"
        ws.cell(6, 20).value = "สถานะการทดสอบรอบนี้"
        row = 7
        while row <= ws.max_row and ws.cell(row, 2).value:
            case_id = ws.cell(row, 2).value
            result = results.get(case_id, {})
            ws.cell(row, 13).value = result.get("status", ws.cell(row, 13).value)
            ws.cell(row, 14).value = result.get("note", ws.cell(row, 14).value)
            ws.cell(row, 18).value = result.get("image", "")
            ws.cell(row, 20).value = result.get("status", "")
            for col in (13, 14, 18, 20):
                ws.cell(row, col).alignment = Alignment(wrap_text=True, vertical="top")
                ws.cell(row, col).font = Font(name="Tahoma", size=10)
            thumb = result.get("thumbnail")
            if thumb and Path(thumb).exists():
                img = XLImage(thumb)
                img.width = 220
                img.height = 140
                anchor = f"S{row}"
                ws.add_image(img, anchor)
                ws.row_dimensions[row].height = 110
            row += 1
        ws.column_dimensions["R"].width = 48
        ws.column_dimensions["S"].width = 32
        ws.column_dimensions["T"].width = 24

    summary = wb["00_สรุป"]
    summary["B8"] = len(results)
    summary["A8"] = "จำนวนเคสที่มีภาพหลักฐาน local UAT แล้ว"
    summary["A9"] = "จำนวนเคสที่เปิดเมนู/หน้าจอจริงรอบนี้สำเร็จ"
    summary["B9"] = sum(1 for r in results.values() if r["status"] != "เปิดเมนูไม่ผ่าน")
    wb.save(OUTPUT_XLSX)


def build_docx(cases: List[dict], results: Dict[str, dict]):
    doc = Document()
    doc.add_heading("หลักฐาน UAT Manufacturing Actual Flow - local UAT", level=1)
    doc.add_paragraph("เอกสารนี้แสดงผลการทดสอบจริงและภาพประกอบต่อเคสของชุด actual flow ที่มีเมนูและฟังก์ชันจริงใน local UAT")

    current_sheet = None
    for case in cases:
        if case["sheet"] != current_sheet:
            current_sheet = case["sheet"]
            doc.add_heading(current_sheet, level=2)
        doc.add_heading(f"{case['case_id']} - {case['scenario']}", level=3)
        table = doc.add_table(rows=0, cols=2)
        table.style = "Table Grid"
        fields = [
            ("ลำดับเหตุการณ์", case["event"]),
            ("Backlog IDs", case["backlog_ids"]),
            ("บทบาท / หน่วยงาน", case["role"]),
            ("Menu Path", case["menu_path"]),
            ("สถานะรอบนี้", results.get(case["case_id"], {}).get("status", "")),
            ("หมายเหตุ", results.get(case["case_id"], {}).get("note", "")),
        ]
        for label, value in fields:
            cells = table.add_row().cells
            cells[0].text = label
            cells[1].text = value or ""
        img_path = results.get(case["case_id"], {}).get("image")
        if img_path and Path(img_path).exists():
            doc.add_paragraph("ภาพหลักฐาน")
            doc.add_picture(img_path, width=Inches(6.5))
        else:
            doc.add_paragraph("ไม่มีภาพหลักฐาน")
    doc.save(OUTPUT_DOCX)


def write_reports(cases: List[dict], results: Dict[str, dict]):
    payload = {
        "total_cases": len(cases),
        "captured_cases": sum(1 for r in results.values() if r.get("image")),
        "failed_cases": [cid for cid, r in results.items() if r["status"] == "เปิดเมนูไม่ผ่าน"],
        "cases": results,
    }
    OUTPUT_JSON.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")

    lines = [
        "# หลักฐาน UAT Manufacturing Actual Flow - local UAT",
        "",
        f"- จำนวนเคสทั้งหมด: `{len(cases)}`",
        f"- จำนวนเคสที่มีภาพหลักฐาน: `{payload['captured_cases']}`",
        f"- จำนวนเคสที่เปิดไม่ผ่าน: `{len(payload['failed_cases'])}`",
        "",
        "## เคสที่เปิดไม่ผ่าน",
    ]
    if payload["failed_cases"]:
        for cid in payload["failed_cases"]:
            lines.append(f"- `{cid}`: {results[cid]['note']}")
    else:
        lines.append("- ไม่มี")
    OUTPUT_MD.write_text("\n".join(lines), encoding="utf-8")


def main():
    cases = load_cases()
    results = capture_cases(cases)
    update_workbook(cases, results)
    build_docx(cases, results)
    write_reports(cases, results)
    print(f"Workbook: {OUTPUT_XLSX}")
    print(f"Word: {OUTPUT_DOCX}")
    print(f"Images: {IMAGE_DIR}")
    print(f"JSON: {OUTPUT_JSON}")


if __name__ == "__main__":
    main()

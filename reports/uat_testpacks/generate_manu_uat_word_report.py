from __future__ import annotations

import ast
import json
import re
from collections import defaultdict
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[2]
REPORTS = ROOT / "reports"
OUT_DOCX = REPORTS / "manu_uat_master_evidence_20260406.docx"

TOPICS_SCRIPT = ROOT / "reports" / "uat_testpacks" / "generate_manu_uat_teststeps.py"
NEW_TOPICS_SCRIPT = ROOT / "reports" / "uat_testpacks" / "enhance_manu_uat_teststeps_v2.py"

FULLFLOW_JSON = REPORTS / "db11_live_fullflow_20260330085840.json"
BACKORDER_JSON = REPORTS / "late_backorder_recovery_uat_test_20260403_final.json"
MOLD_VALIDATION_JSON = REPORTS / "db11_mold_validation_live_20260331022308.json"
MOLD_UI_JSON = REPORTS / "mold_shopfloor_uat_full_20260406.json"
MOLD_UI_CASE_JSON = REPORTS / "mold_shopfloor_uat_test_20260406.json"


def repair_text(value: Any) -> Any:
    if not isinstance(value, str):
        return value
    if any(ch in value for ch in ("à¸", "à¹")):
        try:
            return value.encode("latin1").decode("utf-8")
        except Exception:
            return value
    return value


def load_list_constant(path: Path, variable_name: str) -> list[dict[str, Any]]:
    text = path.read_text(encoding="utf-8", errors="replace")
    match = re.search(rf"^{variable_name}\s*=\s*(\[.*?\n\])", text, re.S | re.M)
    if not match:
        raise RuntimeError(f"Could not find {variable_name} in {path}")
    data = ast.literal_eval(match.group(1))
    return deep_repair(data)


def deep_repair(value: Any) -> Any:
    if isinstance(value, dict):
        return {repair_text(k): deep_repair(v) for k, v in value.items()}
    if isinstance(value, list):
        return [deep_repair(v) for v in value]
    return repair_text(value)


@dataclass
class CaseStatus:
    status: str
    test_env: str
    note: str
    evidence_files: list[str] = field(default_factory=list)


STATUS_MAP: dict[str, CaseStatus] = {
    "MU01-01": CaseStatus(
        status="Passed",
        test_env="UAT",
        note="ทดสอบจริงบน UAT: เปิด Product, ตรวจ On Hand และ Forecasted report ของ FG-PNC-TH-01001 ได้ครบตาม flow review",
        evidence_files=[
            "reports/manu_uat_20260406_mu01_mu04_images/MU01_01_inventory_overview.png",
            "reports/manu_uat_20260406_mu01_mu04_images/MU01_01_product_search.png",
            "reports/manu_uat_20260406_mu01_mu04_images/MU01_01_product_form.png",
            "reports/manu_uat_20260406_mu01_mu04_images/MU01_01_forecast_report.png",
        ],
    ),
    "MU01-02": CaseStatus(
        status="Passed with note",
        test_env="UAT",
        note="ทดสอบจริงบน UAT: เปิด Replenishment และ review route/product สำหรับ FG-PSS-TH-01005 ได้จริง แต่รอบนี้ไม่ได้กด trigger procurement ใหม่เพื่อหลีกเลี่ยงการสร้างเอกสารทับข้อมูลจริง",
        evidence_files=[
            "reports/manu_uat_20260406_mu01_mu04_images/MU01_02_replenishment_fg_pss.png",
            "reports/manu_uat_20260406_mu01_mu04_images/MU01_02_product_route_fg_pss.png",
        ],
    ),
    "MU01-03": CaseStatus(
        status="Passed with note",
        test_env="UAT",
        note="ทดสอบจริงบน UAT: trace เอกสารปลายทางฝั่ง Manufacturing Pharma ของ FG-PNC-TH-01001 ได้จริงจาก MO ตัวอย่างที่มีอยู่แล้วในระบบ แต่ไม่ได้ยิง replenishment ใหม่ในรอบนี้",
        evidence_files=[
            "reports/manu_uat_20260406_mu01_mu04_images/MU01_03_mo_created_fg_pnc.png",
        ],
    ),
    "MU04-03": CaseStatus(
        status="Passed",
        test_env="UAT",
        note="ทดสอบจริงบน UAT: Transfer Pharma partial แล้วเลือก No Backorder จากนั้นสร้าง late backorder กลับได้",
        evidence_files=[str(BACKORDER_JSON.relative_to(ROOT))],
    ),
    "MU05-01": CaseStatus(
        status="Passed",
        test_env="DB11",
        note="ทดสอบจริงจาก sales flow แบบ MTO แล้ว chain ไป MO, delivery, invoice จบ",
        evidence_files=[str(FULLFLOW_JSON.relative_to(ROOT))],
    ),
    "MU06-01": CaseStatus(
        status="Passed",
        test_env="DB11",
        note="ทดสอบจริงจาก Min/Max / orderpoint แล้วสร้าง MO เติม stock ได้",
        evidence_files=[str(FULLFLOW_JSON.relative_to(ROOT))],
    ),
    "MU07-01": CaseStatus(
        status="Passed",
        test_env="DB11",
        note="มี MO done จริงหลายชุดใน full flow และ mold validation",
        evidence_files=[str(FULLFLOW_JSON.relative_to(ROOT))],
    ),
    "MU07-03": CaseStatus(
        status="Passed",
        test_env="UAT",
        note="ทดสอบจริงบน UAT: Manufacturing partial แล้วกด No Backorder จากนั้นสร้าง late backorder กลับได้",
        evidence_files=[str(BACKORDER_JSON.relative_to(ROOT))],
    ),
    "MU08-01": CaseStatus(
        status="Passed",
        test_env="DB11",
        note="มี scrap จริงจากงานผลิตและ trace เอกสาร scrap ได้",
        evidence_files=[str(FULLFLOW_JSON.relative_to(ROOT))],
    ),
    "MU09-01": CaseStatus(
        status="Passed",
        test_env="DB11",
        note="ทดสอบจริง: ระบบ match mold/workcenter ให้อัตโนมัติในหลายเคส plastic production",
        evidence_files=[str(MOLD_VALIDATION_JSON.relative_to(ROOT))],
    ),
    "MU09-02": CaseStatus(
        status="Passed",
        test_env="DB11",
        note="ทดสอบจริง: parallel workorder ไม่ใช้ mold ซ้ำกันแล้ว โดย sibling ที่เกินถูก cancel",
        evidence_files=[str(MOLD_VALIDATION_JSON.relative_to(ROOT))],
    ),
    "MU11-04": CaseStatus(
        status="Passed",
        test_env="UAT",
        note="ทดสอบจริงบน GMP Shop Floor ฝั่ง Plastic ที่มี mold/workcenter พร้อมภาพ UI จริง",
        evidence_files=[str(MOLD_UI_CASE_JSON.relative_to(ROOT))],
    ),
    "MU14-01": CaseStatus(
        status="Passed",
        test_env="UAT",
        note="Card แสดง mold name และ shot counter",
        evidence_files=[
            str(MOLD_UI_CASE_JSON.relative_to(ROOT)),
            "reports/mold_shopfloor_uat_20260406_images/02_mold_card.png",
        ],
    ),
    "MU14-02": CaseStatus(
        status="Passed",
        test_env="UAT",
        note="เปลี่ยน mold จาก card ได้จริงและ backend update",
        evidence_files=[
            str(MOLD_UI_CASE_JSON.relative_to(ROOT)),
            "reports/mold_shopfloor_uat_20260406_images/06_change_mold_dialog.png",
            "reports/mold_shopfloor_uat_20260406_images/07_change_mold_result.png",
        ],
    ),
    "MU14-03": CaseStatus(
        status="Passed",
        test_env="UAT",
        note="mold เต็มแล้วระบบขึ้น warning และกด Continue Anyway ได้",
        evidence_files=[
            str(MOLD_UI_CASE_JSON.relative_to(ROOT)),
            "reports/mold_shopfloor_uat_20260406_images/03_full_mold_warning.png",
            "reports/mold_shopfloor_uat_20260406_images/04_continue_anyway_started.png",
        ],
    ),
    "MU14-04": CaseStatus(
        status="Passed",
        test_env="UAT",
        note="เปลี่ยน mold จาก warning flow แล้ว start งานต่อได้",
        evidence_files=[
            str(MOLD_UI_CASE_JSON.relative_to(ROOT)),
            "reports/mold_shopfloor_uat_20260406_images/08_change_from_warning_result.png",
            "reports/mold_shopfloor_uat_20260406_images/09_changed_mold_started.png",
        ],
    ),
    "MU14-05": CaseStatus(
        status="Passed",
        test_env="UAT",
        note="Reset mold life ใช้งานได้จริง",
        evidence_files=[
            str(MOLD_UI_CASE_JSON.relative_to(ROOT)),
            "reports/mold_shopfloor_uat_20260406_images/10_reset_life_result.png",
        ],
    ),
    "MU14-06": CaseStatus(
        status="Passed with note",
        test_env="UAT",
        note="breakdown recovery path ใช้งานได้ผ่าน More > Report Issue แต่ยังไม่มีปุ่ม breakdown เฉพาะบน card",
        evidence_files=[
            str(MOLD_UI_CASE_JSON.relative_to(ROOT)),
            "reports/mold_shopfloor_uat_20260406_images/11_more_menu_report_issue.png",
        ],
    ),
    "MU14-07": CaseStatus(
        status="Not closed",
        test_env="UAT",
        note="ยังไม่มี sample ที่เหมาะสมพอใน UAT สำหรับ sign-off done/cancel sibling behavior end-to-end ผ่าน UI",
        evidence_files=[str(MOLD_UI_CASE_JSON.relative_to(ROOT))],
    ),
}


def default_status_for_case(case_id: str) -> CaseStatus:
    return STATUS_MAP.get(
        case_id,
        CaseStatus(
            status="Not closed",
            test_env="-",
            note="ยังไม่มีหลักฐานการ rerun แบบ step-by-step พร้อมภาพจริงในชุด evidence ปัจจุบัน",
            evidence_files=[],
        ),
    )


def load_all_topics() -> list[dict[str, Any]]:
    topics = load_list_constant(TOPICS_SCRIPT, "TOPICS")
    new_topics = load_list_constant(NEW_TOPICS_SCRIPT, "NEW_TOPICS")
    mu14 = {
        "sheet": "MU14",
        "title": "Mold / GMP Shop Floor / Recovery",
        "objective": "ทดสอบการแสดงผล mold บน card, การเปลี่ยน mold, warning เมื่อ mold เต็ม, reset life และ breakdown path",
        "coverage": "shop floor mold UI, full mold warning, continue anyway, change mold, reset life, breakdown, queue behavior",
        "cases": [
            {"case_id": "MU14-01", "backlogs": "MA11, MA14", "name": "Card shows mold name and shot counter", "role": "Operator / Supervisor", "path": "Manufacturing > GMP Shop Floor", "pre": "มี WO ฝั่ง plastic ที่มี mold mapping", "data": "MO GMP/MOPL/00012 / WO 364", "steps": "1) เปิด GMP Shop Floor\n2) เลือกการ์ด WO ฝั่ง plastic\n3) ตรวจชื่อ mold และจำนวน shots ที่แสดงบน card", "expected": "card แสดง mold และ shot counter ได้", "fix": "ถ้าไม่ขึ้นให้ตรวจ workcenter, mold mapping และ flag show mold UI", "evidence": "ภาพ card", "note": "อิงผลเทสจริง UAT"},
            {"case_id": "MU14-02", "backlogs": "MA11, MA14", "name": "Change mold from card before start", "role": "Operator", "path": "GMP Shop Floor > Mold button", "pre": "มี mold สำรองที่ compatible", "data": "WO 364, mold 650", "steps": "1) เปิด card\n2) กดปุ่ม Mold\n3) เลือก mold สำรอง\n4) ตรวจ card หลังเปลี่ยน", "expected": "card อัปเดต mold ใหม่และ backend เปลี่ยนจริง", "fix": "ถ้าเลือกไม่ได้ให้ตรวจ compatible molds", "evidence": "ภาพ dialog + card หลังเปลี่ยน", "note": "อิงผลเทสจริง UAT"},
            {"case_id": "MU14-03", "backlogs": "MA11, MA14", "name": "Full mold warning with Continue Anyway", "role": "Operator", "path": "GMP Shop Floor > Start", "pre": "mold life เต็ม", "data": "MO GMP/MOPL/00011 / WO 363 / mold 590", "steps": "1) เปิด WO ที่ mold life เต็ม\n2) กด Start\n3) ตรวจ popup warning\n4) กด Continue Anyway", "expected": "ระบบเตือนแต่ไม่ block และเริ่มงานได้", "fix": "ถ้าไม่เตือนให้ตรวจ mold life limit/current", "evidence": "ภาพ warning + started", "note": "อิงผลเทสจริง UAT"},
            {"case_id": "MU14-04", "backlogs": "MA11, MA14", "name": "Full mold warning with Change Mold then Start", "role": "Operator", "path": "GMP Shop Floor > Start > Change Mold", "pre": "mold life เต็มและมี mold สำรอง", "data": "WO 364", "steps": "1) กด Start\n2) ใน popup เลือก Change Mold\n3) เลือก mold สำรอง\n4) กด Start ต่อ", "expected": "เปลี่ยน mold แล้วเริ่มงานต่อได้", "fix": "ถ้า warning ยังวนซ้ำให้ตรวจ mold current/limit ของ mold ใหม่", "evidence": "ภาพ warning flow + started", "note": "อิงผลเทสจริง UAT"},
            {"case_id": "MU14-05", "backlogs": "MA11", "name": "Reset mold life after maintenance", "role": "Supervisor / Maintenance", "path": "Manufacturing > Configuration > Work Centers", "pre": "mold มี current shots ไม่เป็นศูนย์", "data": "mold 590", "steps": "1) เปิด form ของ mold\n2) กด Reset Life\n3) ยืนยันการ reset\n4) ตรวจ current shots", "expected": "shots ถูก reset เป็น 0", "fix": "ถ้า reset ไม่ได้ให้ตรวจสิทธิ์และปุ่ม action", "evidence": "ภาพ form หลัง reset", "note": "อิงผลเทสจริง UAT"},
            {"case_id": "MU14-06", "backlogs": "MA14", "name": "Breakdown recovery path", "role": "Operator / Supervisor", "path": "GMP Shop Floor > More > Report Issue", "pre": "มีงานที่กำลังทำบน shop floor", "data": "More menu", "steps": "1) เปิด card งาน\n2) เปิดเมนู More\n3) กด Report Issue\n4) ใช้เป็น entry point สำหรับ breakdown handling", "expected": "มีทางเข้ารายงานปัญหาและใช้หยุด/recover งานต่อได้", "fix": "ถ้าเมนูไม่ขึ้นให้ตรวจ custom menu/action", "evidence": "ภาพ More menu", "note": "ไม่มีปุ่ม breakdown แยกเฉพาะ"},
            {"case_id": "MU14-07", "backlogs": "MA14", "name": "Done/Cancel workorders do not reappear as active queue items", "role": "Supervisor / IT", "path": "GMP Shop Floor", "pre": "ต้องมี sample sibling done/cancel บน queue เดียวกัน", "data": "parallel sibling WO", "steps": "1) เตรียม sample queue\n2) ตรวจว่า done/cancel ไม่กลับมา active", "expected": "queue แสดงเฉพาะงานที่ active", "fix": "ถ้ายังเห็นงาน done/cancel ให้ตรวจ queue filter/guard", "evidence": "ต้องเก็บภาพเมื่อมี sample", "note": "ยังไม่ปิดเคส"},
        ],
    }
    return topics + new_topics + [mu14]


def set_cell_shading(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    tc_pr.append(shd)


def style_document(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Aptos"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    style.font.size = Pt(10)
    for name in ("Heading 1", "Heading 2", "Heading 3"):
        s = doc.styles[name]
        s.font.name = "Aptos"
        s._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")


def add_title(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Manufacturing UAT Master Evidence Report")
    r.bold = True
    r.font.size = Pt(20)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.add_run("สรุปทุก MU/case จาก source workbook เดิม พร้อมสถานะจากหลักฐานทดสอบจริงที่มีอยู่ในปัจจุบัน")
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.add_run("Environment: UAT / DB11 | Date: 2026-04-06")
    doc.add_paragraph()
    note = doc.add_paragraph()
    note.add_run("หมายเหตุสำคัญ: ").bold = True
    note.add_run(
        "รายงานฉบับนี้ยึดเฉพาะหลักฐานที่มีอยู่จริงในระบบและไฟล์รายงานปัจจุบันเท่านั้น "
        "เคสที่ยังไม่มี rerun พร้อมภาพจริงจะถูกระบุเป็น Not closed อย่างชัดเจน ไม่ overclaim."
    )


def add_summary_tables(doc: Document, topics: list[dict[str, Any]]) -> None:
    doc.add_heading("1. Coverage Summary", level=1)
    total_cases = sum(len(t["cases"]) for t in topics)
    status_counts: dict[str, int] = defaultdict(int)
    for topic in topics:
        for case in topic["cases"]:
            status_counts[default_status_for_case(case["case_id"]).status] += 1

    p = doc.add_paragraph()
    p.add_run(f"Total MU groups: {len(topics)}\n").bold = True
    p.add_run(f"Total cases: {total_cases}\n")
    for status in ("Passed", "Passed with note", "Not closed"):
        p.add_run(f"{status}: {status_counts.get(status, 0)}\n")

    doc.add_paragraph("Summary by MU")
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for idx, text in enumerate(["MU", "Title", "Cases", "Passed", "Not closed / pending"]):
        hdr[idx].text = text
        set_cell_shading(hdr[idx], "1F4E78")
        for para in hdr[idx].paragraphs:
            for run in para.runs:
                run.font.color.rgb = None
                run.bold = True
    for topic in topics:
        passed = 0
        not_closed = 0
        for case in topic["cases"]:
            if default_status_for_case(case["case_id"]).status == "Passed":
                passed += 1
            else:
                not_closed += 1
        row = table.add_row().cells
        row[0].text = topic["sheet"]
        row[1].text = topic["title"]
        row[2].text = str(len(topic["cases"]))
        row[3].text = str(passed)
        row[4].text = str(not_closed)
        for cell in row:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def add_case_table(doc: Document, case: dict[str, Any], status: CaseStatus) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    fields = [
        ("Case ID", case["case_id"]),
        ("Scenario", case["name"]),
        ("Role", case.get("role", "-")),
        ("Menu Path", case.get("path", "-")),
        ("Pre-condition", case.get("pre", "-")),
        ("Test Data", case.get("data", "-")),
        ("Expected Result", case.get("expected", "-")),
        ("Fix / Recovery", case.get("fix", "-")),
        ("Current Status", f"{status.status} ({status.test_env})"),
        ("Evidence Note", status.note),
        ("Evidence Files", "\n".join(status.evidence_files) if status.evidence_files else "-"),
    ]
    for label, value in fields:
        row = table.add_row().cells
        row[0].text = str(label)
        row[1].text = str(value)
        set_cell_shading(row[0], "D9EAD3")
        row[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        row[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def add_steps(doc: Document, steps_text: str) -> None:
    doc.add_paragraph("Detailed Test Steps", style="Heading 3")
    normalized = str(steps_text).replace("\\n", "\n")
    for line in normalized.splitlines():
        line = line.strip()
        if not line:
            continue
        if re.match(r"^\d+\)", line):
            doc.add_paragraph(line, style="List Number")
        else:
            doc.add_paragraph(line)


def add_images(doc: Document, evidence_files: list[str]) -> None:
    image_paths = [ROOT / path for path in evidence_files if path.lower().endswith(".png")]
    if not image_paths:
        return
    doc.add_paragraph("Screenshot Evidence", style="Heading 3")
    for img_path in image_paths:
        if not img_path.exists():
            continue
        doc.add_paragraph(img_path.name)
        doc.add_picture(str(img_path), width=Cm(15.5))
        doc.add_paragraph()


def build_document(topics: list[dict[str, Any]]) -> None:
    doc = Document()
    style_document(doc)
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    add_title(doc)
    add_summary_tables(doc, topics)

    for topic in topics:
        doc.add_page_break()
        doc.add_heading(f"{topic['sheet']} - {topic['title']}", level=1)
        intro = doc.add_paragraph()
        intro.add_run("Objective: ").bold = True
        intro.add_run(topic.get("objective", "-"))
        intro.add_run("\nCoverage: ").bold = True
        intro.add_run(topic.get("coverage", "-"))

        for case in topic["cases"]:
            status = default_status_for_case(case["case_id"])
            doc.add_heading(f"{case['case_id']} - {case['name']}", level=2)
            add_case_table(doc, case, status)
            add_steps(doc, case.get("steps", "-"))
            add_images(doc, status.evidence_files)
            doc.add_paragraph()

    doc.save(OUT_DOCX)


def main() -> None:
    topics = load_all_topics()
    build_document(topics)
    print(f"Saved: {OUT_DOCX}")


if __name__ == "__main__":
    main()

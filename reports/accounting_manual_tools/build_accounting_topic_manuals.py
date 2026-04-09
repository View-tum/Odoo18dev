from __future__ import annotations

import json
import subprocess
from dataclasses import dataclass, field
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\365_project\TheCool18e\Dev")
TOOLS_DIR = ROOT / "reports" / "accounting_manual_tools"
OUTPUT_DIR = ROOT / "manual" / "Accouting_Manual" / "generated_20260408"
DOCX_DIR = OUTPUT_DIR / "docx"
IMAGE_DIR = OUTPUT_DIR / "images"
CFG_DIR = OUTPUT_DIR / "capture_configs"
PROBE_DIR = TOOLS_DIR / "output" / "probe"
TMP_DIR = TOOLS_DIR / "output" / "tmp_test"
SAMPLES_PATH = TOOLS_DIR / "output" / "manual_live_samples_20260408.json"

BASE_URL = "http://localhost:8811"
DB = "uat"
LOGIN = "admin"
PASSWORD = "365@gmp"

VENDOR_BILLS = {
    "out_confirmed": {"id": 68481, "name": "APD/26/04/00006"},
    "out_paid": {"id": 68482, "name": "APD/26/04/00007"},
    "out_void": {"id": 68483, "name": "APD/26/04/00008"},
}

CUSTOMER_INVOICES = {
    "in_confirmed": {"id": 68484, "name": "INV-E/26/04/00003"},
    "in_paid": {"id": 68485, "name": "INV-E/26/04/00004"},
}


@dataclass
class CaptureSpec:
    key: str
    filename: str
    target_url: str | None = None
    actions: list[dict] = field(default_factory=list)
    highlight_selectors: list[dict] = field(default_factory=list)
    fallback_meta: Path | None = None
    post_nav_wait_ms: int = 3000
    post_click_wait_ms: int = 2000


@dataclass
class TopicSpec:
    code: str
    title: str
    subtitle: str
    menu_paths: list[str]
    objectives: list[str]
    overview: list[str]
    functions_table: list[tuple[str, str, str]]
    steps: list[tuple[str, list[str]]]
    images: list[tuple[str, str]]
    field_tables: list[tuple[str, list[tuple[str, str, str]]]]
    scenarios: list[tuple[str, str, str]]
    journal_tables: list[tuple[str, list[tuple[str, str, float, float]]]]
    cautions: list[str]
    scope_note: str | None = None


def ensure_dirs() -> None:
    for path in (OUTPUT_DIR, DOCX_DIR, IMAGE_DIR, CFG_DIR):
        path.mkdir(parents=True, exist_ok=True)


def load_samples() -> dict:
    return json.loads(SAMPLES_PATH.read_text(encoding="utf-8"))


def set_cell_text(cell, text: str) -> None:
    cell.text = str(text)
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = "TH Sarabun New"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "TH Sarabun New")
            run.font.size = Pt(14)


def shade_cell(cell, color: str = "D9EAF7") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    tc_pr.append(shd)


def add_title(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(24)
    run.font.name = "TH Sarabun New"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "TH Sarabun New")
    run.font.color.rgb = RGBColor(31, 78, 121)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(subtitle)
    run2.font.size = Pt(18)
    run2.font.name = "TH Sarabun New"
    run2._element.rPr.rFonts.set(qn("w:eastAsia"), "TH Sarabun New")


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.name = "TH Sarabun New"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "TH Sarabun New")
    run.font.size = Pt(18 if level == 1 else 16)
    if level == 1:
        run.font.color.rgb = RGBColor(31, 78, 121)


def add_paragraph(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "TH Sarabun New"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "TH Sarabun New")
    run.font.size = Pt(16)


def add_bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        run.font.name = "TH Sarabun New"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "TH Sarabun New")
        run.font.size = Pt(16)


def add_numbered_steps(doc: Document, heading: str, items: list[str]) -> None:
    add_paragraph(doc, heading, bold=True)
    for idx, item in enumerate(items, start=1):
        p = doc.add_paragraph()
        run = p.add_run(f"{idx}. {item}")
        run.font.name = "TH Sarabun New"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "TH Sarabun New")
        run.font.size = Pt(16)


def add_table(doc: Document, headers: list[str], rows: list[tuple]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        set_cell_text(hdr_cells[idx], header)
        shade_cell(hdr_cells[idx])
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)


def add_image(doc: Document, image_path: Path, caption: str) -> None:
    if not image_path.exists():
        add_paragraph(doc, f"[ไม่พบภาพประกอบ: {image_path.name}]")
        return
    doc.add_picture(str(image_path), width=Inches(6.6))
    p = doc.paragraphs[-1]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    run.italic = True
    run.font.name = "TH Sarabun New"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "TH Sarabun New")
    run.font.size = Pt(14)


def configure_doc(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "TH Sarabun New"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "TH Sarabun New")
    style.font.size = Pt(16)
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)


def build_capture_config(spec: CaptureSpec) -> Path:
    cfg = {
        "base_url": BASE_URL,
        "db": DB,
        "login": LOGIN,
        "password": PASSWORD,
        "output_dir": str(IMAGE_DIR),
        "filename": spec.filename,
        "post_nav_wait_ms": spec.post_nav_wait_ms,
        "post_click_wait_ms": spec.post_click_wait_ms,
    }
    if spec.target_url:
        cfg["target_url"] = spec.target_url
    if spec.actions:
        cfg["actions"] = spec.actions
    if spec.highlight_selectors:
        cfg["highlight_selectors"] = spec.highlight_selectors
    path = CFG_DIR / f"{spec.key}.json"
    path.write_text(json.dumps(cfg, ensure_ascii=False, indent=2), encoding="utf-8")
    return path


def generate_capture(spec: CaptureSpec) -> Path:
    cfg_path = build_capture_config(spec)
    meta_path = IMAGE_DIR / f"{spec.filename}.json"
    try:
        subprocess.run(
            ["node", str(TOOLS_DIR / "capture_odoo_page.js"), str(cfg_path)],
            check=True,
            cwd=str(ROOT),
        )
        subprocess.run(
            ["python", str(TOOLS_DIR / "annotate_capture.py"), str(meta_path)],
            check=True,
            cwd=str(ROOT),
        )
        annotated = IMAGE_DIR / f"{Path(spec.filename).stem}_annotated.png"
        if annotated.exists():
            return annotated
    except subprocess.CalledProcessError:
        pass
    if spec.fallback_meta and spec.fallback_meta.exists():
        subprocess.run(
            ["python", str(TOOLS_DIR / "annotate_capture.py"), str(spec.fallback_meta)],
            check=True,
            cwd=str(ROOT),
        )
        return spec.fallback_meta.with_name(spec.fallback_meta.stem + "_annotated.png")
    raise FileNotFoundError(f"Unable to produce image for {spec.key}")


def build_capture_specs(samples: dict) -> dict[str, CaptureSpec]:
    return {
        "group_draft": CaptureSpec(
            key="group_draft",
            filename="group_draft_manual.png",
            target_url=f"{BASE_URL}/odoo/action-1528/{samples['group_payment']['draft_id']}",
            highlight_selectors=[
                {"selector": "[name='member_partner_ids']", "label": "1"},
                {"selector": "button[name='action_search_moves']", "label": "2"},
                {"selector": "[name='line_ids']", "label": "3"},
                {"selector": "button[name='action_confirm_payment']", "label": "4"},
            ],
            fallback_meta=TMP_DIR / "group_draft_test.png.json",
        ),
        "group_wizard": CaptureSpec(
            key="group_wizard",
            filename="group_payment_wizard_manual.png",
            target_url=f"{BASE_URL}/odoo/action-1528/{samples['group_payment']['draft_id']}",
            actions=[
                {"type": "click", "selector": "button[name='action_confirm_payment']", "wait_ms": 2500},
            ],
            highlight_selectors=[
                {"selector": "[name='journal_id']", "label": "1"},
                {"selector": "[name='payment_method_line_id']", "label": "2"},
                {"selector": ".modal-footer .btn-primary", "label": "3"},
            ],
            fallback_meta=PROBE_DIR / "group_payment_wizard.png.json",
        ),
        "group_done": CaptureSpec(
            key="group_done",
            filename="group_done_manual.png",
            target_url=f"{BASE_URL}/odoo/action-1528/{samples['group_payment']['done_id']}",
            highlight_selectors=[
                {"selector": "button[name='action_print_receipt']", "label": "1"},
                {"selector": "button[name='action_view_payments']", "label": "2"},
                {"selector": "[name='line_ids']", "label": "3"},
            ],
            fallback_meta=PROBE_DIR / "group_done_form.png.json",
        ),
        "settings": CaptureSpec(
            key="settings",
            filename="settings_cheque_manual.png",
            target_url=f"{BASE_URL}/odoo/action-1361",
            highlight_selectors=[
                {"selector": "#is_reverse_cheque_entry", "label": "1"},
            ],
            fallback_meta=PROBE_DIR / "settings_direct_wait.png.json",
        ),
        "journal_general": CaptureSpec(
            key="journal_general",
            filename="journal_general_manual.png",
            target_url=f"{BASE_URL}/odoo/action-278/{samples['bank_journal']['id']}",
            highlight_selectors=[
                {"selector": "[name='is_cheque_incoming']", "label": "1"},
                {"selector": "[name='is_cheque_outgoing']", "label": "2"},
                {"selector": "[name='dynamic_cheque_id']", "label": "3"},
            ],
            fallback_meta=PROBE_DIR / "journal_form.png.json",
        ),
        "journal_incoming": CaptureSpec(
            key="journal_incoming",
            filename="journal_incoming_manual.png",
            target_url=f"{BASE_URL}/odoo/action-278/{samples['bank_journal']['id']}",
            highlight_selectors=[
                {"selector": "[name='is_cheque_incoming']", "label": "1"},
                {"selector": "[name='inbound_payment_method_line_ids']", "label": "2"},
            ],
            fallback_meta=PROBE_DIR / "journal_incoming.png.json",
        ),
        "journal_outgoing": CaptureSpec(
            key="journal_outgoing",
            filename="journal_outgoing_manual.png",
            target_url=f"{BASE_URL}/odoo/action-278/{samples['bank_journal']['id']}",
            highlight_selectors=[
                {"selector": "[name='is_cheque_outgoing']", "label": "1"},
                {"selector": "[name='outbound_payment_method_line_ids']", "label": "2"},
            ],
            fallback_meta=PROBE_DIR / "journal_outgoing.png.json",
        ),
        "template": CaptureSpec(
            key="template",
            filename="template_manual.png",
            target_url=f"{BASE_URL}/odoo/action-1367/{samples['template']['id']}",
            highlight_selectors=[
                {"selector": "[name='name']", "label": "1"},
                {"selector": "[name='cheque_hight']", "label": "2"},
                {"selector": "[name='cheque_width']", "label": "3"},
                {"selector": "[name='top_margin']", "label": "4"},
            ],
            fallback_meta=PROBE_DIR / "template_form.png.json",
        ),
        "book": CaptureSpec(
            key="book",
            filename="cheque_book_manual.png",
            target_url=f"{BASE_URL}/odoo/action-1362/{samples['cheque_book']['id']}",
            highlight_selectors=[
                {"selector": "[name='bank_account_journal_id']", "label": "1"},
                {"selector": "[name='first_cheque_no_char']", "label": "2"},
                {"selector": "[name='cheque_book_lines']", "label": "3"},
            ],
            fallback_meta=PROBE_DIR / "book_form.png.json",
        ),
        "bill_register": CaptureSpec(
            key="bill_register",
            filename="bill_register_payment_manual.png",
            target_url=f"{BASE_URL}/odoo/vendor-bills/{VENDOR_BILLS['out_confirmed']['id']}",
            actions=[
                {"type": "click", "selector": "button[name='action_register_payment']", "wait_ms": 2500},
            ],
            highlight_selectors=[
                {"selector": "[name='journal_id']", "label": "1"},
                {"selector": "[name='payment_method_line_id']", "label": "2"},
                {"selector": "[name='wizard_outbound_cheque_lines']", "label": "3"},
                {"selector": ".modal-footer .btn-primary", "label": "4"},
            ],
            fallback_meta=TMP_DIR / "bill_register_payment_test.png.json",
        ),
        "invoice_register": CaptureSpec(
            key="invoice_register",
            filename="invoice_register_payment_manual.png",
            target_url=f"{BASE_URL}/odoo/customer-invoice/{CUSTOMER_INVOICES['in_confirmed']['id']}",
            actions=[
                {"type": "click", "selector": "button[name='action_register_payment']", "wait_ms": 2500},
            ],
            highlight_selectors=[
                {"selector": "[name='journal_id']", "label": "1"},
                {"selector": "[name='payment_method_line_id']", "label": "2"},
                {"selector": "[name='wizard_inbound_cheque_lines']", "label": "3"},
                {"selector": ".modal-footer .btn-primary", "label": "4"},
            ],
            fallback_meta=TMP_DIR / "invoice_register_payment_test.png.json",
        ),
        "cheque_out_confirmed": CaptureSpec(
            key="cheque_out_confirmed",
            filename="cheque_out_confirmed_manual.png",
            target_url=f"{BASE_URL}/odoo/action-1372/{samples['cheques']['outbound_confirmed']['id']}",
            highlight_selectors=[
                {"selector": "[name='pay_partner_id']", "label": "1"},
                {"selector": "[name='amount']", "label": "2"},
                {"selector": "button[name='action_bank_deposit']", "label": "3"},
                {"selector": "button[name='action_view_payment_entries']", "label": "4"},
            ],
            fallback_meta=PROBE_DIR / "cheque_out_cf_form.png.json",
        ),
        "cheque_out_paid": CaptureSpec(
            key="cheque_out_paid",
            filename="cheque_out_paid_manual.png",
            target_url=f"{BASE_URL}/odoo/action-1372/{samples['cheques']['outbound_paid']['id']}",
            highlight_selectors=[
                {"selector": "button[name='action_view_payment_cheque']", "label": "1"},
                {"selector": "button[name='action_view_payment_entries']", "label": "2"},
                {"selector": "[name='clearing_date']", "label": "3"},
            ],
            fallback_meta=PROBE_DIR / "cheque_out_pd_form.png.json",
        ),
        "cheque_out_void": CaptureSpec(
            key="cheque_out_void",
            filename="cheque_out_void_manual.png",
            target_url=f"{BASE_URL}/odoo/action-1372/{samples['cheques']['outbound_void']['id']}",
            highlight_selectors=[
                {"selector": "button[name='action_view_reversed_entries']", "label": "1"},
                {"selector": "[name='void_reason']", "label": "2"},
                {"selector": "button[name='action_reset_to_draft']", "label": "3"},
            ],
            fallback_meta=PROBE_DIR / "cheque_out_void_form.png.json",
        ),
        "cheque_in_confirmed": CaptureSpec(
            key="cheque_in_confirmed",
            filename="cheque_in_confirmed_manual.png",
            target_url=f"{BASE_URL}/odoo/action-1373/{samples['cheques']['inbound_confirmed']['id']}",
            highlight_selectors=[
                {"selector": "[name='pay_partner_id']", "label": "1"},
                {"selector": "[name='amount']", "label": "2"},
                {"selector": "button[name='action_bank_deposit']", "label": "3"},
                {"selector": "button[name='print_dynamic_cheque_report']", "label": "4"},
            ],
            fallback_meta=PROBE_DIR / "cheque_in_cf_form.png.json",
        ),
        "cheque_in_paid": CaptureSpec(
            key="cheque_in_paid",
            filename="cheque_in_paid_manual.png",
            target_url=f"{BASE_URL}/odoo/action-1373/{samples['cheques']['inbound_paid']['id']}",
            highlight_selectors=[
                {"selector": "button[name='action_view_payment_cheque']", "label": "1"},
                {"selector": "button[name='action_view_payment_entries']", "label": "2"},
                {"selector": "[name='clearing_date']", "label": "3"},
            ],
            fallback_meta=PROBE_DIR / "cheque_in_pd_form.png.json",
        ),
    }


def cheque_journal_rows(samples: dict) -> list[tuple[str, str, str]]:
    return [
        ("Cheque Incoming", "เปิดให้ Journal นี้รองรับการรับเช็คจากลูกค้า", "เปิดใช้งานที่ PBAY1"),
        ("Cheque Outgoing", "เปิดให้ Journal นี้รองรับการจ่ายเช็คให้ผู้ขาย", "เปิดใช้งานที่ PBAY1"),
        ("Cheque Form", "เลือกเทมเพลตฟอร์มเช็คที่ใช้พิมพ์เช็คขาออก", samples["template"]["name"]),
        ("Incoming Payment Line", "กำหนดวิธีรับชำระและบัญชีพักเช็ครับ", f"Line ID {samples['bank_journal']['cheque_in_line_id']}"),
        ("Outgoing Payment Line", "กำหนดวิธีจ่ายเช็คและบัญชีพักเช็คจ่าย", f"Line ID {samples['bank_journal']['cheque_out_line_id']}"),
    ]


def move_lines_to_rows(lines: list[dict]) -> list[tuple[str, str, float, float]]:
    rows: list[tuple[str, str, float, float]] = []
    for line in lines:
        rows.append(
            (
                f"{line['account_code']} {line['account_name']}",
                line["label"],
                float(line["debit"]),
                float(line["credit"]),
            )
        )
    return rows


def build_topics(samples: dict, images: dict[str, Path]) -> list[TopicSpec]:
    out_cf = samples["cheques"]["outbound_confirmed"]
    out_pd = samples["cheques"]["outbound_paid"]
    out_void = samples["cheques"]["outbound_void"]
    in_cf = samples["cheques"]["inbound_confirmed"]
    in_pd = samples["cheques"]["inbound_paid"]

    group_rows = [
        ("member_partner_ids", "รายชื่อคู่ค้าภายใต้กลุ่มบริษัทที่ต้องการดึงเอกสารคงค้าง", "เลือกคู่ค้าที่อยู่ในกลุ่มเดียวกันเท่านั้น"),
        ("summary_info", "สรุปยอดคงค้างรวมของคู่ค้าที่เลือก", "ใช้เช็กก่อนเลือกเอกสาร"),
        ("filter_date_from / filter_date_to", "กรองช่วงวันที่เอกสารที่จะดึงเข้ามา", "ช่วยตัดรายการเก่าหรืออนาคตออก"),
        ("filter_min_amount", "กรองรายการที่ยอดค้างต่ำกว่าเกณฑ์", "ช่วยลดรายการย่อยที่ไม่ต้องการรับชำระ"),
        ("line_ids", "ตารางเอกสารที่พบ พร้อมยอดรวม ยอดค้าง และยอดที่จะรับชำระ", "ช่อง is_selected ใช้เลือกเอกสารที่จะรับจริง"),
    ]

    cheque_common_rows = [
        ("name", "เลขที่เช็คหรือเลขอ้างอิงเช็ค", "ใช้ค้นหาและติดตามสถานะเช็ค"),
        ("pay_partner_id", "คู่ค้าผู้จ่ายหรือผู้รับเช็ค", "อ้างอิงคู่สัญญาที่เกี่ยวข้อง"),
        ("amount", "ยอดเช็ค", "ใช้เทียบกับยอดใน payment entry"),
        ("date / cheque_date / clearing_date", "วันที่เอกสาร, วันที่เช็ค, วันที่นำฝาก/ตัดผ่าน", "ใช้ควบคุมช่วงเวลาเชิงบัญชี"),
        ("payment_method_line_id", "วิธีการชำระเงินที่ผูกกับ Journal และบัญชีพักเช็ค", "จุดนี้กำหนดบัญชีเดบิต/เครดิตหลักของ flow"),
    ]

    topics: list[TopicSpec] = []

    topics.append(
        TopicSpec(
            code="3.8",
            title="3.8 การจัดการการรับชำระแบบกลุ่มบริษัท",
            subtitle="Account Receivable",
            menu_paths=["Accounting > Customers > รับชำระเงินกลุ่มลูกค้า"],
            objectives=[
                "เพื่อให้ผู้ใช้สามารถรวมลูกหนี้หลายรายที่อยู่ภายใต้กลุ่มบริษัทเดียวกันมารับชำระในรอบเดียว",
                "เพื่อให้ผู้ใช้รู้วิธีค้นหาเอกสารคงค้าง เลือกยอดที่จะรับ และเปิด wizard รับชำระเงินได้ถูกต้อง",
                "เพื่อให้ตรวจสอบผลลัพธ์หลังรับชำระได้ทั้งระดับใบรับชำระและ Journal Items",
            ],
            overview=[
                "เมนูนี้เป็น custom สำหรับงานรับชำระแบบกลุ่มบริษัท โดยผู้ใช้เลือกสมาชิกในกลุ่มก่อน แล้วให้ระบบค้นหาใบแจ้งหนี้ที่ค้างชำระของสมาชิกที่เลือก จากนั้นจึงเปิดหน้ารับชำระมาตรฐานของ Odoo เพื่อสร้างรายการรับเงินจริง.",
                "ตัวอย่างใน local UAT มีเอกสาร draft ที่ยังไม่สร้าง payment คือเลขที่ 157 และมีเอกสาร done ที่สร้าง payment แล้วคือเลขที่ 142. ในคู่มือนี้จะใช้งานจากสองตัวอย่างนี้เพื่อให้ user เห็นทั้งก่อนและหลังบันทึกรับชำระ.",
            ],
            functions_table=[
                ("ค้นหาเอกสาร", "ดึง invoice/out_invoice และใบลดหนี้ของลูกค้าที่เลือกซึ่งยังค้างชำระอยู่เข้ามาในตาราง line_ids", "ใช้หลังจากเลือกสมาชิกหรือเปลี่ยน filter"),
                ("ดำเนินการรับชำระ", "เปิดหน้าต่าง Account Payment Register ของ Odoo เพื่อให้กำหนด Journal, Payment Method และสร้าง payment จริง", "ใช้เมื่อเลือกรายการที่จะรับชำระครบแล้ว"),
                ("พิมพ์ใบสำคัญ", "พิมพ์เอกสารยืนยันการรับชำระหลังรายการอยู่สถานะ Done", "ใช้หลังบันทึกรับชำระสำเร็จ"),
                ("ดูรายการชำระเงิน", "เปิดรายการ payment ที่ระบบสร้างจาก group payment ใบนั้น", "ใช้ตรวจสอบผลการบันทึก"),
            ],
            steps=[
                ("ขั้นตอนการเปิดและค้นหาเอกสารคงค้าง", [
                    "เข้าเมนู Accounting > Customers > รับชำระเงินกลุ่มลูกค้า แล้วกด New เพื่อสร้างเอกสารใหม่ หรือเปิดเอกสาร draft ตัวอย่างเลขที่ 157 เพื่อดู flow ที่เตรียมไว้.",
                    "ในช่อง ลูกค้า/สมาชิกกลุ่ม ให้เลือกคู่ค้าที่อยู่ภายใต้กลุ่มเดียวกัน จากตัวอย่าง local UAT ใช้ UAT Manual Customer A 20260408 และ UAT Manual Customer B 20260408.",
                    "ถ้าต้องการกรองข้อมูลเพิ่มเติม ให้ระบุช่วงวันที่หรือยอดขั้นต่ำ แล้วกดปุ่ม ค้นหาเอกสาร ระบบจะเติมรายการคงค้างลงในตาราง เอกสารที่พบ.",
                    "ตรวจสอบคอลัมน์ยอดรวม ยอดค้าง และยอดที่จะรับชำระ หากต้องการรับบางรายการให้ติ๊ก is_selected เฉพาะบรรทัดที่ต้องการ.",
                ]),
                ("ขั้นตอนการสร้างรายการรับชำระจริง", [
                    "เมื่อเลือกเอกสารครบแล้ว ให้กดปุ่ม ดำเนินการรับชำระ ระบบจะเปิดหน้าต่าง Payment Register ของ Odoo.",
                    "ใน wizard ให้เลือก Journal ที่ต้องการรับเงินจริง เช่น Journal รับเงินสดหรือบัญชีพักรับชำระตามนโยบายของบริษัท.",
                    "ตรวจสอบ Payment Method และ Payment Date ให้ถูกต้อง จากนั้นกด Create Payment เพื่อสร้างรายการรับเงินจริง.",
                    "เมื่อระบบสร้าง payment สำเร็จ เอกสาร Group Payment จะเปลี่ยนเป็นสถานะ Done และปุ่ม พิมพ์ใบสำคัญ/ดูรายการชำระเงิน จะใช้งานได้.",
                ]),
            ],
            images=[
                ("รูป 3.8.1 หน้าจอ Group Payment สถานะ Draft พร้อมจุดที่ต้องคลิกก่อนค้นหาเอกสาร", "group_draft"),
                ("รูป 3.8.2 หน้าต่าง Payment Register ที่เปิดจากปุ่ม ดำเนินการรับชำระ", "group_wizard"),
                ("รูป 3.8.3 หน้าจอ Group Payment สถานะ Done หลังสร้าง payment แล้ว", "group_done"),
            ],
            field_tables=[("ตารางอธิบายฟิลด์สำคัญของหน้าจอ Group Payment", group_rows)],
            scenarios=[
                ("Scenario A: ค้นหาเอกสารก่อนรับชำระ", "ใช้ record draft ID 157 ซึ่งมีใบ INV-E/26/04/00001 และ INV-E/26/04/00002 ถูกเลือกไว้แล้ว", "ผู้ใช้เห็นรายการคงค้าง 2 ใบ และยังสามารถปรับยอดที่จะรับได้"),
                ("Scenario B: ตรวจสอบรายการหลังสร้าง payment", "ใช้ record done ID 142 แล้วกด ดูรายการชำระเงิน", "ผู้ใช้สามารถไล่ดู payment ที่ระบบสร้างจาก group payment ได้ทันที"),
            ],
            journal_tables=[
                ("ตัวอย่าง Journal Items จาก payment ที่สร้างโดย Group Payment", [
                    ("111103 บัญชีพักเงินสด", "ยอดเงินรับเข้าบัญชีพักก่อนกระทบยอด", 26837.50, 0.0),
                    ("113001 ลูกหนี้การค้า - ในประเทศ", "ตัดยอดลูกหนี้ของใบแจ้งหนี้ในกลุ่ม", 0.0, 26837.50),
                ]),
            ],
            cautions=[
                "ระบบจะใช้ invoice ที่สถานะ Posted และยังมี amount_residual เหลืออยู่เท่านั้น",
                "ถ้าเลือก Journal ผิด รายการเดบิต/เครดิตจะไปลงบัญชีพักคนละชุดทันที จึงควรตรวจสอบใน Payment Register ก่อนกด Create Payment",
            ],
        )
    )

    return topics


def _legacy_unused_block() -> list[TopicSpec]:
    topics: list[TopicSpec] = []

    topics.append(
        TopicSpec(
            code="5.6",
            title="5.6 รับเช็คจากลูกค้า",
            subtitle="Module Cheque",
            menu_paths=[
                "Accounting > Customers > Invoices",
                "Cheque > Cheque > Cheque Receiving",
            ],
            objectives=[
                "เพื่อให้ผู้ใช้รับชำระจากลูกค้าด้วยเช็คได้จากหน้า Invoice",
                "เพื่อให้เข้าใจการลงบัญชีเช็ครับคงค้างและเช็ครับที่ตัดผ่านแล้ว",
            ],
            overview=[
                "การรับเช็คจากลูกค้าเริ่มจาก Invoice แล้วใช้ปุ่ม Register Payment เช่นเดียวกับวิธีรับชำระอื่น แต่เมื่อเลือก Payment Method เป็น Cheque Incoming ระบบจะเปิดส่วนกรอกข้อมูลเช็ค เช่น เลขเช็ค ธนาคาร สาขา และวันที่เช็ค จากนั้นจะสร้าง Cheque Receiving record ให้อัตโนมัติ.",
                f"ใน local UAT ใช้ invoice {CUSTOMER_INVOICES['in_confirmed']['name']} และ {CUSTOMER_INVOICES['in_paid']['name']} เป็นตัวอย่างจริง.",
            ],
            functions_table=[
                ("Register Payment", "จุดเริ่มของการสร้างรายการรับเช็คจาก Invoice", "ใช้กับ Invoice ที่ Posted และยังไม่ชำระเต็ม"),
                ("Incoming Cheque Section", "ส่วนกรอกข้อมูลเช็คที่รับจากลูกค้า", "แสดงเฉพาะเมื่อ payment method line เป็น cheque incoming"),
                ("Print", "พิมพ์ข้อมูลเช็คขาเข้าในกรณีระบบกำหนดให้พิมพ์", "ใช้กับ inbound cheque ที่สถานะ confirmed"),
                ("Bank Deposit / Done", "ใช้เมื่อต้องการเคลียร์เช็คเข้าธนาคาร", "ขั้นตอนเหมือนการนำฝากเช็ค"),
            ],
            steps=[
                ("ขั้นตอนการรับเช็คจากหน้า Invoice", [
                    f"เปิด Invoice {CUSTOMER_INVOICES['in_confirmed']['name']} จากเมนู Accounting > Customers > Invoices แล้วกด Register Payment.",
                    "เลือก Journal = PBAY1 และ Payment Method = Cheque Payment (Inbound).",
                    "กรอกเลขเช็ค, ธนาคาร, สาขา, วันที่เช็ค, หมายเหตุ และยอดเช็คให้ครบ จากนั้นกด Create Payment.",
                    "ระบบจะสร้าง cheque record ในเมนู Cheque Receiving และสร้างรายการบัญชีรับชำระไปยังบัญชีพักเช็ครับ.",
                ]),
            ],
            images=[
                ("รูป 5.6.1 หน้าต่าง Register Payment ของ Invoice ฝั่งรับเช็ค", "invoice_register"),
                ("รูป 5.6.2 เช็ครับสถานะ Confirmed หลังสร้างรายการแล้ว", "cheque_in_confirmed"),
            ],
            field_tables=[("ฟิลด์สำคัญของการรับเช็คขาเข้า", cheque_common_rows + [
                ("bank_account_id / branch", "ธนาคารและสาขาของเช็คที่ลูกค้านำมาชำระ", "ใช้สำหรับติดตามที่มาของเช็ค"),
                ("dynamic_io_cheque_id", "Template ที่ใช้พิมพ์หรือแสดงผลเช็คขาเข้า", "ขึ้นกับการตั้งค่า Journal"),
            ])],
            scenarios=[
                (f"Scenario A: รับเช็คแต่ยังไม่เคลียร์ธนาคาร", f"ใช้เช็ค {in_cf['name']} สถานะ {in_cf['state']}", "ระบบเดบิตบัญชีพักเช็ครับและเครดิตลูกหนี้การค้า"),
                (f"Scenario B: รับเช็คและเคลียร์แล้ว", f"ใช้เช็ค {in_pd['name']} สถานะ {in_pd['state']}", "ระบบย้ายยอดจากบัญชีพักเช็ครับไปยังบัญชีธนาคารจริง"),
            ],
            journal_tables=[
                ("Journal Items ตอนสร้างเช็ครับ (Confirmed)", move_lines_to_rows(in_cf["payment_moves"][0]["lines"])),
            ],
            cautions=[
                "ใน local UAT ตัวอย่าง inbound cheque บางใบมีบริบท settlement เดิม ทำให้ยอดบน cheque record และยอดใน payment entry อาจไม่ตรงกันทุกกรณี ผู้ใช้ต้องตรวจยอดใน Payment Entry ประกอบเสมอ",
                "ควรกรอกเลขเช็คและธนาคารให้ตรงกับหน้าเช็คจริงเพื่อใช้ติดตามภายหลัง",
            ],
        )
    )

    topics.append(
        TopicSpec(
            code="5.7",
            title="5.7 เคลียร์ (กระทบยอด) เช็ค",
            subtitle="Module Cheque",
            menu_paths=[
                "Cheque > Cheque > Cheque Paying",
                "Cheque > Cheque > Cheque Receiving",
                "Cheque > Cheque > Cheque Transactions",
            ],
            objectives=[
                "เพื่อให้ผู้ใช้ทำขั้นตอน Bank Deposit และ Done ของเช็คได้ถูกต้อง",
                "เพื่อให้เข้าใจผลทางบัญชีของการเคลียร์เช็คทั้งฝั่งจ่ายและฝั่งรับ",
            ],
            overview=[
                "ในโมดูลเช็ค การเคลียร์เช็คคือการเปลี่ยนสถานะเช็คจาก Confirmed ไป Bank Deposit และ Done เพื่อสะท้อนว่าเช็คได้เคลื่อนไปถึงธนาคารและถูกตัดผ่านแล้ว. การเคลียร์นี้ไม่ได้หมายถึงการจับคู่ Bank Statement แบบทั่วไปอย่างเดียว แต่เป็น flow ในโมดูลเช็คที่สร้าง Journal Entry ให้ครบ.",
            ],
            functions_table=[
                ("Bank Deposit", "สร้างรายการเคลื่อนย้ายจากบัญชีพักเช็คไปสู่ขั้นตอนการเคลียร์ธนาคาร", "ใช้ก่อน Done"),
                ("Done", "ยืนยันว่าการเคลียร์เช็คเสร็จสิ้น", "เช็คจะเปลี่ยนเป็น Paid"),
                ("Payment Entry", "ใช้ตรวจรายการบัญชีของเช็คในแต่ละช่วง", "ควรตรวจทุกครั้งก่อนปิดงวด"),
            ],
            steps=[
                ("ขั้นตอนการเคลียร์เช็คขาออก", [
                    f"เปิดเช็ค {out_cf['name']} หรือ {out_pd['name']} จากเมนู Cheque Paying.",
                    "ถ้าเช็คอยู่สถานะ Confirmed ให้กด Bank Deposit ก่อน.",
                    "เมื่อเช็คเข้าสถานะ Bank Deposit แล้ว ให้กด Done เพื่อปิด flow และให้ระบบสร้างรายการบัญชีที่เกี่ยวกับธนาคารจริง.",
                ]),
                ("ขั้นตอนการเคลียร์เช็คขาเข้า", [
                    f"เปิดเช็ค {in_cf['name']} หรือ {in_pd['name']} จากเมนู Cheque Receiving.",
                    "กด Bank Deposit เมื่อเช็ครับถูกนำฝากธนาคาร แล้วกด Done เมื่อธนาคารตัดผ่านจริง.",
                ]),
            ],
            images=[
                ("รูป 5.7.1 เช็คขาออกสถานะ Paid พร้อม Smart Buttons สำหรับดูรายการบัญชี", "cheque_out_paid"),
                ("รูป 5.7.2 เช็ครับสถานะ Paid หลังเคลียร์เข้าธนาคารแล้ว", "cheque_in_paid"),
            ],
            field_tables=[("ฟิลด์ที่ต้องตรวจเมื่อเคลียร์เช็ค", [
                ("clearing_date", "วันที่ธนาคารเคลียร์เช็ค", "ควรตรงกับวันที่รับรู้รายการธนาคาร"),
                ("cheque_journal_entry_id", "Journal Entry ที่เกิดจากขั้นเคลียร์เช็ค", "ใช้ตรวจเดบิต/เครดิตธนาคารและบัญชีพัก"),
                ("count_payment_move / count_reversed_move", "จำนวนรายการบัญชีและกลับรายการที่ผูกกับเช็ค", "ใช้ควบคุมการตรวจสอบย้อนหลัง"),
            ])],
            scenarios=[
                ("Scenario A: เคลียร์เช็คจ่าย", f"ใช้ {out_pd['name']} เป็นตัวอย่างเช็คที่ถูกเคลียร์ครบแล้ว", "ตรวจ Payment Entry และ Deposit Move ประกอบ"),
                ("Scenario B: เคลียร์เช็ครับ", f"ใช้ {in_pd['name']} เป็นตัวอย่างเช็ครับที่ตัดผ่านแล้ว", "ตรวจว่าบัญชีพักเช็ครับถูกปิดยอดแล้ว"),
            ],
            journal_tables=[
                ("Journal Items ตอนเคลียร์เช็คขาออกเข้าธนาคาร", move_lines_to_rows(out_pd["deposit_move"]["lines"])),
                ("Journal Items ตอนเคลียร์เช็ครับเข้าธนาคาร", move_lines_to_rows(in_pd["deposit_move"]["lines"])),
            ],
            cautions=[
                "อย่ากด Done ข้ามขั้นโดยไม่ระบุวันที่เคลียร์ที่ถูกต้อง เพราะจะกระทบ period ของบัญชีธนาคาร",
                "ถ้าต้องย้อนรายการ ควรใช้ flow Void/Cancel ของโมดูลเช็ค ไม่ควรแก้ Journal Entry ตรงโดยไม่มีเหตุผล",
            ],
        )
    )

    topics.append(
        TopicSpec(
            code="5.8",
            title="5.8 การยกเลิกเช็คหรือแปลงสถานะเช็ค (Void or Transform a Cheque)",
            subtitle="Module Cheque",
            menu_paths=[
                "Cheque > Cheque > Cheque Paying",
                "Cheque > Cheque > Void Cheque",
            ],
            objectives=[
                "เพื่อให้ผู้ใช้ยกเลิกเช็คที่จ่ายหรือเคลียร์ผิดได้อย่างถูกวิธี",
                "เพื่อให้เข้าใจขอบเขตการใช้งานจริงใน local UAT ว่าปัจจุบันรองรับ Void, Cancel และ Reset To Draft",
            ],
            overview=[
                "ใน local UAT ปุ่มและหน้าจอที่ใช้งานได้จริงของหัวข้อนี้คือ Void, Cancel และ Reset To Draft. ส่วนหน้าจอ Transform Detail ยังถูกคอมเมนต์ซ่อนไว้ใน view จึงไม่ควรอ้างเป็นขั้นตอนใช้งานจริงในคู่มือนี้.",
                f"ตัวอย่างเช็คที่ใช้ประกอบคือเลขที่ {out_void['name']} ซึ่งอยู่สถานะ {out_void['state']} หลังผ่าน flow void แล้ว.",
            ],
            functions_table=[
                ("Void", "เปิด wizard ให้ผู้ใช้ระบุเหตุผลและกลับรายการเช็ค", "ใช้ได้เมื่อเช็คอยู่สถานะ bank_deposit หรือ paid"),
                ("Cancel", "ใช้กับเช็คที่อยู่สถานะ return เพื่อปิดรายการยกเลิก", "มักเกิดจาก wizard Void เรียกต่อ"),
                ("Reset To Draft", "ใช้ย้อนสถานะบางรายการกลับไป draft เพื่อเริ่มต้นใหม่", "ใช้ด้วยความระมัดระวังและควรตรวจ reversed entries ทุกครั้ง"),
            ],
            steps=[
                ("ขั้นตอนการ Void เช็ค", [
                    f"เปิดเช็ค {out_void['name']} จากเมนู Cheque Paying หรือเมนู Void Cheque.",
                    "ถ้าเช็คอยู่สถานะ Paid หรือ Bank Deposit ให้กด Void ระบบจะเปิด wizard ให้กรอกเหตุผล.",
                    "กรอกเหตุผลการยกเลิกแล้วกดยืนยัน ระบบจะสร้าง reversed entries และเปลี่ยนสถานะเช็คตาม flow ของโมดูล.",
                    "เปิด Smart Button Reversed Entry เพื่อตรวจสอบรายการกลับรายการที่เกิดขึ้น.",
                ]),
                ("ขั้นตอน Reset To Draft", [
                    "ถ้าธุรกิจอนุญาตให้ย้อนรายการ ให้ใช้ Reset To Draft เฉพาะกรณีที่ตรวจแล้วว่ายังไม่กระทบเอกสารปลายทางอื่น.",
                    "หลัง Reset ต้องตรวจทั้ง Payment Entry, Reversed Entry และสถานะของ cheque book line ว่ากลับมาใช้งานได้ถูกต้อง.",
                ]),
            ],
            images=[("รูป 5.8.1 ตัวอย่างเช็คที่ถูก Void แล้ว พร้อมจุดตรวจ Reversed Entry และ Reset To Draft", "cheque_out_void")],
            field_tables=[("สิ่งที่ต้องตรวจเมื่อยกเลิกเช็ค", [
                ("void_reason", "เหตุผลการยกเลิกเช็ค", "ควรกรอกให้ชัดเพื่อใช้ตรวจสอบย้อนหลัง"),
                ("reversed_entry_names", "เลขที่รายการกลับรายการที่ระบบสร้าง", "ใช้ติดตามผลกระทบทางบัญชี"),
                ("state / cheque_optinal", "สถานะเช็คและสถานะพิเศษ เช่น return", "ใช้ตัดสินว่าควร Cancel ต่อหรือไม่"),
            ])],
            scenarios=[
                ("Scenario A: ยกเลิกเช็คที่เคลียร์แล้ว", f"ใช้เช็ค {out_void['name']}", "ระบบสร้าง reversed entries และทำให้ตรวจสอบร่องรอยย้อนหลังได้"),
                ("Scenario B: เริ่มต้นใหม่หลังยกเลิก", "ใช้ปุ่ม Reset To Draft เมื่อมีนโยบายอนุญาต", "ต้องตรวจว่า cheque line กลับมาใช้ได้จริงก่อนจ่ายใหม่"),
            ],
            journal_tables=[
                ("Journal Items ตอนสร้างเช็คเดิมก่อนยกเลิก", move_lines_to_rows(out_void["payment_moves"][0]["lines"])),
                ("Journal Items ตอนเคลียร์เช็คก่อนถูก Void", move_lines_to_rows(out_void["deposit_move"]["lines"])),
            ],
            cautions=[
                "Transform Detail ยังไม่เปิดใช้ใน UI ปัจจุบัน จึงไม่ควรใช้อ้างอิงเป็นขั้นตอนในงานจริง",
                "การ Void เช็คกระทบรายการบัญชีหลายชุด ควรให้ผู้มีสิทธิ์ด้านบัญชีหรือหัวหน้างานเป็นผู้ดำเนินการ",
            ],
            scope_note="หมายเหตุ: หัวข้อในเอกสารใช้ชื่อเดิมตามสารบัญ แต่เนื้อหาจะอธิบายเฉพาะฟังก์ชันที่มีจริงใน local UAT เท่านั้น.",
        )
    )

    topics.extend(build_additional_topics(samples, cheque_common_rows, out_cf, out_pd))
    return topics


def render_topic(topic: TopicSpec, images: dict[str, Path]) -> Path:
    doc = Document()
    configure_doc(doc)
    add_title(doc, topic.title, topic.subtitle)
    doc.add_paragraph()

    add_heading(doc, "วัตถุประสงค์")
    add_bullets(doc, topic.objectives)

    add_heading(doc, "ภาพรวมการทำงาน")
    add_bullets(doc, topic.overview)
    if topic.scope_note:
        add_paragraph(doc, topic.scope_note, bold=True)

    add_heading(doc, "เมนูที่ใช้")
    add_bullets(doc, topic.menu_paths)

    add_heading(doc, "ฟังก์ชันที่ทำได้จริง")
    add_table(doc, ["ฟังก์ชัน", "คำอธิบาย", "ใช้เมื่อไร"], topic.functions_table)

    add_heading(doc, "ขั้นตอนการใช้งาน")
    for heading, items in topic.steps:
        add_numbered_steps(doc, heading, items)

    add_heading(doc, "ภาพประกอบการทำงาน")
    for caption, key in topic.images:
        add_image(doc, images[key], caption)

    add_heading(doc, "คำอธิบายฟิลด์และส่วนสำคัญบนหน้าจอ")
    for title, rows in topic.field_tables:
        add_paragraph(doc, title, bold=True)
        add_table(doc, ["Field / Section", "ความหมาย", "วิธีใช้หรือค่าตัวอย่าง"], rows)

    add_heading(doc, "Scenario การใช้งานจากข้อมูลจริงใน local UAT")
    add_table(doc, ["Scenario", "ข้อมูลตัวอย่าง", "ผลลัพธ์ที่ต้องเห็น"], topic.scenarios)

    if topic.journal_tables:
        add_heading(doc, "การอธิบาย Journal Items และขาบัญชี")
        for title, rows in topic.journal_tables:
            add_paragraph(doc, title, bold=True)
            add_table(doc, ["บัญชี", "คำอธิบาย", "Debit", "Credit"], rows)

    add_heading(doc, "ข้อควรระวัง")
    add_bullets(doc, topic.cautions)

    out_path = DOCX_DIR / f"{topic.code}_{topic.title.replace('/', '-').replace(':', '')}.docx"
    doc.save(out_path)
    return out_path


def main() -> None:
    ensure_dirs()
    samples = load_samples()
    capture_specs = build_capture_specs(samples)
    images = {key: generate_capture(spec) for key, spec in capture_specs.items()}
    topics = build_topics(samples, images)
    generated = [str(render_topic(topic, images)) for topic in topics]
    summary = {
        "generated_count": len(generated),
        "docx": generated,
        "images": {key: str(path) for key, path in images.items()},
    }
    (OUTPUT_DIR / "summary.json").write_text(json.dumps(summary, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"Generated {len(generated)} documents")


def build_additional_topics(
    samples: dict,
    cheque_common_rows: list[tuple[str, str, str]],
    out_cf: dict,
    out_pd: dict,
) -> list[TopicSpec]:
    topics: list[TopicSpec] = []

    topics.append(
        TopicSpec(
            code="5.3",
            title="5.3 สร้างสมุดเช็ค (Cheque-Book)",
            subtitle="Module Cheque",
            menu_paths=["Cheque > Cheque Book"],
            objectives=[
                "เพื่อให้ผู้ใช้สร้างสมุดเช็คและชุดเลขเช็คสำหรับใช้จ่ายเช็คในระบบ",
                "เพื่อให้ทราบลำดับสถานะ Draft > Submit > Done ของสมุดเช็ค",
            ],
            overview=[
                "สมุดเช็คเป็น master data ระดับปฏิบัติการที่ใช้เก็บช่วงเลขเช็คของธนาคารแต่ละเล่ม เมื่อระบบจ่ายเช็คขาออก จะดึงเลขเช็คจาก Cheque Book Lines ที่สถานะ draft มาใช้.",
                f"ใน local UAT มีตัวอย่างสมุดเช็คจริงชื่อ {samples['cheque_book']['name']} สถานะ {samples['cheque_book']['state']} สำหรับใช้ประกอบคำอธิบาย.",
            ],
            functions_table=[
                ("Submit", "ยืนยันข้อมูลหัวเล่มเช็คก่อน generate ใบเช็คย่อย", "ใช้เมื่อกรอกธนาคาร จำนวนเช็ค และเลขเริ่มต้นครบแล้ว"),
                ("Generate Cheque", "สร้าง cheque lines ตามจำนวนใบและช่วงเลขที่กำหนด", "หลังจาก generate แล้วจะเห็นรายการในแท็บ Cheque List"),
                ("Clear Cheque", "ล้างรายการ cheque lines ที่สร้างไว้ในสถานะ submit", "ใช้เมื่อกรอกเลขเริ่มต้นผิดและยังไม่ confirm"),
                ("Confirm", "ยืนยันให้สมุดเช็คพร้อมใช้งานจริง", "หลัง confirm แล้ว cheque lines draft จะถูกนำไปใช้ตอนจ่ายเช็ค"),
            ],
            steps=[
                ("ขั้นตอนการสร้างสมุดเช็ค", [
                    "เข้าเมนู Cheque > Cheque Book แล้วกด New.",
                    "กรอก Bank Account, จำนวนเช็ค และ First Cheque No ให้ครบ จากนั้นกด Submit.",
                    "กด Generate Cheque เพื่อให้ระบบสร้างเช็ครายใบในแท็บ Cheque List และตรวจสอบช่วงเลขเริ่มต้นถึงเลขสุดท้าย.",
                    "ถ้าข้อมูลถูกต้อง ให้กด Confirm เพื่อเปิดใช้สมุดเช็คเล่มนี้จริง.",
                ]),
            ],
            images=[("รูป 5.3.1 หน้าจอสมุดเช็ค พร้อมฟิลด์หลักและแท็บ Cheque List", "book")],
            field_tables=[("ฟิลด์สำคัญของสมุดเช็ค", [
                ("bank_account_journal_id", "Journal ธนาคารที่เชื่อมกับเช็คเล่มนี้", samples["bank_journal"]["name"]),
                ("cheque_qty", "จำนวนเช็คที่จะสร้างในเล่ม", "ใช้สร้าง cheque lines"),
                ("first_cheque_no_char / last_cheque_no_char", "ช่วงเลขเช็คของเล่ม", f"ตัวอย่าง draft เหลือ {', '.join(samples['cheque_book']['draft_leaves'][:3])} ..."),
                ("cheque_book_lines", "รายการเช็คย่อยในเล่ม", "แต่ละบรรทัดมีสถานะ draft/paid/cancelled/return"),
            ])],
            scenarios=[
                ("Scenario A: สร้างเล่มใหม่เพื่อจ่ายเช็คชุดใหม่", "เหมาะเมื่อธนาคารออกเช็คเล่มใหม่", "ผู้ใช้ต้องกรอกช่วงเลขเช็คให้ตรงเล่มจริงก่อน confirm"),
                ("Scenario B: ตรวจสอบเช็คคงเหลือในเล่มเดิม", f"เปิดเล่ม {samples['cheque_book']['name']}", "ใช้ดูว่ามีเลขเช็ค draft เหลือกี่ใบก่อนจะสร้างเล่มใหม่"),
            ],
            journal_tables=[],
            cautions=[
                "เมื่อ Confirm แล้ว ควรหลีกเลี่ยงการแก้เลขเช็คเริ่มต้น/สิ้นสุดย้อนหลัง",
                "ถ้ากรอกช่วงเลขเช็คผิด ควรแก้ตั้งแต่สถานะ submit ก่อนนำไปใช้จริง",
            ],
        )
    )

    topics.append(
        TopicSpec(
            code="5.4",
            title="5.4 ชำระบิลผู้ขายด้วยเช็ค",
            subtitle="Module Cheque",
            menu_paths=[
                "Accounting > Vendors > Bills",
                "Cheque > Cheque > Cheque Paying",
            ],
            objectives=[
                "เพื่อให้ผู้ใช้จ่าย Vendor Bill ด้วยเช็คจากหน้าบิลได้ถูกต้อง",
                "เพื่อให้เข้าใจว่าเช็คขาออกสร้างรายการบัญชีอะไรบ้างในแต่ละสถานะ",
            ],
            overview=[
                "การจ่ายบิลผู้ขายด้วยเช็คเริ่มจากปุ่ม Register Payment ในบิลผู้ขาย ผู้ใช้เลือก Journal ธนาคารและ Payment Method เป็นเช็ค แล้วกรอกข้อมูลเลขเช็คและรายละเอียดธนาคาร จากนั้นระบบจะสร้าง Cheque Outbound และ Payment Entry ให้อัตโนมัติ.",
                f"คู่มือนี้ใช้อ้างอิงบิลจริง {VENDOR_BILLS['out_confirmed']['name']} และ {VENDOR_BILLS['out_paid']['name']} เพื่ออธิบายทั้งสถานะ Confirmed และ Paid.",
            ],
            functions_table=[
                ("Register Payment", "เปิด wizard รับ/จ่ายเงินมาตรฐานของ Odoo", "เป็นจุดเริ่มต้นของการสร้างเช็ค"),
                ("Outgoing Cheque Section", "ส่วนกรอก cheque leaf, วันที่เช็ค, หมายเหตุ และ A/C Payee", "จะแสดงเฉพาะเมื่อเลือก Payment Method Line ที่เป็นเช็คขาออก"),
                ("Confirm", "ยืนยันเช็คให้อยู่สถานะ Confirmed", "ใช้เมื่อเช็คออกแล้วแต่ยังไม่ตัดผ่านธนาคาร"),
                ("Bank Deposit / Done", "ตัดยอดเช็คผ่านธนาคารจนเป็น Paid", "ใช้เมื่อเช็คออกจากบัญชีธนาคารจริง"),
            ],
            steps=[
                ("ขั้นตอนการสร้างเช็คจากบิลผู้ขาย", [
                    f"เปิดบิลผู้ขาย {VENDOR_BILLS['out_confirmed']['name']} จากเมนู Accounting > Vendors > Bills แล้วกด Register Payment.",
                    "ใน wizard ให้เลือก Journal = PBAY1 และ Payment Method เป็น Cheque Payment (Outbound).",
                    "กรอกเลขเช็คจาก Cheque Book, วันที่เช็ค, หมายเหตุ และตรวจยอดให้ตรงกับบิล จากนั้นกด Create Payment.",
                    "ระบบจะสร้าง Cheque record ในเมนู Cheque Paying และสร้าง Payment Entry ที่ตัดยอดเจ้าหนี้ไปยังบัญชีพักเช็คจ่าย.",
                ]),
                ("ขั้นตอนการตรวจสอบหลังจ่าย", [
                    f"เปิดเช็คเลขที่ {out_cf['name']} หรือ {out_pd['name']} ในเมนู Cheque > Cheque > Cheque Paying.",
                    "ตรวจสอบสถานะเช็ค ปุ่ม Bank Deposit/Done และ Smart Buttons เช่น Payments Cheque และ Payment Entry.",
                ]),
            ],
            images=[
                ("รูป 5.4.1 หน้าต่าง Register Payment จากบิลผู้ขาย พร้อมจุดที่ต้องกรอก", "bill_register"),
                ("รูป 5.4.2 เช็คขาออกสถานะ Confirmed หลังสร้างรายการแล้ว", "cheque_out_confirmed"),
                ("รูป 5.4.3 เช็คขาออกสถานะ Paid หลังตัดผ่านธนาคาร", "cheque_out_paid"),
            ],
            field_tables=[
                ("ฟิลด์สำคัญของการจ่ายเช็คขาออก", cheque_common_rows + [
                    ("cheque_book_id / cheque_id", "สมุดเช็คและเลขเช็คที่ถูกดึงมาใช้จริง", "ผูกกับ Cheque Book ที่สถานะพร้อมใช้"),
                    ("memo", "ข้อความอ้างอิงบนเช็คและในรายการบัญชี", "ช่วยค้นหารายการย้อนหลัง"),
                    ("ac_payee", "ระบุว่าเช็คเป็น A/C Payee หรือไม่", "มีผลกับข้อมูลที่พิมพ์บนเช็ค"),
                ]),
            ],
            scenarios=[
                (f"Scenario A: จ่ายแล้วแต่ยังไม่ตัดผ่านธนาคาร", f"ใช้เช็ค {out_cf['name']} สถานะ {out_cf['state']}", "เช็คยังอยู่ในบัญชีพักเช็คจ่าย และสามารถติดตามเป็น Outstanding ได้"),
                (f"Scenario B: จ่ายและตัดผ่านแล้ว", f"ใช้เช็ค {out_pd['name']} สถานะ {out_pd['state']}", "ระบบสร้างรายการย้ายจากบัญชีพักเช็คจ่ายไปยังบัญชีธนาคารจริง"),
            ],
            journal_tables=[
                ("Journal Items ตอนสร้างเช็คขาออก (Confirmed)", move_lines_to_rows(out_cf["payment_moves"][0]["lines"])),
                ("Journal Items ตอนเช็คขาออกตัดผ่านธนาคาร (Deposit Move)", move_lines_to_rows(out_pd["deposit_move"]["lines"])),
            ],
            cautions=[
                "ยอดใน wizard ต้องเท่ากับยอดที่ต้องการชำระจริง มิฉะนั้นยอดค้างของบิลจะเหลือหรือจ่ายเกิน",
                "ถ้าเลือกเลขเช็คผิด จะกระทบการติดตาม outstanding cheque และการพิมพ์เช็คทันที",
            ],
        )
    )

    topics.append(
        TopicSpec(
            code="5.5",
            title="5.5 ติดตามเช็คคงค้าง (Outstanding Cheques)",
            subtitle="Module Cheque",
            menu_paths=[
                "Cheque > Cheque > Cheque Paying",
                "Cheque > Cheque > Cheque Transactions",
                "Cheque > Cheque > Paid Cheque",
            ],
            objectives=[
                "เพื่อให้ผู้ใช้ติดตามเช็คที่ยังไม่ตัดผ่านธนาคารได้ถูกต้อง",
                "เพื่อให้แยกสถานะ Confirmed, Bank Deposit และ Paid ของเช็คขาออกได้",
            ],
            overview=[
                "เช็คคงค้างในบริบทของโมดูลนี้หมายถึงเช็คที่สร้างและยืนยันแล้ว แต่ยังไม่เคลียร์ผ่านธนาคาร หรือเช็คที่อยู่ระหว่าง bank deposit. ผู้ใช้สามารถติดตามได้จากเมนู Cheque Paying และ Cheque Transactions โดยดูสถานะและบัญชีพักเช็คจ่ายประกอบ.",
                f"ตัวอย่างเช็ค outstanding จริงใน local UAT คือเลขที่ {out_cf['name']} ซึ่งอยู่สถานะ {out_cf['state']}.",
            ],
            functions_table=[
                ("Payments Cheque", "แสดง payment ที่ผูกกับเช็คใบนั้น", "ใช้ตรวจสอบยอดที่ถูกชำระ"),
                ("Payment Entry", "แสดง Journal Entry ฝั่ง payment", "ใช้ตรวจสอบการลงบัญชีเจ้าหนี้และบัญชีพักเช็คจ่าย"),
                ("Bank Deposit", "ย้ายสถานะจาก Confirmed ไปสู่ขั้นตอนการเคลียร์ธนาคาร", "ใช้เมื่อธนาคารนำเช็คไปเรียกเก็บแล้ว"),
            ],
            steps=[
                ("ขั้นตอนติดตามเช็คคงค้าง", [
                    "เข้าเมนู Cheque > Cheque > Cheque Paying แล้วค้นหาเลขเช็คหรือคู่ค้าที่ต้องการ.",
                    f"เปิดเช็คตัวอย่าง {out_cf['name']} เพื่อตรวจสอบสถานะ Confirmed และใช้ Smart Buttons เพื่อตรวจสอบ Payment Entry.",
                    "ถ้าธนาคารเริ่มนำเช็คไปเรียกเก็บ ให้กด Bank Deposit เพื่อเปลี่ยนเช็คเข้าสถานะขั้นถัดไป.",
                ]),
            ],
            images=[("รูป 5.5.1 ตัวอย่างเช็คขาออกสถานะ Outstanding/Confirmed ที่ยังไม่ตัดผ่านธนาคาร", "cheque_out_confirmed")],
            field_tables=[("สิ่งที่ต้องตรวจบนเช็คคงค้าง", [
                ("state", "สถานะปัจจุบันของเช็ค", "Confirmed = คงค้าง, Bank Deposit = อยู่ระหว่างเคลียร์, Paid = เคลียร์แล้ว"),
                ("payment_ids", "รายการ payment ที่ผูกกับเช็ค", "ใช้ตรวจสอบยอดที่ชำระจริง"),
                ("cheque_journal_entry_id", "รายการ Journal Entry ที่เกิดจากขั้น Bank Deposit/Done", "ใช้ติดตามการเคลียร์เช็ค"),
            ])],
            scenarios=[
                ("Scenario A: เช็คยังไม่ออกจากธนาคาร", f"เช็ค {out_cf['name']} ยังไม่กด Bank Deposit", "ยอดยังอยู่ในบัญชีพักเช็คจ่าย"),
                ("Scenario B: เช็คเคลียร์แล้ว", f"เช็ค {out_pd['name']} อยู่สถานะ {out_pd['state']}", "ยอดออกจากบัญชีพักเช็คจ่ายและตัดเข้าธนาคารแล้ว"),
            ],
            journal_tables=[("ตัวอย่าง Journal Items ที่ใช้ตรวจเช็คคงค้าง", move_lines_to_rows(out_cf["payment_moves"][0]["lines"]))],
            cautions=[
                "เมนู Paid Cheque เป็นมุมมองผลลัพธ์หลังเช็คจบแล้ว ไม่ใช่ที่ใช้ติดตาม outstanding หลัก",
                "ผู้ใช้ควรดู Payment Entry ประกอบเสมอ เพื่อยืนยันว่าบัญชีพักเช็คจ่ายยังมียอดคงอยู่จริง",
            ],
        )
    )

    topics.append(
        TopicSpec(
            code="5.1",
            title="5.1 เปิดใช้งานเช็คเป็นวิธีการชำระเงิน",
            subtitle="Module Cheque",
            menu_paths=[
                "Cheque > Configuration > Settings",
                "Accounting > Configuration > Journals",
            ],
            objectives=[
                "เพื่อให้ผู้ใช้เปิดใช้งานการรับเช็คและจ่ายเช็คบน Journal ธนาคารได้ถูกต้อง",
                "เพื่อให้เข้าใจว่าการตั้งค่า Payment Method Line มีผลต่อบัญชีพักเช็ครับและเช็คจ่ายอย่างไร",
            ],
            overview=[
                "โมดูลเช็คของ Gold Mints ใช้ Journal ธนาคารเป็นตัวตั้งต้น เมื่อเปิดใช้ฟังก์ชัน Cheque Incoming และ Cheque Outgoing แล้ว ระบบจะเพิ่มการตั้งค่าเฉพาะของเช็คใน Journal นั้น เช่น เทมเพลตฟอร์มเช็ค บัญชีพักเช็ครับ และบัญชีพักเช็คจ่าย.",
                f"ใน local UAT ใช้ Journal จริงรหัส {samples['bank_journal']['code']} ชื่อ {samples['bank_journal']['name']} เป็นตัวอย่างหลักของหัวข้อนี้.",
            ],
            functions_table=[
                ("Is Reverse Cheque Entry?", "กำหนดให้ระบบสร้างรายการ reverse สำหรับเช็คที่ถูก void/return ตามนโยบายบริษัท", "อยู่ในหน้า Settings ของโมดูลเช็ค"),
                ("Cheque Incoming", "เปิดให้ Journal รับเช็คจากลูกค้าได้", "เมื่อเปิดแล้ว Payment Method Line ฝั่งรับจะมี flag is_cheque_incoming_line"),
                ("Cheque Outgoing", "เปิดให้ Journal จ่ายเช็คให้ผู้ขายได้", "เมื่อเปิดแล้ว Payment Method Line ฝั่งจ่ายจะมี flag is_cheque_outgoing_line"),
                ("Cheque Form", "ผูกเทมเพลตฟอร์มเช็คที่ใช้พิมพ์เช็คขาออก", "เลือกได้หลาย template ผ่าน many2many"),
            ],
            steps=[
                ("ขั้นตอนการเปิดใช้ฟังก์ชันเช็ค", [
                    "เข้าเมนู Cheque > Configuration > Settings แล้วเปิดตัวเลือก Is Reverse Cheque Entry? หากบริษัทต้องการให้ระบบสร้างรายการกลับรายการให้อัตโนมัติเมื่อมีการ void เช็ค.",
                    "บันทึกการตั้งค่า จากนั้นเข้าเมนู Accounting > Configuration > Journals แล้วเปิด Journal ธนาคารที่ต้องการใช้งานเช็ค.",
                    f"ใน local UAT ให้ใช้ Journal {samples['bank_journal']['code']} ซึ่งตั้งไว้เป็น Journal ธนาคารจริงสำหรับตัวอย่างทั้งหมดของเอกสารชุดนี้.",
                    "ที่หน้า Journal ให้ติ๊ก Cheque Incoming และ Cheque Outgoing แล้วตรวจสอบ Payment Method Lines ทั้งฝั่งรับและฝั่งจ่ายว่าตั้ง flag และบัญชีพักถูกต้อง.",
                ]),
            ],
            images=[
                ("รูป 5.1.1 การเปิดใช้งานตัวเลือก Is Reverse Cheque Entry? ใน Settings", "settings"),
                ("รูป 5.1.2 การตั้งค่าเช็คในหน้า Journal ส่วน General", "journal_general"),
                ("รูป 5.1.3 การตั้งค่า Payment Method Line ฝั่งรับเช็ค", "journal_incoming"),
                ("รูป 5.1.4 การตั้งค่า Payment Method Line ฝั่งจ่ายเช็ค", "journal_outgoing"),
            ],
            field_tables=[
                ("ตารางอธิบายฟังก์ชันและฟิลด์สำคัญใน Journal", cheque_journal_rows(samples)),
                ("บัญชีที่เกี่ยวข้องกับการรับและจ่ายเช็คใน local UAT", [
                    ("payment_account_id ฝั่งรับเช็ค", "บัญชีพักเช็ครับ", "113005 เช็ครับลงวันที่ล่วงหน้า"),
                    ("payment_account_id ฝั่งจ่ายเช็ค", "บัญชีพักเช็คจ่าย", "212004 เช็คจ่ายลงวันที่ล่วงหน้า"),
                    ("Bank Account ของ Journal", "บัญชีเงินฝากธนาคารจริง", "111201 BAY CA สามแยก #046-0-14721-8"),
                ]),
            ],
            scenarios=[
                ("Scenario A: เปิดใช้งานเฉพาะรับเช็ค", "ใช้กับ Journal ที่บริษัทต้องการรับเช็คจากลูกค้า แต่ไม่ใช้จ่ายเช็ค", "ให้เปิด Cheque Incoming อย่างเดียว และกำหนดบัญชีพักเช็ครับ"),
                ("Scenario B: เปิดใช้งานครบทั้งรับและจ่าย", "ใช้กับ Journal ธนาคารกลางที่ต้องรองรับทั้งสองฝั่ง", "ต้องตรวจทั้ง template, incoming line และ outgoing line ให้ครบ"),
            ],
            journal_tables=[],
            cautions=[
                "ถ้าไม่เปิด flag ที่ Journal ต่อให้มี Payment Method Line เป็น cheque อยู่ ผู้ใช้ก็จะไม่เห็นส่วนกรอกข้อมูลเช็คใน wizard",
                "บัญชีพักเช็ครับและเช็คจ่ายต้องเป็นคนละบัญชี เพื่อให้ติดตาม outstanding ได้ชัดเจน",
            ],
        )
    )

    topics.append(
        TopicSpec(
            code="5.2",
            title="5.2 เลือกเทมเพลตฟอร์มเช็ค (Cheque Form Template)",
            subtitle="Module Cheque",
            menu_paths=["Cheque > Configuration > Cheque Lists"],
            objectives=[
                "เพื่อให้ผู้ใช้เลือกหรือปรับเทมเพลตการพิมพ์เช็คให้ตรงกับกระดาษเช็คของธนาคาร",
                "เพื่อเข้าใจว่ากลุ่มฟิลด์ใน template ส่งผลต่อการพิมพ์ตำแหน่งใดบนเช็ค",
            ],
            overview=[
                "เมนู Cheque Lists ใช้เก็บแม่แบบการพิมพ์เช็ค ซึ่ง Journal ฝั่งจ่ายเช็คจะเรียกไปใช้ผ่านฟิลด์ Cheque Form. Template นี้ไม่ได้เกี่ยวกับการบันทึกบัญชีโดยตรง แต่มีผลโดยตรงต่อความถูกต้องของตำแหน่งตัวอักษรบนเช็ค.",
                f"ใน local UAT มี template จริงชื่อ {samples['template']['name']} ใช้เป็นตัวอย่างในคู่มือนี้.",
            ],
            functions_table=[
                ("สร้าง/แก้ไข Template", "กำหนดขนาดเช็ค ระยะขอบ และตำแหน่งตัวอักษรที่ระบบจะพิมพ์", "ใช้ก่อนนำ Journal ไปพิมพ์เช็คจริง"),
                ("กำหนด Payee / Date / Amount / Amount in Word", "คุมตำแหน่งและขนาดตัวอักษรของแต่ละองค์ประกอบ", "ต้องเทียบกับกระดาษเช็คจริงของธนาคาร"),
                ("กำหนด Signature Box", "กำหนดกรอบลงนามหรือพื้นที่เซ็น", "ใช้กับฟอร์มที่มีตำแหน่งเซ็นตายตัว"),
            ],
            steps=[
                ("ขั้นตอนการเลือกและตรวจสอบ Template", [
                    "เข้าเมนู Cheque > Configuration > Cheque Lists แล้วเปิด template ที่ต้องการ เช่น Standard Cheque.",
                    "ตรวจสอบค่าขนาดเช็ค ความสูง ความกว้าง และกลุ่มฟิลด์ตำแหน่งต่าง ๆ เช่น Cheque Date, Payee Name, Amount in Figure และ Amount in Word.",
                    "บันทึก template จากนั้นกลับไปที่ Journal และตรวจว่าฟิลด์ Cheque Form ผูก template นี้แล้ว.",
                ]),
            ],
            images=[("รูป 5.2.1 หน้าจอ Template ของเช็ค พร้อมจุดหลักที่ใช้ปรับตำแหน่งพิมพ์", "template")],
            field_tables=[
                ("กลุ่มฟิลด์สำคัญใน Cheque Form Template", [
                    ("name", "ชื่อ template", "เช่น Standard Cheque"),
                    ("cheque_hight / cheque_width", "กำหนดขนาดกระดาษเช็ค", "ต้องเทียบกับเช็คจริง"),
                    ("top_margin / left_margin / font_size", "กำหนดตำแหน่งและขนาดของวันที่เช็ค", "มีผลกับบรรทัดวันที่ด้านบน"),
                    ("payee_top_margin / payee_left_margin / payee_width", "กำหนดตำแหน่งชื่อผู้รับเงิน", "มีผลกับบรรทัด Payee"),
                    ("af_top_margin / af_left_margin / af_width", "กำหนด Amount in Figure", "ใช้กับยอดตัวเลข"),
                    ("fl_* / sc_*", "กำหนด Amount in Word และการตัดคำ", "ช่วยไม่ให้ข้อความล้นบรรทัด"),
                ]),
            ],
            scenarios=[
                ("Scenario A: ใช้ Template เดิมของธนาคาร", "เปิด Standard Cheque และตรวจว่า Journal PBAY1 ถูกผูกไว้แล้ว", "ผู้ใช้พิมพ์เช็คได้ทันทีโดยไม่ต้องสร้าง template ใหม่"),
                ("Scenario B: เปลี่ยนตำแหน่งตัวอักษรเพราะกระดาษเช็คเปลี่ยนรุ่น", "ปรับ margin และ font size แล้วทดสอบพิมพ์กับเช็คใบตัวอย่าง", "ตำแหน่งตัวอักษรตรงกระดาษโดยไม่กระทบการลงบัญชี"),
            ],
            journal_tables=[],
            cautions=[
                "การแก้ template มีผลกับเช็คที่จะพิมพ์หลังจากนั้นทั้งหมดใน Journal ที่ผูก template นี้",
                "ควรทดสอบพิมพ์กับกระดาษเปล่าก่อนใช้กับเช็คจริง",
            ],
        )
    )

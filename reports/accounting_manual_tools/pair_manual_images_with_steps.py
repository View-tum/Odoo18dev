from __future__ import annotations

from copy import deepcopy
from pathlib import Path
import re
import subprocess

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph


ROOT = Path(r"C:\365_project\TheCool18e\Dev")
DOCX_DIR = ROOT / "manual" / "Accouting_Manual" / "generated_20260408" / "docx"
PDF_DIR = ROOT / "manual" / "Accouting_Manual" / "generated_20260408" / "pdf_review"
IMAGE_DIR = ROOT / "manual" / "Accouting_Manual" / "generated_20260408" / "images"
SOFFICE = Path(r"C:\Program Files\LibreOffice\program\soffice.com")
TARGET_PREFIXES = ("3.8_", "5.", "6.", "7.")

DETAIL_HEADING = "ลำดับการทำงานแบบละเอียด"
STOP_HEADINGS = (
    "4. ตรวจสอบ Journal Entry",
    "คำอธิบายหน้าจอและช่องสำคัญ",
    "คำอธิบายฟิลด์และส่วนสำคัญบนหน้าจอ",
    "ตัวอย่างการใช้งานในสถานการณ์จริง",
    "การอธิบาย Journal Items และขาบัญชี",
    "ข้อควรระวัง",
)


IMAGE_MAP: dict[str, list[tuple[str, str]]] = {
    "3.8": [
        ("nav_dashboard_accounting_real_annotated.png", "ภาพหน้า Dashboard สำหรับเข้าโมดูล Accounting"),
        ("nav_dashboard_accounting_real_annotated.png", "ภาพหน้า Dashboard จุดที่ใช้คลิกเข้า Accounting"),
        ("nav_accounting_group_payment_real_annotated.png", "ภาพเมนูรับชำระเงินกลุ่มลูกค้า"),
        ("group_draft_manual_annotated.png", "ภาพเอกสารรับชำระแบบกลุ่มบริษัทก่อนเริ่มรับเงิน"),
        ("group_payment_wizard_manual_annotated.png", "ภาพหน้าต่างรับชำระเงินของกลุ่มลูกค้า"),
        ("group_done_manual_annotated.png", "ภาพรายการหลังสร้างรับชำระแล้ว"),
        ("journal_group_payment_real_annotated.png", "ภาพ Journal Entry ที่ใช้ตรวจสอบผลการรับชำระ"),
    ],
    "5.1": [
        ("nav_dashboard_cheque_real_annotated.png", "ภาพหน้า Dashboard สำหรับเข้าโมดูล Cheque"),
        ("nav_dashboard_cheque_real_annotated.png", "ภาพหน้า Dashboard จุดที่ใช้คลิกเข้า Cheque"),
        ("nav_cheque_configuration_real_annotated.png", "ภาพเมนูตั้งค่าเช็ค"),
        ("settings_cheque_manual_annotated.png", "ภาพหน้าจอตั้งค่าเปิดใช้งานเช็คในระบบ"),
    ],
    "5.2": [
        ("nav_dashboard_cheque_real_annotated.png", "ภาพหน้า Dashboard สำหรับเข้าโมดูล Cheque"),
        ("nav_cheque_configuration_real_annotated.png", "ภาพเมนู Cheque Form Template"),
        ("template_manual_annotated.png", "ภาพหน้าจอเทมเพลตเช็คที่ใช้กำหนดตำแหน่งพิมพ์"),
    ],
    "5.3": [
        ("nav_dashboard_cheque_real_annotated.png", "ภาพหน้า Dashboard สำหรับเข้าโมดูล Cheque"),
        ("nav_cheque_configuration_real_annotated.png", "ภาพเมนูสร้างสมุดเช็ค"),
        ("cheque_book_manual_annotated.png", "ภาพหน้าจอสมุดเช็คและรายการเลขเช็ค"),
    ],
    "5.4": [
        ("nav_dashboard_accounting_real_annotated.png", "ภาพหน้า Dashboard สำหรับเข้าโมดูล Accounting"),
        ("nav_dashboard_accounting_real_annotated.png", "ภาพหน้า Dashboard จุดที่ใช้คลิกเข้า Accounting"),
        ("nav_accounting_vendors_bills_real_annotated.png", "ภาพเมนู Vendor Bills"),
        ("nav_cheque_operations_real_annotated.png", "ภาพเมนู Cheque Paying สำหรับติดตามเช็คจ่าย"),
        ("vendor_bill_form_real_annotated.png", "ภาพบิลผู้ขายจริงก่อนเริ่มชำระ"),
        ("vendor_bill_form_real_annotated.png", "ภาพบิลผู้ขายจริงที่ใช้เปิดเพื่อเริ่มจ่ายด้วยเช็ค"),
        ("bill_register_payment_manual_annotated.png", "ภาพหน้าต่าง Register Payment ที่ใช้กรอกข้อมูลเช็ค"),
        ("bill_register_payment_manual_annotated.png", "ภาพจุดที่กรอกเลขเช็คและยอดเช็คใน Register Payment"),
        ("cheque_out_confirmed_manual_annotated.png", "ภาพเช็คขาออกสถานะ Confirmed หลังสร้างรายการ"),
    ],
    "5.5": [
        ("nav_dashboard_cheque_real_annotated.png", "ภาพหน้า Dashboard สำหรับเข้าโมดูล Cheque"),
        ("nav_cheque_operations_real_annotated.png", "ภาพเมนู Cheque Paying"),
        ("cheque_out_confirmed_manual_annotated.png", "ภาพเช็คสถานะคงค้างที่ยังไม่ผ่านธนาคาร"),
        ("cheque_out_paid_manual_annotated.png", "ภาพเช็คหลังผ่านธนาคารแล้ว"),
        ("journal_cheque_out_confirmed_real_annotated.png", "ภาพ Journal Entry ของเช็คคงค้าง"),
        ("journal_cheque_out_paid_real_annotated.png", "ภาพ Journal Entry หลังเช็คผ่านธนาคาร"),
    ],
    "5.6": [
        ("nav_dashboard_accounting_real_annotated.png", "ภาพหน้า Dashboard สำหรับเข้าโมดูล Accounting"),
        ("nav_dashboard_accounting_real_annotated.png", "ภาพหน้า Dashboard จุดที่ใช้คลิกเข้า Accounting"),
        ("nav_accounting_customers_invoices_real_annotated.png", "ภาพเมนู Customer Invoices"),
        ("nav_cheque_operations_real_annotated.png", "ภาพเมนู Cheque Receiving ที่ใช้ติดตามเช็ครับ"),
        ("customer_invoice_form_real_annotated.png", "ภาพใบแจ้งหนี้ลูกค้าก่อนรับชำระ"),
        ("customer_invoice_form_real_annotated.png", "ภาพใบแจ้งหนี้ลูกค้าที่ต้องการรับเช็ค"),
        ("invoice_register_payment_real_annotated.png", "ภาพหน้าต่าง Register Payment สำหรับรับเช็ค"),
        ("invoice_register_payment_real_annotated.png", "ภาพจุดที่กรอกเลขเช็คและยอดเช็คใน Register Payment"),
        ("cheque_in_confirmed_manual_annotated.png", "ภาพเช็ครับสถานะ Confirmed หลังรับชำระแล้ว"),
    ],
    "5.7": [
        ("nav_dashboard_cheque_real_annotated.png", "ภาพหน้า Dashboard สำหรับเข้าโมดูล Cheque"),
        ("nav_cheque_operations_real_annotated.png", "ภาพเมนูเช็คที่ใช้เปิดรายการคงค้าง"),
        ("cheque_in_confirmed_manual_annotated.png", "ภาพเช็ครับก่อนเคลียร์"),
        ("cheque_in_paid_manual_annotated.png", "ภาพเช็ครับหลังเคลียร์แล้ว"),
        ("journal_cheque_in_paid_real_annotated.png", "ภาพ Journal Entry หลังเคลียร์เช็ครับ"),
        ("journal_cheque_out_paid_real_annotated.png", "ภาพ Journal Entry หลังเคลียร์เช็คจ่าย"),
    ],
    "5.8": [
        ("nav_dashboard_cheque_real_annotated.png", "ภาพหน้า Dashboard สำหรับเข้าโมดูล Cheque"),
        ("nav_cheque_operations_real_annotated.png", "ภาพเมนูรายการเช็ค"),
        ("cheque_out_void_manual_annotated.png", "ภาพรายการเช็คที่ถูกยกเลิกหรือเปลี่ยนสถานะ"),
        ("journal_cheque_void_reverse_real_annotated.png", "ภาพ Journal Entry ที่เกิดจากการยกเลิกหรือกลับรายการเช็ค"),
    ],
    "6.1": [
        ("nav_dashboard_accounting_real_annotated.png", "ภาพหน้า Dashboard สำหรับเข้าโมดูล Accounting"),
        ("nav_accounting_assets_real_annotated.png", "ภาพเมนู Assets"),
        ("asset_draft_form_real_annotated.png", "ภาพทรัพย์สินสถานะ Draft"),
        ("asset_running_form_real_annotated.png", "ภาพทรัพย์สินที่เริ่มใช้งานแล้ว"),
        ("asset_sell_form_real_annotated.png", "ภาพทรัพย์สินที่ปิดด้วยการขาย"),
        ("asset_dispose_form_real_annotated.png", "ภาพทรัพย์สินที่ตัดจำหน่ายออกจากระบบ"),
    ],
    "6.2": [
        ("nav_dashboard_accounting_real_annotated.png", "ภาพหน้า Dashboard สำหรับเข้าโมดูล Accounting"),
        ("nav_dashboard_accounting_real_annotated.png", "ภาพหน้า Dashboard จุดที่ใช้คลิกเข้า Accounting"),
        ("nav_accounting_asset_models_real_annotated.png", "ภาพเมนู Asset Models"),
        ("nav_accounting_assets_real_annotated.png", "ภาพเมนู Assets สำหรับสร้างทรัพย์สินใหม่"),
        ("asset_model_form_real_annotated.png", "ภาพแบบทรัพย์สินที่ใช้กำหนดบัญชี"),
        ("asset_model_form_real_annotated.png", "ภาพตำแหน่งบัญชีใน Asset Model ที่ต้องตรวจสอบก่อนสร้าง"),
        ("nav_accounting_assets_real_annotated.png", "ภาพรายการ Assets ที่ใช้กด New"),
        ("asset_draft_form_real_annotated.png", "ภาพทรัพย์สินจริงหลังบันทึกเป็น Draft"),
        ("asset_draft_form_real_annotated.png", "ภาพหน้าจอ Draft ที่ใช้ตรวจผลหลังบันทึก"),
    ],
    "6.3": [
        ("nav_dashboard_accounting_real_annotated.png", "ภาพหน้า Dashboard สำหรับเข้าโมดูล Accounting"),
        ("nav_accounting_assets_real_annotated.png", "ภาพเมนู Assets"),
        ("asset_running_form_real_annotated.png", "ภาพทรัพย์สินที่อยู่สถานะ Running"),
        ("journal_asset_depreciation_real_annotated.png", "ภาพ Journal Entry ค่าเสื่อมราคาจริง"),
    ],
    "6.4": [
        ("nav_dashboard_accounting_real_annotated.png", "ภาพหน้า Dashboard สำหรับเข้าโมดูล Accounting"),
        ("nav_accounting_customers_invoices_real_annotated.png", "ภาพเมนู Customer Invoices"),
        ("asset_sale_invoice_real_annotated.png", "ภาพใบขายทรัพย์สินจริง"),
        ("asset_sell_form_real_annotated.png", "ภาพทรัพย์สินหลังขายแล้ว"),
        ("journal_asset_sale_real_annotated.png", "ภาพ Journal Entry ของการขายทรัพย์สิน"),
    ],
    "6.5": [
        ("nav_dashboard_accounting_real_annotated.png", "ภาพหน้า Dashboard สำหรับเข้าโมดูล Accounting"),
        ("nav_accounting_assets_real_annotated.png", "ภาพเมนู Assets"),
        ("asset_dispose_form_real_annotated.png", "ภาพทรัพย์สินที่ตัดจำหน่ายแล้ว"),
        ("journal_asset_disposal_real_annotated.png", "ภาพ Journal Entry ของการตัดจำหน่ายทรัพย์สิน"),
    ],
    "6.6": [
        ("nav_dashboard_accounting_real_annotated.png", "ภาพหน้า Dashboard สำหรับเข้าโมดูล Accounting"),
        ("nav_accounting_fixed_asset_report_real_annotated.png", "ภาพเมนูรายงานทรัพย์สินถาวร"),
        ("fixed_asset_report_page_real_annotated.png", "ภาพหน้ารายงานทรัพย์สินถาวรจริงในระบบ"),
        ("asset_running_form_real_annotated.png", "ภาพหน้าทรัพย์สินที่ใช้เทียบกับรายงาน"),
    ],
    "7.1": [
        ("nav_dashboard_manufacturing_real_annotated.png", "ภาพหน้า Dashboard สำหรับเข้าโมดูล Manufacturing"),
        ("nav_manufacturing_orders_real_annotated.png", "ภาพเมนู Manufacturing Orders"),
        ("manufacturing_order_form_real_annotated.png", "ภาพใบสั่งผลิตจริง"),
        ("nav_manufacturing_bom_real_annotated.png", "ภาพเมนู Bills of Materials"),
        ("manufacturing_bom_form_real_annotated.png", "ภาพ BoM ที่ใช้กับงานผลิตจริง"),
        ("inventory_valuation_page_real_annotated.png", "ภาพหน้ารายงาน Valuation ที่ใช้ตามรอยต้นทุน"),
    ],
    "7.2": [
        ("nav_dashboard_manufacturing_real_annotated.png", "ภาพหน้า Dashboard สำหรับเข้าโมดูล Manufacturing"),
        ("nav_dashboard_manufacturing_real_annotated.png", "ภาพหน้า Dashboard จุดที่ใช้คลิกเข้า Manufacturing"),
        ("nav_manufacturing_orders_real_annotated.png", "ภาพเมนู Manufacturing Orders"),
        ("nav_inventory_valuation_real_annotated.png", "ภาพเมนู Valuation"),
        ("manufacturing_order_form_real_annotated.png", "ภาพใบสั่งผลิตที่ใช้ดูงานจริง"),
        ("manufacturing_order_form_real_annotated.png", "ภาพใบสั่งผลิตที่ใช้ตรวจข้อมูลก่อนเปิดรายการบัญชี"),
        ("inventory_valuation_page_real_annotated.png", "ภาพหน้า Valuation ที่ใช้ไล่รายการบัญชี"),
        ("journal_mfg_raw_fg02001_real_annotated.png", "ภาพ Journal Entry ตอนตัดกึ่งสำเร็จรูป"),
        ("journal_mfg_finished_real_annotated.png", "ภาพ Journal Entry ตอนรับสินค้าสำเร็จรูป"),
    ],
    "7.3": [
        ("nav_dashboard_inventory_real_annotated.png", "ภาพหน้า Dashboard สำหรับเข้าโมดูล Inventory"),
        ("product_category_fg_form_real_annotated.png", "ภาพ Product Category ของสินค้าสำเร็จรูป"),
        ("product_category_rm_form_real_annotated.png", "ภาพ Product Category ของวัตถุดิบ"),
        ("nav_manufacturing_bom_real_annotated.png", "ภาพเมนู Bills of Materials"),
        ("manufacturing_bom_form_real_annotated.png", "ภาพ BoM ที่ใช้กับการตั้งค่าทางบัญชี"),
        ("nav_inventory_valuation_real_annotated.png", "ภาพเมนู Valuation ที่ใช้ตรวจผลของการตั้งค่า"),
    ],
    "7.4": [
        ("nav_dashboard_manufacturing_real_annotated.png", "ภาพหน้า Dashboard สำหรับเข้าโมดูล Manufacturing"),
        ("manufacturing_order_form_real_annotated.png", "ภาพใบสั่งผลิตจริงที่ใช้เป็นตัวอย่าง"),
        ("manufacturing_bom_form_real_annotated.png", "ภาพ BoM ของตัวอย่างที่ใช้งาน"),
        ("inventory_valuation_page_real_annotated.png", "ภาพรายงาน Valuation ของตัวอย่างเดียวกัน"),
        ("journal_mfg_finished_real_annotated.png", "ภาพ Journal Entry ฝั่งรับสินค้าสำเร็จรูป"),
    ],
    "7.5": [
        ("nav_dashboard_manufacturing_real_annotated.png", "ภาพหน้า Dashboard สำหรับเข้าโมดูล Manufacturing"),
        ("nav_manufacturing_scrap_real_annotated.png", "ภาพเมนู Scrap"),
        ("manufacturing_scrap_form_real_annotated.png", "ภาพรายการ Scrap จริงในระบบ"),
        ("inventory_valuation_page_real_annotated.png", "ภาพหน้า Valuation ที่ใช้ตรวจผลของการตัดของ"),
        ("journal_mfg_raw_packaging_real_annotated.png", "ภาพ Journal Entry ที่ใช้ดูผลบัญชีของการตัดวัตถุดิบ"),
    ],
    "7.6": [
        ("nav_dashboard_inventory_real_annotated.png", "ภาพหน้า Dashboard สำหรับเข้าโมดูล Inventory"),
        ("nav_inventory_valuation_real_annotated.png", "ภาพเมนู Valuation"),
        ("inventory_valuation_page_real_annotated.png", "ภาพหน้ารายงาน Valuation จริงในระบบ"),
        ("journal_mfg_raw_fg02001_real_annotated.png", "ภาพ Journal Entry ฝั่งตัดวัตถุดิบ"),
        ("journal_mfg_finished_real_annotated.png", "ภาพ Journal Entry ฝั่งรับสินค้าสำเร็จรูป"),
    ],
    "7.7": [
        ("nav_dashboard_accounting_real_annotated.png", "ภาพหน้า Dashboard สำหรับเข้าโมดูล Accounting"),
        ("nav_accounting_rng8_real_annotated.png", "ภาพเมนูรายงาน รง.8"),
        ("rng8_report_page_real_annotated.png", "ภาพหน้ารายงาน รง.8 จริงในระบบ"),
        ("inventory_valuation_page_real_annotated.png", "ภาพหน้ารายงาน Valuation ที่ใช้เทียบตัวเลขย้อนหลัง"),
    ],
}


def set_run_font(run, size: int = 16, bold: bool = False, italic: bool = False) -> None:
    run.bold = bold
    run.italic = italic
    run.font.name = "TH Sarabun New"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "TH Sarabun New")
    run.font.size = Pt(size)


def make_empty_after(paragraph: Paragraph) -> Paragraph:
    new_p = deepcopy(paragraph._p)
    for child in list(new_p):
        new_p.remove(child)
    paragraph._p.addnext(new_p)
    return Paragraph(new_p, paragraph._parent)


def insert_text_after(paragraph: Paragraph, text: str, *, size: int = 16, bold: bool = False, italic: bool = False) -> Paragraph:
    new_para = make_empty_after(paragraph)
    run = new_para.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic)
    return new_para


def insert_picture_after(paragraph: Paragraph, image_path: Path, caption: str) -> Paragraph:
    pic_para = make_empty_after(paragraph)
    run = pic_para.add_run()
    run.add_picture(str(image_path), width=Inches(6.2))
    cap_para = insert_text_after(pic_para, caption, size=14, italic=True)
    return cap_para


def find_detail_heading(paragraphs: list[Paragraph]) -> int | None:
    for idx, paragraph in enumerate(paragraphs):
        if paragraph.text.strip() == DETAIL_HEADING:
            return idx
    return None


def is_numbered_step(text: str) -> bool:
    return bool(re.match(r"^\d+\.\s+", text.strip()))


def step_number(text: str) -> int | None:
    match = re.match(r"^(\d+)\.\s+", text.strip())
    return int(match.group(1)) if match else None


def find_detail_steps(paragraphs: list[Paragraph]) -> list[Paragraph]:
    heading_idx = find_detail_heading(paragraphs)
    if heading_idx is None:
        return []
    steps: list[Paragraph] = []
    started = False
    previous = 0
    for paragraph in paragraphs[heading_idx + 1 :]:
        text = paragraph.text.strip()
        if not text:
            continue
        if text.startswith(STOP_HEADINGS):
            break
        if text.startswith("รูป "):
            continue
        if is_numbered_step(text):
            number = step_number(text) or 0
            if started and number <= previous:
                break
            started = True
            previous = number
            steps.append(paragraph)
    return steps


def remove_existing_blocks(doc: Document) -> None:
    paragraphs = doc.paragraphs
    remove_indices: set[int] = set()
    block_headers = {"ภาพประกอบของขั้นตอนนี้", "ภาพประกอบตามลำดับขั้นตอน"}
    step_paragraphs = find_detail_steps(paragraphs)
    if step_paragraphs:
        step_indices = [paragraphs.index(p) for p in step_paragraphs]
        for pos, step_idx in enumerate(step_indices):
            next_step_idx = step_indices[pos + 1] if pos + 1 < len(step_indices) else len(paragraphs)
            for j in range(step_idx + 1, next_step_idx):
                text = paragraphs[j].text.strip()
                has_pic = bool(paragraphs[j]._p.xpath('.//pic:pic'))
                if has_pic or text in block_headers or text.startswith("ภาพ") or text.startswith("รูป "):
                    remove_indices.add(j)
    for idx, paragraph in enumerate(paragraphs):
        text = paragraph.text.strip()
        if text in block_headers:
            remove_indices.add(idx)
    for idx in sorted(set(i for i in remove_indices if i < len(paragraphs)), reverse=True):
        element = paragraphs[idx]._element
        element.getparent().remove(element)


def doc_code_from_name(name: str) -> str | None:
    match = re.match(r"^(\d+\.\d+)_", name)
    return match.group(1) if match else None


def enrich_doc(path: Path) -> bool:
    code = doc_code_from_name(path.name)
    if not code or code not in IMAGE_MAP:
        return False
    doc = Document(path)
    remove_existing_blocks(doc)
    step_paragraphs = find_detail_steps(doc.paragraphs)
    if not step_paragraphs:
        return False
    image_specs = IMAGE_MAP[code]
    updated = False
    for idx, step_para in enumerate(step_paragraphs):
        image_name, caption = image_specs[min(idx, len(image_specs) - 1)]
        image_path = IMAGE_DIR / image_name
        if not image_path.exists():
            continue
        insert_picture_after(step_para, image_path, caption)
        updated = True
    if updated:
        doc.save(path)
    return updated


def render_pdfs(paths: list[Path]) -> None:
    if not paths:
        return
    PDF_DIR.mkdir(parents=True, exist_ok=True)
    cmd = [str(SOFFICE), "--headless", "--convert-to", "pdf", "--outdir", str(PDF_DIR), *map(str, paths)]
    subprocess.run(cmd, check=True)


def main() -> None:
    updated: list[Path] = []
    for path in DOCX_DIR.glob("*.docx"):
        if path.name.startswith(TARGET_PREFIXES) and enrich_doc(path):
            updated.append(path)
    render_pdfs(updated)
    print(f"updated {len(updated)} docs")


if __name__ == "__main__":
    main()

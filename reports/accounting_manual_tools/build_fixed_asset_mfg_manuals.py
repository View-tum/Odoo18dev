from __future__ import annotations

import json
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\365_project\TheCool18e\Dev")
TOOLS_DIR = ROOT / "reports" / "accounting_manual_tools"
OUTPUT_DIR = ROOT / "manual" / "Accouting_Manual" / "generated_20260408"
DOCX_DIR = OUTPUT_DIR / "docx"
PDF_DIR = OUTPUT_DIR / "pdf_review"
IMAGE_DIR = OUTPUT_DIR / "images"

MENU_PATH = TOOLS_DIR / "output" / "menu_probe_extended_20260409.json"
LIVE_PATH = TOOLS_DIR / "output" / "fixed_asset_mfg_live_samples_20260409.json"
EXTRA_PATH = TOOLS_DIR / "output" / "fixed_asset_mfg_extra_probe_20260409.json"
DEP_LINES_PATH = TOOLS_DIR / "output" / "fixed_asset_dep_move_lines_20260409.json"
DOC_SAMPLE_PATH = TOOLS_DIR / "output" / "fixed_asset_mfg_doc_samples_20260409.json"
SOFFICE = Path(r"C:\Program Files\LibreOffice\program\soffice.com")


def load_json(path: Path):
    text = path.read_text(encoding="utf-8")
    fixed = text.replace("Ã‚Â ", " ").replace("\xa0", " ")
    candidates = [text, fixed]
    try:
        candidates.append(fixed.encode("latin1").decode("utf-8"))
    except Exception:
        pass
    for candidate in candidates:
        try:
            return json.loads(candidate)
        except Exception:
            continue
    raise ValueError(f"Unable to parse {path}")


def font_run(run, size=16, bold=False, color: RGBColor | None = None):
    run.bold = bold
    run.font.name = "TH Sarabun New"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "TH Sarabun New")
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color


def configure_doc(doc: Document):
    style = doc.styles["Normal"]
    style.font.name = "TH Sarabun New"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "TH Sarabun New")
    style.font.size = Pt(16)
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)


def add_title(doc: Document, title: str, subtitle: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font_run(p.add_run(title), size=24, bold=True, color=RGBColor(31, 78, 121))
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font_run(p2.add_run(subtitle), size=18)


def add_heading(doc: Document, text: str, level=1):
    p = doc.add_paragraph()
    font_run(
        p.add_run(text),
        size=18 if level == 1 else 16,
        bold=True,
        color=RGBColor(31, 78, 121) if level == 1 else None,
    )


def add_para(doc: Document, text: str, bold=False):
    p = doc.add_paragraph()
    font_run(p.add_run(text), bold=bold)


def add_bullets(doc: Document, items: list[str]):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        font_run(p.add_run(item))


def add_steps(doc: Document, items: list[str]):
    for idx, item in enumerate(items, start=1):
        p = doc.add_paragraph()
        font_run(p.add_run(f"{idx}. {item}"))


def shade_cell(cell, color="D9EAF7"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    tc_pr.append(shd)


def set_cell(cell, text):
    cell.text = str(text)
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            font_run(run, size=14)


def add_table(doc: Document, headers: list[str], rows: list[tuple]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        set_cell(table.rows[0].cells[i], header)
        shade_cell(table.rows[0].cells[i])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell(cells[i], value)


def add_image(doc: Document, image: str, caption: str):
    path = IMAGE_DIR / image
    if not path.exists():
        add_para(doc, f"[ไม่พบภาพ: {image}]")
        return
    doc.add_picture(str(path), width=Inches(6.35))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font_run(p.add_run(caption), size=14)


def money(value) -> str:
    try:
        return f"{float(value):,.2f}"
    except Exception:
        return str(value)


def journal_rows(lines: list[dict]) -> list[tuple[str, str, str, str]]:
    return [
        (
            f"{line['account_code']} {line['account_name']}",
            line.get("label", ""),
            money(line.get("debit", 0.0)),
            money(line.get("credit", 0.0)),
        )
        for line in lines
    ]


def write_doc(spec: dict):
    doc = Document()
    configure_doc(doc)
    add_title(doc, spec["title"], spec["subtitle"])
    add_heading(doc, "วัตถุประสงค์")
    add_bullets(doc, spec["objectives"])
    add_heading(doc, "ภาพรวมการทำงาน")
    add_bullets(doc, spec["overview"])
    add_heading(doc, "1. เริ่มจากหน้า Dashboard")
    add_para(doc, "เริ่มจากหน้า Dashboard ของระบบ แล้วกดเข้าโมดูลตามภาพ ก่อนจะไปยังเมนูย่อยที่ใช้ทำรายการจริง")
    add_image(doc, spec["dashboard_image"], spec["dashboard_caption"])
    add_heading(doc, "2. ไปที่เมนูที่ใช้")
    add_para(doc, "เมื่อเข้ามาในโมดูลแล้ว ให้ไปตามเมนูด้านล่างนี้ทีละรายการ")
    add_bullets(doc, spec["menu_paths"])
    add_image(doc, spec["menu_image"], spec["menu_caption"])
    add_heading(doc, "3. ขั้นตอนการใช้งาน")
    add_para(doc, spec["usage_intro"])
    add_steps(doc, spec["steps"])
    for caption, image in spec.get("usage_images", []):
        add_image(doc, image, caption)
    add_heading(doc, "คำอธิบายหน้าจอและช่องสำคัญ")
    add_table(doc, ["หัวข้อ", "ใช้ทำอะไร", "ตัวอย่างจากระบบ"], spec["fields"])
    add_heading(doc, "ตัวอย่างการใช้งานในสถานการณ์จริง")
    add_table(doc, ["สถานการณ์", "ควรทำอย่างไร", "ผลที่ควรได้"], spec["scenarios"])
    add_heading(doc, "4. ตรวจสอบ Journal Entry")
    add_para(doc, spec["je_intro"])
    for caption, image in spec["je_images"]:
        add_image(doc, image, caption)
    for heading, rows, explain in spec["je_tables"]:
        add_para(doc, heading, bold=True)
        add_table(doc, ["บัญชี", "คำอธิบาย", "เดบิต", "เครดิต"], rows)
        add_para(doc, explain)
    add_heading(doc, "ข้อควรระวัง")
    add_bullets(doc, spec["cautions"])
    out = DOCX_DIR / spec["filename"]
    doc.save(out)
    return out


menu = load_json(MENU_PATH)
live = load_json(LIVE_PATH)
extra = load_json(EXTRA_PATH)
dep_lines = load_json(DEP_LINES_PATH)
doc_samples = load_json(DOC_SAMPLE_PATH) if DOC_SAMPLE_PATH.exists() else {}

asset_menu = menu["xmlids"]["account_asset.menu_action_account_asset_form"]["path"]
asset_model_menu = menu["xmlids"]["account_asset.menu_action_account_asset_model_form"]["path"]
fixed_asset_report_menu = menu["xmlids"]["account_fixed_asset_report.menu_accounting_fixed_asset_report"]["path"]
mfg_menu = menu["xmlids"]["mrp.menu_mrp_production_action"]["path"]
bom_menu = menu["xmlids"]["mrp.menu_mrp_bom_form_action"]["path"]
valuation_menu = menu["xmlids"]["stock_account.menu_valuation"]["path"]
rng8_menu = menu["xmlids"]["account_stock_card_rng8.menu_account_stock_card_rng8"]["path"]
product_category_menu = menu["searches"]["Product Categories"][0]["path"]
scrap_menu = menu["searches"]["Scrap"][0]["path"]

fa = live["fixed_asset"]
mfg = live["manufacturing"]
dep_rows = journal_rows(dep_lines["lines"])
sale_rows = journal_rows(fa["sale_move"]["lines"])
disposal_rows = journal_rows(fa["disposal_move"]["lines"])
valuation_rows = [journal_rows(x["lines"]) for x in extra["valuation_moves"]]
scrap = doc_samples.get("scrap", {})


def fixed_asset_specs() -> list[dict]:
    draft, running, sell, dispose = fa["draft_asset"], fa["running_asset"], fa["sell_asset"], fa["dispose_asset"]
    return [
        {
            "filename": "6.1_6.1 การจัดการทรัพย์สิน.docx",
            "title": "6.1 การจัดการทรัพย์สิน",
            "subtitle": "Fixed Asset",
            "objectives": [
                "เพื่อให้ผู้ใช้งานเห็นภาพรวมของวงจรทรัพย์สินตั้งแต่สร้าง ใช้งาน ขาย และตัดจำหน่าย",
                "เพื่อให้ค้นหารายการทรัพย์สินและดูสถานะของแต่ละใบได้ง่าย",
                "เพื่อให้ตรวจสอบผลทางบัญชีของทรัพย์สินได้จากรายการจริงในระบบ",
            ],
            "overview": [
                f"ตัวอย่างจริงในระบบมีทั้ง {draft['name']} สถานะ Draft, {running['name']} สถานะ Running, {sell['name']} ที่ขายแล้ว และ {dispose['name']} ที่ตัดจำหน่ายแล้ว",
                "หน้าจอ Assets เป็นหน้าจอกลางที่ใช้เปิดดูข้อมูลทรัพย์สินแต่ละใบ เช่น มูลค่าเริ่มต้น มูลค่าคงเหลือ แบบทรัพย์สิน และประวัติการลงบัญชี",
            ],
            "dashboard_image": "nav_dashboard_accounting_real_annotated.png",
            "dashboard_caption": "รูป 6.1.1 หน้า Dashboard สำหรับเข้าโมดูล Accounting",
            "menu_paths": [asset_menu],
            "menu_image": "nav_accounting_assets_real_annotated.png",
            "menu_caption": "รูป 6.1.2 เมนู Assets ที่ใช้เปิดรายการทรัพย์สิน",
            "usage_intro": "หัวข้อนี้ใช้รายการทรัพย์สินจริงในระบบเพื่อให้ผู้ใช้งานเปิดตามชื่อรายการและเข้าใจความต่างของแต่ละสถานะได้ทันที",
            "steps": [
                "เข้าเมนู Assets แล้วใช้ช่องค้นหาตามชื่อทรัพย์สิน",
                f"เปิด {draft['name']} เพื่อดูรายการที่ยังอยู่ระหว่างเตรียมข้อมูล",
                f"เปิด {running['name']} เพื่อดูรายการที่เริ่มใช้งานแล้วและมีตารางค่าเสื่อมราคา",
                f"เปิด {sell['name']} หรือ {dispose['name']} เพื่อดูรายการที่ปิดแล้วและตรวจสอบ Journal Entry ต่อได้",
            ],
            "usage_images": [
                ("รูป 6.1.3 หน้าจอทรัพย์สินสถานะ Draft", "asset_draft_form_real_annotated.png"),
                ("รูป 6.1.4 หน้าจอทรัพย์สินสถานะ Running", "asset_running_form_real_annotated.png"),
                ("รูป 6.1.5 หน้าจอทรัพย์สินที่ขายแล้ว", "asset_sell_form_real_annotated.png"),
                ("รูป 6.1.6 หน้าจอทรัพย์สินที่ตัดจำหน่ายแล้ว", "asset_dispose_form_real_annotated.png"),
            ],
            "fields": [
                ("Name", "ชื่อทรัพย์สินที่ใช้ค้นหาและอ้างอิงในรายงาน", running["name"]),
                ("Original Value", "มูลค่าเริ่มต้นของทรัพย์สิน", money(running["original_value"])),
                ("Book Value", "มูลค่าคงเหลือหลังหักค่าเสื่อม", money(running["book_value"])),
                ("Asset Model", "แบบทรัพย์สินที่กำหนดบัญชีและวิธีคิดค่าเสื่อม", running["model"]),
                ("Status", "บอกว่ารายการยังร่าง กำลังใช้งาน หรือปิดแล้ว", running["state"]),
            ],
            "scenarios": [
                ("ตรวจทรัพย์สินก่อนเริ่มใช้งาน", "เปิดรายการสถานะ Draft แล้วตรวจข้อมูลให้ครบ", "ยังไม่มีค่าเสื่อมและยังแก้ข้อมูลหลักได้"),
                ("ติดตามทรัพย์สินที่ใช้งานอยู่", "เปิดรายการสถานะ Running แล้วดูตารางค่าเสื่อม", "เห็นมูลค่าคงเหลือและประวัติค่าเสื่อม"),
                ("ตรวจทรัพย์สินที่ปิดแล้ว", "เปิดรายการที่ขายหรือตัดจำหน่ายแล้วแล้วกดดู Journal Entries", "เห็นผลทางบัญชีที่ปิดรายการนั้นครบ"),
            ],
            "je_intro": "เมื่อทรัพย์สินเริ่มใช้งานหรือถูกปิดรายการ ผู้ใช้งานควรเปิด Journal Entry เพื่อยืนยันว่าระบบลงบัญชีถูกต้อง ตัวอย่างด้านล่างเป็นข้อมูลจริงจากระบบ",
            "je_images": [("รูป 6.1.7 ตัวอย่าง Journal Entry ค่าเสื่อมราคาจริง", "journal_asset_depreciation_real_annotated.png")],
            "je_tables": [("รายการค่าเสื่อมราคาที่เกิดขึ้นจริง", dep_rows, "เดบิตอยู่ฝั่งค่าใช้จ่ายค่าเสื่อมราคา และเครดิตอยู่ฝั่งค่าเสื่อมสะสม เพื่อให้มูลค่าตามบัญชีของทรัพย์สินลดลงตามอายุการใช้งาน")],
            "cautions": [
                "ถ้าปรับข้อมูลสำคัญหลังเริ่มใช้งานแล้ว ควรให้ฝ่ายบัญชีตรวจผลกระทบก่อน",
                "อย่าใช้สถานะอย่างเดียวในการสรุปผล ควรเปิด Journal Entry ประกอบทุกครั้ง",
            ],
        },
        {
            "filename": "6.2_6.2 การสร้างทรัพย์สิน.docx",
            "title": "6.2 การสร้างทรัพย์สิน",
            "subtitle": "Fixed Asset",
            "objectives": [
                "เพื่อให้สร้างทรัพย์สินใหม่ได้จากหน้าจอ Assets",
                "เพื่อให้เข้าใจว่าช่วงสร้างรายการยังไม่ใช่ช่วงที่ระบบตัดค่าเสื่อมราคา",
            ],
            "overview": [
                f"ตัวอย่างจริงในระบบคือ {draft['name']} ซึ่งยังอยู่สถานะ Draft และอ้างอิงแบบทรัพย์สิน {draft['model']}",
                "ก่อนสร้างทรัพย์สินควรเปิด Asset Model เพื่อตรวจบัญชีที่ระบบจะนำไปใช้ต่อ",
            ],
            "dashboard_image": "nav_dashboard_accounting_real_annotated.png",
            "dashboard_caption": "รูป 6.2.1 หน้า Dashboard สำหรับเข้าโมดูล Accounting",
            "menu_paths": [asset_model_menu, asset_menu],
            "menu_image": "nav_accounting_asset_models_real_annotated.png",
            "menu_caption": "รูป 6.2.2 เมนู Asset Models และ Assets",
            "usage_intro": "หัวข้อนี้ใช้แบบทรัพย์สินและรายการร่างจริงในระบบเพื่อให้ผู้ใช้งานตรวจข้อมูลก่อนสร้างรายการใหม่ได้ง่าย",
            "steps": [
                "เข้าเมนู Asset Models แล้วเปิดแบบทรัพย์สินที่ต้องการใช้",
                "ตรวจบัญชีทรัพย์สิน บัญชีค่าเสื่อมสะสม และบัญชีค่าใช้จ่ายค่าเสื่อมให้ครบ",
                "กลับไปที่เมนู Assets แล้วกด New",
                "กรอกชื่อทรัพย์สิน มูลค่าเริ่มต้น วันที่ซื้อ และเลือก Asset Model ให้ตรงประเภท",
                "บันทึกรายการเพื่อให้ทรัพย์สินอยู่ในสถานะ Draft รอเริ่มใช้งาน",
            ],
            "usage_images": [
                ("รูป 6.2.3 หน้าจอ Asset Model ที่ใช้กำหนดบัญชี", "asset_model_form_real_annotated.png"),
                ("รูป 6.2.4 หน้าจอทรัพย์สินสถานะ Draft หลังสร้างรายการ", "asset_draft_form_real_annotated.png"),
            ],
            "fields": [
                ("Asset Model", "บอกระบบว่าจะใช้บัญชีและวิธีคิดค่าเสื่อมแบบใด", draft["model"]),
                ("Original Value", "มูลค่าทรัพย์สินที่จะบันทึกเข้าระบบ", money(draft["original_value"])),
                ("Acquisition Date", "วันที่ซื้อหรือวันที่รับทรัพย์สิน", draft["acquisition_date"]),
                ("Fixed Asset Account", "บัญชีทรัพย์สินหลัก", draft["fixed_asset_account"]),
                ("Depreciation Account", "บัญชีค่าเสื่อมสะสม", draft["depreciation_account"]),
            ],
            "scenarios": [
                ("สร้างทรัพย์สินจากข้อมูลซื้อจริง", "เลือก Asset Model ให้ถูกแล้วจึงบันทึก", "ทรัพย์สินถูกสร้างเป็น Draft และพร้อมเริ่มใช้งาน"),
                ("ตรวจ Asset Model ก่อนสร้าง", "เปิดดูบัญชีใน Asset Model ก่อนทุกครั้ง", "ลดความเสี่ยงการลงบัญชีผิดตั้งแต่ต้น"),
            ],
            "je_intro": "การสร้างทรัพย์สินในสถานะ Draft ยังไม่สร้าง Journal Entry ค่าเสื่อมราคาในทันที ด้านล่างจึงแสดงตัวอย่างรายการค่าเสื่อมที่จะเกิดขึ้นหลังเริ่มใช้งาน",
            "je_images": [("รูป 6.2.5 ตัวอย่าง Journal Entry หลังเริ่มใช้งานทรัพย์สิน", "journal_asset_depreciation_real_annotated.png")],
            "je_tables": [("ตัวอย่างบัญชีที่จะถูกใช้หลังเริ่มใช้งาน", dep_rows, "เมื่อเริ่มใช้งานและถึงรอบตัดค่าเสื่อม ระบบจะเดบิตค่าใช้จ่ายค่าเสื่อมและเครดิตค่าเสื่อมสะสมตามแบบทรัพย์สินที่เลือกไว้ตอนสร้าง")],
            "cautions": [
                "ถ้าเลือก Asset Model ผิดตั้งแต่ตอนสร้าง บัญชีที่จะใช้ตัดค่าเสื่อมจะผิดตามไปด้วย",
                "ก่อนบันทึกควรตรวจชื่อทรัพย์สิน มูลค่า และวันที่ซื้อให้ครบ",
            ],
        },
        {
            "filename": "6.3_6.3 การเริ่มคิดค่าเสื่อมราคา (Activation & Depreciation).docx",
            "title": "6.3 การเริ่มคิดค่าเสื่อมราคา (Activation & Depreciation)",
            "subtitle": "Fixed Asset",
            "objectives": [
                "เพื่อให้ผู้ใช้งานเริ่มใช้งานทรัพย์สินและติดตามค่าเสื่อมราคาได้",
                "เพื่อให้ผู้ใช้งานอ่าน Journal Entry ค่าเสื่อมราคาเป็น",
            ],
            "overview": [
                f"ตัวอย่างทรัพย์สินจริงที่เริ่มใช้งานแล้วคือ {running['name']} สถานะ {running['state']}",
                f"ระบบมี Journal Entry ค่าเสื่อมราคาจริงวันที่ {dep_lines['date']} ให้ใช้เป็นตัวอย่างอ้างอิง",
            ],
            "dashboard_image": "nav_dashboard_accounting_real_annotated.png",
            "dashboard_caption": "รูป 6.3.1 หน้า Dashboard สำหรับเข้าโมดูล Accounting",
            "menu_paths": [asset_menu],
            "menu_image": "nav_accounting_assets_real_annotated.png",
            "menu_caption": "รูป 6.3.2 เมนู Assets สำหรับเปิดรายการที่เริ่มใช้งานแล้ว",
            "usage_intro": "หัวข้อนี้ใช้ทรัพย์สินจริงที่อยู่ในสถานะ Running เพื่อให้ผู้ใช้งานเห็นภาพหลังเริ่มใช้งานและตรวจรายการค่าเสื่อมที่เกิดขึ้นจริง",
            "steps": [
                "เข้าเมนู Assets แล้วเปิดทรัพย์สินที่อยู่ในสถานะ Running",
                "ตรวจมูลค่าเริ่มต้น มูลค่าคงเหลือ และตารางค่าเสื่อมราคา",
                "เมื่อถึงรอบตัดค่าเสื่อม ให้เปิด Journal Entry ของรอบนั้นเพื่อตรวจบัญชีและยอดเงิน",
            ],
            "usage_images": [("รูป 6.3.3 หน้าจอทรัพย์สินสถานะ Running", "asset_running_form_real_annotated.png")],
            "fields": [
                ("Status", "บอกว่าทรัพย์สินเริ่มใช้งานแล้ว", running["state"]),
                ("Original Value", "มูลค่าตั้งต้นของทรัพย์สิน", money(running["original_value"])),
                ("Book Value", "มูลค่าที่เหลืออยู่หลังหักค่าเสื่อม", money(running["book_value"])),
                ("Depreciation Board", "ตารางรอบการตัดค่าเสื่อมราคา", "เปิดดูจากแท็บค่าเสื่อมในรายการทรัพย์สิน"),
            ],
            "scenarios": [
                ("ติดตามค่าเสื่อมตามงวด", "เปิดรายการสถานะ Running และดูตารางค่าเสื่อม", "เห็นทั้งมูลค่าคงเหลือและรายการบัญชีที่ถูกตัดจริง"),
                ("ตรวจงวดค่าเสื่อมผิดปกติ", "เปิด Journal Entry ของรอบนั้นแล้วเทียบบัญชี", "ช่วยหาสาเหตุได้ว่าผิดที่ข้อมูลหรือผิดที่แบบทรัพย์สิน"),
            ],
            "je_intro": "ตัวอย่าง Journal Entry ด้านล่างมาจากรายการค่าเสื่อมราคาจริงในระบบ ผู้ใช้งานควรดูชื่อบัญชีและยอดเดบิตเครดิตทุกครั้ง",
            "je_images": [("รูป 6.3.4 Journal Entry ค่าเสื่อมราคาจริง", "journal_asset_depreciation_real_annotated.png")],
            "je_tables": [("รายการบัญชีที่เกิดขึ้นเมื่อตัดค่าเสื่อมราคา", dep_rows, "ค่าใช้จ่ายค่าเสื่อมราคาอยู่ฝั่งเดบิต เพราะเป็นต้นทุนของงวด ส่วนค่าเสื่อมราคาสะสมอยู่ฝั่งเครดิตเพื่อสะสมยอดที่หักออกจากทรัพย์สิน")],
            "cautions": [
                "ถ้าทรัพย์สินยังอยู่ใน Draft ระบบจะยังไม่สร้าง Journal Entry ค่าเสื่อม",
                "ถ้าต้องแก้ข้อมูลเริ่มต้น ควรให้ฝ่ายบัญชีตรวจผลกระทบของตารางค่าเสื่อมก่อน",
            ],
        },
        {
            "filename": "6.4_6.4 การขายสินทรัพย์ (Selling Assets & Gain_Loss).docx",
            "title": "6.4 การขายสินทรัพย์ (Selling Assets & Gain/Loss)",
            "subtitle": "Fixed Asset",
            "objectives": [
                "เพื่อให้ผู้ใช้งานขายทรัพย์สินและตรวจผลกำไรหรือขาดทุนจากการขายได้",
                "เพื่อให้เชื่อมข้อมูลระหว่างใบขายทรัพย์สินกับ Journal Entry ได้",
            ],
            "overview": [
                f"ตัวอย่างจริงในระบบคือทรัพย์สิน {sell['name']} และใบขาย {fa['sale_invoice']['name']}",
                "หลังขายทรัพย์สิน ระบบจะปิดรายการทรัพย์สินและสร้าง Journal Entry เพื่อตัดบัญชีทรัพย์สิน ตัดค่าเสื่อมสะสม และรับรู้ผลต่างจากการขาย",
            ],
            "dashboard_image": "nav_dashboard_accounting_real_annotated.png",
            "dashboard_caption": "รูป 6.4.1 หน้า Dashboard สำหรับเข้าโมดูล Accounting",
            "menu_paths": [asset_menu, "Accounting > Customers > Invoices"],
            "menu_image": "nav_accounting_customers_invoices_real_annotated.png",
            "menu_caption": "รูป 6.4.2 เมนู Assets และ Invoices ที่ใช้ติดตามการขายทรัพย์สิน",
            "usage_intro": "หัวข้อนี้ใช้ข้อมูลจริงของการขายทรัพย์สินในระบบเพื่อให้ผู้ใช้งานตรวจขั้นตอนตั้งแต่ใบขายจนถึง Journal Entry ได้ครบ",
            "steps": [
                f"เปิดใบขาย {fa['sale_invoice']['name']} เพื่อดูยอดขายและเลขอ้างอิง",
                f"เปิดทรัพย์สิน {sell['name']} เพื่อตรวจว่าสถานะปิดรายการแล้ว",
                "เปิด Journal Entry ที่เกี่ยวข้องเพื่อดูการตัดทรัพย์สิน ตัดค่าเสื่อมสะสม และบันทึกผลต่างจากการขาย",
            ],
            "usage_images": [
                ("รูป 6.4.3 ใบขายทรัพย์สินจริงในระบบ", "asset_sale_invoice_real_annotated.png"),
                ("รูป 6.4.4 หน้าจอทรัพย์สินที่ขายแล้ว", "asset_sell_form_real_annotated.png"),
            ],
            "fields": [
                ("Sale Invoice", "เอกสารขายที่ใช้เรียกเก็บเงินจากผู้ซื้อ", fa["sale_invoice"]["name"]),
                ("Sale Reference", "เลขอ้างอิงที่เชื่อมโยงการขายทรัพย์สิน", fa["sale_invoice"]["ref"]),
                ("Asset State", "บอกว่าทรัพย์สินปิดรายการแล้วหลังการขาย", sell["state"]),
                ("Book Value", "มูลค่าตามบัญชีก่อนขาย", money(sell["book_value"])),
            ],
            "scenarios": [
                ("ขายทรัพย์สินที่ยังมีมูลค่าคงเหลือ", "เปิดใบขายและทรัพย์สินคู่กัน แล้วตรวจ Journal Entry หลังขาย", "เห็นว่าทรัพย์สินถูกตัดออกจากบัญชีและมีผลต่างจากการขาย"),
                ("ตรวจผลกำไรหรือขาดทุน", "ดูบัญชีรายได้อื่นหรือบัญชีต้นทุนที่รับผลต่าง", "รู้ได้ทันทีว่าขายแล้วได้กำไรหรือขาดทุน"),
            ],
            "je_intro": "ตัวอย่างด้านล่างเป็น Journal Entry จริงของการขายทรัพย์สิน ผู้ใช้ควรไล่ดูทีละบรรทัดว่าระบบตัดบัญชีทรัพย์สิน ตัดค่าเสื่อมสะสม และลงผลต่างอย่างไร",
            "je_images": [("รูป 6.4.5 Journal Entry ของการขายทรัพย์สินจริง", "journal_asset_sale_real_annotated.png")],
            "je_tables": [("รายการบัญชีที่เกิดขึ้นเมื่อขายทรัพย์สิน", sale_rows, "บัญชีทรัพย์สินถูกเครดิตออกจากระบบ บัญชีค่าเสื่อมสะสมถูกเดบิตเพื่อล้างยอดสะสมเดิม ส่วนผลต่างจากการขายจะไปอยู่ในบัญชีรายได้หรือกำไรขาดทุนตามที่ระบบคำนวณได้")],
            "cautions": [
                "ควรตรวจว่าทรัพย์สินไม่มีรายการใช้งานค้างอยู่ก่อนขาย",
                "ราคาขายในเอกสารต้องตรงกับข้อมูลจริง เพราะมีผลต่อกำไรหรือขาดทุน",
            ],
        },
        {
            "filename": "6.5_6.5 การตัดจำหน่ายสินทรัพย์ (Disposal _ Scrap).docx",
            "title": "6.5 การตัดจำหน่ายสินทรัพย์ (Disposal / Scrap)",
            "subtitle": "Fixed Asset",
            "objectives": [
                "เพื่อให้ผู้ใช้งานปิดรายการทรัพย์สินที่เลิกใช้ ชำรุด หรือไม่ใช้งานแล้วได้ถูกต้อง",
                "เพื่อให้ตรวจสอบ Journal Entry ของการตัดจำหน่ายได้",
            ],
            "overview": [
                f"ตัวอย่างจริงคือทรัพย์สิน {dispose['name']} ที่ถูกปิดรายการแล้ว",
                "การตัดจำหน่ายจะตัดมูลค่าทรัพย์สินออกจากบัญชีและล้างค่าเสื่อมสะสมตามยอดที่เกิดขึ้นจริง",
            ],
            "dashboard_image": "nav_dashboard_accounting_real_annotated.png",
            "dashboard_caption": "รูป 6.5.1 หน้า Dashboard สำหรับเข้าโมดูล Accounting",
            "menu_paths": [asset_menu],
            "menu_image": "nav_accounting_assets_real_annotated.png",
            "menu_caption": "รูป 6.5.2 เมนู Assets ที่ใช้เปิดรายการเพื่อตัดจำหน่าย",
            "usage_intro": "หัวข้อนี้ใช้ทรัพย์สินจริงที่ถูกตัดจำหน่ายแล้วเพื่ออธิบายขั้นตอนตรวจสอบหลังปิดรายการ",
            "steps": [
                f"เข้าเมนู Assets แล้วเปิดทรัพย์สิน {dispose['name']}",
                "ตรวจว่าสถานะปิดแล้วและไม่มีรอบค่าเสื่อมใหม่เกิดขึ้นต่อ",
                "เปิด Journal Entry ที่เกี่ยวข้องเพื่อดูการตัดบัญชีทรัพย์สินและค่าเสื่อมสะสม",
            ],
            "usage_images": [("รูป 6.5.3 หน้าจอทรัพย์สินที่ตัดจำหน่ายแล้ว", "asset_dispose_form_real_annotated.png")],
            "fields": [
                ("Asset State", "สถานะหลังตัดจำหน่ายแล้ว", dispose["state"]),
                ("Book Value", "มูลค่าคงเหลือก่อนปิดรายการ", money(dispose["book_value"])),
                ("Disposal Date", "วันที่ปิดรายการ", fa["disposal_move"]["date"]),
            ],
            "scenarios": [
                ("ตัดจำหน่ายทรัพย์สินที่เลิกใช้", "เปิดรายการที่ปิดแล้วและตรวจ Journal Entry", "ทรัพย์สินและค่าเสื่อมสะสมถูกล้างออกจากบัญชี"),
                ("ตรวจผลต่างจากการตัดจำหน่าย", "ดูบรรทัดบัญชีที่รับผลต่าง", "เห็นผลกระทบทางบัญชีชัดเจน"),
            ],
            "je_intro": "Journal Entry ด้านล่างเป็นข้อมูลจริงของการตัดจำหน่าย ผู้ใช้งานควรดูว่ามีการล้างบัญชีทรัพย์สินและค่าเสื่อมสะสมครบหรือไม่",
            "je_images": [("รูป 6.5.4 Journal Entry ของการตัดจำหน่ายสินทรัพย์", "journal_asset_disposal_real_annotated.png")],
            "je_tables": [("รายการบัญชีที่เกิดขึ้นเมื่อตัดจำหน่ายทรัพย์สิน", disposal_rows, "ระบบจะเครดิตบัญชีทรัพย์สินเพื่อตัดยอดออก เดบิตค่าเสื่อมสะสมเพื่อล้างยอดเดิม และบันทึกผลต่างในบัญชีต้นทุนหรือขาดทุนจากการตัดจำหน่าย")],
            "cautions": [
                "ไม่ควรตัดจำหน่ายทรัพย์สินที่ยังใช้งานจริงหรือยังมีเอกสารค้างอยู่",
                "ควรตรวจ Journal Entry ทุกครั้งเพื่อยืนยันว่าบัญชีผลต่างไปลงที่บัญชีที่กิจการต้องการ",
            ],
        },
        {
            "filename": "6.6_6.6 รายงานทรัพย์สินถาวร (Fixed Asset Report).docx",
            "title": "6.6 รายงานทรัพย์สินถาวร (Fixed Asset Report)",
            "subtitle": "Fixed Asset",
            "objectives": [
                "เพื่อให้ผู้ใช้งานเปิดรายงานทรัพย์สินถาวรและตรวจข้อมูลทรัพย์สินแต่ละใบได้",
                "เพื่อใช้เทียบรายงานกับ Journal Entry ของค่าเสื่อมราคาและการปิดรายการทรัพย์สิน",
            ],
            "overview": [
                "รายงานสินทรัพย์ถาวรใช้สรุปมูลค่าเริ่มต้น มูลค่าคงเหลือ และความเคลื่อนไหวของทรัพย์สินในช่วงวันที่ที่เลือก",
                "แม้ตัวรายงานจะไม่สร้างรายการบัญชีเอง แต่เป็นหน้าจอหลักที่ผู้ใช้งานใช้ตรวจว่ารายการบัญชีที่เกิดขึ้นสอดคล้องกับทรัพย์สินแต่ละใบหรือไม่",
            ],
            "dashboard_image": "nav_dashboard_accounting_real_annotated.png",
            "dashboard_caption": "รูป 6.6.1 หน้า Dashboard สำหรับเข้าโมดูล Accounting",
            "menu_paths": [fixed_asset_report_menu],
            "menu_image": "nav_accounting_fixed_asset_report_real_annotated.png",
            "menu_caption": "รูป 6.6.2 เมนูรายงานทรัพย์สินถาวร",
            "usage_intro": "หัวข้อนี้ใช้หน้ารายงานจริงในระบบเพื่อให้ผู้ใช้งานเปิดรายงานและตรวจรายการทรัพย์สินที่ต้องการได้ทันที",
            "steps": [
                "เข้าเมนูรายงานทรัพย์สินถาวร",
                "เลือกช่วงวันที่หรือกรองตามแบบทรัพย์สินตามที่ต้องการ",
                "เปิดดูรายการแต่ละใบและเทียบกับหน้าจอทรัพย์สินหรือ Journal Entry ที่เกี่ยวข้อง",
            ],
            "usage_images": [("รูป 6.6.3 หน้ารายงานทรัพย์สินถาวรจริงในระบบ", "fixed_asset_report_page_real_annotated.png")],
            "fields": [
                ("Date Filter", "กำหนดช่วงวันที่ที่ต้องการดูรายการ", "ใช้ช่วงเดียวกับงวดที่ต้องการปิดบัญชี"),
                ("Asset Lines", "แสดงทรัพย์สินแต่ละใบที่เข้าเงื่อนไขรายงาน", "ตรวจชื่อทรัพย์สินและมูลค่าคงเหลือได้"),
                ("Open Asset", "กดเจาะเข้าไปดูรายการทรัพย์สินจริง", "ใช้ตรวจทานข้อมูลต่อจากรายงาน"),
            ],
            "scenarios": [
                ("ตรวจทรัพย์สินก่อนปิดงวด", "เปิดรายงานแล้วกรองช่วงวันที่ของงวด", "เห็นทรัพย์สินที่มียอดเคลื่อนไหวในงวดนั้น"),
                ("เทียบรายงานกับค่าเสื่อม", "เปิดทรัพย์สินและ Journal Entry ของค่าเสื่อมจากรายงานต่อ", "ตรวจสอบความถูกต้องได้ครบทั้งรายงานและบัญชี"),
            ],
            "je_intro": "ตัวอย่างด้านล่างใช้ Journal Entry ค่าเสื่อมราคาจริงเพื่อให้ผู้ใช้เห็นวิธีตรวจเทียบระหว่างรายงานกับบัญชี เมื่อยอดในรายงานเปลี่ยน ควรมีรายการบัญชีรองรับเสมอ",
            "je_images": [("รูป 6.6.4 ตัวอย่าง Journal Entry ค่าเสื่อมราคาที่ใช้เทียบกับรายงาน", "journal_asset_depreciation_real_annotated.png")],
            "je_tables": [("รายการบัญชีที่ใช้ตรวจเทียบกับรายงานทรัพย์สิน", dep_rows, "เมื่อรายงานแสดงว่ามูลค่าคงเหลือลดลง ผู้ใช้ควรเห็นค่าใช้จ่ายค่าเสื่อมและค่าเสื่อมสะสมใน Journal Entry ตรงกับงวดเดียวกัน")],
            "cautions": [
                "รายงานเป็นเครื่องมือสรุปข้อมูล ไม่ใช่หน้าจอแก้ไขรายการโดยตรง",
                "หากรายงานไม่ตรงกับความคาดหวัง ควรย้อนกลับไปตรวจทรัพย์สินแต่ละใบและ Journal Entry ที่เกี่ยวข้อง",
            ],
        },
    ]


def manufacturing_specs() -> list[dict]:
    mo = mfg["sample_mo"]
    fg_cat, rm_cat = mfg["categories"]["fg"], mfg["categories"]["rm"]
    v_rows = mfg["sample_valuation_rows"]
    return [
        {
            "filename": "7.1_7.1 ภาพรวมการไหลของข้อมูล (Overview Flow).docx",
            "title": "7.1 ภาพรวมการไหลของข้อมูล (Overview Flow)",
            "subtitle": "Manufacturing",
            "objectives": [
                "เพื่อให้ผู้ใช้งานเห็นลำดับการทำงานตั้งแต่ BoM, Manufacturing Order, ตัดวัตถุดิบ และรับสินค้าสำเร็จรูป",
                "เพื่อให้เชื่อมหน้าจอการผลิตกับผลทางบัญชีได้",
            ],
            "overview": [
                f"ตัวอย่างจริงใช้ใบสั่งผลิต {mo['name']} ของสินค้า {mo['product']}",
                "ลำดับงานจริงคือ BoM กำหนดส่วนประกอบ จากนั้น MO จะตัดวัตถุดิบจากคลังลอยเข้า Production และรับสินค้าสำเร็จรูปกลับเข้าคลัง",
            ],
            "dashboard_image": "nav_dashboard_manufacturing_real_annotated.png",
            "dashboard_caption": "รูป 7.1.1 หน้า Dashboard สำหรับเข้าโมดูล Manufacturing",
            "menu_paths": [mfg_menu, bom_menu],
            "menu_image": "nav_manufacturing_orders_real_annotated.png",
            "menu_caption": "รูป 7.1.2 เมนู Manufacturing Orders และ Bills of Materials",
            "usage_intro": "หัวข้อนี้ใช้ MO จริงในระบบเพื่อให้ผู้ใช้งานตามเส้นทางการทำงานจากเอกสารหลักไปจนถึงผลทางบัญชีได้ครบ",
            "steps": [
                "เปิด Bills of Materials ของสินค้าที่ต้องการผลิตเพื่อดูส่วนประกอบ",
                f"เปิด Manufacturing Order {mo['name']} เพื่อดูสินค้า วันที่เริ่ม และวันที่เสร็จ",
                "หลังงานเสร็จ ให้เปิด Journal Entry ของการตัดวัตถุดิบและรับสินค้าสำเร็จรูปเพื่อตรวจผลทางบัญชี",
            ],
            "usage_images": [
                ("รูป 7.1.3 หน้าจอ Manufacturing Order จริงในระบบ", "manufacturing_order_form_real_annotated.png"),
                ("รูป 7.1.4 หน้าจอ BoM ของสินค้าที่ใช้ในตัวอย่าง", "manufacturing_bom_form_real_annotated.png"),
            ],
            "fields": [
                ("Manufacturing Order", "เอกสารหลักสำหรับสั่งผลิตและติดตามสถานะงาน", mo["name"]),
                ("Product", "สินค้าสำเร็จรูปที่กำลังผลิต", mo["product"]),
                ("BoM", "โครงสร้างส่วนประกอบที่ระบบใช้เบิกวัตถุดิบ", extra["bom"]["display_name"]),
                ("Picking Type", "บอกสายการผลิตหรือประเภทใบงาน", mo["picking_type"]),
            ],
            "scenarios": [
                ("ติดตามใบสั่งผลิตจริง", "เปิด MO แล้วดู BoM และ Journal Entry ต่อเนื่องกัน", "เข้าใจได้ว่าหน้าจอการผลิตเชื่อมกับบัญชีอย่างไร"),
                ("หาต้นทุนของงานผลิต", "ดูบรรทัดวัตถุดิบและบรรทัดรับสินค้าสำเร็จรูป", "เห็นได้ว่าต้นทุนไหลผ่านบัญชีงานระหว่างทำอย่างไร"),
            ],
            "je_intro": "Journal Entry ด้านล่างแสดงการตัดวัตถุดิบและการรับสินค้าสำเร็จรูปจากใบผลิตจริงเดียวกัน เพื่อให้ผู้ใช้เห็นเส้นทางต้นทุนอย่างง่าย",
            "je_images": [
                ("รูป 7.1.5 Journal Entry ตอนตัดวัตถุดิบเข้าสู่งานระหว่างทำ", "journal_mfg_raw_fg02001_real_annotated.png"),
                ("รูป 7.1.6 Journal Entry ตอนรับสินค้าสำเร็จรูปเข้าคลัง", "journal_mfg_finished_real_annotated.png"),
            ],
            "je_tables": [
                ("ตัวอย่างตัดวัตถุดิบเข้าบัญชีงานระหว่างทำ", valuation_rows[0], "บัญชีงานระหว่างทำถูกเดบิตเพื่อรับต้นทุนเข้าไว้ชั่วคราว และบัญชีสินค้าถูกเครดิตออกจากคลัง"),
                ("ตัวอย่างรับสินค้าสำเร็จรูปเข้าคลัง", valuation_rows[2], "เมื่อผลิตเสร็จ ระบบย้ายต้นทุนออกจากงานระหว่างทำไปยังบัญชีสินค้าสำเร็จรูป"),
            ],
            "cautions": [
                "ถ้า BoM ไม่ครบหรือปริมาณไม่ถูกต้อง ต้นทุนที่เกิดใน Journal Entry ก็จะผิดตามไปด้วย",
                "ควรเปิดดูทั้ง MO และ Journal Entry ทุกครั้งเมื่อมีคำถามเรื่องต้นทุนผลิต",
            ],
        },
        {
            "filename": "7.2_7.2 การบันทึกบัญชีในแต่ละขั้นตอน (Journal Entries Steps).docx",
            "title": "7.2 การบันทึกบัญชีในแต่ละขั้นตอน (Journal Entries Steps)",
            "subtitle": "Manufacturing",
            "objectives": [
                "เพื่อให้ผู้ใช้งานแยกความเข้าใจว่าระบบลงบัญชีตอนไหนบ้างระหว่างการผลิต",
                "เพื่อให้ผู้ใช้งานอ่านเดบิตเครดิตของการผลิตแต่ละขั้นตอนเป็น",
            ],
            "overview": [
                f"ตัวอย่างจริงใช้ Journal Entry ของ {mo['name']} จำนวน 3 รายการ คือ ตัดกึ่งสำเร็จรูป ตัดบรรจุภัณฑ์ และรับสินค้าสำเร็จรูป",
                "ผู้ใช้งานควรดูทีละขั้นว่าอะไรออกจากคลัง อะไรเข้าบัญชีงานระหว่างทำ และอะไรกลับเข้าคลังสินค้าสำเร็จรูป",
            ],
            "dashboard_image": "nav_dashboard_manufacturing_real_annotated.png",
            "dashboard_caption": "รูป 7.2.1 หน้า Dashboard สำหรับเข้าโมดูล Manufacturing",
            "menu_paths": [mfg_menu, valuation_menu],
            "menu_image": "nav_inventory_valuation_real_annotated.png",
            "menu_caption": "รูป 7.2.2 เมนู Manufacturing Orders และ Valuation",
            "usage_intro": "หัวข้อนี้สรุปเฉพาะมุมบัญชีของการผลิต โดยใช้ Journal Entry จริงของระบบเป็นหลักฐานตรง",
            "steps": [
                "เปิดใบสั่งผลิตจริงที่ต้องการตรวจสอบ",
                "เข้าเมนู Valuation หรือเปิด Journal Entries ที่ผูกกับใบผลิตนั้น",
                "ดูบรรทัดตัดวัตถุดิบทีละรายการ",
                "ดูบรรทัดรับสินค้าสำเร็จรูปเมื่อผลิตเสร็จ",
            ],
            "usage_images": [],
            "fields": [
                ("Raw Material Issue", "บรรทัดตัดวัตถุดิบหรือกึ่งสำเร็จรูปออกจากคลัง", extra["valuation_moves"][0]["name"]),
                ("Packaging Issue", "บรรทัดตัดบรรจุภัณฑ์ออกจากคลัง", extra["valuation_moves"][1]["name"]),
                ("Finished Receipt", "บรรทัดรับสินค้าสำเร็จรูปเข้าคลัง", extra["valuation_moves"][2]["name"]),
                ("WIP Account", "บัญชีงานระหว่างทำที่รับต้นทุนระหว่างผลิต", "116021 งานระหว่างทำ"),
            ],
            "scenarios": [
                ("ตรวจการลงบัญชีระหว่างผลิต", "เปิด Journal Entry ทั้ง 3 ใบและไล่ดูทีละขั้น", "เห็นชัดว่าต้นทุนถูกย้ายจาก RM ไป WIP และจาก WIP ไป FG"),
                ("หาสาเหตุที่ต้นทุน FG สูงหรือต่ำ", "ดูยอดวัตถุดิบที่ถูกตัดและยอดรับเข้าของ FG", "ช่วยแยกได้ว่าปัญหาอยู่ที่ BoM การเบิก หรือการรับผลผลิต"),
            ],
            "je_intro": "ด้านล่างคือ Journal Entry จริงทีละขั้นตอนของงานผลิตเดียวกัน ผู้ใช้งานควรอ่านจากบนลงล่างตามลำดับการเกิดรายการ",
            "je_images": [
                ("รูป 7.2.3 Journal Entry ตัดกึ่งสำเร็จรูปเข้าสู่งานระหว่างทำ", "journal_mfg_raw_fg02001_real_annotated.png"),
                ("รูป 7.2.4 Journal Entry ตัดบรรจุภัณฑ์เข้าสู่งานระหว่างทำ", "journal_mfg_raw_packaging_real_annotated.png"),
                ("รูป 7.2.5 Journal Entry รับสินค้าสำเร็จรูปเข้าคลัง", "journal_mfg_finished_real_annotated.png"),
            ],
            "je_tables": [
                ("ขั้นที่ 1 ตัดกึ่งสำเร็จรูป", valuation_rows[0], "ระบบเครดิตบัญชีสินค้าเดิมและเดบิตบัญชีงานระหว่างทำ เพื่อเก็บต้นทุนของชิ้นส่วนที่นำไปผลิต"),
                ("ขั้นที่ 2 ตัดบรรจุภัณฑ์", valuation_rows[1], "บรรจุภัณฑ์ก็ถูกย้ายเข้าไปอยู่ในบัญชีงานระหว่างทำเช่นเดียวกัน"),
                ("ขั้นที่ 3 รับสินค้าสำเร็จรูป", valuation_rows[2], "เมื่อผลิตเสร็จ ระบบเครดิตงานระหว่างทำและเดบิตบัญชีสินค้าสำเร็จรูป"),
            ],
            "cautions": [
                "ถ้าไม่เห็น Journal Entry บางขั้น อาจเกิดจากงานยังไม่ Done หรือหมวดสินค้าไม่ได้ตั้งเป็น real-time valuation",
                "ควรเทียบเลขอ้างอิงของ Journal Entry กับเลข MO ทุกครั้งเพื่อป้องกันเปิดดูผิดใบ",
            ],
        },
        {
            "filename": "7.3_7.3 สรุปผังการตั้งค่าทางบัญชี (Configuration Guide).docx",
            "title": "7.3 สรุปผังการตั้งค่าทางบัญชี (Configuration Guide)",
            "subtitle": "Manufacturing",
            "objectives": [
                "เพื่อให้รู้ว่าควรตรวจหมวดสินค้าและ BoM ที่จุดใดเมื่ออยากเช็กการลงบัญชีการผลิต",
                "เพื่อให้เห็นว่าบัญชีที่ถูกใช้จริงมาจากการตั้งค่าใด",
            ],
            "overview": [
                f"หมวด {fg_cat['name']} และ {rm_cat['name']} ใช้ cost method แบบ {fg_cat['cost_method']} และ valuation แบบ {fg_cat['valuation']}",
                "การตั้งค่าหมวดสินค้าเป็นจุดสำคัญ เพราะกำหนดบัญชีสินค้าขาเข้า ขาออก และบัญชีมูลค่าสินค้าคงคลังที่ใช้ใน Journal Entry",
            ],
            "dashboard_image": "nav_dashboard_inventory_real_annotated.png",
            "dashboard_caption": "รูป 7.3.1 หน้า Dashboard สำหรับเข้าโมดูล Inventory",
            "menu_paths": [product_category_menu, bom_menu],
            "menu_image": "nav_manufacturing_bom_real_annotated.png",
            "menu_caption": "รูป 7.3.2 เมนู Product Categories และ Bills of Materials",
            "usage_intro": "หัวข้อนี้ใช้ข้อมูลจริงของหมวด FG, RM และ BoM เพื่ออธิบายว่าจุดตั้งค่าทางบัญชีอยู่ตรงไหน",
            "steps": [
                "เปิด Product Categories แล้วตรวจหมวด FG ว่าตั้ง cost method และ valuation ถูกต้อง",
                "เปิดหมวด RM เพื่อดูบัญชีมูลค่าวัตถุดิบที่ระบบใช้เวลาตัดของเข้าสู่การผลิต",
                "เปิด BoM เพื่อยืนยันว่าส่วนประกอบที่ใช้ในงานผลิตเป็นรายการที่ถูกต้อง",
            ],
            "usage_images": [
                ("รูป 7.3.3 หมวดสินค้าสำเร็จรูปที่ใช้จริงในระบบ", "product_category_fg_form_real_annotated.png"),
                ("รูป 7.3.4 หมวดวัตถุดิบที่ใช้จริงในระบบ", "product_category_rm_form_real_annotated.png"),
                ("รูป 7.3.5 BoM ที่ใช้จริงในงานผลิตตัวอย่าง", "manufacturing_bom_form_real_annotated.png"),
            ],
            "fields": [
                ("FG Stock Valuation", "บัญชีมูลค่าสินค้าสำเร็จรูป", fg_cat["stock_valuation"]),
                ("RM Stock Valuation", "บัญชีมูลค่าวัตถุดิบ", rm_cat["stock_valuation"]),
                ("Stock Input", "บัญชีรับสินค้าเข้า", fg_cat["stock_input"]),
                ("Stock Output", "บัญชีจ่ายสินค้าออก", fg_cat["stock_output"]),
            ],
            "scenarios": [
                ("ตรวจ setup ก่อนถามเรื่องบัญชี", "เปิดหมวดสินค้าและ BoM ก่อนดู Journal Entry", "ช่วยบอกได้ว่าปัญหาเกิดจากตั้งค่าหรือจากการใช้งาน"),
                ("เทียบบัญชีที่ตั้งกับบัญชีที่ลงจริง", "ดูหมวดสินค้าแล้วเปิด Journal Entry ของ MO เดียวกัน", "ยืนยันได้ว่าระบบใช้บัญชีจาก setup ที่กำหนดไว้จริง"),
            ],
            "je_intro": "ตัวอย่างด้านล่างใช้ Journal Entry รับสินค้าสำเร็จรูปจริง เพื่อชี้ให้เห็นว่าบัญชีในหมวดสินค้าถูกนำมาใช้ตอนผลิตเสร็จอย่างไร",
            "je_images": [("รูป 7.3.6 Journal Entry รับสินค้าสำเร็จรูปที่อ้างอิงจากหมวดสินค้า", "journal_mfg_finished_real_annotated.png")],
            "je_tables": [("รายการบัญชีที่ใช้เทียบกับหมวดสินค้า", valuation_rows[2], "บัญชีงานระหว่างทำและบัญชีสินค้าสำเร็จรูปที่เห็นใน Journal Entry ควรตรงกับสิ่งที่ตั้งไว้ในหมวดสินค้าและโครงสร้างการผลิต")],
            "cautions": [
                "ถ้าหมวดสินค้าใช้ standard cost หรือ manual valuation ผลทางบัญชีจะไม่เหมือนตัวอย่างนี้",
                "การแก้หมวดสินค้าในระหว่างมีธุรกรรมแล้วควรทำโดยฝ่ายที่ดูแลบัญชีและคลังร่วมกัน",
            ],
        },
        {
            "filename": "7.4_7.4 ตัวอย่าง.docx",
            "title": "7.4 ตัวอย่าง",
            "subtitle": "Manufacturing",
            "objectives": [
                "เพื่อให้ผู้ใช้งานเห็นตัวอย่างจริงทั้งเส้นของงานผลิตหนึ่งใบ",
                "เพื่อให้เปิดตามเลขเอกสารจริงแล้วเข้าใจลำดับเหตุการณ์ได้ทันที",
            ],
            "overview": [
                f"ตัวอย่างจริงใช้ใบผลิต {mo['name']} ซึ่งผลิตสินค้า {mo['product']} จำนวน {money(mo['qty'])} หน่วยและผลิตเสร็จแล้ว {money(mo['qty_produced'])} หน่วย",
                "จะเชื่อมให้เห็นทั้งหน้าจอ MO, BoM, รายการตัดวัตถุดิบ และรายการรับสินค้าสำเร็จรูป",
            ],
            "dashboard_image": "nav_dashboard_manufacturing_real_annotated.png",
            "dashboard_caption": "รูป 7.4.1 หน้า Dashboard สำหรับเข้าโมดูล Manufacturing",
            "menu_paths": [mfg_menu, valuation_menu],
            "menu_image": "nav_manufacturing_orders_real_annotated.png",
            "menu_caption": "รูป 7.4.2 เมนู Manufacturing Orders และ Valuation ที่ใช้ตามตัวอย่างจริง",
            "usage_intro": "หัวข้อนี้เขียนเป็นกรณีตัวอย่างเดียวต่อเนื่อง เพื่อให้ผู้ใช้งานเปิดตามแล้วเข้าใจเร็วที่สุด",
            "steps": [
                f"เปิด MO {mo['name']} และตรวจสินค้า วันที่เริ่มผลิต และวันที่ผลิตเสร็จ",
                "เปิด BoM ของสินค้าเดียวกันเพื่อตรวจส่วนประกอบที่ใช้",
                "เปิด Journal Entry ของวัตถุดิบที่ถูกตัดออกจากคลัง",
                "เปิด Journal Entry ของสินค้าสำเร็จรูปที่รับเข้าคลังหลังผลิตเสร็จ",
            ],
            "usage_images": [
                ("รูป 7.4.3 ใบสั่งผลิตจริงที่ใช้เป็นตัวอย่าง", "manufacturing_order_form_real_annotated.png"),
                ("รูป 7.4.4 Journal Entry ตัดวัตถุดิบจริง", "journal_mfg_raw_fg02001_real_annotated.png"),
                ("รูป 7.4.5 Journal Entry รับสินค้าสำเร็จรูปจริง", "journal_mfg_finished_real_annotated.png"),
            ],
            "fields": [
                ("MO Number", "เลขใบสั่งผลิต", mo["name"]),
                ("MO State", "สถานะงานผลิต", mo["state"]),
                ("Start / Finish", "เวลาเริ่มและเวลาจบการผลิต", f"{mo['date_start']} / {mo['date_finished']}"),
                ("Valuation Reference", "เลขอ้างอิงใน Journal Entry และ Stock Valuation", extra["valuation_moves"][2]["name"]),
            ],
            "scenarios": [
                ("ตัวอย่างงานผลิตเสร็จสมบูรณ์", "เปิด MO แล้วตามไปดู JE ขาวัตถุดิบและขา FG", "เห็นต้นทุนไหลครบจากวัตถุดิบไปสินค้าสำเร็จรูป"),
                ("ใช้เป็นต้นแบบเวลาตรวจงานจริง", "ถ้ามีคำถามเรื่องบัญชีของ MO อื่น ให้ไล่แบบเดียวกับตัวอย่างนี้", "ช่วยให้ตรวจงานเป็นลำดับและไม่ตกหล่น"),
            ],
            "je_intro": "ตัวอย่างนี้แสดง Journal Entry ที่สำคัญที่สุด 2 จุด คือ ตัดวัตถุดิบ และรับสินค้าสำเร็จรูป ผู้ใช้งานควรอ่านคู่กันเสมอ",
            "je_images": [
                ("รูป 7.4.6 Journal Entry ตัดวัตถุดิบจากตัวอย่างจริง", "journal_mfg_raw_fg02001_real_annotated.png"),
                ("รูป 7.4.7 Journal Entry รับสินค้าสำเร็จรูปจากตัวอย่างจริง", "journal_mfg_finished_real_annotated.png"),
            ],
            "je_tables": [
                ("ตัดวัตถุดิบของตัวอย่าง", valuation_rows[0], "เห็นชัดว่าวัตถุดิบถูกเครดิตออกจากคลังและถูกนำเข้าไปเก็บไว้ในบัญชีงานระหว่างทำ"),
                ("รับ FG ของตัวอย่าง", valuation_rows[2], "หลังผลิตเสร็จต้นทุนเดียวกันถูกย้ายจากงานระหว่างทำกลับเข้าบัญชีสินค้าสำเร็จรูป"),
            ],
            "cautions": [
                "ถ้า MO ยังไม่เสร็จ รายการรับสินค้าสำเร็จรูปจะยังไม่สมบูรณ์",
                "ถ้าเปิดคนละ MO กับ valuation reference คนละใบ จะทำให้เทียบข้อมูลผิดทันที",
            ],
        },
        {
            "filename": "7.5_7.5 ของเสีย (Scrap & Loss).docx",
            "title": "7.5 ของเสีย (Scrap & Loss)",
            "subtitle": "Manufacturing",
            "objectives": [
                "เพื่อให้ผู้ใช้งานบันทึกของเสียและติดตามผลกระทบต่อสต็อกและบัญชีได้",
                "เพื่อให้เข้าใจความต่างระหว่าง scrap ที่ยังไม่ยืนยันกับผลทางบัญชีหลังมีการตัดของจริง",
            ],
            "overview": [
                f"ตัวอย่างข้อมูลจริงในระบบมีรายการ scrap ชื่อ {scrap.get('origin', '-') or '-'} อยู่ในสถานะ {scrap.get('state', '-') or '-'}",
                "ในวันที่ยังเป็น Draft ระบบจะยังไม่สร้างผลทางบัญชีทันที ผู้ใช้จึงต้องดูทั้งสถานะเอกสารและ Journal Entry ที่เกี่ยวกับการตัดวัตถุดิบประกอบกัน",
            ],
            "dashboard_image": "nav_dashboard_manufacturing_real_annotated.png",
            "dashboard_caption": "รูป 7.5.1 หน้า Dashboard สำหรับเข้าโมดูล Manufacturing",
            "menu_paths": [scrap_menu, valuation_menu],
            "menu_image": "nav_manufacturing_scrap_real_annotated.png",
            "menu_caption": "รูป 7.5.2 เมนู Scrap และ Valuation ที่ใช้ติดตามของเสีย",
            "usage_intro": "หัวข้อนี้ใช้ทั้งหน้าจอ Scrap จริงและตัวอย่าง Valuation ที่เกิดจากการตัดวัตถุดิบออกจากคลัง เพื่ออธิบายผลกระทบแบบที่ผู้ใช้เห็นได้จริงในระบบ",
            "steps": [
                "เข้าเมนู Scrap แล้วเปิดรายการของเสียที่ต้องการตรวจสอบ",
                "ดูสินค้า ปริมาณ ตำแหน่งต้นทาง และตำแหน่ง Scrap",
                "ถ้ารายการยังอยู่ Draft ให้ทราบว่ายังไม่ใช่จุดที่ระบบสรุปผลบัญชีขั้นสุดท้าย",
                "เปิด Valuation และ Journal Entry ที่เกี่ยวข้องกับการตัดวัตถุดิบเพื่อดูผลกระทบทางบัญชีประกอบ",
            ],
            "usage_images": [
                ("รูป 7.5.3 หน้าจอ Scrap จริงในระบบ", "manufacturing_scrap_form_real_annotated.png"),
                ("รูป 7.5.4 หน้าจอ Valuation ที่ใช้ตามรายการของเสียและความสูญเสีย", "inventory_valuation_page_real_annotated.png"),
            ],
            "fields": [
                ("Scrap Document", "เอกสารบันทึกของเสีย", f"{scrap.get('name','-') or '-'} / Qty {scrap.get('qty', 0)}"),
                ("Product", "สินค้าที่ถูกบันทึกเป็นของเสีย", scrap.get("product", "-") or "-"),
                ("Source Location", "ตำแหน่งต้นทางก่อนตัดเป็นของเสีย", scrap.get("location", "-") or "-"),
                ("Scrap Location", "ตำแหน่งที่รับของเสีย", scrap.get("scrap_location", "-") or "-"),
            ],
            "scenarios": [
                ("บันทึกของเสียระหว่างผลิต", "เปิด Scrap แล้วกรอกสินค้า ปริมาณ และตำแหน่งให้ครบ", "มีเอกสารของเสียพร้อมติดตามต่อ"),
                ("ตรวจผลกระทบทางบัญชี", "ถ้ายังไม่เห็น JE จาก scrap โดยตรง ให้เปิด Valuation และ JE ของวัตถุดิบที่เกี่ยวข้อง", "ช่วยอธิบายได้ว่าต้นทุนหรือมูลค่าสต็อกลดลงที่จุดใด"),
            ],
            "je_intro": "ตัวอย่างด้านล่างใช้ Journal Entry จริงของการตัดบรรจุภัณฑ์ออกจากคลังเข้าสู่งานระหว่างทำ เป็นตัวอย่างให้ผู้ใช้งานเข้าใจว่าความสูญเสียจะกระทบมูลค่าสต็อกอย่างไรเมื่อมีการบันทึกและยืนยันรายการจริง",
            "je_images": [("รูป 7.5.5 Journal Entry ตัวอย่างที่ใช้ตรวจผลของของเสียและความสูญเสีย", "journal_mfg_raw_packaging_real_annotated.png")],
            "je_tables": [("ตัวอย่างรายการบัญชีที่ใช้ติดตามผลกระทบของความสูญเสีย", valuation_rows[1], "เมื่อต้นทุนวัสดุถูกตัดออกจากคลัง ระบบจะเครดิตบัญชีสินค้าคงเหลือและเดบิตบัญชีงานระหว่างทำ ผู้ใช้จึงควรติดตามต้นทุนที่หายไปจากสองจุดนี้")],
            "cautions": [
                "ถ้ารายการ scrap ยังเป็น Draft จะยังไม่ใช่หลักฐานว่าระบบลงบัญชีเสร็จแล้ว",
                "ควรตรวจทั้งสถานะ scrap, Valuation และ Journal Entry ก่อนสรุปว่าของเสียกระทบบัญชีครบหรือไม่",
            ],
        },
        {
            "filename": "7.6_7.6 Stock Valuation Report.docx",
            "title": "7.6 Stock Valuation Report",
            "subtitle": "Manufacturing",
            "objectives": [
                "เพื่อให้ผู้ใช้งานเปิดรายงานมูลค่าสต็อกและตรวจรายการรับเข้า-จ่ายออกของงานผลิตได้",
                "เพื่อใช้เทียบรายงานกับ Journal Entry ของการผลิตได้อย่างเป็นลำดับ",
            ],
            "overview": [
                "รายงาน Valuation เป็นหน้าจอที่ใช้ดูว่าของใดถูกตัดออกจากคลังหรือรับเข้าคลังด้วยมูลค่าเท่าใด",
                f"ในตัวอย่างนี้มีรายการอ้างอิง {mo['name']} ที่ใช้ตรวจทั้งวัตถุดิบและสินค้าสำเร็จรูป",
            ],
            "dashboard_image": "nav_dashboard_inventory_real_annotated.png",
            "dashboard_caption": "รูป 7.6.1 หน้า Dashboard สำหรับเข้าโมดูล Inventory",
            "menu_paths": [valuation_menu],
            "menu_image": "nav_inventory_valuation_real_annotated.png",
            "menu_caption": "รูป 7.6.2 เมนู Valuation ที่ใช้เปิดรายงานมูลค่าสต็อก",
            "usage_intro": "หัวข้อนี้ใช้ทั้งหน้ารายงานและ Journal Entry จริงของงานผลิตใบเดียวกัน เพื่อให้ผู้ใช้ไล่ตรวจได้ครบ",
            "steps": [
                "เข้าเมนู Valuation แล้วกรองตามเลขอ้างอิงของ MO หรือสินค้า",
                "เปิดดูบรรทัดวัตถุดิบที่ถูกตัดออกและบรรทัดสินค้าสำเร็จรูปที่ถูกรับเข้า",
                "ถ้าต้องการรู้ผลบัญชี ให้กดเข้า Journal Entry ของบรรทัดนั้นต่อทันที",
            ],
            "usage_images": [("รูป 7.6.3 หน้ารายงาน Stock Valuation จริงในระบบ", "inventory_valuation_page_real_annotated.png")],
            "fields": [
                ("Reference", "เลขอ้างอิงที่เชื่อมรายงานกับ MO", mo["name"]),
                ("Product", "สินค้าที่มีการเคลื่อนไหวและมีผลต่อมูลค่าสต็อก", v_rows[0]["product"]),
                ("Value", "มูลค่าที่เพิ่มหรือลดในสต็อก", money(v_rows[0]["value"])),
                ("Account Move", "เลข Journal Entry ที่เปิดต่อไปตรวจได้", extra["valuation_moves"][0]["name"]),
            ],
            "scenarios": [
                ("เช็กต้นทุนวัตถุดิบของ MO", "กรองที่เลข MO แล้วเปิดบรรทัดวัตถุดิบ", "เห็นว่าของอะไรถูกตัดออกจากคลังเท่าไร"),
                ("เช็กรับ FG เข้าคลัง", "เปิดบรรทัดสินค้าสำเร็จรูปของ MO เดียวกัน", "เห็นมูลค่าที่รับกลับเข้าคลังและเทียบกับ JE ได้"),
            ],
            "je_intro": "เพื่อให้ผู้ใช้เข้าใจความสัมพันธ์ของรายงานกับบัญชี ด้านล่างจะแสดง Journal Entry จริงทั้งฝั่งตัดวัตถุดิบและรับสินค้าสำเร็จรูป",
            "je_images": [
                ("รูป 7.6.4 Journal Entry ของบรรทัดวัตถุดิบในรายงาน Valuation", "journal_mfg_raw_fg02001_real_annotated.png"),
                ("รูป 7.6.5 Journal Entry ของบรรทัดรับ FG ในรายงาน Valuation", "journal_mfg_finished_real_annotated.png"),
            ],
            "je_tables": [
                ("ตัวอย่าง JE ฝั่งวัตถุดิบ", valuation_rows[0], "รายงาน Valuation จะแสดงมูลค่าที่ลดจากคลัง และ JE จะบอกว่ามูลค่านั้นถูกย้ายไปบัญชีใด"),
                ("ตัวอย่าง JE ฝั่งสินค้าสำเร็จรูป", valuation_rows[2], "เมื่อดูรายงานฝั่งรับเข้า ควรเห็น JE ที่ย้ายต้นทุนจากงานระหว่างทำกลับไปที่บัญชี FG"),
            ],
            "cautions": [
                "ถ้ารายงาน Valuation ไม่มีข้อมูล อาจเกิดจากช่วงวันที่ไม่ตรงหรือเอกสารยังไม่ Done",
                "ควรใช้เลขอ้างอิงเดียวกันทั้งในรายงานและ Journal Entry เพื่อป้องกันการอ่านข้ามรายการ",
            ],
        },
        {
            "filename": "7.7_7.7 รายงานส่วนราชการ (รง.8).docx",
            "title": "7.7 รายงานส่วนราชการ (รง.8)",
            "subtitle": "Manufacturing",
            "objectives": [
                "เพื่อให้ผู้ใช้งานเปิดรายงาน รง.8 และตั้งค่าตัวกรองได้ถูกต้อง",
                "เพื่อให้เข้าใจว่ารายงานนี้ใช้ข้อมูลจากสต็อกและเอกสารการเคลื่อนไหว ไม่ได้สร้าง Journal Entry เอง",
            ],
            "overview": [
                f"ค่าตั้งต้นจริงในระบบของรายงานนี้คือ Product Category = {mfg['rng8_defaults']['product_category']} และ Location = {mfg['rng8_defaults']['location']}",
                "ผู้ใช้งานควรใช้รายงานนี้ร่วมกับ Valuation หรือ Journal Entry เพื่อยืนยันความถูกต้องของข้อมูลย้อนหลัง",
            ],
            "dashboard_image": "nav_dashboard_accounting_real_annotated.png",
            "dashboard_caption": "รูป 7.7.1 หน้า Dashboard สำหรับเข้าโมดูล Accounting",
            "menu_paths": [rng8_menu],
            "menu_image": "nav_accounting_rng8_real_annotated.png",
            "menu_caption": "รูป 7.7.2 เมนูรายงาน รง.8",
            "usage_intro": "หัวข้อนี้ใช้หน้ารายงานจริงในระบบและค่าตั้งต้นจริงที่พบในวันที่จัดทำคู่มือ",
            "steps": [
                "เข้าเมนู รง.8",
                "เลือกหมวดสินค้า ตำแหน่งจัดเก็บ และช่วงวันที่ที่ต้องการ",
                "กดดูรายงานแล้วตรวจรายการที่แสดง",
                "ถ้าต้องการยืนยันตัวเลขต่อ ให้เปิด Valuation หรือ Journal Entry ที่เกี่ยวข้องตามเลขอ้างอิง",
            ],
            "usage_images": [("รูป 7.7.3 หน้ารายงาน รง.8 จริงในระบบ", "rng8_report_page_real_annotated.png")],
            "fields": [
                ("Product Category", "หมวดสินค้าที่ใช้กรองรายงาน", mfg["rng8_defaults"]["product_category"]),
                ("Location", "ตำแหน่งคลังที่ใช้กรองรายงาน", mfg["rng8_defaults"]["location"]),
                ("Date From / Date To", "ช่วงวันที่ของรายงาน", f"{mfg['rng8_defaults']['date_from']} ถึง {mfg['rng8_defaults']['date_to']}"),
            ],
            "scenarios": [
                ("ตรวจรายงานส่งหน่วยงานภายนอก", "เลือกช่วงวันที่และหมวดสินค้าให้ตรงกับงวดที่ต้องการ", "ได้ข้อมูลพร้อมตรวจทานและพิมพ์รายงาน"),
                ("เทียบรายงานกับเอกสารบัญชี", "เปิด JE ของรายการรับ FG หรือรายการคลังที่เกี่ยวข้องประกอบ", "ยืนยันความถูกต้องของยอดได้มากขึ้น"),
            ],
            "je_intro": "รายงาน รง.8 เป็นรายงานสรุป ไม่ได้สร้าง Journal Entry เอง ด้านล่างจึงใช้ Journal Entry รับสินค้าสำเร็จรูปจริงเป็นตัวอย่างการตรวจเทียบ",
            "je_images": [("รูป 7.7.4 ตัวอย่าง Journal Entry ที่ใช้ตรวจเทียบกับรายงาน รง.8", "journal_mfg_finished_real_annotated.png")],
            "je_tables": [("ตัวอย่างรายการบัญชีที่ใช้ตรวจเทียบกับรายงาน", valuation_rows[2], "เมื่อรายงานอ้างอิงข้อมูลการผลิตเสร็จ ผู้ใช้งานสามารถเปิด JE นี้เพื่อยืนยันว่ามูลค่ารับเข้าสินค้าสำเร็จรูปมีอยู่จริงในระบบ")],
            "cautions": [
                "รายงาน รง.8 จะถูกต้องก็ต่อเมื่อหมวดสินค้า ตำแหน่งจัดเก็บ และช่วงวันที่เลือกถูกต้อง",
                "ถ้าต้องยืนยันยอดย้อนหลัง ควรเทียบกับ Valuation หรือ Journal Entry เพิ่มเสมอ",
            ],
        },
    ]


def render_pdfs(files: list[Path]):
    PDF_DIR.mkdir(parents=True, exist_ok=True)
    executable = str(SOFFICE) if SOFFICE.exists() else "soffice"
    for doc_path in files:
        subprocess.run(
            [executable, "--headless", "--convert-to", "pdf", "--outdir", str(PDF_DIR), str(doc_path)],
            check=True,
            cwd=str(ROOT),
        )


def main():
    DOCX_DIR.mkdir(parents=True, exist_ok=True)
    specs = fixed_asset_specs() + manufacturing_specs()
    files = [write_doc(spec) for spec in specs]
    render_pdfs(files)
    print(f"generated {len(files)} docs")


if __name__ == "__main__":
    main()

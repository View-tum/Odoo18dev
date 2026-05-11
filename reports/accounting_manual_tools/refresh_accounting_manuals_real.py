from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Inches


ROOT = Path(r"C:\365_project\TheCool18e\Dev")
DOCX_DIR = ROOT / "manual" / "Accouting_Manual" / "generated_20260408" / "docx"
IMAGE_DIR = ROOT / "manual" / "Accouting_Manual" / "generated_20260408" / "images"


GLOBAL_REPLACEMENTS = {
    "local UAT": "ระบบ",
    "local manual": "ระบบ",
    "Scenario การใช้งานจากข้อมูลจริงในระบบ": "Scenario การใช้งานจากข้อมูลจริงในระบบ",
    "ตัวอย่าง Scenario จากข้อมูลจริงในระบบ": "ตัวอย่าง Scenario จากข้อมูลจริงในระบบ",
    "ในระบบ ตัวอย่าง": "ในระบบตัวอย่าง",
    "ใน ระบบ": "ในระบบ",
    "จากตัวอย่าง ระบบ": "จากตัวอย่างระบบ",
    "ตัวอย่างในระบบ มี": "ตัวอย่างในระบบมี",
    "ผู้ใช้ เห็น": "ผู้ใช้เห็น",
    "user": "ผู้ใช้",
}


DOC_SPEC = {
    "3.8_": {
        "images": [
            ("รูป 3.8A.1 หน้า Dashboard สำหรับเข้าโมดูล Accounting", "nav_dashboard_accounting_real_annotated.png"),
            ("รูป 3.8A.2 เมนู Accounting > Customers > รับชำระเงินกลุ่มลูกค้า", "nav_accounting_group_payment_real_annotated.png"),
            ("รูป 3.8A.3 Journal Entry ของการรับชำระแบบกลุ่มบริษัทจากข้อมูลจริงในระบบ", "journal_group_payment_real_annotated.png"),
        ],
        "specific_replacements": {
            "ตัวอย่างในระบบมีเอกสาร draft ที่ยังไม่สร้าง payment คือเลขที่ 157 และมีเอกสาร done ที่สร้าง payment แล้วคือเลขที่ 142. ในคู่มือนี้จะใช้งานจากสองตัวอย่างนี้เพื่อให้ผู้ใช้เห็นทั้งก่อนและหลังบันทึกรับชำระ.": "ในระบบมีทั้งเอกสารรับชำระแบบกลุ่มบริษัทที่อยู่ระหว่างเตรียมรายการ และเอกสารที่บันทึกรับชำระแล้ว โดยคู่มือนี้จะอธิบายให้ผู้ใช้เห็นทั้งขั้นตอนก่อนสร้าง payment และการตรวจสอบผลหลังสร้าง payment จากข้อมูลจริงในระบบ.",
            "ตัวอย่างในระบบมีเอกสาร draft ที่ยังไม่สร้าง payment คือเลขที่ 157 และมีเอกสาร done ที่สร้าง payment แล้วคือเลขที่ 142. ในคู่มือนี้จะใช้งานจากสองตัวอย่างนี้เพื่อให้ผู้ใช้ เห็นทั้งก่อนและหลังบันทึกรับชำระ.": "ในระบบมีทั้งเอกสารรับชำระแบบกลุ่มบริษัทที่อยู่ระหว่างเตรียมรายการ และเอกสารที่บันทึกรับชำระแล้ว โดยคู่มือนี้จะอธิบายให้ผู้ใช้เห็นทั้งขั้นตอนก่อนสร้าง payment และการตรวจสอบผลหลังสร้าง payment จากข้อมูลจริงในระบบ.",
            "ตัวอย่างในระบบมีเอกสาร draft ที่ยังไม่สร้าง payment คือเลขที่ 157 และมีเอกสาร done ที่สร้าง payment แล้วคือเลขที่ 142. ในคู่มือนี้จะใช้งานจากสองตัวอย่างนี้เพื่อให้ user เห็นทั้งก่อนและหลังบันทึกรับชำระ.": "ในระบบมีทั้งเอกสารรับชำระแบบกลุ่มบริษัทที่อยู่ระหว่างเตรียมรายการ และเอกสารที่บันทึกรับชำระแล้ว โดยคู่มือนี้จะอธิบายให้ผู้ใช้เห็นทั้งขั้นตอนก่อนสร้าง payment และการตรวจสอบผลหลังสร้าง payment จากข้อมูลจริงในระบบ.",
            "2. ในช่อง ลูกค้า/สมาชิกกลุ่ม ให้เลือกคู่ค้าที่อยู่ภายใต้กลุ่มเดียวกัน จากตัวอย่างระบบ ใช้ UAT Manual Customer A 20260408 และ UAT Manual Customer B 20260408.": "2. ในช่อง ลูกค้า/สมาชิกกลุ่ม ให้เลือกคู่ค้าที่อยู่ภายใต้กลุ่มเดียวกันตามข้อมูลจริงของรายการที่ต้องการรับชำระ ระบบจะดึงเอกสารคงค้างของสมาชิกที่เลือกมาแสดงในตาราง.",
            "2. ในช่อง ลูกค้า/สมาชิกกลุ่ม ให้เลือกคู่ค้าที่อยู่ภายใต้กลุ่มเดียวกัน จากตัวอย่างระบบใช้ UAT Manual Customer A 20260408 และ UAT Manual Customer B 20260408.": "2. ในช่อง ลูกค้า/สมาชิกกลุ่ม ให้เลือกคู่ค้าที่อยู่ภายใต้กลุ่มเดียวกันตามข้อมูลจริงของรายการที่ต้องการรับชำระ ระบบจะดึงเอกสารคงค้างของสมาชิกที่เลือกมาแสดงในตาราง.",
            "1. เข้าเมนู Accounting > Customers > รับชำระเงินกลุ่มลูกค้า แล้วกด New เพื่อสร้างเอกสารใหม่ หรือเปิดเอกสาร draft ตัวอย่างเลขที่ 157 เพื่อดู flow ที่เตรียมไว้.": "1. เข้าเมนู Accounting > Customers > รับชำระเงินกลุ่มลูกค้า แล้วกด New เพื่อสร้างเอกสารใหม่ หรือเปิดเอกสารจริงที่ผู้ใช้งานต้องการตรวจสอบเพื่อดูรายการที่ค้างชำระอยู่ในระบบ.",
        },
    },
    "5.1_": {
        "images": [
            ("รูป 5.1A.1 หน้า Dashboard สำหรับเข้าโมดูล Cheque", "nav_dashboard_cheque_real_annotated.png"),
            ("รูป 5.1A.2 เมนู Cheque > Configuration", "nav_cheque_configuration_real_annotated.png"),
            ("รูป 5.1A.3 ตัวอย่าง Journal Entry ฝั่งจ่ายเช็คเพื่ออธิบายผลของการตั้งค่า", "journal_cheque_out_confirmed_real_annotated.png"),
            ("รูป 5.1A.4 ตัวอย่าง Journal Entry ฝั่งรับเช็คเพื่ออธิบายผลของการตั้งค่า", "journal_cheque_in_confirmed_real_annotated.png"),
        ],
        "specific_replacements": {
            "ฟังก์ชันที่มีจริงในระบบเท่านั้น.": "ฟังก์ชันที่มีจริงในระบบเท่านั้น.",
        },
    },
    "5.2_": {
        "images": [
            ("รูป 5.2A.1 หน้า Dashboard สำหรับเข้าโมดูล Cheque", "nav_dashboard_cheque_real_annotated.png"),
            ("รูป 5.2A.2 เมนู Cheque > Configuration เพื่อเข้าเลือกเทมเพลตฟอร์มเช็ค", "nav_cheque_configuration_real_annotated.png"),
        ],
        "specific_replacements": {},
    },
    "5.3_": {
        "images": [
            ("รูป 5.3A.1 หน้า Dashboard สำหรับเข้าโมดูล Cheque", "nav_dashboard_cheque_real_annotated.png"),
            ("รูป 5.3A.2 เมนู Cheque > Configuration ที่ใช้เข้าสู่หน้าสมุดเช็ค", "nav_cheque_configuration_real_annotated.png"),
        ],
        "specific_replacements": {
            "ในระบบมีตัวอย่างสมุดเช็คจริงชื่อ UAT-MANUAL-CB-20260408-02 สถานะ done สำหรับใช้ประกอบคำอธิบาย.": "ในระบบมีสมุดเช็คที่ถูกยืนยันใช้งานจริงอยู่แล้ว และมีเลขเช็คคงเหลือสำหรับนำไปใช้ในขั้นตอนการจ่ายเช็ค ผู้ใช้สามารถเปิดสมุดเช็คจริงในระบบเพื่อดูโครงสร้างข้อมูลและสถานะได้ทันที.",
        },
    },
    "5.4_": {
        "images": [
            ("รูป 5.4A.1 หน้า Dashboard สำหรับเข้าโมดูล Accounting", "nav_dashboard_accounting_real_annotated.png"),
            ("รูป 5.4A.2 เมนู Accounting > Vendors > Bills", "nav_accounting_vendors_bills_real_annotated.png"),
            ("รูป 5.4A.3 Journal Entry ตอนสร้างเช็คจ่ายสถานะ Confirmed", "journal_cheque_out_confirmed_real_annotated.png"),
            ("รูป 5.4A.4 Journal Entry หลังเช็คจ่ายถูกตัดผ่านธนาคาร", "journal_cheque_out_paid_real_annotated.png"),
        ],
        "specific_replacements": {},
    },
    "5.5_": {
        "images": [
            ("รูป 5.5A.1 หน้า Dashboard สำหรับเข้าโมดูล Cheque", "nav_dashboard_cheque_real_annotated.png"),
            ("รูป 5.5A.2 เมนู Cheque > Operations", "nav_cheque_operations_real_annotated.png"),
            ("รูป 5.5A.3 Journal Entry ที่ใช้ตรวจยอด Outstanding Cheque", "journal_cheque_out_confirmed_real_annotated.png"),
            ("รูป 5.5A.4 Journal Entry หลังเช็คถูกเคลียร์และย้ายออกจากบัญชีพัก", "journal_cheque_out_paid_real_annotated.png"),
        ],
        "specific_replacements": {},
    },
    "5.6_": {
        "images": [
            ("รูป 5.6A.1 หน้า Dashboard สำหรับเข้าโมดูล Accounting", "nav_dashboard_accounting_real_annotated.png"),
            ("รูป 5.6A.2 เมนู Accounting > Customers > Invoices", "nav_accounting_customers_invoices_real_annotated.png"),
            ("รูป 5.6A.3 หน้าต่าง Register Payment ของ Invoice ฝั่งรับเช็คจากข้อมูลจริงในระบบ", "invoice_register_payment_real_annotated.png"),
            ("รูป 5.6A.4 Journal Entry ตอนสร้างเช็ครับสถานะ Confirmed", "journal_cheque_in_confirmed_real_annotated.png"),
            ("รูป 5.6A.5 Journal Entry หลังเช็ครับถูกตัดผ่านธนาคาร", "journal_cheque_in_paid_real_annotated.png"),
        ],
        "specific_replacements": {
            "[ไม่พบภาพประกอบ: invoice_register_payment_test.png_annotated.png]": "ดูภาพหน้าต่าง Register Payment จากข้อมูลจริงในระบบในหัวข้อภาพพาเข้าเมนูและ Journal Entry ด้านล่าง",
            "ในระบบ ตัวอย่าง inbound cheque บางใบมีบริบท settlement เดิม ทำให้ยอดบน cheque record และยอดใน payment entry อาจไม่ตรงกันทุกกรณี ผู้ใช้ต้องตรวจยอดใน Payment Entry ประกอบเสมอ": "ตัวอย่างเช็ครับบางรายการในระบบมีบริบทการกระทบยอดเดิมร่วมอยู่ด้วย ผู้ใช้จึงต้องเปิด Payment Entry และ Journal Items ประกอบทุกครั้งเพื่อยืนยันยอดและขาบัญชีที่เกิดขึ้นจริง",
        },
    },
    "5.7_": {
        "images": [
            ("รูป 5.7A.1 หน้า Dashboard สำหรับเข้าโมดูล Cheque", "nav_dashboard_cheque_real_annotated.png"),
            ("รูป 5.7A.2 เมนู Cheque > Operations", "nav_cheque_operations_real_annotated.png"),
            ("รูป 5.7A.3 Journal Entry ฝั่งจ่ายเช็คหลังเคลียร์เช็คแล้ว", "journal_cheque_out_paid_real_annotated.png"),
            ("รูป 5.7A.4 Journal Entry ฝั่งรับเช็คหลังเคลียร์เช็คแล้ว", "journal_cheque_in_paid_real_annotated.png"),
        ],
        "specific_replacements": {},
    },
    "5.8_": {
        "images": [
            ("รูป 5.8A.1 หน้า Dashboard สำหรับเข้าโมดูล Cheque", "nav_dashboard_cheque_real_annotated.png"),
            ("รูป 5.8A.2 เมนู Cheque > Operations ที่ใช้เข้าหน้าเอกสารเช็ค", "nav_cheque_operations_real_annotated.png"),
            ("รูป 5.8A.3 Journal Entry ฝั่ง Reverse หลังยกเลิกเช็ค", "journal_cheque_void_reverse_real_annotated.png"),
        ],
        "specific_replacements": {
            "ในระบบ ปุ่มและหน้าจอที่ใช้งานได้จริงของหัวข้อนี้คือ Void, Cancel และ Reset To Draft. ส่วนหน้าจอ Transform Detail ยังถูกคอมเมนต์ซ่อนไว้ใน view จึงไม่ควรอ้างเป็นขั้นตอนใช้งานจริงในคู่มือนี้.": "ปุ่มและหน้าจอที่ใช้งานได้จริงของหัวข้อนี้ในระบบคือ Void, Cancel และ Reset To Draft ส่วนหน้าจอ Transform Detail ยังไม่เปิดใช้งานจริง จึงไม่ถูกนำมาอธิบายในคู่มือนี้.",
        },
    },
}


def iter_paragraphs_and_cells(doc: Document):
    for paragraph in doc.paragraphs:
        yield paragraph
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph


def replace_text(doc: Document, replacements: dict[str, str]) -> None:
    for paragraph in iter_paragraphs_and_cells(doc):
        text = paragraph.text
        if not text:
            continue
        new_text = text
        for old, new in replacements.items():
            new_text = new_text.replace(old, new)
        if new_text != text:
            paragraph.text = new_text


def move_paragraph_before(paragraph, anchor) -> None:
    anchor._p.addprevious(paragraph._p)


def add_paragraph_before(doc: Document, anchor, text: str, style: str | None = None, bold: bool = False):
    paragraph = doc.add_paragraph(style=style)
    run = paragraph.add_run(text)
    run.bold = bold
    move_paragraph_before(paragraph, anchor)
    return paragraph


def add_image_before(doc: Document, anchor, image_path: Path, caption: str) -> None:
    if not image_path.exists():
        return
    image_para = doc.add_paragraph()
    image_para.alignment = 1
    image_para.add_run().add_picture(str(image_path), width=Inches(6.4))
    move_paragraph_before(image_para, anchor)
    caption_para = doc.add_paragraph()
    caption_para.alignment = 1
    caption_para.add_run(caption).bold = False
    move_paragraph_before(caption_para, anchor)


def find_paragraph(doc: Document, text: str):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == text:
            return paragraph
    return None


def ensure_extra_section(doc: Document, doc_name: str, images: list[tuple[str, str]]) -> None:
    section_title = "ภาพพาเข้าเมนูและ Journal Entry จากข้อมูลจริงในระบบ"
    if find_paragraph(doc, section_title):
        return
    anchor = find_paragraph(doc, "คำอธิบายฟิลด์และส่วนสำคัญบนหน้าจอ")
    if anchor is None:
        anchor = find_paragraph(doc, "ข้อควรระวัง")
    if anchor is None:
        anchor = doc.add_paragraph()
    add_paragraph_before(doc, anchor, section_title, style="Heading 1")
    intro = (
        "ภาพในหัวข้อนี้เป็นภาพจากข้อมูลจริงที่เปิดใช้งานอยู่ในระบบ เพื่อให้ผู้ใช้เห็นเส้นทางการคลิกจากหน้าหลัก "
        "เข้าสู่เมนูที่เกี่ยวข้อง รวมถึงภาพตัวอย่าง Journal Entry ที่ต้องเปิดตรวจสอบหลังทำรายการสำเร็จ."
    )
    add_paragraph_before(doc, anchor, intro)
    for caption, image_name in images:
        add_image_before(doc, anchor, IMAGE_DIR / image_name, caption)


def fix_56_placeholder(doc: Document) -> None:
    for paragraph in doc.paragraphs:
        if "[ไม่พบภาพประกอบ:" in paragraph.text:
            paragraph.text = "ดูภาพหน้าต่าง Register Payment จากข้อมูลจริงในระบบในหัวข้อภาพพาเข้าเมนูและ Journal Entry จากข้อมูลจริงในระบบ"


def normalize_doc_specific_text(path_name: str, doc: Document) -> None:
    for paragraph in iter_paragraphs_and_cells(doc):
        text = paragraph.text.strip()
        if not text:
            continue
        if path_name.startswith("3.8_") and "เอกสาร draft ที่ยังไม่สร้าง payment" in text:
            paragraph.text = (
                "ในระบบมีทั้งเอกสารรับชำระแบบกลุ่มบริษัทที่อยู่ระหว่างเตรียมรายการ "
                "และเอกสารที่บันทึกรับชำระแล้ว โดยคู่มือนี้จะอธิบายให้ผู้ใช้เห็นทั้งขั้นตอนก่อนสร้าง payment "
                "และการตรวจสอบผลหลังสร้าง payment จากข้อมูลจริงในระบบ."
            )
        elif path_name.startswith("3.8_") and "UAT Manual Customer A 20260408" in text:
            paragraph.text = (
                "2. ในช่อง ลูกค้า/สมาชิกกลุ่ม ให้เลือกคู่ค้าที่อยู่ภายใต้กลุ่มเดียวกันตามข้อมูลจริงของรายการที่ต้องการรับชำระ "
                "ระบบจะดึงเอกสารคงค้างของสมาชิกที่เลือกมาแสดงในตาราง."
            )
        elif path_name.startswith("5.3_") and "UAT-MANUAL-CB-20260408-02" in text:
            paragraph.text = (
                "ในระบบมีสมุดเช็คที่ถูกยืนยันใช้งานจริงอยู่แล้ว และมีเลขเช็คคงเหลือสำหรับนำไปใช้ในขั้นตอนการจ่ายเช็ค "
                "ผู้ใช้สามารถเปิดสมุดเช็คจริงในระบบเพื่อดูโครงสร้างข้อมูลและสถานะได้ทันที."
            )


def main() -> None:
    for path in sorted(DOCX_DIR.glob("*.docx")):
        spec = None
        for prefix, candidate in DOC_SPEC.items():
            if path.name.startswith(prefix):
                spec = candidate
                break
        if spec is None:
            continue

        doc = Document(path)
        replacements = dict(GLOBAL_REPLACEMENTS)
        replacements.update(spec["specific_replacements"])
        replace_text(doc, replacements)
        normalize_doc_specific_text(path.name, doc)

        if path.name.startswith("5.6_"):
            fix_56_placeholder(doc)

        ensure_extra_section(doc, path.name, spec["images"])
        doc.save(path)
        print(path.name)


if __name__ == "__main__":
    main()

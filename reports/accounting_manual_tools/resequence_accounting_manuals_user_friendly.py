from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.shared import Inches
from docx.table import Table
from docx.text.paragraph import Paragraph


ROOT = Path(r"C:\365_project\TheCool18e\Dev")
DOCX_DIR = ROOT / "manual" / "Accouting_Manual" / "generated_20260408" / "docx"
IMAGE_DIR = ROOT / "manual" / "Accouting_Manual" / "generated_20260408" / "images"


SECTION_RENAMES = {
    "เมนูที่ใช้": "1. เข้าใช้งานจากหน้าหลัก",
    "ขั้นตอนการใช้งาน": "2. ขั้นตอนการทำรายการ",
    "การอธิบาย Journal Items และขาบัญชี": "3. ตรวจสอบรายการบัญชี (Journal Entry)",
    "ฟังก์ชันที่ทำได้จริง": "สิ่งที่ทำได้ในหน้าจอนี้",
    "คำอธิบายฟิลด์และส่วนสำคัญบนหน้าจอ": "คำอธิบายช่องสำคัญบนหน้าจอ",
    "ข้อควรระวัง": "สิ่งที่ควรตรวจสอบก่อนบันทึก",
    "Scenario การใช้งานจากข้อมูลจริงในระบบ": "ตัวอย่างข้อมูลจริงที่ใช้ในคู่มือ",
}


TEXT_REPLACEMENTS = {
    "custom": "เมนูเสริม",
    "wizard": "หน้าต่าง",
    "flow": "ลำดับการทำงาน",
    "filter": "ตัวกรอง",
    "Journal Items และขาบัญชี": "รายการบัญชี (Journal Entry)",
}


DOC_IMAGE_SPEC = {
    "3.8_": {
        "nav": [
            ("รูป 3.8A.1 หน้า Dashboard สำหรับเข้าโมดูล Accounting", "nav_dashboard_accounting_real_annotated.png"),
            ("รูป 3.8A.2 เมนู Accounting > Customers > รับชำระเงินกลุ่มลูกค้า", "nav_accounting_group_payment_real_annotated.png"),
        ],
        "je": [
            ("รูป 3.8A.3 Journal Entry ของการรับชำระแบบกลุ่มบริษัทจากข้อมูลจริงในระบบ", "journal_group_payment_real_annotated.png"),
        ],
    },
    "5.1_": {
        "nav": [
            ("รูป 5.1A.1 หน้า Dashboard สำหรับเข้าโมดูล Cheque", "nav_dashboard_cheque_real_annotated.png"),
            ("รูป 5.1A.2 เมนู Cheque > Configuration", "nav_cheque_configuration_real_annotated.png"),
        ],
        "je": [
            ("รูป 5.1A.3 ตัวอย่าง Journal Entry ฝั่งจ่ายเช็คเพื่ออธิบายผลของการตั้งค่า", "journal_cheque_out_confirmed_real_annotated.png"),
            ("รูป 5.1A.4 ตัวอย่าง Journal Entry ฝั่งรับเช็คเพื่ออธิบายผลของการตั้งค่า", "journal_cheque_in_confirmed_real_annotated.png"),
        ],
    },
    "5.2_": {
        "nav": [
            ("รูป 5.2A.1 หน้า Dashboard สำหรับเข้าโมดูล Cheque", "nav_dashboard_cheque_real_annotated.png"),
            ("รูป 5.2A.2 เมนู Cheque > Configuration เพื่อเข้าเลือกเทมเพลตฟอร์มเช็ค", "nav_cheque_configuration_real_annotated.png"),
        ],
        "je": [],
    },
    "5.3_": {
        "nav": [
            ("รูป 5.3A.1 หน้า Dashboard สำหรับเข้าโมดูล Cheque", "nav_dashboard_cheque_real_annotated.png"),
            ("รูป 5.3A.2 เมนู Cheque > Configuration ที่ใช้เข้าสู่หน้าสมุดเช็ค", "nav_cheque_configuration_real_annotated.png"),
        ],
        "je": [],
    },
    "5.4_": {
        "nav": [
            ("รูป 5.4A.1 หน้า Dashboard สำหรับเข้าโมดูล Accounting", "nav_dashboard_accounting_real_annotated.png"),
            ("รูป 5.4A.2 เมนู Accounting > Vendors > Bills", "nav_accounting_vendors_bills_real_annotated.png"),
        ],
        "je": [
            ("รูป 5.4A.3 Journal Entry ตอนสร้างเช็คจ่ายสถานะ Confirmed", "journal_cheque_out_confirmed_real_annotated.png"),
            ("รูป 5.4A.4 Journal Entry หลังเช็คจ่ายถูกตัดผ่านธนาคาร", "journal_cheque_out_paid_real_annotated.png"),
        ],
    },
    "5.5_": {
        "nav": [
            ("รูป 5.5A.1 หน้า Dashboard สำหรับเข้าโมดูล Cheque", "nav_dashboard_cheque_real_annotated.png"),
            ("รูป 5.5A.2 เมนู Cheque > Operations", "nav_cheque_operations_real_annotated.png"),
        ],
        "je": [
            ("รูป 5.5A.3 Journal Entry ที่ใช้ตรวจยอด Outstanding Cheque", "journal_cheque_out_confirmed_real_annotated.png"),
            ("รูป 5.5A.4 Journal Entry หลังเช็คถูกเคลียร์และย้ายออกจากบัญชีพัก", "journal_cheque_out_paid_real_annotated.png"),
        ],
    },
    "5.6_": {
        "nav": [
            ("รูป 5.6A.1 หน้า Dashboard สำหรับเข้าโมดูล Accounting", "nav_dashboard_accounting_real_annotated.png"),
            ("รูป 5.6A.2 เมนู Accounting > Customers > Invoices", "nav_accounting_customers_invoices_real_annotated.png"),
        ],
        "je": [
            ("รูป 5.6A.3 หน้าต่าง Register Payment ของ Invoice ฝั่งรับเช็คจากข้อมูลจริงในระบบ", "invoice_register_payment_real_annotated.png"),
            ("รูป 5.6A.4 Journal Entry ตอนสร้างเช็ครับสถานะ Confirmed", "journal_cheque_in_confirmed_real_annotated.png"),
            ("รูป 5.6A.5 Journal Entry หลังเช็ครับถูกตัดผ่านธนาคาร", "journal_cheque_in_paid_real_annotated.png"),
        ],
    },
    "5.7_": {
        "nav": [
            ("รูป 5.7A.1 หน้า Dashboard สำหรับเข้าโมดูล Cheque", "nav_dashboard_cheque_real_annotated.png"),
            ("รูป 5.7A.2 เมนู Cheque > Operations", "nav_cheque_operations_real_annotated.png"),
        ],
        "je": [
            ("รูป 5.7A.3 Journal Entry ฝั่งจ่ายเช็คหลังเคลียร์เช็คแล้ว", "journal_cheque_out_paid_real_annotated.png"),
            ("รูป 5.7A.4 Journal Entry ฝั่งรับเช็คหลังเคลียร์เช็คแล้ว", "journal_cheque_in_paid_real_annotated.png"),
        ],
    },
    "5.8_": {
        "nav": [
            ("รูป 5.8A.1 หน้า Dashboard สำหรับเข้าโมดูล Cheque", "nav_dashboard_cheque_real_annotated.png"),
            ("รูป 5.8A.2 เมนู Cheque > Operations ที่ใช้เข้าหน้าเอกสารเช็ค", "nav_cheque_operations_real_annotated.png"),
        ],
        "je": [
            ("รูป 5.8A.3 Journal Entry ฝั่ง Reverse หลังยกเลิกเช็ค", "journal_cheque_void_reverse_real_annotated.png"),
        ],
    },
}


SECTION_ORDER = [
    "วัตถุประสงค์",
    "ภาพรวมการทำงาน",
    "สิ่งที่ทำได้ในหน้าจอนี้",
    "1. เข้าใช้งานจากหน้าหลัก",
    "2. ขั้นตอนการทำรายการ",
    "ภาพประกอบการทำงาน",
    "ตัวอย่างข้อมูลจริงที่ใช้ในคู่มือ",
    "3. ตรวจสอบรายการบัญชี (Journal Entry)",
    "คำอธิบายช่องสำคัญบนหน้าจอ",
    "สิ่งที่ควรตรวจสอบก่อนบันทึก",
]


def body_children(doc: Document):
    body = doc._element.body
    for child in list(body.iterchildren()):
        if child.tag.endswith("}sectPr"):
            continue
        yield child


def paragraph_from_element(doc: Document, element):
    if isinstance(element, CT_P):
        return Paragraph(element, doc)
    return None


def is_heading(doc: Document, element) -> bool:
    paragraph = paragraph_from_element(doc, element)
    if paragraph is None:
        return False
    text = paragraph.text.strip()
    if not text:
        return False
    return text in SECTION_ORDER or text == "ภาพพาเข้าเมนูและ Journal Entry จากข้อมูลจริงในระบบ"


def rename_headings_and_text(doc: Document) -> None:
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text in SECTION_RENAMES:
            paragraph.text = SECTION_RENAMES[text]
        else:
            new_text = paragraph.text
            for old, new in TEXT_REPLACEMENTS.items():
                new_text = new_text.replace(old, new)
            if new_text != paragraph.text:
                paragraph.text = new_text
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    new_text = paragraph.text
                    for old, new in TEXT_REPLACEMENTS.items():
                        new_text = new_text.replace(old, new)
                    if new_text != paragraph.text:
                        paragraph.text = new_text


def segment_sections(doc: Document):
    title_blocks = []
    sections: dict[str, list] = {}
    current_name = None
    current_blocks = []
    before_first_heading = True

    for element in body_children(doc):
        if is_heading(doc, element):
            paragraph = paragraph_from_element(doc, element)
            heading_text = paragraph.text.strip()
            if heading_text == "ภาพพาเข้าเมนูและ Journal Entry จากข้อมูลจริงในระบบ":
                current_name = heading_text
                sections[current_name] = [element]
                before_first_heading = False
                continue
            if before_first_heading:
                before_first_heading = False
            if current_name is not None:
                sections[current_name] = current_blocks
            current_name = heading_text
            current_blocks = [element]
        else:
            if before_first_heading:
                title_blocks.append(element)
            elif current_name is not None:
                current_blocks.append(element)

    if current_name is not None:
        sections[current_name] = current_blocks

    return title_blocks, sections


def add_paragraph_at_end(doc: Document, text: str, style: str | None = None, bold: bool = False):
    paragraph = doc.add_paragraph(style=style)
    run = paragraph.add_run(text)
    run.bold = bold
    return paragraph._p


def add_image_blocks_at_end(doc: Document, caption: str, image_path: Path):
    blocks = []
    if not image_path.exists():
        return blocks
    image_para = doc.add_paragraph()
    image_para.alignment = 1
    image_para.add_run().add_picture(str(image_path), width=Inches(6.4))
    blocks.append(image_para._p)
    caption_para = doc.add_paragraph()
    caption_para.alignment = 1
    caption_para.add_run(caption)
    blocks.append(caption_para._p)
    return blocks


def create_extra_blocks(doc: Document, kind: str, items: list[tuple[str, str]]):
    blocks = []
    if not items:
        return blocks
    if kind == "nav":
        blocks.append(
            add_paragraph_at_end(
                doc,
                "ให้เริ่มจากหน้า Dashboard แล้วกดเข้าโมดูลตามภาพ จากนั้นกดเมนูย่อยตามหมายเลขในภาพเพื่อเข้าสู่หน้าที่ใช้ทำรายการจริง.",
            )
        )
    elif kind == "je":
        blocks.append(
            add_paragraph_at_end(
                doc,
                "หลังบันทึกรายการสำเร็จ ให้เปิด Journal Entry เพื่อตรวจดูว่าระบบบันทึกเงินเข้า เงินออก หรือบัญชีพักไว้ที่บัญชีใดบ้างตามตัวอย่างจริงด้านล่าง.",
            )
        )
    for caption, image_name in items:
        blocks.extend(add_image_blocks_at_end(doc, caption, IMAGE_DIR / image_name))
    return blocks


def detach_blocks(doc: Document, blocks: list) -> list:
    body = doc._element.body
    detached = []
    for block in blocks:
        if block.getparent() is body:
            body.remove(block)
            detached.append(block)
    return detached


def rebuild_document(doc: Document, title_blocks: list, sections: dict[str, list], nav_blocks: list, je_blocks: list):
    body = doc._element.body
    sect_pr = None
    for child in list(body.iterchildren()):
        if child.tag.endswith("}sectPr"):
            sect_pr = child
        else:
            body.remove(child)

    def append_blocks(blocks: list):
        for block in blocks:
            body.append(block)

    append_blocks(title_blocks)

    for section_name in SECTION_ORDER:
        if section_name not in sections:
            continue
        if section_name == "1. เข้าใช้งานจากหน้าหลัก":
            append_blocks(sections[section_name])
            append_blocks(nav_blocks)
        elif section_name == "3. ตรวจสอบรายการบัญชี (Journal Entry)":
            append_blocks(sections[section_name])
            append_blocks(je_blocks)
        else:
            append_blocks(sections[section_name])

    if sect_pr is not None and sect_pr.getparent() is None:
        body.append(sect_pr)


def process_doc(path: Path) -> None:
    doc = Document(path)
    rename_headings_and_text(doc)
    title_blocks, sections = segment_sections(doc)

    old_extra_blocks = sections.pop("ภาพพาเข้าเมนูและ Journal Entry จากข้อมูลจริงในระบบ", [])

    spec = None
    for prefix, candidate in DOC_IMAGE_SPEC.items():
        if path.name.startswith(prefix):
            spec = candidate
            break
    if spec is None:
        spec = {"nav": [], "je": []}

    nav_blocks = create_extra_blocks(doc, "nav", spec["nav"])
    je_blocks = create_extra_blocks(doc, "je", spec["je"])

    nav_blocks = detach_blocks(doc, nav_blocks)
    je_blocks = detach_blocks(doc, je_blocks)
    detach_blocks(doc, old_extra_blocks)

    rebuild_document(doc, title_blocks, sections, nav_blocks, je_blocks)
    doc.save(path)


def main() -> None:
    for path in sorted(DOCX_DIR.glob("*.docx")):
        process_doc(path)
        print(path.name)


if __name__ == "__main__":
    main()

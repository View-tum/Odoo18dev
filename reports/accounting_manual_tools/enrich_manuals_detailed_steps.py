from __future__ import annotations

import re
import subprocess
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph


ROOT = Path(r"C:\365_project\TheCool18e\Dev")
DOCX_DIR = ROOT / "manual" / "Accouting_Manual" / "generated_20260408" / "docx"
PDF_DIR = ROOT / "manual" / "Accouting_Manual" / "generated_20260408" / "pdf_review"
SOFFICE = Path(r"C:\Program Files\LibreOffice\program\soffice.com")
TARGET_PREFIXES = ("3.8_", "5.", "6.", "7.")


def set_run_font(run, size: int = 16, bold: bool = False) -> None:
    run.bold = bold
    run.font.name = "TH Sarabun New"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "TH Sarabun New")
    run.font.size = Pt(size)


def insert_paragraph_after(paragraph, text: str = ""):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        run = new_para.add_run(text)
        set_run_font(run)
    return new_para


def is_heading(text: str) -> bool:
    text = text.strip()
    return text.startswith(("1. ", "2. ", "3. ", "4. ", "5. ")) or text in {
        "เมนูที่ใช้",
        "ขั้นตอนการใช้งาน",
        "ภาพประกอบการทำงาน",
        "คำอธิบายหน้าจอและช่องสำคัญ",
        "คำอธิบายฟิลด์และส่วนสำคัญบนหน้าจอ",
        "ตัวอย่างการใช้งานในสถานการณ์จริง",
        "การอธิบาย Journal Items และขาบัญชี",
        "ข้อควรระวัง",
    }


def strip_number(text: str) -> str:
    return re.sub(r"^\d+\.\s*", "", text.strip())


def collect_menu_paths(paragraphs: list) -> list[str]:
    start = None
    for idx, p in enumerate(paragraphs):
        if "ไปที่เมนูที่ใช้" in p.text or p.text.strip() == "เมนูที่ใช้":
            start = idx
            break
    if start is None:
        return []
    results: list[str] = []
    for p in paragraphs[start + 1 :]:
        text = p.text.strip()
        if not text:
            continue
        if text.startswith(("3. ", "ขั้นตอนการใช้งาน", "รูป ", "คำอธิบาย", "ภาพประกอบ", "ตัวอย่างการใช้งาน", "การอธิบาย Journal", "ข้อควรระวัง")):
            break
        if ">" in text:
            results.append(text)
    return results


def collect_existing_steps(paragraphs: list) -> tuple[int | None, int | None, list[str]]:
    heading_idx = None
    for idx, p in enumerate(paragraphs):
        if "ขั้นตอนการใช้งาน" in p.text:
            heading_idx = idx
            break
    if heading_idx is None:
        return None, None, []
    insert_after_idx = heading_idx
    if heading_idx + 1 < len(paragraphs):
        next_text = paragraphs[heading_idx + 1].text.strip()
        if next_text and not re.match(r"^\d+\.\s*", next_text) and not next_text.startswith(("รูป ", "คำอธิบาย", "ภาพประกอบ", "ตัวอย่างการใช้งาน", "4. ", "ข้อควรระวัง")):
            insert_after_idx = heading_idx + 1
    steps: list[str] = []
    for p in paragraphs[heading_idx + 1 :]:
        text = p.text.strip()
        if not text:
            continue
        if text == "ลำดับการทำงานแบบละเอียด":
            break
        if text.startswith(("รูป ", "คำอธิบาย", "ภาพประกอบ", "ตัวอย่างการใช้งาน", "4. ", "ข้อควรระวัง")):
            break
        if re.match(r"^\d+\.\s*", text):
            steps.append(strip_number(text))
    return heading_idx, insert_after_idx, steps


def build_detailed_steps(doc: Document) -> list[str]:
    paragraphs = doc.paragraphs
    menu_paths = collect_menu_paths(paragraphs)
    _, _, steps = collect_existing_steps(paragraphs)
    module = "Accounting"
    if menu_paths:
        module = menu_paths[0].split(">")[0].strip()
    detailed: list[str] = [
        "เริ่มจากหน้า Dashboard ของระบบ แล้วตรวจว่าผู้ใช้เข้าบริษัทและเมนูได้ถูกต้องก่อนเริ่มทำรายการ",
        f"คลิกเข้าโมดูล {module} จากหน้า Dashboard",
    ]
    if menu_paths:
        first_menu = menu_paths[0]
        detailed.append(f"เมื่อเข้ามาในโมดูลแล้ว ให้ไปที่เมนู {first_menu}")
        for extra_menu in menu_paths[1:]:
            detailed.append(f"ถ้าขั้นตอนนี้ต้องเปิดหน้าจอที่เกี่ยวข้องเพิ่มเติม ให้ไปที่เมนู {extra_menu}")
    detailed.append("เมื่อหน้าจอเปิดแล้ว ให้ตรวจชื่อเอกสาร วันที่ คู่ค้า และจำนวนเงินหรือจำนวนสินค้าให้ตรงกับรายการที่ต้องการทำ")
    for step in steps:
        detailed.append(step)
    detailed.append("หลังบันทึกหรือยืนยันรายการแล้ว ให้กลับมาตรวจผลที่หน้าจอเอกสารและหน้ารายการบัญชีทุกครั้ง")
    return detailed


def enrich_doc(path: Path) -> bool:
    doc = Document(path)
    paragraphs = doc.paragraphs
    texts = [p.text.strip() for p in paragraphs]
    if "ลำดับการทำงานแบบละเอียด" in texts:
        return False
    heading_idx, insert_after_idx, _ = collect_existing_steps(paragraphs)
    if heading_idx is None or insert_after_idx is None:
        return False
    detailed_steps = build_detailed_steps(doc)
    anchor = paragraphs[insert_after_idx]
    subheading = insert_paragraph_after(anchor, "ลำดับการทำงานแบบละเอียด")
    if subheading.runs:
        set_run_font(subheading.runs[0], 16, bold=True)
    current = subheading
    for idx, item in enumerate(detailed_steps, start=1):
        current = insert_paragraph_after(current, f"{idx}. {item}")
        if current.runs:
            set_run_font(current.runs[0], 16)
    doc.save(path)
    return True


def render_pdfs(paths: list[Path]) -> None:
    if not paths:
        return
    PDF_DIR.mkdir(parents=True, exist_ok=True)
    cmd = [str(SOFFICE), "--headless", "--convert-to", "pdf", "--outdir", str(PDF_DIR), *map(str, paths)]
    subprocess.run(cmd, check=True)


def main() -> None:
    targets = [p for p in DOCX_DIR.glob("*.docx") if p.name.startswith(TARGET_PREFIXES)]
    updated: list[Path] = []
    for path in targets:
        if enrich_doc(path):
            updated.append(path)
    render_pdfs(updated)
    print(f"updated {len(updated)} docs")


if __name__ == "__main__":
    main()

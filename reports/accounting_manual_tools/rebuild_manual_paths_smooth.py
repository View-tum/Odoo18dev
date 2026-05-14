from __future__ import annotations

from copy import deepcopy
from importlib.util import module_from_spec, spec_from_file_location
from pathlib import Path
import re
import subprocess

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph

ROOT = Path(r"C:\365_project\TheCool18e\Dev")
DOCX_DIR = ROOT / "manual" / "Accouting_Manual" / "generated_20260408" / "docx"
PDF_DIR = ROOT / "manual" / "Accouting_Manual" / "generated_20260408" / "pdf_review"
SOFFICE = Path(r"C:\Program Files\LibreOffice\program\soffice.com")
PAIR_SCRIPT = ROOT / "reports" / "accounting_manual_tools" / "pair_manual_images_with_steps.py"

SECTION_STOP = "4. ตรวจสอบ Journal Entry"
SECTION_MENU = "เมนูที่ใช้"
SECTION_STEPS = "3. ขั้นตอนการใช้งาน"

START_MARKERS = {
    SECTION_MENU,
    "ฟังก์ชันที่ทำได้จริง",
    "สิ่งที่ทำได้ในหน้าจอนี้",
    "1. เริ่มจากหน้า Dashboard",
    "2. ไปที่เมนูที่ใช้",
    SECTION_STEPS,
    "ขั้นตอนการใช้งาน",
    "ลำดับการทำงานแบบละเอียด",
}

REMOVE_TEXTS = {
    SECTION_MENU,
    "ฟังก์ชันที่ทำได้จริง",
    "สิ่งที่ทำได้ในหน้าจอนี้",
    "1. เริ่มจากหน้า Dashboard",
    "2. ไปที่เมนูที่ใช้",
    SECTION_STEPS,
    "ขั้นตอนการใช้งาน",
    "ลำดับการทำงานแบบละเอียด",
    "ภาพประกอบการใช้งาน",
    "ภาพประกอบตามลำดับขั้นตอน",
    "คำอธิบายหน้าจอและช่องสำคัญ",
    "คำอธิบายฟิลด์และส่วนสำคัญบนหน้าจอ",
    "ตารางอธิบายฟิลด์สำคัญใน Settings และ Journal",
    "ตารางอธิบายฟิลด์สำคัญของการจ่ายเช็คขาออก",
    "ตัวอย่างการใช้งานในสถานการณ์จริง",
}

SKIP_PREFIXES = (
    "รูป ",
    "ภาพ ",
    "หัวข้อนี้",
    "เริ่มจากหน้า Dashboard",
    "จากหน้า Dashboard",
    "คลิกเข้าโมดูล",
    "เมื่อเข้ามาในโมดูลแล้ว",
)


def load_pair_module():
    spec = spec_from_file_location("pair_manual_images_with_steps", PAIR_SCRIPT)
    module = module_from_spec(spec)
    assert spec.loader is not None
    spec.loader.exec_module(module)
    return module


PAIR = load_pair_module()
IMAGE_MAP = PAIR.IMAGE_MAP
IMAGE_DIR = PAIR.IMAGE_DIR

OVERRIDES: dict[str, dict[str, list[str]]] = {
    "3.8": {
        "menus": ["Accounting > Customers > รับชำระเงินกลุ่มลูกค้า"],
        "steps": [
            "เริ่มจากหน้า Dashboard แล้วคลิกเข้าโมดูล Accounting",
            "เมื่อเข้ามาในโมดูล Accounting แล้ว ให้ไปที่เมนู Accounting > Customers > รับชำระเงินกลุ่มลูกค้า",
            "เปิดเอกสารรับชำระแบบกลุ่มบริษัทที่ต้องการทำรายการ หรือกด New เพื่อสร้างรายการใหม่ พร้อมตรวจกลุ่มลูกค้า สมาชิก วันที่รับชำระ และยอดรวมให้ตรงกับข้อมูลจริง",
            "เลือกกลุ่มลูกค้าและสมาชิกในกลุ่มที่ต้องการรับเงิน แล้วกดค้นหาเอกสารค้างชำระ",
            "ตรวจสอบใบแจ้งหนี้ที่ดึงขึ้นมา เลือกเฉพาะบรรทัดที่ต้องการรับเงิน แล้วบันทึกหรือยืนยันรายการ จากนั้นกลับมาตรวจสถานะและรายการบัญชีของเอกสาร",
        ],
    },
    "5.7": {
        "menus": [
            "Cheque > Cheque > Cheque Paying",
            "Cheque > Cheque > Cheque Receiving",
        ],
        "steps": [
            "เริ่มจากหน้า Dashboard แล้วคลิกเข้าโมดูล Cheque",
            "เมื่อเข้ามาในโมดูล Cheque แล้ว ให้ใช้เมนู Cheque > Cheque > Cheque Paying และ Cheque > Cheque > Cheque Receiving ตามลำดับของงาน",
            "เปิดเช็คที่ต้องการเคลียร์ เช่นเช็คจ่าย 860200001 หรือเช็ครับ RCV-INV-E/26/04/00003 แล้วตรวจเลขเอกสาร วันที่ คู่ค้า วิธีชำระ และยอดเงินให้ตรงกับรายการจริง",
            "หากเช็คยังอยู่สถานะ Confirmed ให้กด Bank Deposit เมื่อนำเช็คฝากธนาคารหรือส่งตัดผ่าน และเมื่อธนาคารตัดผ่านจริงให้กด Done จากนั้นกลับมาตรวจสถานะและรายการบัญชีอีกครั้ง",
        ],
    },
    "5.8": {
        "menus": [
            "Cheque > Cheque > Cheque Paying",
            "Cheque > Cheque > Void Cheque",
        ],
        "steps": [
            "เริ่มจากหน้า Dashboard แล้วคลิกเข้าโมดูล Cheque",
            "เมื่อเข้ามาในโมดูล Cheque แล้ว ให้ใช้เมนู Cheque > Cheque > Cheque Paying และ Cheque > Cheque > Void Cheque ตามลำดับของงาน",
            "เปิดเช็คจริงที่ต้องการยกเลิก เช่น 860200003 ตรวจเลขเอกสาร วันที่ คู่ค้า วิธีชำระ และสถานะให้ถูกต้อง จากนั้นกด Void กรอกเหตุผล และยืนยันรายการก่อนกลับมาตรวจรายการกลับทางและสถานะสุดท้าย",
        ],
    },
}


def set_run_font(run, size: int = 16, bold: bool = False, italic: bool = False) -> None:
    run.bold = bold
    run.italic = italic
    run.font.name = "TH Sarabun New"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "TH Sarabun New")
    run.font.size = Pt(size)


def make_empty_after(paragraph: Paragraph) -> Paragraph:
    new_p = deepcopy(paragraph._p)
    for child in list(new_p):
        new_p.remove(child)
    paragraph._p.addnext(new_p)
    return Paragraph(new_p, paragraph._parent)


def insert_text_after(
    paragraph: Paragraph,
    text: str,
    *,
    size: int = 16,
    bold: bool = False,
    italic: bool = False,
    style: str | None = None,
) -> Paragraph:
    new_para = make_empty_after(paragraph)
    if style:
        try:
            new_para.style = style
        except Exception:
            pass
    run = new_para.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic)
    return new_para


def insert_picture_after(paragraph: Paragraph, image_path: Path, caption: str) -> Paragraph:
    pic_para = make_empty_after(paragraph)
    run = pic_para.add_run()
    run.add_picture(str(image_path), width=Inches(6.2))
    cap_para = insert_text_after(pic_para, caption, size=14, italic=True)
    return cap_para


def remove_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def doc_code_from_name(name: str) -> str | None:
    match = re.match(r"^(\d+\.\d+)_", name)
    return match.group(1) if match else None


def clean_text(text: str) -> str:
    return re.sub(r"\s+", " ", text.strip())


def find_journal_idx(paragraphs: list[Paragraph]) -> int | None:
    for idx, paragraph in enumerate(paragraphs):
        if clean_text(paragraph.text) == SECTION_STOP:
            return idx
    return None


def find_rebuild_start(paragraphs: list[Paragraph], journal_idx: int) -> int | None:
    for idx, paragraph in enumerate(paragraphs[:journal_idx]):
        text = clean_text(paragraph.text)
        if text in START_MARKERS:
            return idx
    return None


def extract_menu_lines(paragraphs: list[Paragraph], start: int, stop: int) -> list[str]:
    lines: list[str] = []
    in_menu = False
    for paragraph in paragraphs[start:stop]:
        text = clean_text(paragraph.text)
        if not text:
            continue
        if text == SECTION_MENU:
            in_menu = True
            continue
        if in_menu:
            if text in START_MARKERS - {SECTION_MENU}:
                break
            if text.startswith(("รูป ", "ภาพ ")):
                continue
            if re.match(r"^[A-Za-z].*\s>\s", text):
                lines.append(text)
    if not lines:
        for paragraph in paragraphs[start:stop]:
            text = clean_text(paragraph.text)
            if text.startswith(("รูป ", "ภาพ ")):
                continue
            if re.match(r"^[A-Za-z].*\s>\s", text):
                lines.append(text)
                continue
            for match in re.finditer(r"(Accounting|Cheque|Inventory|Manufacturing)(?:\s*>\s*[^>]+?){1,5}", text):
                candidate = match.group(0).strip()
                for token in (" และ ", " แล้ว", " เพื่อ", " ตาม", " พร้อม", " โดย", " จากนั้น"):
                    if token in candidate:
                        candidate = candidate.split(token, 1)[0].strip()
                if " > " in candidate:
                    lines.append(candidate)
    ordered: list[str] = []
    seen: set[str] = set()
    for line in lines:
        if line not in seen:
            seen.add(line)
            ordered.append(line)
    return ordered


def merge_parts(parts: list[str]) -> str:
    cleaned = [clean_text(part).rstrip(" .") for part in parts if clean_text(part)]
    if not cleaned:
        return ""
    merged = cleaned[0]
    for extra in cleaned[1:]:
        if extra.startswith(
            ("กด", "เลือก", "กรอก", "ตรวจ", "เปิด", "กลับ", "เมื่อ", "หาก", "แล้ว", "ในส่วน", "เมื่อต้องการ")
        ):
            merged += f" จากนั้น {extra}"
        else:
            merged += f" และ {extra}"
    return merged


def normalize_group_title(text: str) -> str:
    text = clean_text(text)
    text = re.sub(r"^ขั้นตอนที่\s*\d+\s*", "", text)
    return text.rstrip(" .")


def extract_operation_steps(paragraphs: list[Paragraph], start: int, stop: int) -> list[str]:
    groups: list[dict[str, object]] = []
    current: dict[str, object] | None = None
    for paragraph in paragraphs[start:stop]:
        text = clean_text(paragraph.text)
        if not text or text in REMOVE_TEXTS or text.startswith(("รูป ", "ภาพ ")):
            continue
        if any(text.startswith(prefix) for prefix in SKIP_PREFIXES):
            continue
        if text.startswith("ขั้นตอนที่"):
            current = {"title": normalize_group_title(text), "items": []}
            groups.append(current)
            continue
        if re.match(r"^\d+\.\s+", text):
            item = re.sub(r"^\d+\.\s*", "", text).strip()
            if any(item.startswith(prefix) for prefix in ("เริ่มจากหน้า Dashboard", "คลิกเข้าโมดูล", "เมื่อเข้ามาในโมดูล")):
                continue
            if current is None:
                current = {"title": "", "items": []}
                groups.append(current)
            current["items"].append(item)
            continue
    steps: list[str] = []
    for group in groups:
        title = str(group["title"]).strip()
        items = list(group["items"])
        if title and items:
            steps.append(f"{title} โดย {merge_parts(items)}")
        elif items:
            steps.extend(items)
        elif title:
            steps.append(title)
    ordered: list[str] = []
    seen: set[str] = set()
    for step in steps:
        key = clean_text(step)
        if key and key not in seen:
            seen.add(key)
            ordered.append(step)
    return ordered


def unique_step_images(code: str) -> list[Path]:
    ordered: list[Path] = []
    seen: set[str] = set()
    for image_name, _caption in IMAGE_MAP.get(code, []):
        if image_name.startswith("journal_"):
            continue
        if image_name in seen:
            continue
        image_path = IMAGE_DIR / image_name
        if image_path.exists():
            seen.add(image_name)
            ordered.append(image_path)
    return ordered


def compress_steps_to_image_count(steps: list[str], image_count: int) -> list[str]:
    compressed = steps[:]
    if image_count <= 0:
        return []
    while len(compressed) > image_count and len(compressed) > 1:
        merge_at = max(2, len(compressed) - 2)
        merged = merge_parts([compressed[merge_at], compressed[merge_at + 1]])
        compressed = compressed[:merge_at] + [merged] + compressed[merge_at + 2 :]
    return compressed


def module_name_from_menu(menu_lines: list[str]) -> str:
    if not menu_lines:
        return "โมดูลที่เกี่ยวข้อง"
    return menu_lines[0].split(" > ", 1)[0].strip()


def build_intro_steps(menu_lines: list[str]) -> list[str]:
    module_name = module_name_from_menu(menu_lines)
    intro = [f"เริ่มจากหน้า Dashboard แล้วคลิกเข้าโมดูล {module_name}"]
    if menu_lines:
        if len(menu_lines) == 1:
            intro.append(f"เมื่อเข้ามาในโมดูล {module_name} แล้ว ให้ไปที่เมนู {menu_lines[0]}")
        else:
            intro.append(f"เมื่อเข้ามาในโมดูล {module_name} แล้ว ให้ใช้เมนู {' และ '.join(menu_lines)} ตามลำดับของงาน")
    return intro


def remove_range(paragraphs: list[Paragraph], start: int, stop: int) -> None:
    for idx in range(stop - 1, start - 1, -1):
        remove_paragraph(paragraphs[idx])


def rebuild_doc(path: Path) -> bool:
    code = doc_code_from_name(path.name)
    if not code:
        return False
    doc = Document(path)
    paragraphs = doc.paragraphs
    journal_idx = find_journal_idx(paragraphs)
    if journal_idx is None:
        return False
    start_idx = find_rebuild_start(paragraphs, journal_idx)
    if start_idx is None:
        return False
    if code in OVERRIDES:
        menu_lines = OVERRIDES[code]["menus"]
        steps = OVERRIDES[code]["steps"]
    else:
        menu_lines = extract_menu_lines(paragraphs, start_idx, journal_idx)
        steps = build_intro_steps(menu_lines) + extract_operation_steps(paragraphs, start_idx, journal_idx)
    steps = [step for step in steps if clean_text(step)]
    image_paths = unique_step_images(code)
    steps = compress_steps_to_image_count(steps, len(image_paths))
    if not steps or not image_paths:
        return False

    anchor = paragraphs[start_idx - 1] if start_idx > 0 else paragraphs[0]
    remove_range(paragraphs, start_idx, journal_idx)

    current = anchor
    current = insert_text_after(current, SECTION_MENU, size=16, bold=True)
    for line in menu_lines:
        current = insert_text_after(current, line, size=16, style="List Bullet")
    current = insert_text_after(current, SECTION_STEPS, size=16, bold=True)
    for idx, step in enumerate(steps, start=1):
        current = insert_text_after(current, f"{idx}. {step}", size=16)
        if idx - 1 < len(image_paths):
            current = insert_picture_after(current, image_paths[idx - 1], f"ภาพประกอบขั้นตอนที่ {idx}")

    doc.save(path)
    return True


def render(paths: list[Path]) -> None:
    if not paths:
        return
    PDF_DIR.mkdir(parents=True, exist_ok=True)
    cmd = [str(SOFFICE), "--headless", "--convert-to", "pdf", "--outdir", str(PDF_DIR), *map(str, paths)]
    subprocess.run(cmd, check=True)


def main() -> None:
    updated: list[Path] = []
    for path in sorted(DOCX_DIR.glob("*.docx")):
        if rebuild_doc(path):
            updated.append(path)
    render(updated)
    print(f"rebuilt smooth paths in {len(updated)} docs")


if __name__ == "__main__":
    main()

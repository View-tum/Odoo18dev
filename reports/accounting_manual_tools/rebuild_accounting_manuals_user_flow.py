from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.shared import Inches


ROOT = Path(r"C:\365_project\TheCool18e\Dev")
DOCX_DIR = ROOT / "manual" / "Accouting_Manual" / "generated_20260408" / "docx"
IMAGE_DIR = ROOT / "manual" / "Accouting_Manual" / "generated_20260408" / "images"
PDF_DIR = ROOT / "manual" / "Accouting_Manual" / "generated_20260408" / "pdf_review"


DOC_SPECS = {
    "3.8_": {
        "dashboard_image": ("รูป 3.8A.1 หน้า Dashboard สำหรับเข้าโมดูล Accounting", "nav_dashboard_accounting_real_annotated.png"),
        "menu_image": ("รูป 3.8A.2 เมนู Accounting > Customers > รับชำระเงินกลุ่มลูกค้า", "nav_accounting_group_payment_real_annotated.png"),
        "menu_paths": ["Accounting > Customers > รับชำระเงินกลุ่มลูกค้า"],
        "usage_steps": [
            "เปิดเอกสารรับชำระแบบกลุ่มบริษัทที่ต้องการทำรายการ หรือกด New เพื่อสร้างรายการใหม่",
            "เลือกกลุ่มลูกค้าและสมาชิกในกลุ่มที่ต้องการรับเงิน แล้วกดค้นหาเอกสารค้างชำระ",
            "ตรวจสอบใบแจ้งหนี้ที่ดึงขึ้นมา เลือกเฉพาะบรรทัดที่ต้องการรับเงิน และตรวจยอดรวมให้ถูกต้อง",
            "กดปุ่มดำเนินการรับชำระ ระบบจะเปิดหน้าจอ Payment Register ให้เลือก Journal, วันที่รับเงิน และวิธีรับชำระ",
            "ตรวจสอบข้อมูลอีกครั้ง แล้วกด Create Payment เพื่อสร้างรายการรับเงินจริง",
        ],
        "usage_note": "หัวข้อนี้ใช้ข้อมูลจริงของเอกสารรับชำระแบบกลุ่มบริษัทที่มีอยู่ในระบบเพื่อให้ผู้ใช้เห็นขั้นตอนก่อนรับเงินและหลังรับเงินจริงครบในเรื่องเดียว",
        "usage_images": [],
        "je_intro": "หลังรับเงินแล้ว ให้เปิด Payment Entry หรือ Journal Entry เพื่อดูว่าระบบตัดลูกหนี้และรับเงินเข้าบัญชีใดบ้าง ผู้ใช้ควรดูทั้งชื่อบัญชี เดบิต เครดิต และยอดสุทธิทุกครั้ง",
        "je_images": [("รูป 3.8A.3 Journal Entry ของการรับชำระแบบกลุ่มบริษัทจากข้อมูลจริงในระบบ", "journal_group_payment_real_annotated.png")],
    },
    "5.1_": {
        "dashboard_image": ("รูป 5.1A.1 หน้า Dashboard สำหรับเข้าโมดูล Cheque", "nav_dashboard_cheque_real_annotated.png"),
        "menu_image": ("รูป 5.1A.2 เมนู Cheque > Configuration", "nav_cheque_configuration_real_annotated.png"),
        "menu_paths": ["Cheque > Configuration > Settings", "Accounting > Configuration > Journals"],
        "usage_steps": [
            "เข้าเมนู Settings ของโมดูลเช็คเพื่อตรวจว่าระบบเปิดใช้การรับเช็คและจ่ายเช็คแล้ว",
            "เปิด Journal ธนาคารที่ใช้งานจริง เช่น PBAY1 แล้วตรวจ Payment Methods ทั้งฝั่งรับและจ่าย",
            "ตรวจบัญชีพักเช็ครับและบัญชีพักเช็คจ่ายให้ครบ เพราะสองบัญชีนี้เป็นจุดที่ระบบจะบันทึกรายการชั่วคราวก่อนธนาคารตัดผ่าน",
            "บันทึกการตั้งค่า แล้วทดลองเปิดหน้าจอ Register Payment ของเอกสารจริง 1 ใบเพื่อยืนยันว่าช่องเช็คแสดงขึ้นมาแล้ว",
        ],
        "usage_note": "หัวข้อนี้เป็นการตั้งค่าพื้นฐาน จึงยังไม่สร้าง Journal Entry โดยตรง แต่ถ้าตั้งค่าไม่ครบ ผู้ใช้จะไม่สามารถทำหัวข้อรับเช็คและจ่ายเช็คต่อได้",
        "usage_images": [],
        "je_intro": "แม้หน้าตั้งค่านี้จะไม่ลงบัญชีเอง แต่ค่าที่ตั้งจะถูกนำไปใช้ตอนสร้าง Payment Entry จริง ตัวอย่างด้านล่างเป็น Journal Entry จริงจากฝั่งเช็คจ่ายและเช็ครับเพื่อให้ผู้ใช้เห็นผลของการตั้งค่า",
        "je_images": [
            ("รูป 5.1A.3 ตัวอย่าง Journal Entry ฝั่งจ่ายเช็คจากข้อมูลจริงในระบบ", "journal_cheque_out_confirmed_real_annotated.png"),
            ("รูป 5.1A.4 ตัวอย่าง Journal Entry ฝั่งรับเช็คจากข้อมูลจริงในระบบ", "journal_cheque_in_confirmed_real_annotated.png"),
        ],
    },
    "5.2_": {
        "dashboard_image": ("รูป 5.2A.1 หน้า Dashboard สำหรับเข้าโมดูล Cheque", "nav_dashboard_cheque_real_annotated.png"),
        "menu_image": ("รูป 5.2A.2 เมนู Cheque > Configuration เพื่อเข้าเลือกเทมเพลตฟอร์มเช็ค", "nav_cheque_configuration_real_annotated.png"),
        "menu_paths": ["Cheque > Configuration > Cheque Lists", "Accounting > Configuration > Journals"],
        "usage_steps": [
            "เปิดแบบฟอร์มเช็คที่ใช้งานอยู่จริง เช่น Standard Cheque",
            "ตรวจตำแหน่งวันที่ ชื่อผู้รับเงิน จำนวนเงินตัวเลข และจำนวนเงินตัวอักษรให้ตรงกับแบบกระดาษเช็คของธนาคาร",
            "กลับไปที่ Journal ที่ใช้งานจริง แล้วตรวจว่า Journal นั้นเลือกเทมเพลตฟอร์มเช็คถูกใบอยู่แล้ว",
            "พิมพ์ทดสอบกับกระดาษเปล่าก่อนทุกครั้ง เพื่อดูว่าตำแหน่งพิมพ์ตรงกับเช็คจริงหรือไม่",
        ],
        "usage_note": "หน้าจอนี้ใช้สำหรับจัดตำแหน่งการพิมพ์บนเช็คจริง ผู้ใช้ควรดูทั้งแบบฟอร์มเช็คและ Journal ที่จะนำแบบฟอร์มนี้ไปใช้ร่วมกัน",
        "usage_images": [],
        "je_intro": "หน้าจอเทมเพลตเช็คไม่สร้าง Journal Entry โดยตรง แต่จะถูกใช้ตอนพิมพ์เช็คจากรายการจ่ายจริง ตัวอย่างด้านล่างจึงแสดง Journal Entry ของเช็คจ่ายจริงที่อาศัยเทมเพลตนี้ร่วมด้วย",
        "je_images": [("รูป 5.2A.3 ตัวอย่าง Journal Entry ของเช็คจ่ายจากข้อมูลจริงในระบบ", "journal_cheque_out_confirmed_real_annotated.png")],
    },
    "5.3_": {
        "dashboard_image": ("รูป 5.3A.1 หน้า Dashboard สำหรับเข้าโมดูล Cheque", "nav_dashboard_cheque_real_annotated.png"),
        "menu_image": ("รูป 5.3A.2 เมนู Cheque > Configuration ที่ใช้เข้าสู่หน้าสมุดเช็ค", "nav_cheque_configuration_real_annotated.png"),
        "menu_paths": ["Cheque > Cheque Book"],
        "usage_steps": [
            "กด New เพื่อสร้างสมุดเช็คใหม่ หรือเปิดสมุดเช็คจริงที่ใช้งานอยู่เพื่อดูตัวอย่างข้อมูล",
            "เลือก Journal ธนาคาร แล้วกรอกเลขเช็คเริ่มต้น เลขเช็คสิ้นสุด และจำนวนใบให้ตรงกับเล่มจริง",
            "กด Submit และ Confirm เพื่อให้เลขเช็คแต่ละใบพร้อมใช้งาน",
            "เปิดบรรทัดเช็คในเล่มเพื่อตรวจว่าสถานะของแต่ละใบยังว่างและพร้อมถูกดึงไปใช้ตอนจ่ายเช็คจริง",
        ],
        "usage_note": "สมุดเช็คเป็นข้อมูลตั้งต้นของการจ่ายเช็ค ถ้าไม่มีเลขเช็คในเล่ม ระบบจะไม่สามารถให้ผู้ใช้เลือกเลขเช็คตอนจ่ายบิลได้",
        "usage_images": [],
        "je_intro": "หน้าจอสมุดเช็คไม่สร้าง Journal Entry ด้วยตัวเอง แต่เลขเช็คจากเล่มนี้จะถูกใช้ในรายการจ่ายจริง ตัวอย่างด้านล่างแสดงผล Journal Entry หลังนำเลขเช็คไปใช้ใน Payment Entry แล้ว",
        "je_images": [("รูป 5.3A.3 ตัวอย่าง Journal Entry ของเช็คจ่ายที่ใช้เลขเช็คจากสมุดเช็คจริง", "journal_cheque_out_confirmed_real_annotated.png")],
    },
    "5.4_": {
        "dashboard_image": ("รูป 5.4A.1 หน้า Dashboard สำหรับเข้าโมดูล Accounting", "nav_dashboard_accounting_real_annotated.png"),
        "menu_image": ("รูป 5.4A.2 เมนู Accounting > Vendors > Bills", "nav_accounting_vendors_bills_real_annotated.png"),
        "menu_paths": ["Accounting > Vendors > Bills", "Cheque > Cheque > Cheque Paying"],
        "usage_steps": [
            "เปิดบิลผู้ขายจริงที่ต้องการจ่าย เช่น APD/26/04/00006 หรือ APD/26/04/00007",
            "กด Register Payment แล้วเลือก Journal ธนาคารที่ใช้งานจริง และเลือกวิธีจ่ายเป็นเช็คขาออก",
            "กรอกเลขเช็ค ธนาคาร สาขา วันที่เช็ค และยอดเช็คให้ตรงกับเช็คจริง",
            "กด Create Payment ระบบจะสร้าง Payment Entry และสร้างเอกสารเช็คขาออกให้อัตโนมัติ",
            "ถ้าต้องการติดตามต่อ ให้เปิดเมนู Cheque Paying เพื่อดูสถานะเช็คว่าอยู่ที่ Confirmed, Bank Deposit หรือ Paid",
        ],
        "usage_note": "หัวข้อนี้ใช้ข้อมูลจริงของบิลผู้ขายและเช็คจ่ายในระบบ ผู้ใช้สามารถเปิดตามเลขเอกสารจริงแล้วทำตามลำดับได้ทันที",
        "usage_images": [],
        "je_intro": "หลังสร้างเช็คจ่ายแล้ว ให้เปิด Payment Entry และ Journal Items เพื่อตรวจว่าระบบลงบัญชีพักเช็คจ่ายอย่างไร และหลังเช็คผ่านธนาคารแล้วบัญชีใดถูกย้ายออกไป",
        "je_images": [
            ("รูป 5.4A.3 Journal Entry ตอนสร้างเช็คจ่ายสถานะ Confirmed", "journal_cheque_out_confirmed_real_annotated.png"),
            ("รูป 5.4A.4 Journal Entry หลังเช็คจ่ายถูกตัดผ่านธนาคาร", "journal_cheque_out_paid_real_annotated.png"),
        ],
    },
    "5.5_": {
        "dashboard_image": ("รูป 5.5A.1 หน้า Dashboard สำหรับเข้าโมดูล Cheque", "nav_dashboard_cheque_real_annotated.png"),
        "menu_image": ("รูป 5.5A.2 เมนู Cheque > Operations", "nav_cheque_operations_real_annotated.png"),
        "menu_paths": ["Cheque > Cheque > Cheque Paying", "Cheque > Cheque > Cheque Transactions", "Cheque > Cheque > Paid Cheque"],
        "usage_steps": [
            "เปิดเช็คขาออกที่ยังค้างอยู่จริง เช่น 860200001 เพื่อตรวจสถานะปัจจุบัน",
            "ดูวันที่เช็ค เลขเช็ค ผู้รับเงิน และยอดเช็ค แล้วเปิด Smart Button ของ Payment Entry เพื่อตรวจรายการบัญชีที่เกี่ยวข้อง",
            "เมื่อเช็คถูกนำฝากหรือส่งเข้าธนาคารแล้ว ให้กด Bank Deposit",
            "เมื่อธนาคารตัดผ่านจริง ให้กด Done แล้วตรวจว่าเช็คย้ายออกจากเมนูคงค้างไปยัง Paid Cheque",
        ],
        "usage_note": "หัวข้อนี้เน้นการติดตามเช็คที่ยังไม่ตัดผ่านธนาคาร ผู้ใช้ควรใช้เมนู Cheque Paying เป็นจุดหลัก แล้วเปิด Payment Entry ประกอบทุกครั้ง",
        "usage_images": [],
        "je_intro": "ผู้ใช้ควรเทียบ Journal Entry ก่อนและหลังเคลียร์เช็คเสมอ เพื่อดูว่าบัญชีพักเช็คจ่ายถูกล้างออก และเงินออกจากธนาคารจริงแล้วหรือยัง",
        "je_images": [
            ("รูป 5.5A.3 Journal Entry ที่ใช้ตรวจยอด Outstanding Cheque", "journal_cheque_out_confirmed_real_annotated.png"),
            ("รูป 5.5A.4 Journal Entry หลังเช็คถูกเคลียร์และย้ายออกจากบัญชีพัก", "journal_cheque_out_paid_real_annotated.png"),
        ],
    },
    "5.6_": {
        "dashboard_image": ("รูป 5.6A.1 หน้า Dashboard สำหรับเข้าโมดูล Accounting", "nav_dashboard_accounting_real_annotated.png"),
        "menu_image": ("รูป 5.6A.2 เมนู Accounting > Customers > Invoices", "nav_accounting_customers_invoices_real_annotated.png"),
        "menu_paths": ["Accounting > Customers > Invoices", "Cheque > Cheque > Cheque Receiving"],
        "usage_steps": [
            "เปิด Invoice จริง เช่น INV-E/26/04/00003 หรือ INV-E/26/04/00004",
            "กด Register Payment แล้วเลือก Journal = PBAY1 และ Payment Method = Cheque Payment (Inbound)",
            "กรอกเลขเช็ค ธนาคาร สาขา วันที่เช็ค และยอดเช็คให้ตรงกับเช็คที่ลูกค้านำมาชำระ",
            "กด Create Payment ระบบจะสร้าง Cheque Receiving record และสร้าง Payment Entry ให้อัตโนมัติ",
            "ถ้าต้องการติดตามต่อ ให้เปิดเมนู Cheque Receiving เพื่อดูสถานะว่าเช็คยังค้าง นำฝากแล้ว หรือผ่านธนาคารแล้ว",
        ],
        "usage_note": "หัวข้อนี้ใช้ Invoice และ Payment จริงในระบบ ผู้ใช้สามารถเปิดตามเลขเอกสารจริงและดูหน้าต่าง Register Payment ได้ตามภาพ",
        "usage_images": [("รูป 5.6A.3 หน้าต่าง Register Payment ของ Invoice ฝั่งรับเช็คจากข้อมูลจริงในระบบ", "invoice_register_payment_real_annotated.png")],
        "je_intro": "หลังรับเช็คแล้ว ให้เปิด Payment Entry และ Journal Items เพื่อตรวจว่าระบบลงบัญชีพักเช็ครับ และเมื่อเช็คผ่านธนาคารแล้วระบบย้ายยอดไปบัญชีธนาคารถูกต้องหรือไม่",
        "je_images": [
            ("รูป 5.6A.4 Journal Entry ตอนสร้างเช็ครับสถานะ Confirmed", "journal_cheque_in_confirmed_real_annotated.png"),
            ("รูป 5.6A.5 Journal Entry หลังเช็ครับถูกตัดผ่านธนาคาร", "journal_cheque_in_paid_real_annotated.png"),
        ],
    },
    "5.7_": {
        "dashboard_image": ("รูป 5.7A.1 หน้า Dashboard สำหรับเข้าโมดูล Cheque", "nav_dashboard_cheque_real_annotated.png"),
        "menu_image": ("รูป 5.7A.2 เมนู Cheque > Operations", "nav_cheque_operations_real_annotated.png"),
        "menu_paths": ["Cheque > Cheque > Cheque Paying", "Cheque > Cheque > Cheque Receiving"],
        "usage_steps": [
            "เปิดเช็คที่ต้องการเคลียร์ เช่นเช็คจ่าย 860200001 หรือเช็ครับ RCV-INV-E/26/04/00003",
            "ตรวจว่าสถานะปัจจุบันยังอยู่ที่ Confirmed ก่อน ถ้าใช่ให้กด Bank Deposit เมื่อนำเช็คฝากธนาคารหรือส่งตัดผ่าน",
            "เมื่อธนาคารตัดผ่านจริง ให้กด Done เพื่อปิดรายการ",
            "กลับไปตรวจว่าเอกสารถูกย้ายสถานะและไม่ค้างอยู่ในเมนูเดิมแล้ว",
        ],
        "usage_note": "ผู้ใช้ควรเคลียร์เช็คเมื่อมีหลักฐานจากธนาคารแล้วเท่านั้น เพื่อให้สถานะเช็คและยอดในบัญชีตรงกับของจริง",
        "usage_images": [],
        "je_intro": "หลังเช็คถูกเคลียร์แล้ว ให้เปิด Journal Entry เพื่อดูว่าบัญชีพักเช็คถูกล้างออกและยอดไปอยู่บัญชีธนาคารจริงแล้วทั้งฝั่งรับและฝั่งจ่าย",
        "je_images": [
            ("รูป 5.7A.3 Journal Entry ฝั่งจ่ายเช็คหลังเคลียร์เช็คแล้ว", "journal_cheque_out_paid_real_annotated.png"),
            ("รูป 5.7A.4 Journal Entry ฝั่งรับเช็คหลังเคลียร์เช็คแล้ว", "journal_cheque_in_paid_real_annotated.png"),
        ],
    },
    "5.8_": {
        "dashboard_image": ("รูป 5.8A.1 หน้า Dashboard สำหรับเข้าโมดูล Cheque", "nav_dashboard_cheque_real_annotated.png"),
        "menu_image": ("รูป 5.8A.2 เมนู Cheque > Operations ที่ใช้เข้าสู่หน้าเอกสารเช็ค", "nav_cheque_operations_real_annotated.png"),
        "menu_paths": ["Cheque > Cheque > Cheque Paying", "Cheque > Cheque > Void Cheque"],
        "usage_steps": [
            "เปิดเช็คจริงที่ต้องการยกเลิก เช่น 860200003",
            "ตรวจสถานะเช็คก่อน ถ้าอยู่ในสถานะที่ระบบอนุญาต ให้กด Void",
            "กรอกเหตุผลการยกเลิก แล้วกดยืนยันเพื่อให้ระบบสร้างรายการกลับรายการให้อัตโนมัติ",
            "ถ้าธุรกิจอนุญาตและตรวจแล้วว่าเอกสารปลายทางยังไม่กระทบรายการอื่น สามารถใช้ Reset To Draft เพื่อย้อนกลับไปแก้ข้อมูลได้",
            "หลังยกเลิกหรือย้อนสถานะ ให้กลับไปดู Payment Entry, Reversed Entry และสถานะของเลขเช็คในสมุดเช็คทุกครั้ง",
        ],
        "usage_note": "หัวข้อนี้ใช้เช็คยกเลิกจริงในระบบเพื่ออธิบายการย้อนสถานะ ผู้ใช้ควรทำเฉพาะกรณีที่ตรวจเอกสารต้นทางและปลายทางครบแล้ว",
        "usage_images": [],
        "je_intro": "เมื่อยกเลิกเช็ค ระบบจะสร้าง Journal Entry กลับรายการ ผู้ใช้ต้องเปิดดูทั้งรายการเดิมและรายการกลับเพื่อยืนยันว่าบัญชีพักและบัญชีธนาคารถูกย้อนถูกต้อง",
        "je_images": [("รูป 5.8A.3 Journal Entry ฝั่ง Reverse หลังยกเลิกเช็ค", "journal_cheque_void_reverse_real_annotated.png")],
    },
}


def find_paragraph(doc: Document, text: str):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == text:
            return paragraph
    return None


def iter_body_children(doc: Document):
    body = doc._element.body
    for child in list(body.iterchildren()):
        if child.tag.endswith("}sectPr"):
            continue
        yield child


def remove_between(doc: Document, start_paragraph, end_paragraph) -> None:
    body = doc._element.body
    removing = False
    for child in list(iter_body_children(doc)):
        if child is start_paragraph._p:
            removing = True
            continue
        if child is end_paragraph._p:
            break
        if removing and child.getparent() is body:
            body.remove(child)


def insert_paragraph_before(doc: Document, anchor, text: str, style: str | None = None, bullet: bool = False):
    paragraph = doc.add_paragraph(style=style)
    if bullet:
        paragraph.style = "List Bullet"
    paragraph.add_run(text)
    anchor._p.addprevious(paragraph._p)
    return paragraph


def insert_image_before(doc: Document, anchor, caption: str, image_name: str) -> None:
    image_path = IMAGE_DIR / image_name
    if not image_path.exists():
        return
    image_para = doc.add_paragraph()
    image_para.alignment = 1
    image_para.add_run().add_picture(str(image_path), width=Inches(6.3))
    anchor._p.addprevious(image_para._p)
    caption_para = doc.add_paragraph()
    caption_para.alignment = 1
    caption_para.add_run(caption)
    anchor._p.addprevious(caption_para._p)


def insert_user_flow(doc: Document, anchor, spec: dict) -> None:
    insert_paragraph_before(doc, anchor, "1. เริ่มจากหน้า Dashboard", style="Heading 1")
    insert_paragraph_before(doc, anchor, "เริ่มจากหน้า Dashboard ของระบบ แล้วกดเข้าโมดูลตามภาพ ก่อนจะไปยังเมนูย่อยที่ใช้ทำรายการจริง")
    insert_image_before(doc, anchor, *spec["dashboard_image"])

    insert_paragraph_before(doc, anchor, "2. ไปที่เมนูที่ใช้", style="Heading 1")
    insert_paragraph_before(doc, anchor, "เมื่อเข้ามาในโมดูลแล้ว ให้ไปตามเมนูด้านล่างนี้ทีละรายการ")
    for path in spec["menu_paths"]:
        insert_paragraph_before(doc, anchor, path, bullet=True)
    insert_image_before(doc, anchor, *spec["menu_image"])

    insert_paragraph_before(doc, anchor, "3. ขั้นตอนการใช้งาน", style="Heading 1")
    if spec.get("usage_note"):
        insert_paragraph_before(doc, anchor, spec["usage_note"])
    for idx, step in enumerate(spec["usage_steps"], start=1):
        insert_paragraph_before(doc, anchor, f"{idx}. {step}")
    for caption, image_name in spec.get("usage_images", []):
        insert_image_before(doc, anchor, caption, image_name)

    insert_paragraph_before(doc, anchor, "4. ตรวจสอบ Journal Entry", style="Heading 1")
    insert_paragraph_before(doc, anchor, spec["je_intro"])
    for caption, image_name in spec["je_images"]:
        insert_image_before(doc, anchor, caption, image_name)


def normalize_simple_language(doc: Document) -> None:
    replacements = {
        "Wizard": "หน้าต่าง",
        "wizard": "หน้าต่าง",
        "custom": "เมนูเสริม",
        "flow": "ลำดับการทำงาน",
        "filter": "ตัวกรอง",
        "local UAT": "ระบบ",
        "local manual": "คู่มือ",
    }
    for paragraph in doc.paragraphs:
        text = paragraph.text
        if not text:
            continue
        for old, new in replacements.items():
            text = text.replace(old, new)
        if paragraph.text != text:
            paragraph.text = text
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    text = paragraph.text
                    for old, new in replacements.items():
                        text = text.replace(old, new)
                    if paragraph.text != text:
                        paragraph.text = text


def process_doc(path: Path, spec: dict) -> None:
    doc = Document(path)
    normalize_simple_language(doc)

    start_heading = find_paragraph(doc, "สิ่งที่ทำได้ในหน้าจอนี้")
    field_heading = find_paragraph(doc, "คำอธิบายช่องสำคัญบนหน้าจอ")
    if start_heading is None or field_heading is None:
        raise RuntimeError(f"Could not find anchor headings in {path.name}")

    remove_between(doc, start_heading, field_heading)
    insert_user_flow(doc, field_heading, spec)
    doc.save(path)


def main() -> None:
    for prefix, spec in DOC_SPECS.items():
        path = next(DOCX_DIR.glob(f"{prefix}*.docx"))
        process_doc(path, spec)
        print(path.name)


if __name__ == "__main__":
    main()

from __future__ import annotations

import json
import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\365_project\TheCool18e\Dev")
OUTPUT_DIR = ROOT / "manual" / "Accouting_Manual" / "generated_20260408"
DOCX_DIR = OUTPUT_DIR / "docx"
PDF_DIR = OUTPUT_DIR / "pdf_review"
IMAGE_DIR = OUTPUT_DIR / "images"
SAMPLES_PATH = ROOT / "reports" / "accounting_manual_tools" / "output" / "manual_live_samples_20260408.json"
INSPECTION_PATH = ROOT / "reports" / "accounting_manual_tools" / "accounting_manual_inspection.json"
SOFFICE = Path(r"C:\Program Files\LibreOffice\program\soffice.com")

DOCX_PATH = DOCX_DIR / "5_Module Cheque คู่มือรวม.docx"

FONT_NAME = "Angsana New"
BASE_SIZE = 16

IMAGE_CALLOUTS: dict[str, list[tuple[str, str]]] = {
    "nav_dashboard_cheque_real_annotated.png": [
        ("1", "การ์ด Cheque บนหน้า Dashboard ใช้กดเข้าโมดูลเช็ค"),
    ],
    "nav_dashboard_accounting_real_annotated.png": [
        ("1", "การ์ด Accounting บนหน้า Dashboard ใช้กดเข้าโมดูลบัญชี"),
    ],
    "nav_cheque_configuration_real_annotated.png": [
        ("1", "เมนู Cheque หลัก"),
        ("2", "เมนู Configuration"),
        ("3", "เมนู Settings สำหรับตั้งค่าเช็ค"),
    ],
    "nav_cheque_operations_real_annotated.png": [
        ("1", "เมนู Cheque หลัก"),
        ("2", "เมนู Cheque Book"),
        ("3", "เมนู Cheque Receiving หรือรายการเช็คฝั่งรับ"),
        ("4", "เมนู Cheque Transactions สำหรับติดตามเช็ค"),
        ("5", "เมนูย่อยอื่นที่ใช้ตามสถานะเช็ค เช่น Paid หรือ Void"),
    ],
    "nav_accounting_vendors_bills_real_annotated.png": [
        ("1", "เมนู Vendors"),
        ("2", "เมนู Bills"),
    ],
    "nav_accounting_customers_invoices_real_annotated.png": [
        ("1", "เมนู Customers"),
        ("2", "เมนู Invoices"),
    ],
    "cheque_book_manual_annotated.png": [
        ("1", "ส่วนหัวของสมุดเช็ค เช่น ชื่อเล่ม หรือข้อมูลหลักของเล่ม"),
        ("2", "ช่องที่ใช้เลือกธนาคารหรือแบบฟอร์มเช็ค"),
        ("3", "ตารางเลขเช็คที่ระบบสร้างให้พร้อมใช้งาน"),
    ],
    "vendor_bill_open_real_annotated.png": [
        ("1", "แถบปุ่มด้านบน ใช้เริ่มขั้นตอนจ่ายบิล"),
        ("2", "ข้อมูลผู้ขายและวันที่ของบิล"),
        ("3", "ยอดคงค้างที่ต้องตรวจให้ตรงก่อนจ่าย"),
        ("4", "รายการสินค้าและยอดของบิล"),
    ],
    "cheque_out_confirmed_manual_annotated.png": [
        ("1", "เลขเช็คหรือข้อมูลสำคัญของรายการเช็ค"),
        ("2", "ยอดเงินของเช็ค"),
        ("3", "สถานะของเช็ค"),
        ("4", "ปุ่มหรือส่วนงานที่ใช้เปลี่ยนสถานะเช็ค"),
    ],
    "cheque_out_paid_manual_annotated.png": [
        ("1", "สถานะเดิมของเช็ค"),
        ("2", "สถานะใหม่หลังเช็คผ่านธนาคาร"),
        ("3", "ข้อมูลยอดเงินหรือยอดที่เคลียร์แล้ว"),
    ],
    "cheque_in_confirmed_manual_annotated.png": [
        ("1", "เลขเช็ครับหรือข้อมูลอ้างอิง"),
        ("2", "ยอดเงินของเช็ครับ"),
        ("3", "สถานะเช็คที่ยังค้างอยู่"),
    ],
    "cheque_in_paid_manual_annotated.png": [
        ("1", "สถานะเดิมก่อนผ่านธนาคาร"),
        ("2", "สถานะหลังผ่านธนาคาร"),
        ("3", "ยอดเงินที่ระบบเคลียร์เข้าธนาคารแล้ว"),
    ],
    "cheque_out_void_manual_annotated.png": [
        ("1", "สถานะหลังยกเลิกเช็ค"),
        ("3", "ข้อมูลรายการเช็คที่ถูกยกเลิก"),
    ],
    "customer_invoice_open_real_annotated.png": [
        ("1", "แถบปุ่มด้านบน ใช้เริ่มขั้นตอนรับชำระจากลูกค้า"),
        ("2", "ข้อมูลลูกค้าและวันที่ของใบแจ้งหนี้"),
        ("3", "ยอดคงค้างที่ต้องตรวจให้ตรงก่อนรับชำระ"),
        ("4", "รายการสินค้าและยอดในใบแจ้งหนี้"),
    ],
    "journal_cheque_out_confirmed_real_annotated.png": [
        ("1", "ชื่อ Journal และวันที่ของรายการบัญชี"),
        ("2", "แถวชื่อบัญชีที่ระบบลงรายการ"),
        ("3", "ตาราง Journal Items ที่ใช้ดู Debit และ Credit"),
    ],
    "journal_cheque_out_paid_real_annotated.png": [
        ("1", "ชื่อ Journal ตอนเคลียร์เช็ค"),
        ("2", "บัญชีที่ถูกใช้ในรายการเคลียร์"),
        ("3", "ตาราง Debit และ Credit ตอนเงินออกจากธนาคาร"),
    ],
    "journal_cheque_in_confirmed_real_annotated.png": [
        ("1", "ชื่อ Journal ของการรับเช็ค"),
        ("2", "บัญชีลูกหนี้และบัญชีพักเช็ครับ"),
        ("3", "ตาราง Journal Items ที่แสดงการตัดลูกหนี้และพักยอดเช็ครับ"),
    ],
    "journal_cheque_in_paid_real_annotated.png": [
        ("1", "ชื่อ Journal ตอนเคลียร์เช็ครับ"),
        ("2", "บัญชีธนาคารและบัญชีพักเช็ครับ"),
        ("3", "ตาราง Journal Items ที่แสดงเงินเข้าธนาคารจริง"),
    ],
    "journal_cheque_void_reverse_real_annotated.png": [
        ("1", "ชื่อรายการย้อนกลับ"),
        ("2", "บัญชีที่ใช้ย้อนรายการ"),
        ("3", "ตาราง Journal Items ที่แสดงการกลับทิศทางของเดบิตและเครดิต"),
    ],
    "journal_outgoing_manual_annotated.png": [
        ("1", "บรรทัดวิธีจ่ายเงินหรือข้อมูลของ Journal"),
        ("2", "บัญชีที่ใช้กับวิธีชำระแบบเช็ค"),
    ],
}


def set_font(run, *, size=BASE_SIZE, bold=False, italic=False, color: RGBColor | None = None):
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def paragraph_keep(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    keep_lines = OxmlElement("w:keepLines")
    p_pr.append(keep_lines)


def paragraph_keep_with_next(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    keep = OxmlElement("w:keepNext")
    p_pr.append(keep)


def add_para(doc: Document, text: str, *, size=BASE_SIZE, bold=False, italic=False, align=None, space_after=6):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    set_font(r, size=size, bold=bold, italic=italic)
    paragraph_keep(p)
    return p


def add_heading(doc: Document, text: str, *, level=1):
    size = 20 if level == 1 else 18
    color = RGBColor(31, 78, 121) if level == 1 else RGBColor(0, 0, 0)
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_font(r, size=size, bold=True, color=color)
    p.paragraph_format.space_after = Pt(4)
    paragraph_keep(p)
    paragraph_keep_with_next(p)
    return p


def add_bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    set_font(r)
    paragraph_keep(p)


def add_numbered_step(doc: Document, idx: int, text: str):
    p = doc.add_paragraph()
    r = p.add_run(f"{idx}. {text}")
    set_font(r)
    p.paragraph_format.space_after = Pt(4)
    paragraph_keep(p)
    return p


def add_image(doc: Document, image_name: str, caption: str, *, width=6.65):
    path = IMAGE_DIR / image_name
    if not path.exists():
        add_para(doc, f"[ไม่พบภาพประกอบ: {image_name}]", italic=True)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    p.paragraph_format.space_after = Pt(3)
    paragraph_keep(p)
    cap = add_para(doc, caption, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=15, space_after=8)
    paragraph_keep(cap)
    callouts = IMAGE_CALLOUTS.get(image_name, [])
    if callouts:
        title = add_para(doc, "คำอธิบายกรอบสี่เหลี่ยมในภาพ", bold=True, space_after=2)
        paragraph_keep_with_next(title)
        for label, text in callouts:
            add_para(doc, f"{label}. {text}", space_after=2)


def set_cell(cell, text: str, *, bold=False):
    cell.text = str(text)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for p in cell.paragraphs:
        p.paragraph_format.space_after = Pt(2)
        for r in p.runs:
            set_font(r, size=16, bold=bold)


def shade_cell(cell, color="D9EAF7"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    tc_pr.append(shd)


def add_table(doc: Document, headers: list[str], rows: list[tuple]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        set_cell(table.rows[0].cells[idx], header, bold=True)
        shade_cell(table.rows[0].cells[idx])
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell(cells[idx], value)
    doc.add_paragraph()
    return table


def configure_doc(doc: Document):
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    style.font.size = Pt(BASE_SIZE)
    sec = doc.sections[0]
    sec.start_type = WD_SECTION_START.NEW_PAGE
    sec.top_margin = Inches(0.65)
    sec.bottom_margin = Inches(0.65)
    sec.left_margin = Inches(0.75)
    sec.right_margin = Inches(0.75)


def move_rows(lines: list[dict], meaning_by_code: dict[str, str]) -> list[tuple[str, str, str, str]]:
    rows = []
    for line in lines:
        rows.append(
            (
                f"{line['account_code']} {line['account_name']}",
                meaning_by_code.get(line["account_code"], line["label"]),
                f"{line['debit']:,.2f}",
                f"{line['credit']:,.2f}",
            )
        )
    return rows


def build_subsection_header(doc: Document, title: str, purpose: str, menu_path: str):
    add_heading(doc, title, level=1)
    add_heading(doc, "วัตถุประสงค์", level=2)
    add_para(doc, purpose)
    add_heading(doc, "เมนูที่ใช้", level=2)
    add_para(doc, menu_path)


def build_steps_with_images(doc: Document, steps: list[tuple[str, str | None, str | None]]):
    add_heading(doc, "ขั้นตอนการใช้งาน", level=2)
    for idx, (step_text, image_name, caption) in enumerate(steps, start=1):
        add_numbered_step(doc, idx, step_text)
        if image_name and caption:
            add_image(doc, image_name, caption)


def build_fields(doc: Document, rows: list[tuple[str, str]]):
    add_heading(doc, "ช่องสำคัญที่ควรดูบนหน้าจอ", level=2)
    add_table(doc, ["Field Name", "Meaning"], rows)


def build_journal(doc: Document, intro: list[str], image_name: str | None, image_caption: str | None, rows: list[tuple[str, str, str, str]]):
    add_heading(doc, "ตรวจสอบ Journal Entry", level=2)
    for line in intro:
        add_para(doc, line)
    if image_name and image_caption:
        add_image(doc, image_name, image_caption)
    if rows:
        add_table(doc, ["บัญชี", "ความหมาย", "Debit", "Credit"], rows)


def build_common_intro(doc: Document, bank_journal: dict, template: dict, cheque_book: dict):
    add_para(doc, "บทที่ 5 Module Cheque", size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=3)
    add_para(
        doc,
        "คู่มือการใช้งานเช็คสำหรับการรับเงินและการจ่ายเงิน",
        size=18,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=10,
    )
    add_heading(doc, "สารบัญ", level=1)
    for item in [
        "5.3 สร้างสมุดเช็ค",
        "5.4 ชำระบิลผู้ขายด้วยเช็ค",
        "5.5 ติดตามเช็คคงค้าง",
        "5.6 รับเช็คจากลูกค้า",
        "5.7 เคลียร์ (กระทบยอด) เช็ค",
        "5.8 ยกเลิกเช็คหรือเปลี่ยนสถานะเช็ค",
    ]:
        add_para(doc, item, space_after=2)
    doc.add_paragraph()
    add_heading(doc, "ภาพรวมของบทนี้", level=1)
    add_para(
        doc,
        "บทนี้อธิบายการใช้งานเช็คตั้งแต่การสร้างสมุดเช็ค การนำเช็คไปใช้รับและจ่ายเงินจริง "
        "รวมถึงการติดตามเช็คที่ยังไม่ผ่านธนาคาร การเคลียร์เช็ค และการยกเลิกรายการเมื่อมีการเปลี่ยนแปลง",
    )
    add_para(
        doc,
        f"ข้อมูลตัวอย่างทั้งหมดใช้จากระบบจริง เช่น Journal {bank_journal['code']} {bank_journal['name']} "
        f"เทมเพลต {template['name']} และสมุดเช็ค {cheque_book['name']}",
    )
    add_heading(doc, "ฟังก์ชันที่ใช้งานได้จริงในบทนี้", level=1)
    add_table(
        doc,
        ["ฟังก์ชัน", "ใช้ทำอะไร", "เมนูหรือหน้าจอที่ใช้"],
        [
            ("สร้างสมุดเช็ค", "กำหนดช่วงเลขเช็คสำหรับใช้งาน", "Cheque > Cheque Book"),
            ("จ่ายผู้ขายด้วยเช็ค", "ออกเช็คเพื่อตัดหนี้เจ้าหนี้", "Accounting > Vendors > Bills"),
            ("รับเช็คจากลูกค้า", "รับเช็คเพื่อตัดหนี้ลูกหนี้", "Accounting > Customers > Invoices"),
            ("ติดตามเช็คคงค้าง", "ดูว่าเช็คใบใดยังไม่ผ่านธนาคาร", "Cheque > Cheque"),
            ("เคลียร์เช็ค", "เปลี่ยนเช็คจากคงค้างเป็นผ่านธนาคาร", "Cheque > Cheque"),
            ("ยกเลิกเช็ค", "ย้อนรายการเมื่อเช็คถูกยกเลิก", "Cheque > Cheque"),
        ],
    )


def main():
    sys.stdout.reconfigure(encoding="utf-8")
    DOCX_DIR.mkdir(parents=True, exist_ok=True)
    PDF_DIR.mkdir(parents=True, exist_ok=True)
    samples = json.loads(SAMPLES_PATH.read_text(encoding="utf-8"))
    inspection = json.loads(INSPECTION_PATH.read_text(encoding="utf-8"))

    bank_journal = samples["bank_journal"]
    template = samples["template"]
    cheque_book = samples["cheque_book"]
    chq = samples["cheques"]

    out_cf = chq["outbound_confirmed"]
    out_pd = chq["outbound_paid"]
    out_void = chq["outbound_void"]
    in_cf = chq["inbound_confirmed"]
    in_pd = chq["inbound_paid"]

    doc = Document()
    configure_doc(doc)
    build_common_intro(doc, bank_journal, template, cheque_book)

    build_subsection_header(
        doc,
        "5.3 สร้างสมุดเช็ค (Cheque-Book)",
        "หัวข้อนี้ใช้สร้างชุดเลขเช็คที่จะนำไปใช้งานจริง ระบบจะนำเลขเหล่านี้ไปเรียกใช้ตอนจ่ายเช็คให้ผู้ขาย",
        "Dashboard > Cheque > Cheque Book",
    )
    build_steps_with_images(
        doc,
        [
            ("เริ่มจากหน้า Dashboard แล้วคลิก Cheque เพื่อเข้าโมดูล", "nav_dashboard_cheque_real_annotated.png", "เริ่มจากหน้า Dashboard"),
            ("ไปที่เมนู Cheque Book แล้วกด New เพื่อสร้างสมุดเช็คใหม่", "nav_cheque_operations_real_annotated.png", "เมนู Cheque Book สำหรับสร้างสมุดเช็ค"),
            (f"กรอกชื่อสมุดเช็ค เลือก Journal {bank_journal['code']} และเลือกเทมเพลต {template['name']} จากนั้นกำหนดเลขเริ่มต้นและเลขสิ้นสุดของเล่ม", "cheque_book_manual_annotated.png", "หน้าจอสร้างสมุดเช็ค"),
            ("บันทึกและยืนยันสมุดเช็ค เมื่อยืนยันแล้วระบบจะสร้างใบเช็คว่างตามช่วงเลขให้พร้อมใช้งาน", "cheque_book_manual_annotated.png", "ยืนยันสมุดเช็คเพื่อสร้างเลขเช็ค"),
        ],
    )
    build_fields(
        doc,
        [
            ("Name", "ชื่อสมุดเช็ค"),
            ("Journal", "ธนาคารที่สมุดเช็คเล่มนี้จะผูกใช้งาน"),
            ("Cheque Form Template", "แบบฟอร์มที่จะใช้พิมพ์เช็คจากเล่มนี้"),
            ("Start Number", "เลขเช็คใบแรกของเล่ม"),
            ("End Number", "เลขเช็คใบสุดท้ายของเล่ม"),
            ("State", "สถานะของสมุดเช็ค"),
        ],
    )
    build_journal(
        doc,
        [
            f"ตัวอย่างจริงในระบบมีสมุดเช็คชื่อ {cheque_book['name']} สถานะ {cheque_book['state']} และมีเลขเช็คพร้อมใช้ เช่น {', '.join(cheque_book['draft_leaves'][:3])}",
            "การสร้างสมุดเช็คยังไม่ทำให้เกิดรายการบัญชี แต่เป็นการเตรียมเลขเช็คเพื่อใช้ในขั้นตอนการจ่ายเช็คจริง",
        ],
        "cheque_book_manual_annotated.png",
        "ตัวอย่างสมุดเช็คที่สร้างไว้ในระบบจริง",
        [],
    )

    doc.add_page_break()

    build_subsection_header(
        doc,
        "5.4 ชำระบิลผู้ขายด้วยเช็ค",
        "หัวข้อนี้ใช้เมื่อต้องการออกเช็คเพื่อชำระหนี้ให้ผู้ขาย โดยระบบจะสร้างรายการบัญชีตอนบันทึกเช็ค และจะสร้างรายการเพิ่มเติมอีกครั้งตอนเช็คผ่านธนาคาร",
        "Dashboard > Accounting > Vendors > Bills",
    )
    build_steps_with_images(
        doc,
        [
            ("เริ่มจากหน้า Dashboard แล้วคลิกการ์ด Accounting เพื่อเข้าสู่เมนูบัญชีเจ้าหนี้", "nav_dashboard_accounting_real_annotated.png", "เริ่มจากหน้า Dashboard เพื่อไปเมนูบัญชี"),
            ("ไปที่เมนู Vendors > Bills แล้วเปิดบิลผู้ขายที่ต้องการจ่าย", "nav_accounting_vendors_bills_real_annotated.png", "เมนูบิลผู้ขาย"),
            ("เมื่อเปิดบิลแล้ว ให้ตรวจยอดคงค้าง ข้อมูลผู้ขาย และกดปุ่มที่แถบด้านบนเพื่อเริ่มการชำระ จากนั้นกรอก Journal วิธีชำระ เลขเช็ค วันที่เช็ค และจำนวนเงินให้ตรงกับเอกสารจริง", "vendor_bill_open_real_annotated.png", "หน้าจอบิลผู้ขายที่ใช้เริ่มขั้นตอนจ่ายด้วยเช็ค"),
            ("เมื่อบันทึกรายการแล้ว ให้เปิดเมนู Cheque เพื่อตรวจสถานะเช็ค ถ้าเช็คยังไม่ผ่านธนาคารจะอยู่ในสถานะ Confirmed", "cheque_out_confirmed_manual_annotated.png", "ตัวอย่างเช็คจ่ายที่ยังคงค้าง"),
            ("เมื่อเช็คผ่านธนาคารแล้ว ให้เปลี่ยนสถานะเป็น Paid และกลับมาตรวจว่าเอกสารถูกย้ายออกจากเช็คค้างแล้ว", "cheque_out_paid_manual_annotated.png", "ตัวอย่างเช็คจ่ายที่ผ่านธนาคารแล้ว"),
        ],
    )
    build_fields(
        doc,
        [
            ("Journal", "ธนาคารที่ใช้จ่ายเช็ค"),
            ("Payment Method", "วิธีชำระเงินที่เลือกเป็น Cheque Payment"),
            ("Cheque Number", "เลขเช็คที่ดึงจากสมุดเช็ค"),
            ("Cheque Date", "วันที่บนเช็ค"),
            ("Memo", "ข้อความอ้างอิงของรายการ"),
            ("State", "สถานะของเช็ค เช่น Confirmed หรือ Paid"),
        ],
    )
    build_journal(
        doc,
        [
            f"ตัวอย่างเช็คจ่ายสถานะ Confirmed คือเลข {out_cf['name']} อ้างอิงรายการบัญชี {out_cf['payment_moves'][0]['name']}",
            "เมื่อบันทึกเช็คจ่าย ระบบจะตัดเจ้าหนี้ออกก่อน และย้ายยอดไปพักไว้ในบัญชีเช็คจ่ายลงวันที่ล่วงหน้า",
            f"ตัวอย่างเช็คจ่ายสถานะ Paid คือเลข {out_pd['name']} เมื่อเคลียร์เช็คแล้วระบบสร้างรายการธนาคาร {out_pd['deposit_move']['name']} เพิ่มอีกใบ",
        ],
        "journal_cheque_out_confirmed_real_annotated.png",
        "Journal Entry ตอนบันทึกเช็คจ่ายผู้ขาย",
        move_rows(
            out_cf["payment_moves"][0]["lines"],
            {
                "212001": "ด้าน Debit แปลว่าระบบตัดยอดเจ้าหนี้ของบิลผู้ขายออกจากบัญชีเจ้าหนี้",
                "212004": "ด้าน Credit แปลว่าระบบย้ายยอดไปพักไว้ในบัญชีเช็คจ่ายลงวันที่ล่วงหน้า รอเช็คผ่านธนาคาร",
            },
        )
        + move_rows(
            out_pd["deposit_move"]["lines"],
            {
                "111201": "เมื่อเช็คผ่านธนาคาร ระบบเครดิตบัญชีธนาคารจริงเพื่อลดยอดเงินฝากตามเงินที่ออกจากบัญชี",
                "212004": "ระบบเดบิตบัญชีเช็คจ่ายลงวันที่ล่วงหน้าเพื่อล้างยอดพักเช็คจ่ายเดิม",
            },
        ),
    )

    doc.add_page_break()

    build_subsection_header(
        doc,
        "5.5 ติดตามเช็คคงค้าง (Outstanding Cheques)",
        "หัวข้อนี้ใช้ติดตามเช็คที่ออกไปหรือรับมาแล้ว แต่ยังไม่ผ่านธนาคาร เพื่อให้ผู้ใช้รู้ว่าเช็คใบใดยังค้างอยู่และต้องติดตามต่อ",
        "Dashboard > Cheque > Cheque > Cheque Transactions",
    )
    build_steps_with_images(
        doc,
        [
            ("เริ่มจากหน้า Dashboard แล้วคลิก Cheque เพื่อเข้าเมนูติดตามเช็ค", "nav_dashboard_cheque_real_annotated.png", "เริ่มจากหน้า Dashboard ไปโมดูล Cheque"),
            ("ไปที่เมนู Cheque > Cheque Transactions หรือเมนูย่อยที่ใช้ดูเช็คค้าง แล้วใช้ตัวกรองสถานะเพื่อดูเฉพาะเช็คที่ยังค้าง", "nav_cheque_operations_real_annotated.png", "เมนูติดตามเช็คคงค้าง"),
            ("เปิดเช็คแต่ละใบเพื่อตรวจเลขเช็ค ผู้รับหรือผู้จ่าย ยอดเงิน วันเช็ค และสถานะ หากยังเป็น Confirmed แปลว่าเช็คยังไม่ผ่านธนาคาร", "cheque_out_confirmed_manual_annotated.png", "ตัวอย่างเช็คจ่ายค้างในสถานะ Confirmed"),
            ("กรณีเป็นเช็ครับ ให้เปิดรายการเช็ครับค้างเพื่อตรวจว่าลูกค้านำเช็คเข้ามาแล้ว แต่ยังไม่ได้ฝากหรือยังไม่ผ่านธนาคาร", "cheque_in_confirmed_manual_annotated.png", "ตัวอย่างเช็ครับค้างในสถานะ Confirmed"),
        ],
    )
    build_fields(
        doc,
        [
            ("Name", "เลขที่เช็คหรือเลขเอกสารเช็ครับ"),
            ("Partner", "ชื่อผู้ขายหรือลูกค้าที่เกี่ยวข้องกับเช็ค"),
            ("Amount", "จำนวนเงินบนเช็ค"),
            ("Cheque Date", "วันที่หน้าเช็ค"),
            ("State", "สถานะของเช็คในปัจจุบัน"),
            ("Journal", "ธนาคารที่ใช้กับเช็คใบนี้"),
        ],
    )
    build_journal(
        doc,
        [
            "หน้าติดตามเช็คคงค้างเป็นหน้าสำหรับตรวจสอบสถานะเช็คเป็นหลัก จึงไม่ได้สร้าง Journal Entry ใหม่จากการเปิดดูหน้าจอนี้",
            "สิ่งที่ผู้ใช้ควรทำคือเปิดรายการบัญชีของเช็คแต่ละใบเพื่อยืนยันว่าตอนนี้ยอดยังพักอยู่ในบัญชีเช็ครับหรือเช็คจ่ายจริง",
            "ถ้าสถานะยังเป็น Confirmed ยอดจะยังค้างอยู่ในบัญชีพักเช็ค ไม่ได้เข้าออกธนาคารจริง",
        ],
        "journal_outgoing_manual_annotated.png",
        "หน้าจอ Journal ที่ใช้ตรวจสอบรายการของเช็คคงค้าง",
        [],
    )

    doc.add_page_break()

    build_subsection_header(
        doc,
        "5.6 รับเช็คจากลูกค้า",
        "หัวข้อนี้ใช้เมื่อกิจการได้รับเช็คจากลูกค้าเพื่อชำระหนี้ ระบบจะตัดยอดลูกหนี้และพักยอดไว้ในบัญชีเช็ครับลงวันที่ล่วงหน้า ก่อนจะเคลียร์เข้าธนาคารในภายหลัง",
        "Dashboard > Accounting > Customers > Invoices",
    )
    build_steps_with_images(
        doc,
        [
            ("เริ่มจากหน้า Dashboard แล้วคลิกการ์ด Accounting", "nav_dashboard_accounting_real_annotated.png", "เริ่มจากหน้า Dashboard ไปโมดูล Accounting"),
            ("ไปที่เมนู Customers > Invoices แล้วเปิดใบแจ้งหนี้ลูกค้าที่ต้องการรับชำระ", "nav_accounting_customers_invoices_real_annotated.png", "เมนูใบแจ้งหนี้ลูกค้า"),
            ("เมื่อเปิดใบแจ้งหนี้แล้ว ให้ตรวจยอดคงค้าง ข้อมูลลูกค้า และกดปุ่มที่แถบด้านบนเพื่อเริ่มรับชำระ จากนั้นกรอกวิธีรับเงิน เลขเช็ค วันที่เช็ค และยอดที่ลูกค้าชำระจริง", "customer_invoice_open_real_annotated.png", "หน้าจอใบแจ้งหนี้ลูกค้าที่ใช้เริ่มขั้นตอนรับชำระ"),
            ("หลังบันทึกรายการ ให้ไปตรวจที่เมนูเช็คเพื่อดูว่าเช็ครับใบนี้อยู่ในสถานะ Confirmed", "cheque_in_confirmed_manual_annotated.png", "ตัวอย่างเช็ครับหลังบันทึกรายการ"),
            ("เมื่อเช็ครับผ่านธนาคารแล้ว ให้เปลี่ยนสถานะเป็น Paid และตรวจว่าเช็ครับออกจากรายการคงค้างแล้ว", "cheque_in_paid_manual_annotated.png", "ตัวอย่างเช็ครับที่ผ่านธนาคารแล้ว"),
        ],
    )
    build_fields(
        doc,
        [
            ("Journal", "ธนาคารที่จะใช้รับเช็ค"),
            ("Payment Method", "วิธีรับเงินที่เลือกเป็น Cheque Payment"),
            ("Cheque Number", "เลขที่เช็ครับ"),
            ("Cheque Date", "วันที่บนเช็คที่ลูกค้ามอบให้"),
            ("Amount", "ยอดเงินที่รับตามเช็ค"),
            ("Customer", "ชื่อลูกค้าที่ชำระเงิน"),
        ],
    )
    build_journal(
        doc,
        [
            f"ตัวอย่างเช็ครับสถานะ Confirmed คือ {in_cf['name']} และระบบสร้างรายการบัญชี {in_cf['payment_moves'][0]['name']}",
            "ขั้นแรก ระบบจะตัดยอดลูกหนี้ก่อน แล้วนำยอดไปพักไว้ในบัญชีเช็ครับลงวันที่ล่วงหน้า เพราะเช็คยังไม่ผ่านธนาคาร",
            f"เมื่อเคลียร์เช็คแล้ว ระบบจะสร้างรายการธนาคาร {in_pd['deposit_move']['name']} เพื่อย้ายยอดจากบัญชีพักเช็ครับเข้าบัญชีธนาคารจริง",
        ],
        "journal_cheque_in_confirmed_real_annotated.png",
        "Journal Entry ตอนรับเช็คจากลูกค้า",
        move_rows(
            in_cf["payment_moves"][0]["lines"],
            {
                "113001": "ด้าน Credit แปลว่าระบบตัดยอดลูกหนี้ของใบแจ้งหนี้ออก เพราะลูกค้านำเช็คมาชำระแล้ว",
                "113005": "ด้าน Debit แปลว่าระบบพักยอดไว้ในบัญชีเช็ครับลงวันที่ล่วงหน้า รอเช็คผ่านธนาคาร",
            },
        )
        + move_rows(
            in_pd["deposit_move"]["lines"],
            {
                "111201": "เมื่อเช็คผ่านธนาคาร ระบบเดบิตบัญชีธนาคารจริงเพื่อเพิ่มยอดเงินฝาก",
                "113005": "ระบบเครดิตบัญชีเช็ครับลงวันที่ล่วงหน้าเพื่อล้างยอดพักเช็คเดิม",
            },
        ),
    )

    doc.add_page_break()

    build_subsection_header(
        doc,
        "5.7 เคลียร์ (กระทบยอด) เช็ค",
        "หัวข้อนี้ใช้เมื่อเช็คที่รับมาหรือจ่ายออกไปผ่านธนาคารแล้ว ผู้ใช้ต้องเปลี่ยนสถานะของเช็คและตรวจว่าระบบสร้างรายการย้ายยอดจากบัญชีพักเช็คไปบัญชีธนาคารจริง",
        "Dashboard > Cheque > Cheque > Cheque Transactions",
    )
    build_steps_with_images(
        doc,
        [
            ("เข้าโมดูล Cheque แล้วเปิดรายการเช็คที่อยู่ในสถานะ Confirmed ซึ่งหมายถึงเช็คยังค้างและยังไม่ผ่านธนาคาร", "nav_cheque_operations_real_annotated.png", "เมนูติดตามเช็คที่ใช้เคลียร์เช็ค"),
            ("เปิดเช็คแต่ละใบแล้วตรวจว่าจำนวนเงิน ผู้รับหรือผู้จ่าย และวันที่เช็คตรงกับข้อมูลที่ธนาคารแจ้งว่าผ่านแล้ว", "cheque_out_confirmed_manual_annotated.png", "ตัวอย่างเช็คจ่ายก่อนเคลียร์"),
            ("กดปุ่มเปลี่ยนสถานะหรือทำรายการให้เช็คเป็น Paid เพื่อยืนยันว่าเช็คผ่านธนาคารแล้ว", "cheque_out_paid_manual_annotated.png", "ตัวอย่างเช็คจ่ายหลังเคลียร์แล้ว"),
            ("สำหรับเช็ครับ ให้ทำในแนวทางเดียวกัน โดยตรวจเช็ครับที่เป็น Confirmed แล้วเปลี่ยนเป็น Paid เมื่อเงินเข้าธนาคารเรียบร้อย", "cheque_in_paid_manual_annotated.png", "ตัวอย่างเช็ครับหลังเคลียร์แล้ว"),
            ("กลับไปเปิด Journal Entry ของรายการเคลียร์เช็คเพื่อตรวจว่ายอดถูกย้ายออกจากบัญชีพักเช็คไปบัญชีธนาคารจริงแล้ว", "journal_cheque_out_paid_real_annotated.png", "Journal ตอนเคลียร์เช็คจ่าย"),
        ],
    )
    build_fields(
        doc,
        [
            ("State", "ใช้ดูว่าเช็คยังค้างหรือผ่านธนาคารแล้ว"),
            ("Paid Date", "วันที่เช็คผ่านธนาคาร"),
            ("Deposit Move", "รายการบัญชีที่ระบบสร้างตอนเคลียร์เช็ค"),
            ("Partner", "ผู้รับหรือผู้จ่ายที่เกี่ยวข้อง"),
            ("Amount", "ยอดเงินของเช็คที่ต้องตรงกับรายงานธนาคาร"),
        ],
    )
    build_journal(
        doc,
        [
            f"ตัวอย่างเคลียร์เช็คจ่ายใช้รายการบัญชี {out_pd['deposit_move']['name']}",
            f"ตัวอย่างเคลียร์เช็ครับใช้รายการบัญชี {in_pd['deposit_move']['name']}",
            "หลักการอ่านรายการบัญชีคือ บัญชีพักเช็คจะถูกล้างออก และบัญชีธนาคารจริงจะรับผลของรายการแทน",
        ],
        "journal_cheque_out_paid_real_annotated.png",
        "ตัวอย่าง Journal ตอนเคลียร์เช็คจ่าย",
        move_rows(
            out_pd["deposit_move"]["lines"],
            {
                "111201": "เช็คจ่ายผ่านธนาคารแล้ว เงินออกจากบัญชีธนาคารจริง จึงลงฝั่ง Credit ของบัญชีธนาคาร",
                "212004": "ระบบเดบิตบัญชีเช็คจ่ายลงวันที่ล่วงหน้าเพื่อล้างยอดพักเช็คจ่าย",
            },
        )
        + move_rows(
            in_pd["deposit_move"]["lines"],
            {
                "111201": "เช็ครับผ่านธนาคารแล้ว เงินเข้าบัญชีจริง จึงลงฝั่ง Debit ของบัญชีธนาคาร",
                "113005": "ระบบเครดิตบัญชีเช็ครับลงวันที่ล่วงหน้าเพื่อล้างยอดพักเช็ครับ",
            },
        ),
    )

    doc.add_page_break()

    build_subsection_header(
        doc,
        "5.8 การยกเลิกเช็คหรือแปลงสถานะเช็ค (Void or Transform a Cheque)",
        "หัวข้อนี้ใช้เมื่อเช็คที่บันทึกไปแล้วไม่สามารถใช้งานต่อได้ เช่น เช็คเสีย เช็คยกเลิก หรือมีความจำเป็นต้องย้อนรายการ เพื่อให้สถานะเช็คและรายการบัญชีกลับมาถูกต้อง",
        "Dashboard > Cheque > Cheque > Void Cheque",
    )
    build_steps_with_images(
        doc,
        [
            ("เข้าเมนู Cheque แล้วเปิดรายการเช็คที่ต้องการยกเลิก ตรวจให้แน่ใจก่อนว่าเป็นเช็คใบที่ต้องการย้อนจริง", "nav_cheque_operations_real_annotated.png", "เมนูที่ใช้เปิดรายการเช็คเพื่อยกเลิก"),
            ("ตรวจข้อมูลเลขเช็ค คู่ค้า และยอดเงินก่อนกด Void หรือเปลี่ยนสถานะ เพื่อป้องกันการย้อนรายการผิดใบ", "cheque_out_void_manual_annotated.png", "หน้าจอเช็คที่ถูกยกเลิก"),
            ("หลังยกเลิกแล้ว ให้กลับไปตรวจสถานะของเช็คว่าเป็น Cancelled และไม่มีผลค้างในรายการเช็คคงค้างอีกต่อไป", "cheque_out_void_manual_annotated.png", "สถานะเช็คหลังยกเลิก"),
            ("เปิด Journal Entry ของเอกสารย้อนกลับเพื่อตรวจว่าระบบสร้างรายการกลับทิศทางของเช็คเดิมให้แล้ว", "journal_cheque_void_reverse_real_annotated.png", "Journal ที่ระบบสร้างตอนยกเลิกเช็ค"),
        ],
    )
    build_fields(
        doc,
        [
            ("State", "สถานะหลังยกเลิกจะเปลี่ยนเป็น Cancelled"),
            ("Reverse Entry", "รายการบัญชีที่ระบบสร้างเพื่อย้อนผลของเช็คเดิม"),
            ("Reference", "ข้อความอ้างอิงที่ช่วยบอกว่าเป็นการยกเลิกรายการใด"),
            ("Partner", "คู่ค้าที่เกี่ยวข้องกับเช็คที่ยกเลิก"),
        ],
    )
    build_journal(
        doc,
        [
            f"ตัวอย่างจริงใช้เช็คเลข {out_void['name']} และระบบสร้างรายการย้อนกลับชื่อ {out_void['reversed_entry_names'][-1]}",
            f"จากภาพ ระบบย้อนอ้างอิงจากรายการเดิม {out_void['payment_moves'][0]['name']} เพื่อให้ยอดเจ้าหนี้และยอดพักเช็คกลับสู่สภาพก่อนออกเช็ค",
            "การยกเลิกเช็คต้องเปิดดู Journal ทุกครั้ง เพราะเป็นขั้นตอนที่มีผลต่อยอดเจ้าหนี้และยอดพักเช็คโดยตรง",
        ],
        "journal_cheque_void_reverse_real_annotated.png",
        "Journal ที่ระบบสร้างตอนยกเลิกเช็ค",
        [
            ("212004 เช็คจ่ายลงวันที่ล่วงหน้า", "รายการย้อนจะเดบิตบัญชีเช็คจ่ายลงวันที่ล่วงหน้าเพื่อยกเลิกผลของเช็คจ่ายเดิม", "9,000.00", "0.00"),
            ("212001 เจ้าหนี้การค้า - ในประเทศ", "รายการย้อนจะเครดิตบัญชีเจ้าหนี้เพื่อคืนยอดหนี้ให้กลับมาอยู่ในระบบอีกครั้ง", "0.00", "9,000.00"),
        ],
    )

    add_heading(doc, "สรุปการใช้งานบทที่ 5", level=1)
    add_para(
        doc,
        "บทนี้ครอบคลุมตั้งแต่การเตรียมระบบให้รองรับเช็ค ไปจนถึงการใช้งานจริงในทั้งฝั่งรับเงินและจ่ายเงิน "
        "รวมถึงการติดตามเช็คคงค้าง การเคลียร์เช็ค และการย้อนรายการเมื่อต้องยกเลิกเช็ค ผู้ใช้งานควรยึดหลักว่าในทุกธุรกรรมต้องกลับไปเปิดดู Journal Entry เสมอ "
        "เพื่อให้มั่นใจว่าระบบลงบัญชีฝั่งรับเงิน ฝั่งลูกหนี้ เจ้าหนี้ และบัญชีพักเช็คถูกต้องตรงกับเอกสารจริง",
    )

    doc.save(DOCX_PATH)
    subprocess.run(
        [str(SOFFICE), "--headless", "--convert-to", "pdf", "--outdir", str(PDF_DIR), str(DOCX_PATH)],
        check=True,
    )
    print(f"built {DOCX_PATH}")


if __name__ == "__main__":
    main()

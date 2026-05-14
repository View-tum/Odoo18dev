from __future__ import annotations

from pathlib import Path
import subprocess
import sys

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(r"C:\365_project\TheCool18e\Dev")
DOCX_DIR = ROOT / "manual" / "Accouting_Manual" / "generated_20260408" / "docx"
DOCX_PATH = next(DOCX_DIR.glob("3.8_*.docx"))
PDF_DIR = ROOT / "manual" / "Accouting_Manual" / "generated_20260408" / "pdf_review"
IMAGE_DIR = ROOT / "manual" / "Accouting_Manual" / "generated_20260408" / "images"
SOFFICE = Path(r"C:\Program Files\LibreOffice\program\soffice.com")


BASE_FONT = "Angsana New"
BASE_SIZE = 16


def set_font(run, *, size=BASE_SIZE, bold=False):
    run.font.name = BASE_FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), BASE_FONT)
    run.font.size = Pt(size)
    run.bold = bold


def keep_with_next(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    keep = OxmlElement("w:keepNext")
    p_pr.append(keep)


def keep_lines_together(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    keep = OxmlElement("w:keepLines")
    p_pr.append(keep)


def add_para(doc: Document, text: str, *, size=BASE_SIZE, bold=False, align=None, space_after=6):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_font(r, size=size, bold=bold)
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    keep_lines_together(p)
    return p


def add_heading(doc: Document, text: str, *, size=BASE_SIZE):
    p = add_para(doc, text, size=size, bold=True, space_after=4)
    keep_with_next(p)
    return p


def add_pic(doc: Document, filename: str, caption: str, *, width=6.1):
    path = IMAGE_DIR / filename
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    p.add_run().add_picture(str(path), width=Inches(width))
    keep_lines_together(p)
    cap = add_para(doc, caption, size=BASE_SIZE, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
    keep_lines_together(cap)


def add_field_table(doc: Document, rows: list[tuple[str, str]]):
    title = add_para(doc, "ช่องสำคัญที่ควรตรวจบนหน้าจอ", size=18, bold=True, space_after=4)
    keep_with_next(title)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    headers = ["Field Name", "Meaning"]
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
    for name, meaning in rows:
        cells = table.add_row().cells
        cells[0].text = name
        cells[1].text = meaning
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                for run in p.runs:
                    set_font(run, size=BASE_SIZE, bold=(row == table.rows[0]))


def add_journal_explanation_table(doc: Document):
    title = add_para(doc, "แนวคิดการอ่านเดบิตและเครดิตในหัวข้อนี้", size=18, bold=True, space_after=4)
    keep_with_next(title)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = ["บรรทัดที่ต้องดู", "ความหมายแบบเข้าใจง่าย", "สิ่งที่ควรตรวจ"]
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    rows = [
        ("Debit (Dr) เงินรับหรือบัญชีพักรับชำระ", "เป็นฝั่งที่ระบบบันทึกว่าเงินเข้ามาแล้ว หรือมีรายการรับเงินรอเคลียร์", "ตรวจชื่อบัญชีและยอดว่าเท่ากับยอดที่รับจริง"),
        ("Credit (Cr) ลูกหนี้การค้า", "เป็นฝั่งที่ระบบตัดยอดหนี้ของใบแจ้งหนี้ที่ถูกเลือกมารับชำระ", "ตรวจว่าเลขใบแจ้งหนี้และยอดที่ตัดตรงกับเอกสารลูกหนี้จริง"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                for run in p.runs:
                    set_font(run, size=BASE_SIZE, bold=(row == table.rows[0]))


def add_contents(doc: Document):
    title = add_para(doc, "สารบัญ", size=BASE_SIZE, bold=True, space_after=4)
    keep_with_next(title)
    items = [
        "1. วัตถุประสงค์",
        "2. ภาพรวมการทำงาน",
        "3. เมนูที่ใช้",
        "4. ขั้นตอนการใช้งาน",
        "5. ตรวจสอบ Journal Entry",
        "6. ข้อควรระวัง",
    ]
    for item in items:
        add_para(doc, item, size=BASE_SIZE, space_after=2)


def build_document():
    sys.stdout.reconfigure(encoding="utf-8")
    doc = Document()
    sec = doc.sections[0]
    sec.start_type = WD_SECTION_START.NEW_PAGE
    sec.top_margin = Inches(0.6)
    sec.bottom_margin = Inches(0.6)
    sec.left_margin = Inches(0.7)
    sec.right_margin = Inches(0.7)

    add_para(doc, "3.8 การจัดการการรับชำระแบบกลุ่มบริษัท", size=BASE_SIZE, bold=True, space_after=3)
    add_para(
        doc,
        "คู่มือนี้อธิบายการรับชำระเงินแบบกลุ่มบริษัทด้วยภาษาที่ผู้ใช้งานทั่วไปอ่านแล้วทำตามได้ทันที "
        "โดยเริ่มตั้งแต่การเข้าเมนู การเปิดรายการ การเลือกใบแจ้งหนี้ที่ต้องการรับเงิน "
        "ไปจนถึงการกลับมาตรวจผลทางบัญชีหลังบันทึกรายการเสร็จ",
        size=BASE_SIZE,
        space_after=10,
    )
    add_contents(doc)
    doc.add_paragraph()

    add_heading(doc, "วัตถุประสงค์")
    add_para(doc, "เพื่อให้ผู้ใช้งานสามารถรับชำระเงินจากลูกค้าหลายรายในกลุ่มเดียวกันได้จากเอกสารเดียว", size=BASE_SIZE)
    add_para(doc, "เพื่อให้ผู้ใช้งานตรวจสอบได้ว่าระบบตัดยอดลูกหนี้ถูกใบ ถูกยอด และลงบัญชีถูกต้อง", size=BASE_SIZE, space_after=10)

    add_heading(doc, "ภาพรวมการทำงาน")
    add_para(
        doc,
        "หัวข้อนี้ใช้กับกรณีที่ลูกค้าในเครือหรือกลุ่มบริษัทชำระเงินรวมกันครั้งเดียว "
        "ผู้ใช้งานจึงต้องเลือกรายการค้างชำระของสมาชิกแต่ละรายให้ถูกต้องก่อนบันทึกรับเงิน",
        size=BASE_SIZE,
    )
    add_para(
        doc,
        "เมื่อบันทึกรายการแล้ว ระบบจะสร้างข้อมูลรับชำระ และตัดยอดใบแจ้งหนี้ที่เลือกไว้ "
        "ดังนั้นขั้นตอนที่สำคัญที่สุดคือการตรวจรายชื่อลูกค้า เลขใบแจ้งหนี้ และยอดรวมก่อนกดยืนยัน",
        size=BASE_SIZE,
        space_after=10,
    )

    add_heading(doc, "เมนูที่ใช้")
    add_para(doc, "Accounting > Customers > รับชำระเงินกลุ่มลูกค้า", size=BASE_SIZE, space_after=10)

    add_heading(doc, "3. ขั้นตอนการใช้งาน")
    steps = [
        (
            "1. เริ่มจากหน้า Dashboard ของระบบ แล้วมองหาการ์ดของโมดูล Accounting เพื่อเข้าไปทำรายการรับเงิน",
            "nav_dashboard_accounting_real_annotated.png",
            "ภาพขั้นตอนที่ 1 จุดเริ่มต้นจากหน้า Dashboard",
        ),
        (
            "2. เมื่อเข้าสู่โมดูล Accounting แล้ว ให้ไปที่เมนู Accounting > Customers > รับชำระเงินกลุ่มลูกค้า "
            "เมนูนี้เป็นจุดรวมเอกสารรับเงินของลูกค้าหลายรายในกลุ่มเดียวกัน",
            "nav_accounting_group_payment_real_annotated.png",
            "ภาพขั้นตอนที่ 2 เมนูรับชำระเงินกลุ่มลูกค้า",
        ),
        (
            "3. ที่หน้ารายการ ให้เปิดเอกสารที่มีอยู่แล้ว หรือกด New เพื่อสร้างรายการใหม่ "
            "จากนั้นตรวจส่วนหัวของเอกสารก่อน ได้แก่ ชื่อกลุ่มลูกค้า วันที่รับชำระ สมุดรายวัน และวิธีรับเงิน",
            "group_draft_manual_annotated.png",
            "ภาพขั้นตอนที่ 3 หน้าจอเอกสารก่อนเลือกรายการลูกหนี้",
        ),
        (
            "4. เมื่อข้อมูลส่วนหัวครบแล้ว ให้เลือกกลุ่มลูกค้าและสมาชิกในกลุ่มที่ต้องการรับเงิน "
            "แล้วกดค้นหาเอกสารค้างชำระ เพื่อให้ระบบดึงใบแจ้งหนี้ที่ยังมียอดค้างของสมาชิกแต่ละรายขึ้นมาให้เลือก",
            "group_payment_wizard_manual_annotated.png",
            "ภาพขั้นตอนที่ 4 หน้าจอค้นหาและเลือกใบแจ้งหนี้ค้างชำระ",
        ),
        (
            "5. ตรวจแต่ละบรรทัดที่ระบบดึงขึ้นมา โดยดูชื่อสมาชิก เลขใบแจ้งหนี้ ยอดค้าง และยอดที่จะรับจริง "
            "เลือกเฉพาะบรรทัดที่อยู่ในหลักฐานการรับเงินจริง เพื่อป้องกันการตัดยอดผิดใบหรือเกินยอด",
            "group_payment_wizard_manual_annotated.png",
            "ภาพขั้นตอนที่ 5 การตรวจสอบบรรทัดใบแจ้งหนี้ก่อนบันทึกรับเงิน",
        ),
        (
            "6. เมื่อตรวจครบแล้ว ให้บันทึกหรือยืนยันรายการรับชำระ จากนั้นกลับมาตรวจสถานะของเอกสารอีกครั้ง "
            "รวมถึงเปิดรายการบัญชีของเอกสารเพื่อดูว่าระบบตัดลูกหนี้และบันทึกเงินรับถูกต้องแล้ว",
            "group_done_manual_annotated.png",
            "ภาพขั้นตอนที่ 6 เอกสารหลังบันทึกรับชำระสำเร็จ",
        ),
    ]
    for text, image, caption in steps:
        add_para(doc, text, size=BASE_SIZE)
        add_pic(doc, image, caption)

    add_field_table(
        doc,
        [
            ("Customer Group", "ชื่อกลุ่มลูกค้าที่ใช้รวมเอกสารเพื่อรับชำระครั้งเดียว"),
            ("Members", "รายชื่อลูกค้าในกลุ่มที่ระบบจะดึงเอกสารค้างชำระมาให้เลือก"),
            ("Journal", "ชื่อสมุดรายวันที่ใช้บันทึกรายการบัญชีของการรับเงินครั้งนี้"),
            ("Payment Date", "วันที่ที่ระบบจะนำไปใช้เป็นวันที่รับเงินและวันที่ลงบัญชี"),
            ("Payment Method", "วิธีรับเงิน เช่น เงินโอน เงินสด หรือวิธีที่กิจการใช้จริง"),
            ("Amount", "ยอดรวมที่ผู้ใช้งานตั้งใจรับชำระในครั้งนั้น"),
        ],
    )

    add_heading(doc, "4. ตรวจสอบ Journal Entry")
    add_heading(doc, "วิธีอ่าน Journal ในหัวข้อนี้")
    add_para(
        doc,
        "หลังบันทึกรายการรับชำระแล้ว ให้เปิดรายการบัญชีของเอกสารเดิม ไม่ต้องไปเปิดจากเมนูอื่นก่อน "
        "เพราะจะช่วยให้มั่นใจว่า Journal ที่เห็นเป็นของรายการรับเงินครั้งนี้จริง",
        size=BASE_SIZE,
    )
    add_para(
        doc,
        "การอ่าน Journal ของหัวข้อนี้ให้ดู 3 เรื่องพร้อมกัน คือ เลขอ้างอิงเอกสาร ชื่อคู่ค้า "
        "และยอดรวมที่รับเงิน เพื่อยืนยันว่าระบบตัดยอดลูกหนี้ของสมาชิกในกลุ่มถูกคนและถูกใบ",
        size=BASE_SIZE,
    )
    add_para(
        doc,
        "จากภาพตัวอย่าง ระบบสร้างรายการบัญชีเลขที่ PCSH3/2026/00843 "
        "อ้างอิงใบแจ้งหนี้ INV-D/26/02/01422 วันที่บัญชี 23/02/2026 ยอด 26,837.50 บาท",
        size=BASE_SIZE,
    )
    add_para(
        doc,
        "บรรทัดด้าน Debit (Dr) คือบัญชี 111103 บัญชีพักเงินสด จำนวน 26,837.50 บาท "
        "ความหมายคือระบบรับรู้ว่าเงินเข้ามาแล้ว และพักยอดไว้ในบัญชีรับชำระของกิจการ",
        size=BASE_SIZE,
    )
    add_para(
        doc,
        "บรรทัดด้าน Credit (Cr) คือบัญชี 113001 ลูกหนี้การค้า - ในประเทศ จำนวน 26,837.50 บาท "
        "ความหมายคือระบบตัดยอดหนี้ของลูกค้าที่อยู่ในใบแจ้งหนี้รายการนี้ออกจากบัญชีลูกหนี้",
        size=BASE_SIZE,
    )
    add_para(
        doc,
        "ถ้ายอด Debit และ Credit เท่ากัน และเลขอ้างอิงตรงกับเอกสารที่ผู้ใช้งานเพิ่งรับชำระ "
        "แปลว่าระบบลงบัญชีของการรับเงินครั้งนี้ถูกต้องแล้ว",
        size=BASE_SIZE,
        space_after=10,
    )

    add_pic(
        doc,
        "journal_group_payment_real_annotated.png",
        "ภาพ Journal Entry ของการรับชำระแบบกลุ่มบริษัท พร้อมคำอธิบาย Dr/Cr",
    )

    add_journal_explanation_table(doc)

    add_heading(doc, "ข้อควรระวัง")
    add_para(doc, "อย่าเลือกใบแจ้งหนี้เกินกว่าหลักฐานการรับเงินจริง เพราะระบบจะตัดยอดลูกหนี้ทันทีหลังบันทึกรายการ", size=BASE_SIZE)
    add_para(doc, "ก่อนยืนยันรายการทุกครั้ง ควรตรวจวันที่รับเงิน ชื่อกลุ่มลูกค้า และยอดรวมให้ตรงกับเอกสารประกอบ", size=BASE_SIZE)
    add_para(doc, "หลังบันทึกแล้ว ควรเปิด Journal Entry ตรวจซ้ำอย่างน้อยหนึ่งครั้ง เพื่อยืนยันว่าระบบลงบัญชีตรงกับรายการรับเงินจริง", size=BASE_SIZE)

    DOCX_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(DOCX_PATH)
    PDF_DIR.mkdir(parents=True, exist_ok=True)
    subprocess.run(
        [str(SOFFICE), "--headless", "--convert-to", "pdf", "--outdir", str(PDF_DIR), str(DOCX_PATH)],
        check=True,
    )
    print("rebuilt 3.8 manual")


if __name__ == "__main__":
    build_document()

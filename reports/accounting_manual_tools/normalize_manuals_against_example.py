from __future__ import annotations

from copy import deepcopy
from pathlib import Path
import re
import subprocess

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.table import Table
from docx.text.paragraph import Paragraph


ROOT = Path(r"C:\365_project\TheCool18e\Dev")
DOCX_DIR = ROOT / "manual" / "Accouting_Manual" / "generated_20260408" / "docx"
PDF_DIR = ROOT / "manual" / "Accouting_Manual" / "generated_20260408" / "pdf_review"
SOFFICE = Path(r"C:\Program Files\LibreOffice\program\soffice.com")
TARGET_PREFIXES = ("3.8_", "5.", "6.", "7.")

DETAIL_HEADING = "ลำดับการทำงานแบบละเอียด"
JE_HEADING = "4. ตรวจสอบ Journal Entry"
FIELD_HEADING_TEXTS = {"Field Name", "ชื่อช่องที่ควรดูบนหน้าจอ"}
FIELD_TABLE_HEADERS = {"Field Name", "Field / Section", "หัวข้อ", "ชื่อช่องที่ควรดูบนหน้าจอ"}
JOURNAL_GUIDE_HEADING = "วิธีอ่าน Journal ในหัวข้อนี้"

FIELD_MAP: dict[str, list[str]] = {
    "3.8": ["Customer Group", "Members", "Journal", "Payment Date", "Payment Method", "Amount"],
    "5.1": ["Journal", "Incoming Payments", "Outgoing Payments", "Outstanding Receipts Account", "Outstanding Payments Account"],
    "5.2": ["Template Name", "Bank Name", "Payee Position", "Amount Position", "Date Position"],
    "5.3": ["Cheque Book", "Bank Journal", "Start Number", "End Number", "Next Number", "Status"],
    "5.4": ["Vendor Bill", "Register Payment", "Journal", "Payment Method", "Cheque Number", "Amount"],
    "5.5": ["Cheque Number", "Partner", "Due Date", "Amount", "State", "Bank Journal"],
    "5.6": ["Customer Invoice", "Register Payment", "Journal", "Payment Method", "Cheque Number", "Amount"],
    "5.7": ["Cheque Number", "Deposit Date", "Clear Date", "State", "Journal Items"],
    "5.8": ["Cheque Number", "State", "Void Reason", "New Cheque", "Reverse Entry"],
    "6.1": ["Asset Name", "Asset Model", "Original Value", "Acquisition Date", "Book Value", "Status"],
    "6.2": ["Name", "Original Value", "Acquisition Date", "Asset Model", "Quantity", "Company"],
    "6.3": ["Acquisition Date", "First Depreciation Date", "Method Number", "Method Period", "Book Value", "Depreciation Board"],
    "6.4": ["Asset Name", "Invoice", "Sale Amount", "Gain/Loss Account", "Status", "Journal Entry"],
    "6.5": ["Asset Name", "Disposal Date", "Residual Value", "Loss Account", "Status", "Journal Entry"],
    "6.6": ["Asset Name", "Asset Model", "Original Value", "Book Value", "Acquisition Date", "Status"],
    "7.1": ["Product", "BoM", "Components", "Quantity", "Scheduled Date", "Status"],
    "7.2": ["MO Number", "Product", "Quantity", "Components", "Valuation", "Journal Entry"],
    "7.3": ["Product Category", "Valuation", "Costing Method", "Stock Valuation Account", "Production Account", "Stock Journal"],
    "7.4": ["MO Number", "Product", "Quantity", "Components", "Finished Product", "Journal Entry"],
    "7.5": ["Scrap Product", "Quantity", "Source Location", "Scrap Location", "Reason", "Reference"],
    "7.6": ["Date", "Product", "Quantity", "Total Value", "Reference", "Journal Entry"],
    "7.7": ["Date From", "Date To", "Product", "Quantity", "Unit", "Lot/Batch"],
}

FIELD_MEANINGS: dict[str, str] = {
    "Customer Group": "ชื่อกลุ่มลูกค้าที่ใช้รวมเอกสารเพื่อรับชำระครั้งเดียว",
    "Members": "รายชื่อลูกค้าในกลุ่มที่ระบบจะดึงเอกสารค้างชำระมาให้เลือก",
    "Journal": "สมุดรายวันที่ระบบจะใช้บันทึกรายการบัญชีของเอกสารนี้",
    "Payment Date": "วันที่ที่ใช้บันทึกรายการรับหรือจ่ายเงิน",
    "Payment Method": "วิธีรับหรือจ่ายเงิน เช่น เช็ค โอน หรือเงินสด",
    "Amount": "ยอดเงินของรายการที่กำลังทำ",
    "Incoming Payments": "วิธีรับเงินที่อนุญาตให้ใช้กับ Journal นี้",
    "Outgoing Payments": "วิธีจ่ายเงินที่อนุญาตให้ใช้กับ Journal นี้",
    "Outstanding Receipts Account": "บัญชีพักเงินรับที่ใช้ก่อนเงินเข้าธนาคารจริง",
    "Outstanding Payments Account": "บัญชีพักเงินจ่ายที่ใช้ก่อนเช็คหรือเงินออกจริง",
    "Template Name": "ชื่อแบบฟอร์มเช็คที่ใช้เลือกตอนพิมพ์เช็ค",
    "Bank Name": "ชื่อธนาคารที่จะแสดงบนแบบฟอร์มเช็ค",
    "Payee Position": "ตำแหน่งข้อความชื่อผู้รับเงินบนแบบฟอร์มเช็ค",
    "Amount Position": "ตำแหน่งจำนวนเงินบนแบบฟอร์มเช็ค",
    "Date Position": "ตำแหน่งวันที่บนแบบฟอร์มเช็ค",
    "Cheque Book": "ชื่อสมุดเช็คที่ใช้คุมเลขเช็คในชุดเดียวกัน",
    "Bank Journal": "Journal ธนาคารที่ผูกกับสมุดเช็คหรือรายการเช็คนี้",
    "Start Number": "เลขเช็คใบแรกของสมุดเช็ค",
    "End Number": "เลขเช็คใบสุดท้ายของสมุดเช็ค",
    "Next Number": "เลขเช็คใบถัดไปที่ระบบจะหยิบมาใช้",
    "Status": "สถานะปัจจุบันของเอกสารหรือรายการ",
    "Vendor Bill": "บิลผู้ขายที่ใช้เป็นต้นทางของการจ่ายเงิน",
    "Register Payment": "หน้าต่างที่ใช้กรอกรายละเอียดการรับหรือจ่ายเงินก่อนระบบสร้างรายการจริง",
    "Cheque Number": "เลขที่พิมพ์บนเช็คที่ใช้ติดตามรายการ",
    "Partner": "คู่ค้าที่เกี่ยวข้องกับเอกสารนี้",
    "Due Date": "วันที่ครบกำหนดของเช็คหรือเอกสาร",
    "Deposit Date": "วันที่นำเช็คไปฝากหรือส่งเข้าธนาคาร",
    "Clear Date": "วันที่เช็คผ่านธนาคารหรือเคลียร์รายการสำเร็จ",
    "Journal Items": "รายละเอียดบรรทัดบัญชีเดบิตและเครดิตที่ระบบสร้าง",
    "Void Reason": "เหตุผลที่ใช้ยกเลิกเช็คหรือกลับรายการ",
    "New Cheque": "เช็คใบใหม่ที่ใช้แทนเช็คเดิมหลังการแปลงสถานะ",
    "Reverse Entry": "รายการบัญชีกลับทางที่ระบบสร้างเมื่อยกเลิกเอกสาร",
    "Asset Name": "ชื่อทรัพย์สินที่ใช้ค้นหาและติดตามในระบบ",
    "Asset Model": "แบบทรัพย์สินที่กำหนดอายุการใช้งานและบัญชีที่เกี่ยวข้อง",
    "Original Value": "มูลค่าทุนเริ่มต้นของทรัพย์สิน",
    "Acquisition Date": "วันที่ซื้อหรือรับทรัพย์สินเข้าระบบ",
    "Book Value": "มูลค่าคงเหลือทางบัญชีของทรัพย์สิน",
    "Name": "ชื่อรายการหลักของเอกสารหรือแบบฟอร์ม",
    "Quantity": "จำนวนหน่วยของรายการนั้น",
    "Company": "บริษัทเจ้าของรายการในระบบ",
    "First Depreciation Date": "วันที่เริ่มลงค่าเสื่อมงวดแรก",
    "Method Number": "จำนวนงวดที่ระบบจะตัดค่าเสื่อม",
    "Method Period": "ช่วงเวลาระหว่างงวดค่าเสื่อม",
    "Depreciation Board": "ตารางแสดงงวดค่าเสื่อมทั้งหมดของทรัพย์สิน",
    "Invoice": "ใบแจ้งหนี้ที่ใช้เชื่อมโยงกับรายการขายทรัพย์สิน",
    "Sale Amount": "ยอดขายทรัพย์สินที่เกิดขึ้นจริง",
    "Gain/Loss Account": "บัญชีที่ใช้รับรู้กำไรหรือขาดทุนจากการขาย",
    "Disposal Date": "วันที่ตัดจำหน่ายหรือปลดทรัพย์สิน",
    "Residual Value": "มูลค่าคงเหลือก่อนปิดทรัพย์สิน",
    "Loss Account": "บัญชีที่ใช้รับรู้ผลขาดทุนจากการตัดจำหน่าย",
    "Product": "ชื่อสินค้าที่ใช้ในงานผลิตหรือรายงาน",
    "BoM": "สูตรการผลิตที่ระบุส่วนประกอบของสินค้า",
    "Components": "วัตถุดิบหรือชิ้นส่วนที่ใช้ในงานนั้น",
    "Scheduled Date": "วันที่ระบบวางแผนให้เริ่มหรือทำรายการ",
    "MO Number": "เลขที่ใบสั่งผลิตที่ใช้ตามรอยงาน",
    "Valuation": "มูลค่าสินค้าคงคลังที่เกิดจากการเคลื่อนไหวของสินค้า",
    "Product Category": "หมวดสินค้าที่กำหนดวิธีคิดต้นทุนและการลงบัญชี",
    "Costing Method": "วิธีคำนวณต้นทุนสินค้า เช่น FIFO หรือ Average",
    "Stock Valuation Account": "บัญชีสินค้าคงคลังที่รับมูลค่าจากสินค้า",
    "Production Account": "บัญชีงานระหว่างทำที่ใช้สะสมต้นทุนการผลิต",
    "Stock Journal": "สมุดรายวันที่ใช้บันทึกรายการสินค้าคงคลัง",
    "Finished Product": "สินค้าสำเร็จรูปที่ได้จากงานผลิต",
    "Scrap Product": "สินค้าที่ถูกตัดเป็นของเสีย",
    "Source Location": "ตำแหน่งต้นทางที่สินค้าถูกตัดออก",
    "Scrap Location": "ตำแหน่งปลายทางที่เก็บของเสีย",
    "Reason": "เหตุผลหรือคำอธิบายประกอบรายการ",
    "Reference": "เลขอ้างอิงที่ใช้เชื่อมกลับไปยังเอกสารต้นทาง",
    "Date": "วันที่ของรายการหรือรายงาน",
    "Total Value": "มูลค่ารวมของรายการที่แสดงในรายงาน",
    "Date From": "วันเริ่มต้นของช่วงวันที่ที่ใช้กรองรายงาน",
    "Date To": "วันสิ้นสุดของช่วงวันที่ที่ใช้กรองรายงาน",
    "Unit": "หน่วยนับที่ใช้แสดงปริมาณสินค้า",
    "Lot/Batch": "เลขล็อตหรือแบตช์ที่ใช้ติดตามสินค้า",
    "Journal Entry": "รายการบัญชีที่ระบบสร้างขึ้นเพื่อบันทึกผลทางการเงินของเหตุการณ์นั้น",
}

JOURNAL_GUIDE_MAP: dict[str, list[str]] = {
    "3.8": [
        "Journal คือสมุดรายวันที่ระบบใช้บันทึกรายการรับชำระของเอกสารนี้",
        "ให้เปิดดูชื่อ Journal, เลขอ้างอิงเอกสาร และยอดเดบิตเครดิตก่อน เพื่อยืนยันว่าเป็นรายการของกลุ่มลูกค้าที่กำลังตรวจจริง",
        "กรณีรับชำระสำเร็จ ระบบจะเดบิตบัญชีเงินรับหรือบัญชีพักรับชำระ และเครดิตบัญชีลูกหนี้ตามใบแจ้งหนี้ที่ถูกเลือก",
    ],
    "5.1": [
        "หัวข้อนี้เป็นการตั้งค่า จึงยังไม่เกิด Journal Entry ทันที แต่ค่าที่ตั้งใน Journal จะถูกนำไปใช้ตอนรับหรือจ่ายเงินจริง",
        "ถ้าตั้ง Outstanding account ถูกต้อง รายการเช็คจะลงบัญชีพักก่อน แล้วค่อยย้ายเมื่อเช็คผ่านธนาคาร",
    ],
    "5.2": [
        "หัวข้อนี้เป็นการตั้งค่าแบบฟอร์มเช็ค จึงยังไม่สร้างรายการบัญชี",
        "ผลของการตั้งค่านี้จะมีผลเฉพาะรูปแบบการพิมพ์เช็ค ไม่ได้เปลี่ยนเดบิตหรือเครดิตของระบบ",
    ],
    "5.3": [
        "การสร้างสมุดเช็คยังไม่สร้าง Journal Entry เพราะเป็นการกำหนดเลขเช็คที่จะใช้ในอนาคต",
        "เมื่อมีการนำเลขเช็คจากสมุดนี้ไปจ่ายเงินจริง ระบบจึงจะสร้างรายการบัญชีตาม Journal ที่ผูกไว้",
    ],
    "5.4": [
        "Journal คือสมุดรายวันที่ใช้ลงบัญชีของเช็คจ่ายใบนี้ เช่น Bank Journal หรือ Cheque Journal",
        "ตอนเช็คยังไม่ผ่านธนาคาร ระบบมักเดบิตเจ้าหนี้และเครดิตบัญชีพักเช็คจ่ายก่อน เพื่อให้รู้ว่ายังเป็นเช็คคงค้าง",
        "เมื่อเช็คผ่านธนาคาร ระบบจะย้ายยอดจากบัญชีพักเช็คจ่ายไปยังบัญชีธนาคารจริง",
    ],
    "5.5": [
        "ให้ดูชื่อ Journal และสถานะของเช็คก่อนว่าเป็น Confirmed หรือ Paid เพราะความหมายทางบัญชีต่างกัน",
        "สถานะ Confirmed แปลว่ายังอยู่ในบัญชีพักเช็คจ่าย ส่วนสถานะ Paid แปลว่ายอดถูกย้ายไปธนาคารแล้ว",
        "เมื่อเปิด Journal Items ให้ดูว่าชื่อบัญชีและยอดตรงกับสถานะของเช็คที่กำลังติดตามอยู่",
    ],
    "5.6": [
        "สำหรับเช็ครับ ระบบจะใช้ Journal ของการรับเงินเพื่อบันทึกรายการเข้าบัญชีพักเช็ครับก่อน",
        "เมื่อรับเช็คสำเร็จ ระบบจะเดบิตบัญชีพักเช็ครับหรือบัญชีเงินรับ และเครดิตลูกหนี้ของใบแจ้งหนี้ลูกค้า",
        "ถ้าเช็คยังไม่ผ่านธนาคาร ยอดจะยังไม่เข้าบัญชีธนาคารจริงจนกว่าจะมีการเคลียร์เช็ค",
    ],
    "5.7": [
        "ขั้นตอนเคลียร์เช็คคือการย้ายยอดจากบัญชีพักเช็ครับหรือเช็คจ่าย ไปยังบัญชีธนาคารจริง",
        "ให้เทียบเลขเช็ค วันที่เคลียร์ และชื่อบัญชีใน Journal Items ว่าตรงกับเช็คใบที่ต้องการปิดสถานะ",
        "ถ้าดูถูกใบ จะเห็นว่าบัญชีพักลดลงและบัญชีธนาคารเพิ่มหรือลดตามทิศทางของรายการ",
    ],
    "5.8": [
        "เมื่อยกเลิกหรือแปลงสถานะเช็ค ระบบจะสร้าง Reverse Entry เพื่อกลับผลทางบัญชีของเช็คเดิม",
        "ให้ดูเลขอ้างอิงของรายการกลับทางและชื่อบัญชีว่ากลับเดบิตเครดิตตรงข้ามกับรายการเดิมหรือไม่",
        "ถ้ามีการออกเช็คใบใหม่แทน ระบบจะมีทั้งรายการกลับทางของใบเดิมและรายการใหม่ของใบแทน",
    ],
    "6.1": [
        "Journal ของทรัพย์สินจะเริ่มเห็นชัดเมื่อมีการเริ่มใช้งาน ขาย ตัดจำหน่าย หรือบันทึกค่าเสื่อม",
        "ให้ดูว่าบัญชีทรัพย์สิน บัญชีค่าเสื่อมสะสม และบัญชีกำไรขาดทุนถูกใช้ตรงตามเหตุการณ์ของทรัพย์สินหรือไม่",
    ],
    "6.2": [
        "การสร้างทรัพย์สินสถานะ Draft ยังไม่สร้าง Journal Entry ทันที เพราะยังเป็นเพียงการตั้งต้นข้อมูล",
        "บัญชีที่จะถูกใช้จริงถูกกำหนดจาก Asset Model ที่เลือก จึงควรเปิดดูแบบทรัพย์สินก่อนบันทึกทุกครั้ง",
    ],
    "6.3": [
        "เมื่อเริ่มคิดค่าเสื่อม ระบบจะใช้ Journal ของค่าเสื่อมสร้างรายการรายงวดให้อัตโนมัติ",
        "โดยทั่วไประบบจะเดบิตค่าใช้จ่ายค่าเสื่อม และเครดิตค่าเสื่อมสะสมของทรัพย์สินนั้น",
        "ให้ดูงวด วันที่ และเลขอ้างอิงว่าตรงกับทรัพย์สินและรอบตัดค่าเสื่อมที่ต้องการตรวจหรือไม่",
    ],
    "6.4": [
        "การขายทรัพย์สินจะมีทั้งรายการขายและรายการตัดยอดทรัพย์สินออกจากบัญชี",
        "ให้ดูว่าระบบปิดมูลค่าทุนและค่าเสื่อมสะสมออกก่อน แล้วค่อยรับรู้กำไรหรือขาดทุนจากการขาย",
        "ถ้ายอดขายไม่เท่ามูลค่าคงเหลือ ระบบจะแสดงผลต่างไว้ที่บัญชีกำไรหรือขาดทุนจากการขายทรัพย์สิน",
    ],
    "6.5": [
        "การตัดจำหน่ายหรือ Scrap ทรัพย์สินคือการปิดยอดทรัพย์สินออกจากระบบโดยไม่มีการขาย",
        "ให้ดูว่าระบบตัดมูลค่าทุนและค่าเสื่อมสะสมออก แล้วรับรู้ผลขาดทุนคงเหลือในบัญชีที่กำหนดไว้",
    ],
    "6.6": [
        "รายงานทรัพย์สินถาวรใช้เทียบกับ Journal Entry ที่เกิดจากทรัพย์สินแต่ละใบ",
        "ควรอ่านค่าจากรายงานคู่กับบัญชีทรัพย์สินและค่าเสื่อมสะสม เพื่อยืนยันว่ารายงานและรายการบัญชีไปในทิศทางเดียวกัน",
    ],
    "7.1": [
        "ภาพรวมการผลิตจะเกี่ยวกับ Journal ของสินค้าคงคลังและงานระหว่างทำเป็นหลัก",
        "ให้ดูว่าแต่ละขั้นของการผลิตย้ายมูลค่าจากวัตถุดิบไปงานระหว่างทำ และจากงานระหว่างทำไปสินค้าสำเร็จรูปอย่างไร",
    ],
    "7.2": [
        "Journal ของการผลิตใช้บอกว่าต้นทุนกำลังย้ายจากวัตถุดิบไปงานระหว่างทำ และสุดท้ายเข้าไปที่สินค้าสำเร็จรูป",
        "อ่านง่ายที่สุดโดยดูจากเลขอ้างอิง MO เดียวกัน แล้วไล่ตามลำดับเวลาใน Valuation หรือ Journal Entries",
        "บรรทัดเดบิตคือบัญชีที่รับมูลค่าเข้า ส่วนบรรทัดเครดิตคือบัญชีที่มูลค่าถูกย้ายออกไป",
    ],
    "7.3": [
        "หัวข้อนี้เป็นการอธิบายการตั้งค่าที่มีผลกับ Journal ของงานผลิต ไม่ใช่การสร้าง Journal โดยตรง",
        "ถ้ากำหนด Product Category และบัญชีคงคลังถูกต้อง รายการผลิตจริงจะลงบัญชีได้ตามที่ต้องการ",
    ],
    "7.4": [
        "ตัวอย่างงานผลิตควรเปิดดูทั้ง Valuation และ Journal Entry เพื่อเห็นภาพต้นทุนครบทุกขั้น",
        "ให้ดูว่ารายการตัดวัตถุดิบและรับสินค้าเสร็จมีเลขอ้างอิงเดียวกันและใช้บัญชีสอดคล้องกับหมวดสินค้า",
    ],
    "7.5": [
        "Scrap และ Loss จะกระทบต้นทุนโดยทำให้มูลค่าสินค้าหรือวัตถุดิบถูกย้ายออกจากที่เดิม",
        "ให้ดูชื่อสินค้า ปริมาณ และบัญชีที่เกี่ยวข้องว่าเป็นของเสียของงานเดียวกับที่กำลังตรวจจริง",
    ],
    "7.6": [
        "รายงาน Stock Valuation ใช้ดูมูลค่าสินค้าคงคลังที่มาจาก Journal และ Stock Moves จริง",
        "ควรอ่านคู่กับเลขอ้างอิงเอกสารและ Journal Entry เพื่อยืนยันว่ามูลค่าที่เห็นในรายงานมาจากรายการใด",
    ],
    "7.7": [
        "รายงาน รง.8 ไม่ได้สร้าง Journal Entry เอง แต่ใช้ข้อมูลสินค้าและการเคลื่อนไหวจริงมาสรุปผล",
        "เมื่อจำเป็นต้องตรวจย้อนหลัง ให้เทียบกับ Valuation หรือ Journal ของเอกสารต้นทางเพื่อยืนยันตัวเลข",
    ],
}


def set_run_font(run, size: int = 16, bold: bool = False) -> None:
    run.bold = bold
    run.font.name = "TH Sarabun New"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "TH Sarabun New")
    run.font.size = Pt(size)


def make_empty_after(paragraph: Paragraph) -> Paragraph:
    new_p = deepcopy(paragraph._p)
    for child in list(new_p):
        new_p.remove(child)
    paragraph._p.addnext(new_p)
    return Paragraph(new_p, paragraph._parent)


def insert_paragraph_after(paragraph: Paragraph, text: str, *, size: int = 16, bold: bool = False) -> Paragraph:
    new_para = make_empty_after(paragraph)
    run = new_para.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return new_para


def remove_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def remove_table(table: Table) -> None:
    element = table._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def find_index_by_text(paragraphs: list[Paragraph], text: str) -> int | None:
    for idx, paragraph in enumerate(paragraphs):
        if paragraph.text.strip() == text:
            return idx
    return None


def doc_code_from_name(name: str) -> str | None:
    match = re.match(r"^(\d+\.\d+)_", name)
    return match.group(1) if match else None


def split_step(text: str) -> tuple[int, str] | None:
    match = re.match(r"^(\d+)\.\s*(.+)$", text.strip())
    if not match:
        return None
    return int(match.group(1)), match.group(2).strip()


def is_numbered_step(text: str) -> bool:
    return split_step(text) is not None


def insert_table_before(paragraph: Paragraph, fields: list[str]) -> None:
    temp_doc = Document()
    table = temp_doc.add_table(rows=len(fields) + 1, cols=1)
    table.style = "Table Grid"
    table.cell(0, 0).text = "Field Name"
    for row_idx, field_name in enumerate(fields, start=1):
        table.cell(row_idx, 0).text = field_name
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    set_run_font(run, size=16, bold=(para.text.strip() == "Field Name"))
    paragraph._p.addprevious(deepcopy(table._tbl))


def cleanup_field_blocks(doc: Document) -> None:
    for paragraph in list(doc.paragraphs):
        if paragraph.text.strip() in FIELD_HEADING_TEXTS:
            remove_paragraph(paragraph)
    for table in list(doc.tables):
        if not table.rows or not table.rows[0].cells:
            continue
        header = table.rows[0].cells[0].text.strip()
        if header in FIELD_TABLE_HEADERS:
            remove_table(table)


def cleanup_journal_guides(doc: Document) -> None:
    journal_lines = {line for lines in JOURNAL_GUIDE_MAP.values() for line in lines}
    for paragraph in list(doc.paragraphs):
        if paragraph.text.strip() == JOURNAL_GUIDE_HEADING:
            remove_paragraph(paragraph)
    for paragraph in list(doc.paragraphs):
        if paragraph.text.strip() in journal_lines:
            remove_paragraph(paragraph)


def rewrite_generic_text(content: str, code: str) -> str:
    generic_verify = " พร้อมตรวจเลขเอกสาร วันที่ คู่ค้า ยอดเงิน หรือจำนวนสินค้าให้ตรงกับรายการจริง"
    if generic_verify in content:
        content = content.replace(generic_verify, verify_phrase_for_code(code))
    content = re.sub(r"( เพื่อเข้าสู่ชุดเมนูของหัวข้อนี้)+$", " เพื่อเข้าสู่ชุดเมนูของหัวข้อนี้", content)
    content = re.sub(r"( เพื่อเปิดหน้ารายการหลักที่ต้องใช้งาน)+$", " เพื่อเปิดหน้ารายการหลักที่ต้องใช้งาน", content)
    if content.startswith("เริ่มจากหน้า Dashboard ของระบบ"):
        return "เริ่มจากหน้า Dashboard ของระบบ แล้วตรวจสอบชื่อบริษัท สิทธิ์ผู้ใช้งาน และเลขเอกสารที่จะใช้ให้ถูกต้องก่อนเริ่มทำรายการ"
    if content.startswith("คลิกเข้าโมดูล "):
        if "เพื่อเข้าสู่ชุดเมนูของหัวข้อนี้" in content:
            return content
        return content + " เพื่อเข้าสู่ชุดเมนูของหัวข้อนี้"
    if content.startswith("เมื่อเข้ามาในโมดูลแล้ว ให้ไปที่เมนู "):
        if "เพื่อเปิดหน้ารายการหลักที่ต้องใช้งาน" in content:
            return content
        return re.sub(r"^เมื่อเข้ามาในโมดูลแล้ว ให้ไปที่เมนู (.+)$", r"เมื่อเข้ามาในโมดูลแล้ว ให้ไปที่เมนู \1 เพื่อเปิดหน้ารายการหลักที่ต้องใช้งาน", content)
    if content.startswith("เมื่อหน้าจอเปิดแล้ว ให้ตรวจ"):
        return "เปิดหน้าจอรายการแล้วตรวจข้อมูลหลักให้ตรงกับรายการจริงก่อนกดทำรายการ"
    if content.startswith("หลังบันทึกหรือยืนยันรายการแล้ว"):
        return "หลังบันทึกหรือยืนยันรายการแล้ว ให้กลับมาตรวจผลที่หน้าจอเอกสาร สถานะรายการ และหน้ารายการบัญชีทุกครั้ง"
    return content


def verify_phrase_for_code(code: str) -> str:
    if code == "3.8":
        return " พร้อมตรวจกลุ่มลูกค้า สมาชิก วันที่รับชำระ และยอดรวมให้ตรงกับข้อมูลจริง"
    if code.startswith("5."):
        return " พร้อมตรวจเลขเอกสาร วันที่ คู่ค้า วิธีชำระ และยอดเงินให้ตรงกับรายการจริง"
    if code.startswith("6."):
        return " พร้อมตรวจชื่อทรัพย์สิน แบบทรัพย์สิน มูลค่า และวันที่ซื้อให้ตรงกับข้อมูลจริง"
    if code.startswith("7."):
        return " พร้อมตรวจเลขอ้างอิงสินค้า ปริมาณ และสถานะงานให้ตรงกับรายการจริง"
    return " พร้อมตรวจข้อมูลหลักให้ตรงกับรายการจริง"


def merge_optional_menu(step3: str, step4: str) -> str:
    prefix = "ถ้าขั้นตอนนี้ต้องเปิดหน้าจอที่เกี่ยวข้องเพิ่มเติม ให้ไปที่เมนู "
    if step4.startswith(prefix):
        extra_menu = step4[len(prefix):].strip()
        base = step3.rstrip(" .")
        return f"{base} และถ้าต้องเปิดหน้าจอที่เกี่ยวข้องเพิ่มเติม ให้ไปที่เมนู {extra_menu}"
    return step3


def merge_verify_into_specific(next_step: str, code: str) -> str:
    extra = verify_phrase_for_code(code)
    if "ตรวจ" in next_step:
        return next_step
    if next_step.endswith(extra):
        return next_step
    if next_step.startswith(("เปิด", "เข้าเมนู", "กลับไปที่เมนู", "กด", "กรอก", "เลือก", "ดู", "คลิก")):
        return next_step.rstrip(" .") + extra
    return next_step


def smooth_detailed_steps(doc: Document, code: str) -> bool:
    paragraphs = doc.paragraphs
    detail_idx = find_index_by_text(paragraphs, DETAIL_HEADING)
    je_idx = find_index_by_text(paragraphs, JE_HEADING)
    if detail_idx is None or je_idx is None or je_idx <= detail_idx:
        return False

    step_entries: list[tuple[int, Paragraph, str]] = []
    for idx in range(detail_idx + 1, je_idx):
        paragraph = paragraphs[idx]
        text = paragraph.text.strip()
        if is_numbered_step(text):
            step_entries.append((idx, paragraph, split_step(text)[1]))
    if not step_entries:
        return False

    changed = False
    keep_flags = [True] * len(step_entries)
    contents = [content for _, _, content in step_entries]

    menu_idx = next((idx for idx, text in enumerate(contents) if text.startswith("เมื่อเข้ามาในโมดูลแล้ว ให้ไปที่เมนู ")), None)
    if menu_idx is not None:
        next_idx = menu_idx + 1
        if next_idx < len(contents) and contents[next_idx].startswith("ถ้าขั้นตอนนี้ต้องเปิดหน้าจอที่เกี่ยวข้องเพิ่มเติม ให้ไปที่เมนู "):
            contents[menu_idx] = merge_optional_menu(contents[menu_idx], contents[next_idx])
            keep_flags[next_idx] = False
            changed = True

    verify_idx = next(
        (
            idx
            for idx, text in enumerate(contents)
            if text.startswith("เมื่อหน้าจอเปิดแล้ว ให้ตรวจ")
            or text.startswith("เปิดหน้าจอรายการแล้วตรวจ")
        ),
        None,
    )
    if verify_idx is not None and keep_flags[verify_idx]:
        for idx in range(verify_idx + 1, len(contents)):
            if keep_flags[idx]:
                contents[idx] = merge_verify_into_specific(contents[idx], code)
                keep_flags[verify_idx] = False
                changed = True
                break

    new_order: list[tuple[Paragraph, str]] = []
    for keep, (_, paragraph, _), content in zip(keep_flags, step_entries, contents):
        if not keep:
            remove_paragraph(paragraph)
            changed = True
            continue
        rewritten = rewrite_generic_text(content, code)
        if rewritten != content:
            changed = True
        new_order.append((paragraph, rewritten))

    for new_no, (paragraph, content) in enumerate(new_order, start=1):
        new_text = f"{new_no}. {content}"
        if paragraph.text != new_text:
            paragraph.text = new_text
            for run in paragraph.runs:
                set_run_font(run, size=16, bold=False)
            changed = True

    return changed


def insert_single_field_table(doc: Document, code: str) -> bool:
    fields = FIELD_MAP.get(code, [])
    if not fields:
        return False
    paragraphs = doc.paragraphs
    je_idx = find_index_by_text(paragraphs, JE_HEADING)
    if je_idx is None:
        return False
    temp_doc = Document()
    table = temp_doc.add_table(rows=len(fields) + 1, cols=2)
    table.style = "Table Grid"
    table.cell(0, 0).text = "Field Name"
    table.cell(0, 1).text = "Meaning"
    for row_idx, field_name in enumerate(fields, start=1):
        table.cell(row_idx, 0).text = field_name
        table.cell(row_idx, 1).text = FIELD_MEANINGS.get(field_name, f"ใช้ตรวจข้อมูลของช่อง {field_name} ให้ตรงกับรายการจริง")
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    set_run_font(run, size=16, bold=(para.text.strip() in {"Field Name", "Meaning"}))
    paragraphs[je_idx]._p.addprevious(deepcopy(table._tbl))
    return True


def insert_journal_guide(doc: Document, code: str) -> bool:
    guide_lines = JOURNAL_GUIDE_MAP.get(code)
    if not guide_lines:
        return False
    paragraphs = doc.paragraphs
    je_idx = find_index_by_text(paragraphs, JE_HEADING)
    if je_idx is None:
        return False
    anchor = paragraphs[je_idx]
    last = insert_paragraph_after(anchor, JOURNAL_GUIDE_HEADING, size=16, bold=True)
    for line in guide_lines:
        last = insert_paragraph_after(last, line, size=16, bold=False)
    return True


def apply_layout_tuning(doc: Document) -> bool:
    changed = False
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        has_drawing = bool(paragraph._p.xpath(".//w:drawing"))
        if is_numbered_step(text) or text in {DETAIL_HEADING, JE_HEADING, JOURNAL_GUIDE_HEADING}:
            paragraph.paragraph_format.keep_with_next = True
            changed = True
        elif has_drawing:
            paragraph.paragraph_format.keep_with_next = True
            paragraph.paragraph_format.space_after = Pt(0)
            changed = True
        elif text.startswith("ภาพ"):
            paragraph.paragraph_format.keep_with_next = True
            paragraph.paragraph_format.space_after = Pt(6)
            changed = True
        elif text.startswith("รูป "):
            paragraph.paragraph_format.keep_with_next = True
            paragraph.paragraph_format.space_after = Pt(6)
            changed = True
    return changed


def normalize_doc(path: Path) -> bool:
    code = doc_code_from_name(path.name)
    if not code:
        return False

    doc = Document(path)
    changed = False

    cleanup_field_blocks(doc)
    cleanup_journal_guides(doc)
    if smooth_detailed_steps(doc, code):
        changed = True

    if find_index_by_text(doc.paragraphs, JE_HEADING) is not None:
        if insert_single_field_table(doc, code):
            changed = True
        if insert_journal_guide(doc, code):
            changed = True

    if apply_layout_tuning(doc):
        changed = True

    if changed:
        doc.save(path)
    return changed


def render_pdfs(paths: list[Path]) -> None:
    if not paths:
        return
    PDF_DIR.mkdir(parents=True, exist_ok=True)
    cmd = [str(SOFFICE), "--headless", "--convert-to", "pdf", "--outdir", str(PDF_DIR), *map(str, paths)]
    subprocess.run(cmd, check=True)


def main() -> None:
    updated: list[Path] = []
    for path in DOCX_DIR.glob("*.docx"):
        if path.name.startswith(TARGET_PREFIXES) and normalize_doc(path):
            updated.append(path)
    render_pdfs(updated)
    print(f"normalized {len(updated)} docs")


if __name__ == "__main__":
    main()

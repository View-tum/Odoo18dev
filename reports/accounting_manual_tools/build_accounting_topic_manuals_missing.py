from __future__ import annotations

import json
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\365_project\TheCool18e\Dev")
OUTPUT_DIR = ROOT / "manual" / "Accouting_Manual" / "generated_20260408"
DOCX_DIR = OUTPUT_DIR / "docx"
IMAGE_DIR = OUTPUT_DIR / "images"
SAMPLES_PATH = ROOT / "reports" / "accounting_manual_tools" / "output" / "manual_live_samples_20260408.json"
SUMMARY_PATH = OUTPUT_DIR / "missing_topics_summary_20260409.json"

VENDOR_BILLS = {
    "out_confirmed": {"id": 68481, "name": "APD/26/04/00006"},
    "out_paid": {"id": 68482, "name": "APD/26/04/00007"},
}


def repair_text(value: Any) -> Any:
    if isinstance(value, dict):
        return {key: repair_text(val) for key, val in value.items()}
    if isinstance(value, list):
        return [repair_text(item) for item in value]
    if not isinstance(value, str):
        return value
    if "à" not in value:
        return value
    for encoding in ("latin1", "cp1252"):
        try:
            return value.encode(encoding).decode("utf-8")
        except Exception:
            continue
    return value


def set_run_font(run, size: int = 16, bold: bool = False, italic: bool = False, color: RGBColor | None = None) -> None:
    run.bold = bold
    run.italic = italic
    run.font.name = "TH Sarabun New"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "TH Sarabun New")
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color


def set_cell_text(cell, text: str) -> None:
    cell.text = str(text)
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            set_run_font(run, 14)


def shade_cell(cell, color: str = "D9EAF7") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    tc_pr.append(shd)


def setup(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "TH Sarabun New"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "TH Sarabun New")
    style.font.size = Pt(16)
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)


def add_title(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    set_run_font(run, 24, bold=True, color=RGBColor(31, 78, 121))
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(subtitle)
    set_run_font(run2, 18)


def add_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, 18, bold=True, color=RGBColor(31, 78, 121))


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        set_run_font(run, 16)


def add_steps(doc: Document, title: str, steps: list[str]) -> None:
    p = doc.add_paragraph()
    run = p.add_run(title)
    set_run_font(run, 16, bold=True)
    for idx, step in enumerate(steps, start=1):
        p = doc.add_paragraph()
        run = p.add_run(f"{idx}. {step}")
        set_run_font(run, 16)


def add_table(doc: Document, headers: list[str], rows: list[tuple]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = True
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header)
        shade_cell(table.rows[0].cells[idx])
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)


def add_image(doc: Document, image_name: str, caption: str) -> None:
    path = IMAGE_DIR / image_name
    if not path.exists():
        return
    doc.add_picture(str(path), width=Inches(6.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(caption)
    set_run_font(run, 14, italic=True)


def add_table_section(doc: Document, title: str, rows: list[tuple], headers: list[str]) -> None:
    add_heading(doc, title)
    add_table(doc, headers, rows)


def move_lines_to_rows(lines: list[dict], descriptions: dict[str, str] | None = None) -> list[tuple[str, str, float, float]]:
    rows: list[tuple[str, str, float, float]] = []
    descriptions = descriptions or {}
    for line in lines:
        account_code = line["account_code"]
        rows.append(
            (
                f"{account_code} {line['account_name']}",
                descriptions.get(account_code, line["label"]),
                float(line["debit"]),
                float(line["credit"]),
            )
        )
    return rows


def cheque_journal_rows(samples: dict) -> list[tuple[str, str, str]]:
    return [
        ("Cheque Incoming", "เปิดให้ Journal รองรับการรับเช็คจากลูกค้า", f"เปิดใช้งานที่ Journal {samples['bank_journal']['code']}"),
        ("Cheque Outgoing", "เปิดให้ Journal รองรับการจ่ายเช็คให้ผู้ขาย", f"เปิดใช้งานที่ Journal {samples['bank_journal']['code']}"),
        ("Cheque Form", "เลือกเทมเพลตฟอร์มเช็คที่ใช้พิมพ์เช็คขาออก", samples["template"]["name"]),
        ("Incoming Payment Line", "กำหนดวิธีรับชำระและบัญชีพักเช็ครับ", f"Line ID {samples['bank_journal']['cheque_in_line_id']}"),
        ("Outgoing Payment Line", "กำหนดวิธีจ่ายเช็คและบัญชีพักเช็คจ่าย", f"Line ID {samples['bank_journal']['cheque_out_line_id']}"),
    ]


def save_topic(
    filename: str,
    title: str,
    subtitle: str,
    objectives: list[str],
    overview: list[str],
    menus: list[str],
    functions_rows: list[tuple],
    steps: list[tuple[str, list[str]]],
    images: list[tuple[str, str]],
    field_tables: list[tuple[str, list[tuple], list[str]]],
    scenarios: list[tuple],
    journal_tables: list[tuple[str, list[tuple]]],
    cautions: list[str],
) -> None:
    doc = Document()
    setup(doc)
    add_title(doc, title, subtitle)
    doc.add_paragraph()
    add_heading(doc, "วัตถุประสงค์")
    add_bullets(doc, objectives)
    add_heading(doc, "ภาพรวมการทำงาน")
    add_bullets(doc, overview)
    add_heading(doc, "เมนูที่ใช้")
    add_bullets(doc, menus)
    add_heading(doc, "ฟังก์ชันที่ทำได้จริง")
    add_table(doc, ["ฟังก์ชัน", "คำอธิบาย", "ใช้เมื่อใด"], functions_rows)
    add_heading(doc, "ขั้นตอนการใช้งาน")
    for step_title, step_items in steps:
        add_steps(doc, step_title, step_items)
    add_heading(doc, "ภาพประกอบการใช้งาน")
    for caption, image_name in images:
        add_image(doc, image_name, caption)
    add_heading(doc, "คำอธิบายฟิลด์และส่วนสำคัญบนหน้าจอ")
    for table_title, rows, headers in field_tables:
        add_table_section(doc, table_title, rows, headers)
    add_heading(doc, "ตัวอย่าง Scenario จากข้อมูลจริงใน local UAT")
    add_table(doc, ["Scenario", "ข้อมูลตัวอย่าง", "ผลลัพธ์ที่ต้องเห็น"], scenarios)
    if journal_tables:
        add_heading(doc, "การอธิบาย Journal Items และขาบัญชี")
        add_bullets(
            doc,
            [
                "Debit หมายถึงบัญชีฝั่งรับผลประโยชน์หรือรับมูลค่าเพิ่มขึ้นในรายการนั้น",
                "Credit หมายถึงบัญชีฝั่งที่ถูกตัดออกหรือเป็นแหล่งที่มาของมูลค่าในรายการนั้น",
                "ผู้ใช้ควรเปิด Smart Button ชื่อ Payment Entry หรือ Reversed Entry เพื่อตรวจสอบรายการจริงทุกครั้ง",
            ],
        )
        for table_title, rows in journal_tables:
            add_table_section(doc, table_title, rows, ["บัญชี", "คำอธิบายเชิงธุรกิจ", "Debit", "Credit"])
    add_heading(doc, "ข้อควรระวัง")
    add_bullets(doc, cautions)
    doc.save(DOCX_DIR / filename)


def main() -> None:
    DOCX_DIR.mkdir(parents=True, exist_ok=True)
    samples = repair_text(json.loads(SAMPLES_PATH.read_text(encoding="utf-8")))
    out_cf = samples["cheques"]["outbound_confirmed"]
    out_pd = samples["cheques"]["outbound_paid"]
    book = samples["cheque_book"]
    template = samples["template"]
    bank_journal = samples["bank_journal"]

    cheque_common_rows = [
        ("name", "เลขที่เช็คที่ระบบสร้างให้จากเช็คเล่มหรือรูปแบบรับเช็ค", "เช่น 860200001 หรือ RCV-INV-E/26/04/00003"),
        ("state", "สถานะของเช็คในปัจจุบัน", "draft, confirmed, paid, cancelled"),
        ("journal_id", "Journal ธนาคารที่ใช้กับเช็คใบนี้", f"{bank_journal['code']} {bank_journal['name']}"),
        ("partner_id", "คู่ค้าเจ้าของรายการรับหรือจ่าย", "ใช้ตรวจสอบว่าจ่ายให้ใครหรือรับมาจากใคร"),
        ("amount", "ยอดเงินบนเช็ค", "ต้องตรวจสอบให้ตรงกับยอดชำระจริง"),
        ("payment_ids", "รายการ payment ที่เชื่อมกับเช็ค", "ใช้ drill down ไปยัง Payment Entry"),
    ]

    topics_created: list[str] = []

    save_topic(
        filename="5.1_5.1 เปิดใช้งานเช็คเป็นวิธีการชำระเงิน.docx",
        title="5.1 เปิดใช้งานเช็คเป็นวิธีการชำระเงิน",
        subtitle="Module Cheque",
        objectives=[
            "เพื่อให้ผู้ใช้งานสามารถเปิดใช้การรับเช็คและจ่ายเช็คบน Journal ธนาคารได้ถูกต้อง",
            "เพื่อให้เข้าใจว่าฟิลด์ใน Journal และ Payment Method Line มีผลต่อบัญชีพักเช็ครับและเช็คจ่ายอย่างไร",
            "เพื่อให้ผู้ใช้งานตรวจสอบหน้าจอที่เกี่ยวข้องก่อนเริ่มใช้งาน scenario รับเช็คและจ่ายเช็คจริง",
        ],
        overview=[
            "โมดูลเช็คของ Gold Mints ใช้ Journal ธนาคารเป็นแกนหลัก เมื่อเปิดใช้ฟังก์ชันเช็คแล้ว ผู้ใช้จะสามารถเลือกวิธีชำระเป็นเช็คได้จาก Register Payment และเมนูเช็คเฉพาะทางจะเริ่มใช้งานได้",
            f"ตัวอย่างใน local UAT ใช้ Journal รหัส {bank_journal['code']} ชื่อ {bank_journal['name']} ซึ่งมีทั้ง Payment Method แบบรับเช็คและจ่ายเช็คพร้อมใช้งานจริง",
            "การตั้งค่าที่ถูกต้องในขั้นตอนนี้เป็นฐานของหัวข้อ 5.3 ถึง 5.8 ทั้งหมด หากตั้งไม่ครบ ผู้ใช้จะไม่เห็นช่องกรอกข้อมูลเช็คใน wizard",
        ],
        menus=[
            "Cheque > Configuration > Settings",
            "Accounting > Configuration > Journals",
            f"Journal ตัวอย่างที่ใช้ใน local UAT คือ {bank_journal['code']}",
        ],
        functions_rows=[
            ("Is Reverse Cheque Entry?", "กำหนดให้ระบบสร้างรายการกลับรายการให้โดยอัตโนมัติเมื่อมีการ Void เช็ค", "ใช้เมื่อบริษัทต้องการควบคุมการกลับรายการทางบัญชีในแนวเดียวกัน"),
            ("Cheque Incoming", "เปิดการใช้งานรับเช็คจากลูกค้าใน Journal นี้", "ใช้กับ Journal ที่รับชำระจากลูกค้าด้วยเช็ค"),
            ("Cheque Outgoing", "เปิดการใช้งานจ่ายเช็คให้ผู้ขายใน Journal นี้", "ใช้กับ Journal ที่ต้องออกเช็คจ่าย"),
            ("Cheque Form", "ผูกเทมเพลตฟอร์มเช็คสำหรับใช้พิมพ์เช็คขาออก", "ใช้หลังจากสร้างหรือเลือก Template แล้ว"),
            ("Payment Method Line", "กำหนดวิธีชำระเงินและบัญชีพักของเช็คแต่ละฝั่ง", "ใช้ตรวจสอบว่ารายการบัญชีจะลงถูกบัญชีเมื่อรับหรือจ่ายเช็ค"),
        ],
        steps=[
            ("ขั้นตอนที่ 1 เปิดการตั้งค่าระดับระบบ", [
                "เข้าเมนู Cheque > Configuration > Settings แล้วตรวจสอบตัวเลือก Is Reverse Cheque Entry? หากองค์กรต้องการให้ระบบสร้างรายการกลับรายการเมื่อมีการ Void เช็คให้เปิดใช้งานไว้",
                "กด Save เพื่อบันทึกการตั้งค่าระดับระบบก่อนเข้าสู่การตั้งค่าระดับ Journal",
            ]),
            ("ขั้นตอนที่ 2 เปิดใช้งานบน Journal ธนาคาร", [
                "เข้าเมนู Accounting > Configuration > Journals แล้วเปิด Journal รหัส PBAY1",
                "ในแท็บ General ให้ติ๊กเปิด Cheque Incoming และ Cheque Outgoing",
                "ในส่วน Cheque Form ให้เลือก Standard Cheque หรือ Template ที่องค์กรใช้จริง",
            ]),
            ("ขั้นตอนที่ 3 ตรวจสอบ Payment Method Line", [
                "ในส่วน Incoming Payments ให้ตรวจว่า Payment Method Line ที่เป็นเช็ครับมีการตั้งบัญชีพักเช็ครับไว้แล้ว",
                "ในส่วน Outgoing Payments ให้ตรวจว่า Payment Method Line ที่เป็นเช็คจ่ายมีการตั้งบัญชีพักเช็คจ่ายไว้แล้ว",
                "ทดสอบเปิด Register Payment ในเอกสารตัวอย่างเพื่อยืนยันว่าระบบแสดงส่วนกรอกข้อมูลเช็คจริง",
            ]),
        ],
        images=[
            ("รูป 5.1.1 การเปิดใช้งานตัวเลือก Is Reverse Cheque Entry? ในหน้า Settings", "settings_cheque_manual_annotated.png"),
            ("รูป 5.1.2 การตั้งค่าเช็คในหน้า Journal ส่วน General", "journal_general_manual_annotated.png"),
            ("รูป 5.1.3 การตั้งค่า Payment Method Line ฝั่งรับเช็ค", "journal_incoming_manual_annotated.png"),
            ("รูป 5.1.4 การตั้งค่า Payment Method Line ฝั่งจ่ายเช็ค", "journal_outgoing_manual_annotated.png"),
        ],
        field_tables=[
            ("ตารางอธิบายฟิลด์สำคัญใน Settings และ Journal", cheque_journal_rows(samples) + [
                ("Is Reverse Cheque Entry?", "ตัวเลือกใน Settings ระดับระบบ", "เมื่อเปิดใช้ ระบบจะเตรียม flow สำหรับ reverse entry ตอน Void เช็ค"),
                ("Bank Account ของ Journal", "บัญชีธนาคารจริงของ Journal", "ใช้เป็นปลายทางตอนเช็คขาออกเคลียร์ธนาคารหรือเช็ครับถูกฝากเข้าธนาคาร"),
            ], ["Field / Section", "ความหมาย", "วิธีใช้งานหรือค่าตัวอย่าง"]),
            ("บัญชีที่เกี่ยวข้องกับการรับและจ่ายเช็คใน local UAT", [
                ("payment_account_id ฝั่งรับเช็ค", "บัญชีพักเช็ครับ", "113005 เช็ครับลงวันที่ล่วงหน้า"),
                ("payment_account_id ฝั่งจ่ายเช็ค", "บัญชีพักเช็คจ่าย", "212004 เช็คจ่ายลงวันที่ล่วงหน้า"),
                ("Bank Account ของ Journal", "บัญชีเงินฝากธนาคารจริง", "111201 BAY CA สามแยก #046-0-14721-8"),
            ], ["Field / Section", "ความหมาย", "วิธีใช้งานหรือค่าตัวอย่าง"]),
        ],
        scenarios=[
            ("Scenario A: เปิดใช้งานเฉพาะรับเช็ค", "เหมาะกับ Journal ที่ใช้รับเงินลูกค้าอย่างเดียว", "ผู้ใช้จะเห็นเฉพาะส่วนรับเช็คในหน้ารับชำระเงิน"),
            ("Scenario B: เปิดใช้งานครบทั้งรับและจ่าย", f"ใช้กับ Journal {bank_journal['code']} ใน local UAT", "ผู้ใช้สามารถใช้ Journal เดียวกันได้ทั้งหัวข้อ 5.4 และ 5.6"),
        ],
        journal_tables=[],
        cautions=[
            "ถ้าไม่เปิด flag ที่ Journal แม้จะมี Payment Method Line เป็นเช็คอยู่ ระบบจะไม่แสดงส่วนกรอกข้อมูลเช็คใน wizard",
            "บัญชีพักเช็ครับและบัญชีพักเช็คจ่ายต้องเป็นคนละบัญชี เพื่อให้ติดตาม Outstanding และเคลียร์ธนาคารได้ถูกต้อง",
            "หลังเปลี่ยนการตั้งค่า Journal ควรทดสอบด้วยเอกสารตัวอย่าง 1 รายการก่อนเริ่มใช้งานจริง",
        ],
    )
    topics_created.append("5.1")

    save_topic(
        filename="5.2_5.2 เลือกเทมเพลตฟอร์มเช็ค (Cheque Form Template).docx",
        title="5.2 เลือกเทมเพลตฟอร์มเช็ค (Cheque Form Template)",
        subtitle="Module Cheque",
        objectives=[
            "เพื่อให้ผู้ใช้งานเลือกหรือปรับเทมเพลตการพิมพ์เช็คให้ตรงกับกระดาษเช็คของธนาคาร",
            "เพื่อให้เข้าใจว่ากลุ่มฟิลด์ตำแหน่งพิมพ์แต่ละชุดมีผลกับส่วนใดบนหน้าเช็ค",
            "เพื่อให้ผู้ใช้งานย้อนกลับไปผูก Template ที่ Journal ได้อย่างถูกต้อง",
        ],
        overview=[
            "Cheque Form Template เป็นแม่แบบการพิมพ์เช็คขาออก ไม่ได้สร้างบัญชีโดยตรง แต่มีผลกับความถูกต้องของตำแหน่งข้อความบนเช็คจริง",
            f"ใน local UAT มี Template ตัวอย่างชื่อ {template['name']} และถูกใช้ร่วมกับ Journal {bank_journal['code']}",
            "ผู้ใช้งานควรเข้าใจทั้งขนาดกระดาษเช็ค ระยะขอบ ตำแหน่งวันที่ ชื่อผู้รับเงิน จำนวนเงินตัวเลข และจำนวนเงินตัวอักษร ก่อนนำไปพิมพ์กับเช็คจริง",
        ],
        menus=["Cheque > Configuration > Cheque Lists", "Accounting > Configuration > Journals"],
        functions_rows=[
            ("สร้าง/แก้ไข Template", "กำหนดขนาดเช็ค ระยะขอบ และตำแหน่งข้อความที่ระบบจะพิมพ์", "ใช้เมื่อเปลี่ยนแบบฟอร์มเช็คหรือธนาคารออกกระดาษเช็คใหม่"),
            ("กำหนด Payee / Date / Amount", "คุมตำแหน่งข้อความสำคัญบนเช็ค", "ใช้เพื่อให้ข้อความตรงกับพื้นที่พิมพ์บนกระดาษจริง"),
            ("ผูก Template กับ Journal", "ทำให้ Journal ขาออกเรียกแม่แบบนี้ไปใช้ตอนพิมพ์", "ใช้หลังตรวจสอบ Template แล้ว"),
        ],
        steps=[
            ("ขั้นตอนที่ 1 เปิดและตรวจ Template", [
                "เข้าเมนู Cheque > Configuration > Cheque Lists แล้วเปิด Template ชื่อ Standard Cheque",
                "ตรวจสอบขนาดเช็ค ความสูง ความกว้าง และค่าระยะขอบหลัก ๆ ให้ตรงกับกระดาษเช็คของธนาคาร",
            ]),
            ("ขั้นตอนที่ 2 ตรวจฟิลด์ตำแหน่งพิมพ์", [
                "ตรวจกลุ่มฟิลด์ที่เกี่ยวกับวันที่เช็ค ชื่อผู้รับเงิน จำนวนเงินตัวเลข และจำนวนเงินตัวอักษร",
                "หากต้องปรับตำแหน่ง ให้แก้เฉพาะค่าที่เกี่ยวข้องและจดบันทึกค่าเดิมไว้ก่อนทุกครั้ง",
            ]),
            ("ขั้นตอนที่ 3 ผูกกับ Journal", [
                "กลับไปที่เมนู Accounting > Configuration > Journals แล้วเปิด PBAY1",
                "ตรวจว่าฟิลด์ Cheque Form ของ Journal ผูกกับ Standard Cheque แล้ว",
                "ทดสอบพิมพ์กับกระดาษเปล่าก่อนใช้กับเช็คจริง",
            ]),
        ],
        images=[("รูป 5.2.1 หน้าจอ Template ของเช็ค พร้อมจุดหลักที่ใช้ปรับตำแหน่งพิมพ์", "template_manual_annotated.png")],
        field_tables=[
            ("กลุ่มฟิลด์สำคัญใน Cheque Form Template", [
                ("name", "ชื่อ Template", template["name"]),
                ("cheque_hight / cheque_width", "กำหนดขนาดกระดาษเช็ค", "ต้องเทียบกับกระดาษเช็คจริง"),
                ("top_margin / left_margin / font_size", "กำหนดตำแหน่งและขนาดข้อความวันที่เช็ค", "ใช้เมื่อวันที่พิมพ์สูงหรือต่ำเกินไป"),
                ("payee_top_margin / payee_left_margin / payee_width", "กำหนดตำแหน่งชื่อผู้รับเงิน", "ใช้เมื่อชื่อผู้รับเงินไม่ตรงบรรทัด"),
                ("af_top_margin / af_left_margin / af_width", "กำหนด Amount in Figure", "ใช้ควบคุมตำแหน่งยอดตัวเลข"),
                ("fl_* / sc_*", "กำหนด Amount in Word และการตัดคำ", "ใช้เมื่อข้อความตัวอักษรล้นบรรทัดหรือไม่ตรงตำแหน่ง"),
            ], ["Field / Section", "ความหมาย", "วิธีใช้งานหรือค่าตัวอย่าง"]),
        ],
        scenarios=[
            ("Scenario A: ใช้ Template เดิมของธนาคาร", "Template Standard Cheque พร้อมใช้งานอยู่แล้ว", "ผู้ใช้สามารถผูกกับ Journal แล้วพิมพ์เช็คได้ทันที"),
            ("Scenario B: ปรับตำแหน่งข้อความเพราะฟอร์มเช็คเปลี่ยน", "แก้ margin และ font size ก่อนพิมพ์ทดสอบ", "ตำแหน่งตัวอักษรจะตรงกับกระดาษเช็คใหม่โดยไม่กระทบการลงบัญชี"),
        ],
        journal_tables=[],
        cautions=[
            "การแก้ Template มีผลกับเช็คที่จะพิมพ์หลังจากนั้นทั้งหมดใน Journal ที่ผูก Template นี้",
            "ควรพิมพ์ทดสอบกับกระดาษเปล่าก่อนทุกครั้งเพื่อป้องกันการเสียหน้าเช็คจริง",
        ],
    )
    topics_created.append("5.2")

    save_topic(
        filename="5.3_5.3 สร้างสมุดเช็ค (Cheque-Book).docx",
        title="5.3 สร้างสมุดเช็ค (Cheque-Book)",
        subtitle="Module Cheque",
        objectives=[
            "เพื่อให้ผู้ใช้งานสร้างสมุดเช็คและชุดเลขเช็คสำหรับใช้จ่ายเช็คในระบบได้ถูกต้อง",
            "เพื่อให้เข้าใจลำดับสถานะ Draft > Submit > Done ของสมุดเช็ค",
            "เพื่อให้ผู้ใช้งานสามารถตรวจเช็คคงเหลือในเล่มได้จากหน้าจอเดียว",
        ],
        overview=[
            "สมุดเช็คเป็น Master Data ระดับปฏิบัติการสำหรับเก็บช่วงเลขเช็คของธนาคาร เมื่อมีการจ่ายเช็ค ระบบจะดึงเลขจาก Cheque Book Lines ที่พร้อมใช้งานมาใช้",
            f"ใน local UAT มีตัวอย่างเล่มจริงชื่อ {book['name']} สถานะ {book['state']} ผูกกับ Journal {bank_journal['code']}",
            "ผู้ใช้งานต้องสร้างเล่มให้ถูกก่อนใช้งานการจ่ายเช็ค มิฉะนั้น Wizard จ่ายเช็คจะไม่มีเลขเช็คให้เลือก",
        ],
        menus=["Cheque > Cheque Book"],
        functions_rows=[
            ("Submit", "ยืนยันข้อมูลหัวเล่มก่อนสร้างเลขเช็คย่อย", "ใช้เมื่อกรอก Journal, จำนวนเช็ค และเลขเริ่มต้นครบแล้ว"),
            ("Generate Cheque", "สร้าง Cheque Lines ตามจำนวนใบและช่วงเลขที่กำหนด", "ใช้หลังจาก Submit"),
            ("Clear Cheque", "ล้างรายการเช็คย่อยที่สร้างไว้ในสถานะ Submit", "ใช้เมื่อกรอกช่วงเลขผิดและยังไม่ Confirm"),
            ("Confirm", "ยืนยันให้สมุดเช็คพร้อมใช้งานจริง", "ใช้เมื่อรายการเลขเช็คถูกต้องและพร้อมนำไปจ่าย"),
        ],
        steps=[
            ("ขั้นตอนที่ 1 สร้างหัวสมุดเช็ค", [
                "เข้าเมนู Cheque > Cheque Book แล้วกด New",
                "กรอก Bank Account Journal, จำนวนเช็ค และ First Cheque No ให้ครบ",
                "กด Submit เพื่อยืนยันข้อมูลหัวเล่ม",
            ]),
            ("ขั้นตอนที่ 2 สร้างเลขเช็คย่อย", [
                "กด Generate Cheque เพื่อให้ระบบสร้างรายการในแท็บ Cheque List ตามจำนวนใบที่ระบุ",
                "ตรวจช่วงเลขเช็คตั้งแต่ใบแรกถึงใบสุดท้ายว่าตรงกับเล่มจริง",
            ]),
            ("ขั้นตอนที่ 3 ยืนยันให้พร้อมใช้งาน", [
                "หากเลขเช็คถูกต้องให้กด Confirm",
                "หลัง Confirm แล้ว เลขเช็คสถานะ Draft ในเล่มจะถูกเรียกใช้จาก Wizard จ่ายเช็คได้",
            ]),
        ],
        images=[("รูป 5.3.1 หน้าจอสมุดเช็ค พร้อมฟิลด์หลักและแท็บ Cheque List", "cheque_book_manual_annotated.png")],
        field_tables=[
            ("ฟิลด์สำคัญของสมุดเช็ค", [
                ("bank_account_journal_id", "Journal ธนาคารที่เชื่อมกับเช็คเล่มนี้", f"{bank_journal['code']} {bank_journal['name']}"),
                ("cheque_qty", "จำนวนเช็คที่จะสร้างในเล่ม", "ระบบใช้ค่านี้สร้างรายการเช็คย่อย"),
                ("first_cheque_no_char / last_cheque_no_char", "ช่วงเลขเช็คของเล่ม", f"ตัวอย่างเลขที่เหลือ เช่น {', '.join(book['draft_leaves'][:3])}"),
                ("cheque_book_lines", "รายการเช็คย่อยในเล่ม", "ใช้ติดตามว่าใบใด draft, paid, cancelled หรือ return"),
                ("state", "สถานะของสมุดเช็ค", "Draft > Submit > Done"),
            ], ["Field / Section", "ความหมาย", "วิธีใช้งานหรือค่าตัวอย่าง"]),
        ],
        scenarios=[
            ("Scenario A: สร้างเล่มใหม่เพื่อจ่ายเช็คชุดใหม่", "ใช้เมื่อธนาคารออกสมุดเช็คเล่มใหม่", "ผู้ใช้ต้องกำหนดช่วงเลขเช็คให้ตรงกับเล่มจริงก่อน Confirm"),
            ("Scenario B: ตรวจเลขเช็คคงเหลือในเล่มเดิม", f"เปิดเล่ม {book['name']}", "ผู้ใช้สามารถดูเลขเช็คที่ยังเป็น Draft ได้จากแท็บ Cheque List"),
        ],
        journal_tables=[],
        cautions=[
            "เมื่อ Confirm แล้ว ควรหลีกเลี่ยงการแก้เลขเช็คเริ่มต้นหรือสิ้นสุดย้อนหลัง",
            "หากกรอกช่วงเลขผิด ควรแก้ตั้งแต่สถานะ Submit ก่อนนำเลขเช็คไปใช้งานจริง",
        ],
    )
    topics_created.append("5.3")

    save_topic(
        filename="5.4_5.4 ชำระบิลผู้ขายด้วยเช็ค.docx",
        title="5.4 ชำระบิลผู้ขายด้วยเช็ค",
        subtitle="Module Cheque",
        objectives=[
            "เพื่อให้ผู้ใช้งานจ่าย Vendor Bill ด้วยเช็คจากหน้าบิลได้ถูกต้อง",
            "เพื่อให้เข้าใจผลลัพธ์ที่ระบบสร้างทั้งในเมนูเช็คและใน Journal Items",
            "เพื่อให้ผู้ใช้สามารถตรวจสอบเดบิตและเครดิตในแต่ละสถานะของเช็คขาออกได้",
        ],
        overview=[
            "การจ่ายบิลผู้ขายด้วยเช็คเริ่มจากปุ่ม Register Payment ใน Vendor Bill ผู้ใช้เลือก Journal ธนาคารและ Payment Method เป็นเช็ค จากนั้นกรอกเลขเช็คและรายละเอียดในส่วน Outgoing Cheque",
            f"คู่มือนี้ใช้อ้างอิงบิลจริง {VENDOR_BILLS['out_confirmed']['name']} และ {VENDOR_BILLS['out_paid']['name']} เพื่ออธิบายทั้งสถานะ Confirmed และ Paid",
            "เมื่อผู้ใช้กด Create Payment ระบบจะสร้างเช็คขาออกและ Payment Entry ให้อัตโนมัติ จากนั้นผู้ใช้ติดตามเช็คต่อในเมนู Cheque Paying",
        ],
        menus=["Accounting > Vendors > Bills", "Cheque > Cheque > Cheque Paying"],
        functions_rows=[
            ("Register Payment", "เปิด Wizard รับหรือจ่ายเงินมาตรฐานของ Odoo", "ใช้เป็นจุดเริ่มต้นของการจ่ายเช็คจากหน้าบิล"),
            ("Outgoing Cheque Section", "ส่วนกรอกเลขเช็ค วันที่เช็ค หมายเหตุ และ A/C Payee", "แสดงเฉพาะเมื่อเลือก Payment Method เป็นเช็คขาออก"),
            ("Confirm", "ยืนยันเช็คให้อยู่สถานะ Confirmed", "ใช้เมื่อออกเช็คแล้วแต่ยังไม่ตัดผ่านธนาคาร"),
            ("Bank Deposit / Done", "ตัดยอดเช็คผ่านธนาคารจนเป็น Paid", "ใช้เมื่อธนาคารนำเช็คไปตัดบัญชีจริง"),
            ("Payment Entry", "Smart Button สำหรับเปิด Journal Entry ที่ระบบสร้าง", "ใช้ตรวจสอบการลงบัญชีเดบิตและเครดิต"),
        ],
        steps=[
            ("ขั้นตอนที่ 1 เปิดบิลและเข้าสู่ Wizard จ่ายเช็ค", [
                f"เปิดบิลผู้ขาย {VENDOR_BILLS['out_confirmed']['name']} จากเมนู Accounting > Vendors > Bills",
                "กดปุ่ม Register Payment เพื่อเปิด Wizard รับหรือจ่ายเงิน",
                f"เลือก Journal = {bank_journal['code']} และ Payment Method = Cheque Payment (Outbound)",
            ]),
            ("ขั้นตอนที่ 2 กรอกข้อมูลเช็คและสร้างรายการ", [
                "ในส่วน Outgoing Cheque ให้เลือกเลขเช็คจาก Cheque Book",
                "กรอกวันที่เช็ค หมายเหตุ และตรวจยอดให้ตรงกับบิลที่จะชำระ",
                "กด Create Payment เพื่อสร้างเช็คขาออกและ Payment Entry",
            ]),
            ("ขั้นตอนที่ 3 ติดตามสถานะหลังจ่าย", [
                f"เปิดเช็คเลขที่ {out_cf['name']} หรือ {out_pd['name']} ในเมนู Cheque > Cheque > Cheque Paying",
                "ตรวจสอบสถานะเช็ค ปุ่ม Bank Deposit หรือ Done และ Smart Buttons เช่น Payments Cheque, Payment Entry และ Reversed Entry",
                "เมื่อต้องการเคลียร์ธนาคาร ให้กด Bank Deposit หรือ Done ตามสิทธิ์และ flow ของระบบ",
            ]),
        ],
        images=[
            ("รูป 5.4.1 หน้าต่าง Register Payment จากบิลผู้ขาย พร้อมจุดที่ต้องกรอก", "bill_register_payment_manual_annotated.png"),
            ("รูป 5.4.2 เช็คขาออกสถานะ Confirmed หลังสร้างรายการแล้ว", "cheque_out_confirmed_manual_annotated.png"),
            ("รูป 5.4.3 เช็คขาออกสถานะ Paid หลังตัดผ่านธนาคาร", "cheque_out_paid_manual_annotated.png"),
        ],
        field_tables=[
            ("ฟิลด์สำคัญของการจ่ายเช็คขาออก", cheque_common_rows + [
                ("cheque_book_id / cheque_id", "สมุดเช็คและเลขเช็คที่ถูกดึงมาใช้จริง", "ต้องมาจากเล่มที่อยู่สถานะพร้อมใช้งาน"),
                ("memo", "ข้อความอ้างอิงบนเช็คและในรายการบัญชี", "ช่วยค้นหารายการย้อนหลัง"),
                ("ac_payee", "ระบุว่าเช็คเป็น A/C Payee หรือไม่", "ใช้ควบคุมข้อความบนหน้าเช็ค"),
                ("payment_method_line_id", "วิธีการชำระเงินจริง", f"ใน local UAT ใช้ Line ID {bank_journal['cheque_out_line_id']}"),
            ], ["Field / Section", "ความหมาย", "วิธีใช้งานหรือค่าตัวอย่าง"]),
        ],
        scenarios=[
            ("Scenario A: จ่ายแล้วแต่ยังไม่ตัดผ่านธนาคาร", f"ใช้เช็ค {out_cf['name']} สถานะ {out_cf['state']}", "ยอดจะยังคงอยู่ในบัญชีพักเช็คจ่ายและติดตามเป็น Outstanding ได้"),
            ("Scenario B: จ่ายและตัดผ่านแล้ว", f"ใช้เช็ค {out_pd['name']} สถานะ {out_pd['state']}", "ระบบจะสร้างรายการย้ายจากบัญชีพักเช็คจ่ายไปยังบัญชีธนาคารจริง"),
        ],
        journal_tables=[
            ("Journal Items ตอนสร้างเช็คขาออกสถานะ Confirmed", move_lines_to_rows(out_cf["payment_moves"][0]["lines"], {
                "212001": "Debit เจ้าหนี้การค้า เพื่อปิดหนี้ของบิลผู้ขายที่จ่ายด้วยเช็ค",
                "212004": "Credit บัญชีพักเช็คจ่าย เพื่อแสดงว่าออกเช็คแล้วแต่ธนาคารยังไม่ตัดเงินจริง",
            })),
            ("Journal Items ตอนเช็คขาออกถูกตัดผ่านธนาคาร", move_lines_to_rows(out_pd["deposit_move"]["lines"], {
                "212004": "Debit บัญชีพักเช็คจ่ายเพื่อล้างยอดคงค้างของเช็คที่ออกไปแล้ว",
                "111201": "Credit บัญชีธนาคารจริง เพราะเงินถูกตัดออกจากบัญชีธนาคารแล้ว",
            })),
        ],
        cautions=[
            "ยอดใน Wizard ต้องเท่ากับยอดที่ต้องการชำระจริง มิฉะนั้นยอดค้างของบิลจะเหลือหรือจ่ายเกิน",
            "ถ้าเลือกเลขเช็คผิด จะกระทบการติดตาม Outstanding Cheque และการพิมพ์เช็คทันที",
            "ก่อนกด Bank Deposit หรือ Done ควรเปิด Payment Entry เพื่อตรวจสอบบัญชีพักและบัญชีธนาคารทุกครั้ง",
        ],
    )
    topics_created.append("5.4")

    save_topic(
        filename="5.5_5.5 ติดตามเช็คคงค้าง (Outstanding Cheques).docx",
        title="5.5 ติดตามเช็คคงค้าง (Outstanding Cheques)",
        subtitle="Module Cheque",
        objectives=[
            "เพื่อให้ผู้ใช้งานติดตามเช็คที่ยังไม่ตัดผ่านธนาคารได้ถูกต้อง",
            "เพื่อให้แยกสถานะ Confirmed, ระหว่างเคลียร์, และ Paid ได้ชัดเจน",
            "เพื่อให้ผู้ใช้งานตรวจสอบยอดคงค้างจาก Journal Items ได้อย่างเข้าใจ",
        ],
        overview=[
            "Outstanding Cheque ในบริบทของระบบนี้หมายถึงเช็คที่สร้างและยืนยันแล้ว แต่ยังไม่ถูกเคลียร์ผ่านธนาคาร ผู้ใช้สามารถติดตามผ่านเมนู Cheque Paying และ Cheque Transactions",
            f"ตัวอย่างเช็คคงค้างจริงใน local UAT คือเลขที่ {out_cf['name']} สถานะ {out_cf['state']}",
            "หัวข้อนี้เน้นการอ่านสถานะ การเปิด Smart Buttons เพื่อตรวจสอบ Payment Entry และการตัดสินใจว่าเมื่อไรควรกด Bank Deposit หรือ Done",
        ],
        menus=["Cheque > Cheque > Cheque Paying", "Cheque > Cheque > Cheque Transactions", "Cheque > Cheque > Paid Cheque"],
        functions_rows=[
            ("Cheque Paying", "มุมมองหลักสำหรับเช็คขาออกที่ยังทำงานอยู่", "ใช้ค้นหาเช็คที่ยัง Confirmed หรือยังไม่ปิด flow"),
            ("Cheque Transactions", "มุมมองรวมธุรกรรมเช็คทั้งหมด", "ใช้ค้นหาย้อนหลังหรือใช้ filter เพิ่มเติม"),
            ("Paid Cheque", "มุมมองเช็คที่เคลียร์แล้ว", "ใช้ตรวจเช็คผลลัพธ์หลังจบ flow"),
            ("Payment Entry", "Smart Button เปิดรายการบัญชีของเช็ค", "ใช้ยืนยันว่าบัญชีพักเช็คจ่ายยังมียอดคงอยู่หรือถูกล้างไปแล้ว"),
            ("Bank Deposit", "ปุ่มเปลี่ยนสถานะเช็คสู่ขั้นตอนเคลียร์ธนาคาร", "ใช้เมื่อธนาคารเริ่มนำเช็คไปเรียกเก็บ"),
        ],
        steps=[
            ("ขั้นตอนที่ 1 ค้นหาเช็คคงค้าง", [
                "เข้าเมนู Cheque > Cheque > Cheque Paying แล้วค้นหาเลขเช็คหรือชื่อคู่ค้าที่ต้องการ",
                f"เปิดเช็คตัวอย่าง {out_cf['name']} เพื่อดูรายละเอียดเช็คขาออกที่ยังคงค้าง",
            ]),
            ("ขั้นตอนที่ 2 ตรวจสถานะและรายการบัญชี", [
                "ตรวจสอบฟิลด์ state ว่าอยู่ในสถานะ Confirmed หรือสถานะอื่น",
                "กด Smart Button Payment Entry เพื่อตรวจสอบว่ารายการบัญชีปัจจุบันยังคง Credit บัญชีพักเช็คจ่ายอยู่หรือไม่",
            ]),
            ("ขั้นตอนที่ 3 ตัดสินใจเมื่อเช็คเริ่มเคลียร์", [
                "หากธนาคารเริ่มนำเช็คไปเรียกเก็บ ให้กด Bank Deposit เพื่อขยับเช็คเข้าสู่ขั้นตอนการเคลียร์",
                f"เมื่อเช็คถูกตัดผ่านบัญชีแล้ว สามารถตรวจเช็คตัวอย่างเลขที่ {out_pd['name']} ในเมนู Paid Cheque ประกอบได้",
            ]),
        ],
        images=[
            ("รูป 5.5.1 ตัวอย่างเช็คขาออกสถานะ Outstanding/Confirmed ที่ยังไม่ตัดผ่านธนาคาร", "cheque_out_confirmed_manual_annotated.png"),
            ("รูป 5.5.2 ตัวอย่างเช็คขาออกสถานะ Paid สำหรับเปรียบเทียบหลังเคลียร์แล้ว", "cheque_out_paid_manual_annotated.png"),
        ],
        field_tables=[
            ("สิ่งที่ต้องตรวจบนเช็คคงค้าง", [
                ("state", "สถานะปัจจุบันของเช็ค", "Confirmed = คงค้าง, Paid = เคลียร์แล้ว"),
                ("payment_ids", "รายการ Payment ที่ผูกกับเช็ค", "ใช้ตรวจสอบยอดที่จ่ายจริงและชื่อรายการชำระ"),
                ("payment_moves", "Payment Entry ที่ระบบสร้างตอนออกเช็ค", "ใช้ดูว่าปิดเจ้าหนี้แล้วหรือยัง"),
                ("deposit_move", "Journal Entry ตอนเคลียร์ธนาคาร", "จะมีค่าเมื่อเช็คถูกตัดผ่านธนาคารแล้ว"),
            ], ["Field / Section", "ความหมาย", "วิธีใช้งานหรือค่าตัวอย่าง"]),
        ],
        scenarios=[
            ("Scenario A: เช็คยังไม่ออกจากธนาคาร", f"เช็ค {out_cf['name']} ยังอยู่สถานะ {out_cf['state']}", "ยอดยังอยู่ในบัญชีพักเช็คจ่าย"),
            ("Scenario B: เช็คเคลียร์แล้ว", f"เช็ค {out_pd['name']} อยู่สถานะ {out_pd['state']}", "ยอดออกจากบัญชีพักเช็คจ่ายและตัดเข้าบัญชีธนาคารจริงแล้ว"),
        ],
        journal_tables=[
            ("ตัวอย่าง Journal Items ที่ใช้ตรวจเช็คคงค้าง", move_lines_to_rows(out_cf["payment_moves"][0]["lines"], {
                "212001": "Debit เจ้าหนี้การค้า หมายถึงหนี้จากบิลผู้ขายถูกปิดแล้ว",
                "212004": "Credit บัญชีพักเช็คจ่าย หมายถึงเช็คออกแล้วแต่เงินยังไม่ออกจากธนาคารจริง",
            })),
            ("ตัวอย่าง Journal Items หลังเช็คเคลียร์แล้วเพื่อใช้เทียบผลลัพธ์", move_lines_to_rows(out_pd["deposit_move"]["lines"], {
                "212004": "Debit บัญชีพักเช็คจ่ายเพื่อล้างยอด Outstanding",
                "111201": "Credit บัญชีธนาคารจริงเพื่อแสดงว่าเงินถูกตัดออกแล้ว",
            })),
        ],
        cautions=[
            "เมนู Paid Cheque ใช้ดูผลลัพธ์หลังเช็คจบแล้ว ไม่ใช่มุมมองหลักสำหรับ Outstanding",
            "ผู้ใช้ควรเปิด Payment Entry ทุกครั้งเพื่อยืนยันว่าบัญชีพักเช็คจ่ายยังมียอดคงอยู่จริง",
            "หากพบว่าเช็คอยู่สถานะ Confirmed นานผิดปกติ ควรประสานธนาคารและตรวจ Bank Statement ควบคู่กัน",
        ],
    )
    topics_created.append("5.5")

    SUMMARY_PATH.write_text(
        json.dumps(
            {
                "created_topics": topics_created,
                "docx_dir": str(DOCX_DIR),
                "image_dir": str(IMAGE_DIR),
            },
            ensure_ascii=False,
            indent=2,
        ),
        encoding="utf-8",
    )


if __name__ == "__main__":
    main()

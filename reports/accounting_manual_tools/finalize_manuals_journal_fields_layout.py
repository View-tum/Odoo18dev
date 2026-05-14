from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph


DOCX_DIR = Path(r"C:\365_project\TheCool18e\Dev\manual\Accouting_Manual\generated_20260408\docx")
PDF_DIR = Path(r"C:\365_project\TheCool18e\Dev\manual\Accouting_Manual\generated_20260408\pdf_review")
SOFFICE = Path(r"C:\Program Files\LibreOffice\program\soffice.com")
IMAGE_DIR = Path(r"C:\365_project\TheCool18e\Dev\manual\Accouting_Manual\generated_20260408\images")


FIELD_MAP: dict[str, list[str]] = {
    "3.8": ["Customer Group", "Members", "Journal", "Payment Date", "Payment Method", "Amount"],
    "5.1": ["Journal", "Incoming Payments", "Outgoing Payments", "Outstanding Receipts Account", "Outstanding Payments Account"],
    "5.2": ["Template Name", "Bank Name", "Payee Position", "Amount Position", "Date Position"],
    "5.3": ["Cheque Book", "Bank Journal", "Start Number", "End Number", "Next Number", "Status"],
    "5.4": ["Vendor Bill", "Register Payment", "Journal", "Payment Method", "Cheque Number", "Status"],
    "5.5": ["Cheque Number", "Partner", "Journal", "Due Date", "Status", "Journal Items"],
    "5.6": ["Customer Invoice", "Register Payment", "Journal", "Payment Method", "Cheque Number", "Status"],
    "5.7": ["Cheque Number", "Deposit Date", "Clear Date", "Journal", "Status"],
    "5.8": ["Cheque Number", "Status", "Void Reason", "New Cheque", "Reverse Entry"],
    "6.1": ["Asset Name", "Asset Model", "Original Value", "Acquisition Date", "Book Value", "Status"],
    "6.2": ["Name", "Original Value", "Acquisition Date", "Asset Model", "Company", "Status"],
    "6.3": ["Asset Name", "First Depreciation Date", "Method Number", "Method Period", "Journal Entry", "Status"],
    "6.4": ["Asset Name", "Invoice", "Sale Amount", "Gain/Loss Account", "Journal Entry", "Status"],
    "6.5": ["Asset Name", "Disposal Date", "Residual Value", "Loss Account", "Status", "Journal Entry"],
    "6.6": ["Asset Name", "Asset Model", "Original Value", "Book Value", "Acquisition Date", "Status"],
    "7.1": ["Product", "BoM", "Components", "Quantity", "Scheduled Date", "Status"],
    "7.2": ["MO Number", "Product", "Quantity", "Components", "Valuation", "Journal Entry"],
    "7.3": ["Product Category", "Valuation", "Costing Method", "Stock Valuation Account", "Production Account", "Stock Journal"],
    "7.4": ["MO Number", "Product", "Quantity", "Components", "Finished Product", "Journal Entry"],
    "7.5": ["Scrap Product", "Quantity", "Source Location", "Scrap Location", "Reason", "Reference"],
    "7.6": ["Date", "Product", "Quantity", "Total Value", "Reference", "Journal Entry"],
    "7.7": ["Date From", "Date To", "Product", "Quantity", "Unit", "Lot/Batch"],
}


FIELD_MEANINGS: dict[str, str] = {
    "Customer Group": "ชื่อกลุ่มลูกค้าที่ใช้รวมเอกสารเพื่อรับชำระครั้งเดียว",
    "Members": "รายชื่อลูกค้าในกลุ่มที่ระบบจะดึงเอกสารค้างชำระมาให้เลือก",
    "Journal": "สมุดรายวันที่ระบบจะใช้บันทึกรายการบัญชีของเอกสารนี้",
    "Payment Date": "วันที่ที่ใช้บันทึกรายการรับหรือจ่ายเงิน",
    "Payment Method": "วิธีรับหรือจ่ายเงิน เช่น เช็ค โอน หรือเงินสด",
    "Amount": "ยอดเงินของรายการที่กำลังทำ",
    "Incoming Payments": "วิธีรับเงินที่อนุญาตให้ใช้กับ Journal นี้",
    "Outgoing Payments": "วิธีจ่ายเงินที่อนุญาตให้ใช้กับ Journal นี้",
    "Outstanding Receipts Account": "บัญชีพักเงินรับที่ใช้ก่อนเงินเข้าธนาคารจริง",
    "Outstanding Payments Account": "บัญชีพักเงินจ่ายที่ใช้ก่อนเช็คหรือเงินออกจริง",
    "Template Name": "ชื่อแบบฟอร์มเช็คที่ใช้เลือกตอนพิมพ์เช็ค",
    "Bank Name": "ชื่อธนาคารที่จะแสดงบนแบบฟอร์มเช็ค",
    "Payee Position": "ตำแหน่งข้อความชื่อผู้รับเงินบนแบบฟอร์มเช็ค",
    "Amount Position": "ตำแหน่งจำนวนเงินบนแบบฟอร์มเช็ค",
    "Date Position": "ตำแหน่งวันที่บนแบบฟอร์มเช็ค",
    "Cheque Book": "ชื่อสมุดเช็คที่ใช้คุมเลขเช็คในชุดเดียวกัน",
    "Bank Journal": "Journal ธนาคารที่ผูกกับสมุดเช็คหรือรายการเช็คนี้",
    "Start Number": "เลขเช็คใบแรกของสมุดเช็ค",
    "End Number": "เลขเช็คใบสุดท้ายของสมุดเช็ค",
    "Next Number": "เลขเช็คใบถัดไปที่ระบบจะหยิบมาใช้",
    "Status": "สถานะปัจจุบันของเอกสารหรือรายการ",
    "Vendor Bill": "บิลผู้ขายที่ใช้เป็นต้นทางของการจ่ายเงิน",
    "Register Payment": "หน้าต่างที่ใช้กรอกรายละเอียดการรับหรือจ่ายเงินก่อนระบบสร้างรายการจริง",
    "Cheque Number": "เลขที่พิมพ์บนเช็คที่ใช้ติดตามรายการ",
    "Partner": "คู่ค้าที่เกี่ยวข้องกับเอกสารนี้",
    "Due Date": "วันที่ครบกำหนดของเช็คหรือเอกสาร",
    "Deposit Date": "วันที่นำเช็คไปฝากหรือส่งเข้าธนาคาร",
    "Clear Date": "วันที่เช็คผ่านธนาคารหรือเคลียร์รายการสำเร็จ",
    "Journal Items": "รายละเอียดบรรทัดบัญชีเดบิตและเครดิตที่ระบบสร้าง",
    "Void Reason": "เหตุผลที่ใช้ยกเลิกเช็คหรือกลับรายการ",
    "New Cheque": "เช็คใบใหม่ที่ใช้แทนเช็คเดิมหลังการแปลงสถานะ",
    "Reverse Entry": "รายการบัญชีกลับทางที่ระบบสร้างเมื่อยกเลิกเอกสาร",
    "Customer Invoice": "ใบแจ้งหนี้ลูกค้าที่ใช้เป็นต้นทางของการรับเช็ค",
    "Asset Name": "ชื่อทรัพย์สินที่ใช้ค้นหาและติดตามในระบบ",
    "Asset Model": "แบบทรัพย์สินที่กำหนดอายุการใช้งานและบัญชีที่เกี่ยวข้อง",
    "Original Value": "มูลค่าทุนเริ่มต้นของทรัพย์สิน",
    "Acquisition Date": "วันที่ซื้อหรือรับทรัพย์สินเข้าระบบ",
    "Book Value": "มูลค่าคงเหลือทางบัญชีของทรัพย์สิน",
    "Name": "ชื่อรายการหลักของเอกสารหรือแบบฟอร์ม",
    "Quantity": "จำนวนหน่วยของรายการนั้น",
    "Company": "บริษัทเจ้าของรายการในระบบ",
    "First Depreciation Date": "วันที่เริ่มลงค่าเสื่อมงวดแรก",
    "Method Number": "จำนวนงวดที่ระบบจะตัดค่าเสื่อม",
    "Method Period": "ช่วงเวลาระหว่างงวดค่าเสื่อม",
    "Depreciation Board": "ตารางแสดงงวดค่าเสื่อมทั้งหมดของทรัพย์สิน",
    "Invoice": "ใบแจ้งหนี้ที่ใช้เชื่อมโยงกับรายการขายทรัพย์สิน",
    "Sale Amount": "ยอดขายทรัพย์สินที่เกิดขึ้นจริง",
    "Gain/Loss Account": "บัญชีที่ใช้รับรู้กำไรหรือขาดทุนจากการขาย",
    "Disposal Date": "วันที่ตัดจำหน่ายหรือปลดทรัพย์สิน",
    "Residual Value": "มูลค่าคงเหลือก่อนปิดทรัพย์สิน",
    "Loss Account": "บัญชีที่ใช้รับรู้ผลขาดทุนจากการตัดจำหน่าย",
    "Product": "ชื่อสินค้าที่ใช้ในงานผลิตหรือรายงาน",
    "BoM": "สูตรการผลิตที่ระบุส่วนประกอบของสินค้า",
    "Components": "วัตถุดิบหรือชิ้นส่วนที่ใช้ในงานนั้น",
    "Scheduled Date": "วันที่ระบบวางแผนให้เริ่มหรือทำรายการ",
    "MO Number": "เลขที่ใบสั่งผลิตที่ใช้ตามรอยงาน",
    "Valuation": "มูลค่าสินค้าคงคลังที่เกิดจากการเคลื่อนไหวของสินค้า",
    "Product Category": "หมวดสินค้าที่กำหนดวิธีคิดต้นทุนและการลงบัญชี",
    "Costing Method": "วิธีคำนวณต้นทุนสินค้า เช่น FIFO หรือ Average",
    "Stock Valuation Account": "บัญชีสินค้าคงคลังที่รับมูลค่าจากสินค้า",
    "Production Account": "บัญชีงานระหว่างทำที่ใช้สะสมต้นทุนการผลิต",
    "Stock Journal": "สมุดรายวันที่ใช้บันทึกรายการสินค้าคงคลัง",
    "Finished Product": "สินค้าสำเร็จรูปที่ได้จากงานผลิต",
    "Scrap Product": "สินค้าที่ถูกตัดเป็นของเสีย",
    "Source Location": "ตำแหน่งต้นทางที่สินค้าถูกตัดออก",
    "Scrap Location": "ตำแหน่งปลายทางที่เก็บของเสีย",
    "Reason": "เหตุผลหรือคำอธิบายประกอบรายการ",
    "Reference": "เลขอ้างอิงที่ใช้เชื่อมกลับไปยังเอกสารต้นทาง",
    "Date": "วันที่ของรายการหรือรายงาน",
    "Total Value": "มูลค่ารวมของรายการที่แสดงในรายงาน",
    "Date From": "วันเริ่มต้นของช่วงวันที่ที่ใช้กรองรายงาน",
    "Date To": "วันสิ้นสุดของช่วงวันที่ที่ใช้กรองรายงาน",
    "Unit": "หน่วยนับที่ใช้แสดงปริมาณสินค้า",
    "Lot/Batch": "เลขล็อตหรือแบตช์ที่ใช้ติดตามสินค้า",
    "Journal Entry": "รายการบัญชีที่ระบบสร้างขึ้นเพื่อบันทึกผลทางการเงินของเหตุการณ์นั้น",
}


JOURNAL_GUIDE_MAP: dict[str, list[str]] = {
    "3.8": [
        "Journal คือสมุดรายวันที่ระบบใช้บันทึกรายการรับชำระของเอกสารนี้",
        "ให้ดูชื่อ Journal เลขอ้างอิงเอกสาร และยอดเดบิตเครดิตก่อน เพื่อยืนยันว่าเป็นรายการของกลุ่มลูกค้าที่กำลังตรวจจริง",
        "กรณีรับชำระสำเร็จ ระบบจะเดบิตบัญชีเงินรับหรือบัญชีพักรับชำระ และเครดิตบัญชีลูกหนี้ตามใบแจ้งหนี้ที่ถูกเลือก",
    ],
    "5.1": [
        "หัวข้อนี้เป็นการตั้งค่า จึงยังไม่เกิด Journal Entry ทันที แต่ค่าที่ตั้งใน Journal จะถูกนำไปใช้ตอนรับหรือจ่ายเงินจริง",
        "ถ้าตั้ง Outstanding account ถูกต้อง รายการเช็คจะลงบัญชีพักก่อน แล้วค่อยย้ายเมื่อเช็คผ่านธนาคาร",
    ],
    "5.2": [
        "หัวข้อนี้เป็นการตั้งค่าแบบฟอร์มเช็ค จึงยังไม่สร้างรายการบัญชี",
        "ผลของการตั้งค่านี้จะมีผลเฉพาะรูปแบบการพิมพ์เช็ค ไม่ได้เปลี่ยนเดบิตหรือเครดิตของระบบ",
    ],
    "5.3": [
        "การสร้างสมุดเช็คยังไม่สร้าง Journal Entry เพราะเป็นการกำหนดเลขเช็คที่จะใช้ในอนาคต",
        "เมื่อมีการนำเลขเช็คจากสมุดนี้ไปจ่ายเงินจริง ระบบจึงจะสร้างรายการบัญชีตาม Journal ที่ผูกไว้",
    ],
    "5.4": [
        "Journal คือสมุดรายวันที่ใช้ลงบัญชีของเช็คจ่ายใบนี้ เช่น Bank Journal หรือ Cheque Journal",
        "ตอนเช็คยังไม่ผ่านธนาคาร ระบบมักเดบิตเจ้าหนี้และเครดิตบัญชีพักเช็คจ่ายก่อน เพื่อให้รู้ว่ายังเป็นเช็คคงค้าง",
        "เมื่อเช็คผ่านธนาคาร ระบบจะย้ายยอดจากบัญชีพักเช็คจ่ายไปยังบัญชีธนาคารจริง",
        "หลังสร้างเช็คจ่ายแล้ว ให้เปิด Payment Entry และ Journal Items เพื่อตรวจว่าระบบลงบัญชีพักเช็คจ่ายอย่างไร และหลังเช็คผ่านธนาคารแล้วบัญชีใดถูกย้ายออกไป",
    ],
    "5.5": [
        "ให้ดูชื่อ Journal และสถานะของเช็คก่อนว่าเป็น Confirmed หรือ Paid เพราะความหมายทางบัญชีต่างกัน",
        "สถานะ Confirmed แปลว่ายังอยู่ในบัญชีพักเช็คจ่าย ส่วนสถานะ Paid แปลว่ายอดถูกย้ายไปธนาคารแล้ว",
        "เมื่อเปิด Journal Items ให้ดูว่าชื่อบัญชีและยอดตรงกับสถานะของเช็คที่กำลังติดตามอยู่",
    ],
    "5.6": [
        "สำหรับเช็ครับ ระบบจะใช้ Journal ของการรับเงินเพื่อบันทึกรายการเข้าบัญชีพักเช็ครับก่อน",
        "เมื่อรับเช็คสำเร็จ ระบบจะเดบิตบัญชีพักเช็ครับหรือบัญชีเงินรับ และเครดิตลูกหนี้ของใบแจ้งหนี้ลูกค้า",
        "ถ้าเช็คยังไม่ผ่านธนาคาร ยอดจะยังไม่เข้าบัญชีธนาคารจริงจนกว่าจะมีการเคลียร์เช็ค",
    ],
    "5.7": [
        "ขั้นตอนเคลียร์เช็คคือการย้ายยอดจากบัญชีพักเช็ครับหรือเช็คจ่าย ไปยังบัญชีธนาคารจริง",
        "ให้เทียบเลขเช็ค วันที่เคลียร์ และชื่อบัญชีใน Journal Items ว่าตรงกับเช็คใบที่ต้องการปิดสถานะ",
        "ถ้าดูถูกใบ จะเห็นว่าบัญชีพักลดลงและบัญชีธนาคารเพิ่มหรือลดตามทิศทางของรายการ",
    ],
    "5.8": [
        "เมื่อยกเลิกหรือแปลงสถานะเช็ค ระบบจะสร้าง Reverse Entry เพื่อกลับผลทางบัญชีของเช็คเดิม",
        "ให้ดูเลขอ้างอิงของรายการกลับทางและชื่อบัญชีว่ากลับเดบิตเครดิตตรงข้ามกับรายการเดิมหรือไม่",
        "ถ้ามีการออกเช็คใบใหม่แทน ระบบจะมีทั้งรายการกลับทางของใบเดิมและรายการใหม่ของใบแทน",
    ],
    "6.1": [
        "Journal ของทรัพย์สินจะเริ่มเห็นชัดเมื่อมีการเริ่มใช้งาน ขาย ตัดจำหน่าย หรือบันทึกค่าเสื่อม",
        "ให้ดูว่าบัญชีทรัพย์สิน บัญชีค่าเสื่อมสะสม และบัญชีกำไรขาดทุนถูกใช้ตรงตามเหตุการณ์ของทรัพย์สินหรือไม่",
    ],
    "6.2": [
        "การสร้างทรัพย์สินสถานะ Draft ยังไม่สร้าง Journal Entry ทันที เพราะยังเป็นเพียงการตั้งต้นข้อมูล",
        "บัญชีที่จะถูกใช้จริงถูกกำหนดจาก Asset Model ที่เลือก จึงควรเปิดดูแบบทรัพย์สินก่อนบันทึกทุกครั้ง",
    ],
    "6.3": [
        "เมื่อเริ่มคิดค่าเสื่อม ระบบจะใช้ Journal ของค่าเสื่อมสร้างรายการรายงวดให้อัตโนมัติ",
        "โดยทั่วไปจะเดบิตค่าเสื่อมราคาและเครดิตค่าเสื่อมราคาสะสมตามบัญชีที่อยู่ใน Asset Model",
        "ให้เทียบวันที่งวด จำนวนเงิน และชื่อทรัพย์สินใน Journal Entry กับ Depreciation Board",
    ],
    "6.4": [
        "ตอนขายทรัพย์สิน ระบบจะปิดมูลค่าคงเหลือของทรัพย์สิน ตัดค่าเสื่อมสะสม และรับรู้กำไรหรือขาดทุนจากการขาย",
        "ให้เทียบยอดขายจริงกับ Book Value เพื่อดูว่าระบบลงกำไรหรือขาดทุนไว้ในบัญชีใด",
        "บรรทัดสำคัญคือบัญชีลูกหนี้หรือเงินรับ บัญชีทรัพย์สิน บัญชีค่าเสื่อมสะสม และบัญชีกำไรขาดทุนจากการขาย",
    ],
    "6.5": [
        "การตัดจำหน่ายทรัพย์สินจะปิดสินทรัพย์ออกจากงบและรับรู้ผลต่างที่เหลืออยู่",
        "ให้ดูว่าระบบย้ายยอดออกจากบัญชีทรัพย์สินและค่าเสื่อมสะสมครบหรือไม่ และผลต่างถูกปิดที่ Loss Account ถูกต้องหรือไม่",
    ],
    "6.6": [
        "รายงานทรัพย์สินถาวรเป็นหน้าสรุป จึงไม่ได้สร้าง Journal Entry ใหม่",
        "ให้ใช้รายงานนี้เทียบกับยอดใน Journal Entry และ Depreciation Board เพื่อดูว่ามูลค่าคงเหลือสอดคล้องกันหรือไม่",
    ],
    "7.1": [
        "หัวข้อนี้อธิบายภาพรวมของการไหลข้อมูล จึงใช้เพื่อดูว่ารายการใดจะไปสร้าง Journal ในขั้นไหน",
        "จุดที่มักเกิด Journal คือการตัดวัตถุดิบเข้า WIP การรับผลผลิตเข้า FG และการบันทึกของเสียหรือปรับมูลค่า",
    ],
    "7.2": [
        "Journal ของการผลิตใช้บอกว่าต้นทุนกำลังย้ายจากวัตถุดิบไปงานระหว่างทำ และจากงานระหว่างทำไปเป็นสินค้าสำเร็จรูป",
        "อ่านง่ายที่สุดโดยดูจากเลขอ้างอิง MO เดียวกัน แล้วไล่ตามลำดับการตัดวัตถุดิบ รับสินค้า และปรับมูลค่า",
        "บรรทัดเดบิตคือบัญชีที่รับมูลค่าเข้า ส่วนบรรทัดเครดิตคือบัญชีที่ปล่อยต้นทุนออกจากขั้นก่อนหน้า",
    ],
    "7.3": [
        "หัวข้อนี้เป็นการสรุปผังการตั้งค่าทางบัญชี จึงยังไม่สร้าง Journal ใหม่ในตัวเอง",
        "แต่ทุกบัญชีที่ตั้งใน Product Category และ Stock Journal จะถูกใช้จริงตอนผลิต รับเข้า โอน และปรับมูลค่า",
    ],
    "7.4": [
        "หัวข้อนี้ใช้ตัวอย่างจริงเพื่อให้เห็นว่า MO หนึ่งใบสร้าง Journal หลายช่วงได้อย่างไร",
        "ให้ดูเลข MO เดียวกันในทุก Journal แล้วเทียบว่าวัตถุดิบถูกตัดเมื่อไร งานระหว่างทำเพิ่มเมื่อไร และ FG เข้าสินค้าเมื่อไร",
    ],
    "7.5": [
        "ของเสียที่บันทึกในงานผลิตอาจกระทบทั้งปริมาณและมูลค่าต้นทุน ขึ้นกับวิธีตั้งค่าบัญชีและการปรับมูลค่า",
        "ให้ตรวจ Journal ของ Scrap หรือ Landed Cost ที่เกี่ยวข้อง ว่าต้นทุนถูกย้ายออกจากวัตถุดิบหรือบวกกลับเข้า FG อย่างไร",
    ],
    "7.6": [
        "Stock Valuation Report เป็นรายงานสรุปมูลค่าสินค้าคงคลังจากรายการที่บันทึกไว้แล้ว จึงไม่ได้สร้าง Journal ใหม่",
        "ให้ใช้รายงานนี้เทียบกับ Journal Entry และ Stock Moves เพื่อดูว่าปริมาณและมูลค่าไปทางเดียวกันหรือไม่",
    ],
    "7.7": [
        "รายงาน รง.8 เป็นรายงานสรุปเพื่อหน่วยงานภายนอก จึงไม่ได้สร้าง Journal Entry ใหม่",
        "เมื่อตรวจรายงานนี้ ควรเทียบกับ Stock Moves และ Journal ของงานผลิตเพื่อให้แน่ใจว่าปริมาณและต้นทุนสอดคล้องกัน",
    ],
}

JOURNAL_IMAGE_MAP: dict[str, list[tuple[str, str]]] = {
    "3.8": [("journal_group_payment_real_annotated.png", "ภาพ Journal Entry ของการรับชำระแบบกลุ่มบริษัท พร้อมคำอธิบาย Dr/Cr")],
    "5.4": [("journal_cheque_out_confirmed_real_annotated.png", "ภาพ Journal Entry ของเช็คจ่ายสถานะ Confirmed พร้อมคำอธิบาย Dr/Cr")],
    "5.5": [
        ("journal_cheque_out_confirmed_real_annotated.png", "ภาพ Journal Entry ของเช็คคงค้าง พร้อมคำอธิบาย Dr/Cr"),
        ("journal_cheque_out_paid_real_annotated.png", "ภาพ Journal Entry หลังเช็คผ่านธนาคาร พร้อมคำอธิบาย Dr/Cr"),
    ],
    "5.6": [("journal_cheque_in_confirmed_real_annotated.png", "ภาพ Journal Entry ของการรับเช็คจากลูกค้า พร้อมคำอธิบาย Dr/Cr")],
    "5.7": [
        ("journal_cheque_in_paid_real_annotated.png", "ภาพ Journal Entry หลังเคลียร์เช็ครับ พร้อมคำอธิบาย Dr/Cr"),
        ("journal_cheque_out_paid_real_annotated.png", "ภาพ Journal Entry หลังเคลียร์เช็คจ่าย พร้อมคำอธิบาย Dr/Cr"),
    ],
    "5.8": [("journal_cheque_void_reverse_real_annotated.png", "ภาพ Journal Entry ของการยกเลิกหรือกลับรายการเช็ค พร้อมคำอธิบาย Dr/Cr")],
    "6.2": [("journal_asset_depreciation_real_annotated.png", "ภาพ Journal Entry ตัวอย่างที่เกี่ยวข้องกับทรัพย์สิน พร้อมคำอธิบาย Dr/Cr")],
    "6.3": [("journal_asset_depreciation_real_annotated.png", "ภาพ Journal Entry ค่าเสื่อมราคาจริง พร้อมคำอธิบาย Dr/Cr")],
    "6.4": [("journal_asset_sale_real_annotated.png", "ภาพ Journal Entry ของการขายทรัพย์สิน พร้อมคำอธิบาย Dr/Cr")],
    "6.5": [("journal_asset_disposal_real_annotated.png", "ภาพ Journal Entry ของการตัดจำหน่ายทรัพย์สิน พร้อมคำอธิบาย Dr/Cr")],
    "7.2": [
        ("journal_mfg_raw_fg02001_real_annotated.png", "ภาพ Journal Entry ตอนตัดวัตถุดิบเข้าสู่งานระหว่างทำ พร้อมคำอธิบาย Dr/Cr"),
        ("journal_mfg_finished_real_annotated.png", "ภาพ Journal Entry ตอนรับสินค้าสำเร็จรูปเข้าคลัง พร้อมคำอธิบาย Dr/Cr"),
    ],
    "7.3": [("journal_mfg_finished_real_annotated.png", "ภาพ Journal Entry ตัวอย่างที่ผูกกับการตั้งค่าบัญชีการผลิต พร้อมคำอธิบาย Dr/Cr")],
    "7.4": [
        ("journal_mfg_raw_fg02001_real_annotated.png", "ภาพ Journal Entry ตัดวัตถุดิบของตัวอย่างจริง พร้อมคำอธิบาย Dr/Cr"),
        ("journal_mfg_finished_real_annotated.png", "ภาพ Journal Entry รับ FG ของตัวอย่างจริง พร้อมคำอธิบาย Dr/Cr"),
    ],
    "7.5": [("journal_mfg_raw_packaging_real_annotated.png", "ภาพ Journal Entry ตัวอย่างของของเสียและความสูญเสีย พร้อมคำอธิบาย Dr/Cr")],
    "7.6": [
        ("journal_mfg_raw_fg02001_real_annotated.png", "ภาพ Journal Entry ฝั่งวัตถุดิบจากรายงาน Valuation พร้อมคำอธิบาย Dr/Cr"),
        ("journal_mfg_finished_real_annotated.png", "ภาพ Journal Entry ฝั่งรับ FG จากรายงาน Valuation พร้อมคำอธิบาย Dr/Cr"),
    ],
    "7.7": [("journal_mfg_finished_real_annotated.png", "ภาพ Journal Entry ที่ใช้เทียบกับรายงาน รง.8 พร้อมคำอธิบาย Dr/Cr")],
}


def set_keep_with_next(paragraph: Paragraph, value: bool = True) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    keep = ppr.find(qn("w:keepNext"))
    if value and keep is None:
        keep = OxmlElement("w:keepNext")
        ppr.append(keep)
    elif not value and keep is not None:
        ppr.remove(keep)


def set_keep_together(paragraph: Paragraph, value: bool = True) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    keep = ppr.find(qn("w:keepLines"))
    if value and keep is None:
        keep = OxmlElement("w:keepLines")
        ppr.append(keep)
    elif not value and keep is not None:
        ppr.remove(keep)


def paragraph_after(paragraph: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    created = Paragraph(new_p, paragraph._parent)
    if text:
        created.add_run(text)
    if style:
        created.style = style
    return created


def picture_after(paragraph: Paragraph, image_path: Path, caption: str) -> Paragraph:
    pic_para = paragraph_after(paragraph)
    run = pic_para.add_run()
    run.add_picture(str(image_path), width=Inches(6.2))
    cap = paragraph_after(pic_para, caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return cap


def remove_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def remove_table(table) -> None:
    element = table._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def iter_blocks(doc: Document):
    body = doc._body._element
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield "p", Paragraph(child, doc._body)
        elif child.tag == qn("w:tbl"):
            yield "t", child


def code_from_name(name: str) -> str:
    return name.split("_", 1)[0]


def find_paragraph(doc: Document, needle: str) -> Paragraph | None:
    for p in doc.paragraphs:
        if p.text.strip() == needle:
            return p
    return None


def cleanup_duplicate_journal_section(doc: Document) -> None:
    guide_indexes = [i for i, p in enumerate(doc.paragraphs) if p.text.strip() == "วิธีอ่าน Journal ในหัวข้อนี้"]
    for idx in reversed(guide_indexes[1:]):
        remove_paragraph(doc.paragraphs[idx])
    heading_indexes = [i for i, p in enumerate(doc.paragraphs) if p.text.strip() == "4. ตรวจสอบ Journal Entry"]
    for idx in reversed(heading_indexes[1:]):
        remove_paragraph(doc.paragraphs[idx])


def ensure_field_table(doc: Document, code: str) -> None:
    fields = FIELD_MAP.get(code, [])
    if not fields:
        return
    tables = list(doc.tables)
    field_table = None
    for table in tables:
        if table.rows and len(table.rows[0].cells) >= 2:
            hdr = [c.text.strip() for c in table.rows[0].cells[:2]]
            if hdr == ["Field Name", "Meaning"]:
                field_table = table
            elif hdr[0] == "Field Name":
                remove_table(table)
    if field_table is None:
        anchor = find_paragraph(doc, "4. ตรวจสอบ Journal Entry") or find_paragraph(doc, "3. ขั้นตอนการใช้งาน")
        if anchor is None:
            return
        tbl = anchor._parent.add_table(rows=1, cols=2, width=0)
        anchor._p.addprevious(tbl._element)
        field_table = tbl
    while len(field_table.rows) > 1:
        field_table._tbl.remove(field_table.rows[-1]._tr)
    field_table.style = "Table Grid"
    field_table.cell(0, 0).text = "Field Name"
    field_table.cell(0, 1).text = "Meaning"
    for field in fields:
        cells = field_table.add_row().cells
        cells[0].text = field
        cells[1].text = FIELD_MEANINGS.get(field, f"ความหมายของช่อง {field}")
    for row in field_table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.keep_together = True


def ensure_journal_guide(doc: Document, code: str) -> None:
    guide_lines = JOURNAL_GUIDE_MAP.get(code, [])
    if not guide_lines:
        return
    heading = find_paragraph(doc, "4. ตรวจสอบ Journal Entry")
    if heading is None:
        heading = doc.add_paragraph("4. ตรวจสอบ Journal Entry")
    cleanup_duplicate_journal_section(doc)
    existing = find_paragraph(doc, "วิธีอ่าน Journal ในหัวข้อนี้")
    if existing is not None:
        idx = next(
            (i for i, p in enumerate(doc.paragraphs) if p._element is existing._element),
            -1,
        )
        if idx < 0:
            idx = 0
        # remove old guide lines until next heading
        following = []
        for p in doc.paragraphs[idx + 1:]:
            txt = p.text.strip()
            if txt.startswith("5.") or txt.startswith("4.") or txt.startswith("ตาราง") or txt == "Field Name":
                break
            if txt:
                following.append(p)
        for p in reversed(following):
            remove_paragraph(p)
        cursor = existing
    else:
        cursor = paragraph_after(heading, "วิธีอ่าน Journal ในหัวข้อนี้")
        cursor.style = "Heading 3"
    set_keep_with_next(heading, True)
    set_keep_with_next(cursor, True)
    prev = cursor
    for line in guide_lines:
        p = paragraph_after(prev, line)
        p.style = "Normal"
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.space_before = Pt(0)
        set_keep_with_next(p, True)
        prev = p


def ensure_journal_images(doc: Document, code: str) -> None:
    image_specs = JOURNAL_IMAGE_MAP.get(code, [])
    if not image_specs:
        return
    heading = find_paragraph(doc, "4. ตรวจสอบ Journal Entry")
    if heading is None:
        return
    guide = find_paragraph(doc, "วิธีอ่าน Journal ในหัวข้อนี้")
    anchor = guide or heading

    # move anchor to last guide paragraph if present
    if guide is not None:
        guide_idx = next((i for i, p in enumerate(doc.paragraphs) if p._element is guide._element), -1)
        started = False
        for p in doc.paragraphs[guide_idx + 1 :] if guide_idx >= 0 else []:
            txt = p.text.strip()
            if not started:
                started = True
            if txt.startswith("ภาพ Journal Entry") or txt.startswith("ภาพ Journal"):
                continue
            if txt.startswith("Scenario") or txt.startswith("บัญชี") or txt == "Field Name" or txt.startswith("5.") or txt.startswith("รูป "):
                break
            if txt:
                anchor = p

    # remove old inserted JE images/captions immediately after guide block until first table/next section
    cleanup = []
    start_idx = next((i for i, p in enumerate(doc.paragraphs) if p._element is anchor._element), -1)
    if start_idx >= 0:
        for p in doc.paragraphs[start_idx + 1 :]:
            txt = p.text.strip()
            has_pic = bool(p._element.xpath(".//pic:pic"))
            if txt.startswith("Scenario") or txt.startswith("บัญชี") or txt == "Field Name" or txt.startswith("5.") or txt.startswith("ตาราง"):
                break
            if has_pic or txt.startswith("ภาพ Journal Entry") or txt.startswith("ภาพ Journal"):
                cleanup.append(p)
        for p in reversed(cleanup):
            remove_paragraph(p)

    prev = anchor
    for image_name, caption in image_specs:
        image_path = IMAGE_DIR / image_name
        if image_path.exists():
            prev = picture_after(prev, image_path, caption)


def remove_redundant_step_lines(doc: Document) -> None:
    redundant_starts = (
        "เมื่อล็อกอินเข้าสู่ระบบแล้ว",
        "เมื่อเข้ามาในโมดูลแล้ว",
        "ตรวจหน้าจอ",
    )
    for p in list(doc.paragraphs):
        txt = p.text.strip()
        if any(txt.startswith(prefix) for prefix in redundant_starts):
            remove_paragraph(p)
        if txt in {"ภาพประกอบของขั้นตอนนี้", "คำอธิบายหน้าจอ", "ตัวอย่างการใช้งานในหัวข้อนี้"}:
            remove_paragraph(p)


def apply_layout(doc: Document) -> None:
    for i, p in enumerate(doc.paragraphs):
        txt = p.text.strip()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(6)
        if txt.startswith(("1.", "2.", "3.", "4.", "5.", "6.", "7.", "8.", "9.")):
            set_keep_with_next(p, True)
            set_keep_together(p, True)
        if txt in {"3. ขั้นตอนการใช้งาน", "4. ตรวจสอบ Journal Entry", "วิธีอ่าน Journal ในหัวข้อนี้"}:
            set_keep_with_next(p, True)
            set_keep_together(p, True)
        if txt.startswith("ภาพ"):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_keep_with_next(p, True)
        if i + 1 < len(doc.paragraphs):
            nxt = doc.paragraphs[i + 1]
            if txt.startswith(tuple(f"{n}." for n in range(1, 10))) and nxt._element.xpath(".//pic:pic"):
                set_keep_with_next(p, True)
                set_keep_together(nxt, True)
    for table in doc.tables:
        for row in table.rows:
            tr_pr = row._tr.get_or_add_trPr()
            cant_split = tr_pr.find(qn("w:cantSplit"))
            if cant_split is None:
                tr_pr.append(OxmlElement("w:cantSplit"))


def normalize_one(path: Path) -> None:
    doc = Document(path)
    code = code_from_name(path.name)
    remove_redundant_step_lines(doc)
    ensure_field_table(doc, code)
    ensure_journal_guide(doc, code)
    ensure_journal_images(doc, code)
    apply_layout(doc)
    doc.save(path)


def render_all() -> None:
    PDF_DIR.mkdir(parents=True, exist_ok=True)
    import subprocess

    for path in sorted(DOCX_DIR.glob("*.docx")):
        subprocess.run(
            [str(SOFFICE), "--headless", "--convert-to", "pdf", "--outdir", str(PDF_DIR), str(path)],
            check=True,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
        )


def main() -> None:
    for path in sorted(DOCX_DIR.glob("*.docx")):
        normalize_one(path)
    render_all()
    print(f"finalized {len(list(DOCX_DIR.glob('*.docx')))} docs")


if __name__ == "__main__":
    main()

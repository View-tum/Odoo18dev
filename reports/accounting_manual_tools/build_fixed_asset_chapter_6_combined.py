from __future__ import annotations

import json
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\365_project\TheCool18e\Dev")
TOOLS_DIR = ROOT / "reports" / "accounting_manual_tools"
OUTPUT_DIR = ROOT / "manual" / "Accouting_Manual" / "generated_20260408"
DOCX_DIR = OUTPUT_DIR / "docx"
PDF_DIR = OUTPUT_DIR / "pdf_review"
IMAGE_DIR = OUTPUT_DIR / "images"
SAMPLES_PATH = TOOLS_DIR / "output" / "fixed_asset_ch6_live_samples_20260410.json"
TRANSACTIONS_PATH = TOOLS_DIR / "output" / "fixed_asset_ch6_transactions_20260410.json"
DEP_MOVE_PATH = TOOLS_DIR / "output" / "fixed_asset_dep_move_lines_20260409.json"
IMAGE_MAP_PATH = TOOLS_DIR / "output" / "fixed_asset_ch6_image_map_20260410.json"
SOFFICE = Path(r"C:\Program Files\LibreOffice\program\soffice.com")

DOCX_PATH = DOCX_DIR / "6_Fixed Asset คู่มือรวม.docx"
PDF_PATH = PDF_DIR / "6_Fixed Asset คู่มือรวม.pdf"

FONT_NAME = "Angsana New"
BASE_SIZE = 16


IMAGE_CALLOUTS: dict[str, list[tuple[str, str]]] = {
    "ch6_nav_dashboard_accounting_annotated.png": [
        ("1", "กดการ์ด Accounting จากหน้า Dashboard เพื่อเริ่มงานด้านบัญชีและทรัพย์สิน"),
    ],
    "ch6_nav_dashboard_accounting_annotated_crop.png": [
        ("1", "กดการ์ด Accounting จากหน้า Dashboard เพื่อเริ่มงานด้านบัญชีและทรัพย์สิน"),
    ],
    "ch6_nav_assets_annotated_crop.png": [
        ("1", "ส่วนควบคุมด้านบน ใช้ดูตัวกรอง ค้นหา หรือเปิดมุมมองอื่นของรายการทรัพย์สิน"),
        ("2", "ปุ่ม New ใช้สร้างทรัพย์สินรายการใหม่"),
        ("3", "ตารางรายการทรัพย์สิน ใช้เปิดดูรายละเอียดทรัพย์สินแต่ละใบ"),
    ],
    "ch6_nav_asset_models_annotated_crop.png": [
        ("1", "ส่วนควบคุมด้านบนของหน้าต้นแบบทรัพย์สิน"),
        ("2", "ปุ่ม New ใช้สร้างต้นแบบทรัพย์สินใหม่"),
        ("3", "รายการต้นแบบทรัพย์สิน ใช้เลือกต้นแบบที่ต้องการนำไปใช้"),
    ],
    "ch6_asset_model_form_annotated_crop.png": [
        ("1", "วิธีคิดค่าเสื่อมราคา เช่น แบบเส้นตรง"),
        ("2", "บัญชีสินทรัพย์ถาวรที่ใช้รับมูลค่าทรัพย์สิน"),
        ("3", "บัญชีค่าเสื่อมราคาสะสม"),
        ("4", "บัญชีค่าใช้จ่ายค่าเสื่อมราคา"),
        ("5", "สมุดรายวันที่ใช้บันทึกรายการค่าเสื่อมราคา"),
    ],
    "ch6_asset_draft_form_annotated_crop.png": [
        ("1", "ปุ่ม Confirm ใช้เปลี่ยนทรัพย์สินจากฉบับร่างไปเป็นทรัพย์สินที่เริ่มใช้งานจริง"),
        ("2", "ต้นแบบทรัพย์สินที่เลือกใช้"),
        ("3", "มูลค่าทุนของทรัพย์สิน"),
        ("4", "วันที่เริ่มรับรู้ทรัพย์สิน"),
        ("5", "สมุดรายวันที่เกี่ยวข้องกับรายการนี้"),
    ],
    "ch6_asset_open_form_annotated_crop.png": [
        ("1", "ปุ่ม Modify Depreciation ใช้ปรับแผนค่าเสื่อมราคาเมื่อมีการเปลี่ยนแปลง"),
        ("2", "ปุ่ม Open Entries ใช้เปิดดูรายการบัญชีของทรัพย์สินใบนี้"),
        ("3", "มูลค่าคงเหลือสุทธิของทรัพย์สิน ณ วันที่ดูข้อมูล"),
        ("4", "มูลค่าทุนตั้งต้นของทรัพย์สิน"),
        ("5", "บัญชีค่าใช้จ่ายค่าเสื่อมราคาที่ระบบจะใช้ตอนตัดค่าเสื่อม"),
    ],
    "ch6_nav_fixed_asset_report_annotated_crop.png": [
        ("2", "พื้นที่รายงานทรัพย์สินถาวร ใช้ดูยอดรวมและยอดคงเหลือของทรัพย์สิน"),
    ],
    "ch6_asset_sell_form_annotated_crop.png": [
        ("1", "แถบสถานะของทรัพย์สิน หลังขายแล้วสถานะจะปิดรายการ"),
        ("3", "มูลค่าคงเหลือของทรัพย์สินก่อนปิดรายการ"),
        ("4", "รายละเอียดทรัพย์สิน ใช้ตรวจสอบข้อมูลก่อนและหลังการขาย"),
    ],
    "ch6_asset_dispose_form_annotated_crop.png": [
        ("1", "แถบสถานะของทรัพย์สินหลังตัดจำหน่าย"),
        ("3", "มูลค่าคงเหลือก่อนปิดรายการ"),
        ("4", "รายละเอียดทรัพย์สินที่ใช้ตรวจสอบก่อนตัดจำหน่าย"),
    ],
    "ch6_asset_sale_invoice_annotated_crop.png": [
        ("1", "สถานะของใบขายสินทรัพย์"),
        ("2", "ชื่อลูกค้าที่ซื้อสินทรัพย์"),
        ("3", "รายการสินค้าที่ขายในเอกสารนี้"),
        ("4", "ยอดรวมของใบขาย"),
    ],
    "ch6_journal_depreciation_annotated_crop.png": [
        ("1", "หัวเอกสารบัญชี เช่น ชื่อสมุดรายวันและวันที่บันทึก"),
        ("2", "แท็บ Journal Items ใช้เปิดดูบรรทัดบัญชีเดบิตและเครดิต"),
        ("3", "ตารางบรรทัดบัญชีของรายการค่าเสื่อม"),
    ],
    "ch6_journal_sale_annotated_crop.png": [
        ("1", "หัวเอกสารบัญชีจากการขายทรัพย์สิน"),
        ("2", "แท็บ Journal Items ของรายการนี้"),
        ("3", "ตารางเดบิตและเครดิตที่ระบบลงบัญชีจริง"),
    ],
    "ch6_journal_disposal_annotated_crop.png": [
        ("1", "หัวเอกสารบัญชีจากการตัดจำหน่ายทรัพย์สิน"),
        ("2", "แท็บ Journal Items"),
        ("3", "ตารางเดบิตและเครดิตของการตัดทรัพย์สินออกจากระบบ"),
    ],
}


def load_json(path: Path) -> dict:
    text = path.read_text(encoding="utf-8")
    fixed = text.replace("\xa0", " ")
    candidates = [text, fixed]
    try:
        candidates.append(fixed.encode("latin1").decode("utf-8"))
    except Exception:
        pass
    for candidate in candidates:
        try:
            return json.loads(candidate)
        except Exception:
            continue
    raise ValueError(f"Unable to parse {path}")


def set_font(run, *, size=BASE_SIZE, bold=False, italic=False, color: RGBColor | None = None):
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def paragraph_keep(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    keep_lines = OxmlElement("w:keepLines")
    p_pr.append(keep_lines)


def paragraph_keep_with_next(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    keep = OxmlElement("w:keepNext")
    p_pr.append(keep)


def add_para(doc: Document, text: str, *, size=BASE_SIZE, bold=False, italic=False, align=None, space_after=6):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    set_font(r, size=size, bold=bold, italic=italic)
    paragraph_keep(p)
    return p


def add_heading(doc: Document, text: str, *, level=1):
    size = 20 if level == 1 else 18
    color = RGBColor(31, 78, 121) if level == 1 else RGBColor(0, 0, 0)
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_font(r, size=size, bold=True, color=color)
    p.paragraph_format.space_after = Pt(4)
    paragraph_keep(p)
    paragraph_keep_with_next(p)
    return p


def add_bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    set_font(r, size=BASE_SIZE)
    paragraph_keep(p)


def add_numbered_step(doc: Document, idx: int, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(f"{idx}. {text}")
    set_font(r, size=BASE_SIZE)
    paragraph_keep(p)
    return p


def set_cell(cell, text: str, *, bold=False):
    cell.text = str(text)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for p in cell.paragraphs:
        p.paragraph_format.space_after = Pt(2)
        for r in p.runs:
            set_font(r, size=BASE_SIZE, bold=bold)


def shade_cell(cell, color="D9EAF7"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    tc_pr.append(shd)


def add_table(doc: Document, headers: list[str], rows: list[tuple]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        set_cell(hdr[idx], header, bold=True)
        shade_cell(hdr[idx])
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell(cells[idx], value)
    doc.add_paragraph()
    return table


def add_image(doc: Document, image_name: str, caption: str, *, width=6.75):
    path = IMAGE_DIR / image_name
    if not path.exists():
        add_para(doc, f"[ไม่พบภาพประกอบ: {image_name}]", italic=True)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    p.paragraph_format.space_after = Pt(3)
    paragraph_keep(p)
    cap = add_para(doc, caption, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=15, space_after=6)
    paragraph_keep(cap)
    callouts = IMAGE_CALLOUTS.get(image_name, [])
    if callouts:
        title = add_para(doc, "คำอธิบายกรอบสี่เหลี่ยมในภาพ", bold=True, space_after=2)
        paragraph_keep_with_next(title)
        for label, text in callouts:
            add_para(doc, f"{label}. {text}", space_after=2)


def configure_doc(doc: Document):
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    style.font.size = Pt(BASE_SIZE)
    sec = doc.sections[0]
    sec.start_type = WD_SECTION_START.NEW_PAGE
    sec.top_margin = Inches(0.65)
    sec.bottom_margin = Inches(0.65)
    sec.left_margin = Inches(0.75)
    sec.right_margin = Inches(0.75)


def move_rows(lines: list[dict], meaning_by_code: dict[str, str]) -> list[tuple[str, str, str, str]]:
    rows: list[tuple[str, str, str, str]] = []
    for line in lines:
        rows.append(
            (
                f"{line['account_code']} {line['account_name']}",
                meaning_by_code.get(line["account_code"], line.get("label", "")),
                f"{line['debit']:,.2f}",
                f"{line['credit']:,.2f}",
            )
        )
    return rows


def build_subsection_header(doc: Document, title: str, purpose: str, overview: str, functions: list[str], menu_path: str):
    add_heading(doc, title, level=1)
    add_heading(doc, "วัตถุประสงค์", level=2)
    add_para(doc, purpose)
    add_heading(doc, "ภาพรวมการทำงาน", level=2)
    add_para(doc, overview)
    add_heading(doc, "ฟังก์ชันที่ใช้ในหัวข้อนี้", level=2)
    for item in functions:
        add_bullet(doc, item)
    add_heading(doc, "เมนูที่ใช้", level=2)
    add_para(doc, menu_path)


def build_steps_with_images(doc: Document, steps: list[tuple[str, str | None, str | None]]):
    add_heading(doc, "ขั้นตอนการใช้งานแบบละเอียด", level=2)
    for idx, (step_text, image_name, caption) in enumerate(steps, start=1):
        add_numbered_step(doc, idx, step_text)
        if image_name and caption:
            add_image(doc, image_name, caption)


def build_fields(doc: Document, rows: list[tuple[str, str]]):
    add_heading(doc, "ช่องสำคัญที่ควรดูบนหน้าจอ", level=2)
    add_table(doc, ["Field Name", "Meaning"], rows)


def build_journal(doc: Document, intro: list[str], image_name: str | None, image_caption: str | None, rows: list[tuple[str, str, str, str]]):
    add_heading(doc, "ตรวจสอบ Journal Entry", level=2)
    for line in intro:
        add_para(doc, line)
    if image_name and image_caption:
        add_image(doc, image_name, image_caption)
    if rows:
        add_table(doc, ["บัญชี", "ความหมาย", "Debit", "Credit"], rows)


def main():
    samples = load_json(SAMPLES_PATH)
    transactions = load_json(TRANSACTIONS_PATH)
    dep = load_json(DEP_MOVE_PATH)
    image_map = load_json(IMAGE_MAP_PATH)

    data = samples["samples"]
    draft_asset = data["draft_asset"]
    open_asset = data["open_asset"]
    model_asset = data["model_asset"]
    sell_asset = transactions["sell_asset"]
    sale_invoice = transactions["sale_invoice"]
    sale_move = transactions["sale_move"]
    dispose_asset = transactions["dispose_asset"]
    disposal_move = transactions["disposal_move"]

    doc = Document()
    configure_doc(doc)

    add_para(doc, "บทที่ 6 Fixed Asset", size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=3)
    add_para(doc, "คู่มือรวมการจัดการทรัพย์สินถาวร", size=18, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)

    add_heading(doc, "สารบัญ", level=1)
    for item in [
        "6.1 การจัดการทรัพย์สิน",
        "6.2 การสร้างทรัพย์สิน",
        "6.3 การเริ่มคิดค่าเสื่อมราคา (Activation & Depreciation)",
        "6.4 การขายสินทรัพย์ (Selling Assets & Gain/Loss)",
        "6.5 การตัดจำหน่ายสินทรัพย์ (Disposal / Scrap)",
        "6.6 รายงานทรัพย์สินถาวร (Fixed Asset Report)",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "การเข้าเมนูหลักของบทนี้", level=1)
    add_para(
        doc,
        "บทนี้ใช้เมนูหลักในหมวด Accounting ทั้งหมด โดยผู้ใช้งานควรเริ่มจากหน้า Dashboard แล้วกดเข้าโมดูล Accounting จากนั้นจึงเลือกเมนูย่อยตามงานที่ต้องทำ เช่น รายการทรัพย์สิน ต้นแบบทรัพย์สิน หรือรายงานทรัพย์สินถาวร",
    )
    add_image(doc, "ch6_nav_dashboard_accounting_annotated.png", "หน้า Dashboard แบบเต็มหน้าจอที่ใช้เริ่มเข้าโมดูล Accounting", width=6.9)
    add_para(doc, f"เมนูรายการทรัพย์สินใช้เส้นทาง {samples['menus']['assets']}")
    add_image(doc, image_map["ch6_nav_assets"], "หน้าเมนู Assets ที่ใช้ดูและเปิดรายการทรัพย์สิน", width=6.9)
    add_para(doc, f"เมนูต้นแบบทรัพย์สินใช้เส้นทาง {samples['menus']['asset_models']}")
    add_image(doc, image_map["ch6_nav_asset_models"], "หน้าเมนู Asset Models ที่ใช้กำหนดบัญชีและวิธีคิดค่าเสื่อมราคา", width=6.9)

    doc.add_page_break()

    build_subsection_header(
        doc,
        "6.1 การจัดการทรัพย์สิน",
        "หัวข้อนี้ใช้สำหรับดูรายการทรัพย์สินที่มีอยู่แล้ว ตรวจสอบสถานะ มูลค่าทุน มูลค่าคงเหลือ และเข้าไปทำงานต่อ เช่น ปรับแผนค่าเสื่อมราคา ขายทรัพย์สิน ตัดจำหน่าย หรือเปิดดูรายการบัญชีที่เกี่ยวข้อง",
        "เมื่อเข้าหน้ารายการทรัพย์สิน ผู้ใช้งานจะเห็นทรัพย์สินหลายสถานะ เช่น Draft สำหรับรายการที่ยังไม่เริ่มใช้งาน และ Open สำหรับรายการที่เริ่มใช้งานแล้ว หน้าจอเดียวกันนี้จึงเป็นจุดเริ่มต้นของงานประจำด้านทรัพย์สินเกือบทั้งหมด",
        [
            "ค้นหารายการทรัพย์สิน",
            "เปิดดูสถานะและมูลค่าของทรัพย์สิน",
            "เปิดดูรายการบัญชีของทรัพย์สิน",
            "เข้าไปทำงานต่อ เช่น ปรับค่าเสื่อมราคา ขาย หรือตัดจำหน่าย",
        ],
        samples["menus"]["assets"],
    )
    build_steps_with_images(
        doc,
        [
            ("เริ่มจากหน้า Dashboard แล้วกดเข้าโมดูล Accounting เพื่อรวมงานด้านบัญชีและทรัพย์สินไว้ในที่เดียว", "ch6_nav_dashboard_accounting_annotated.png", "เริ่มเข้าโมดูล Accounting จากหน้า Dashboard"),
            ("เลือกเมนู Assets เพื่อเปิดรายการทรัพย์สินทั้งหมดของบริษัท แล้วดูภาพรวมก่อนว่าต้องทำงานกับทรัพย์สินใบใด", image_map["ch6_nav_assets"], "หน้า Assets ใช้ค้นหาและเปิดรายการทรัพย์สิน"),
            (f"เปิดทรัพย์สินที่ยังเป็น Draft เช่น {draft_asset['name'].splitlines()[0]} เพื่อดูว่าข้อมูลยังอยู่ในช่วงเตรียมรายการ และยังไม่เริ่มตัดค่าเสื่อมราคา", image_map["ch6_asset_draft_form"], "ตัวอย่างทรัพย์สินสถานะ Draft"),
            (f"เปิดทรัพย์สินที่ใช้งานอยู่จริง เช่น {open_asset['name']} เพื่อดูมูลค่าตั้งต้น มูลค่าคงเหลือ และปุ่มที่ใช้จัดการต่อ", image_map["ch6_asset_open_form"], "ตัวอย่างทรัพย์สินสถานะ Open"),
        ],
    )
    build_fields(
        doc,
        [
            ("Name", "ชื่อทรัพย์สินที่ใช้แยกแต่ละรายการ"),
            ("State", "สถานะของทรัพย์สิน เช่น Draft, Open หรือ Close"),
            ("Acquisition Date", "วันที่เริ่มรับรู้ทรัพย์สิน"),
            ("Original Value", "มูลค่าตั้งต้นของทรัพย์สิน"),
            ("Book Value", "มูลค่าคงเหลือสุทธิหลังหักค่าเสื่อมราคาที่รับรู้แล้ว"),
            ("Model", "ต้นแบบทรัพย์สินที่ใช้กำหนดบัญชีและวิธีคิดค่าเสื่อมราคา"),
            ("Journal", "สมุดรายวันที่ใช้บันทึกรายการของทรัพย์สิน"),
        ],
    )
    build_journal(
        doc,
        [
            "ในหัวข้อนี้ หน้ารายการทรัพย์สินยังไม่ใช่การสร้าง Journal Entry ใหม่ด้วยตัวเอง แต่เป็นจุดที่ผู้ใช้งานใช้เปิดดูรายการบัญชีของทรัพย์สินที่เกิดขึ้นแล้ว",
            f"ตัวอย่างทรัพย์สิน {open_asset['name']} สามารถกดปุ่ม Open Entries เพื่อดูรายการค่าเสื่อมราคาและรายการบัญชีที่เกี่ยวข้องย้อนหลังได้",
        ],
        image_map["ch6_journal_depreciation"],
        "ตัวอย่าง Journal Entry ที่เปิดดูได้จากทรัพย์สินที่ใช้งานอยู่",
        move_rows(
            dep["lines"],
            {
                "530072": "เป็นค่าใช้จ่ายค่าเสื่อมราคาของงวดนั้น จึงอยู่ฝั่ง Debit",
                "123014": "เป็นค่าเสื่อมราคาสะสมของทรัพย์สิน จึงอยู่ฝั่ง Credit",
            },
        ),
    )

    doc.add_page_break()

    build_subsection_header(
        doc,
        "6.2 การสร้างทรัพย์สิน",
        "หัวข้อนี้ใช้สำหรับสร้างทรัพย์สินรายการใหม่จากข้อมูลจริงของสินทรัพย์ เช่น เครื่องจักร โปรแกรมคอมพิวเตอร์ หรืออุปกรณ์สำนักงาน โดยกำหนดมูลค่าทุน วันที่เริ่มใช้งาน และต้นแบบที่ต้องการใช้",
        "การสร้างทรัพย์สินแบ่งได้เป็นสองส่วน คือ เลือกต้นแบบที่เหมาะสมก่อน แล้วจึงสร้างรายการทรัพย์สินจริง เมื่อตัวรายการยังเป็น Draft ระบบยังไม่ตัดค่าเสื่อมราคาและยังไม่สร้าง Journal Entry ค่าเสื่อมราคา",
        [
            "เลือกต้นแบบทรัพย์สินให้ตรงประเภท",
            "สร้างทรัพย์สินใหม่และกรอกข้อมูลหลัก",
            "บันทึกรายการในสถานะ Draft",
            "ตรวจสอบว่าข้อมูลพร้อมสำหรับเริ่มใช้งาน",
        ],
        f"{samples['menus']['asset_models']} และ {samples['menus']['assets']}",
    )
    build_steps_with_images(
        doc,
        [
            ("เข้าเมนู Asset Models ก่อนเพื่อเลือกต้นแบบที่เหมาะกับทรัพย์สินที่จะสร้าง เพราะต้นแบบจะเป็นตัวกำหนดบัญชีและวิธีคิดค่าเสื่อมราคาให้กับรายการใหม่", image_map["ch6_nav_asset_models"], "หน้า Asset Models ใช้เลือกต้นแบบทรัพย์สิน"),
            (f"เปิดต้นแบบ {model_asset['name']} แล้วตรวจสอบว่าวิธีคิดค่าเสื่อมราคา บัญชีทรัพย์สิน บัญชีค่าเสื่อมสะสม และบัญชีค่าใช้จ่ายถูกต้องก่อนนำไปใช้", image_map["ch6_asset_model_form"], "ตัวอย่างต้นแบบทรัพย์สินที่ใช้ในระบบจริง"),
            ("กลับมาที่เมนู Assets แล้วกด New เพื่อสร้างรายการทรัพย์สินใหม่ จากนั้นกรอกชื่อทรัพย์สิน วันที่รับรู้ มูลค่าทุน และเลือกต้นแบบให้ตรงกับสินทรัพย์จริง", image_map["ch6_nav_assets"], "หน้า Assets ที่ใช้สร้างทรัพย์สินใหม่"),
            (f"ตรวจสอบรายการ Draft ที่สร้างแล้ว เช่น {draft_asset['name'].splitlines()[0]} ว่าข้อมูลหลักครบ ก่อนจะกด Confirm ในขั้นตอนถัดไป", image_map["ch6_asset_draft_form"], "ตัวอย่างทรัพย์สินใหม่ในสถานะ Draft"),
        ],
    )
    build_fields(
        doc,
        [
            ("Name", "ชื่อทรัพย์สินที่ต้องการบันทึก"),
            ("Model", "ต้นแบบทรัพย์สินที่ใช้ดึงบัญชีและวิธีคิดค่าเสื่อมราคา"),
            ("Original Value", "มูลค่าทุนของทรัพย์สิน"),
            ("Acquisition Date", "วันที่เริ่มรับรู้ทรัพย์สิน"),
            ("Journal", "สมุดรายวันที่จะใช้ลงรายการที่เกี่ยวกับทรัพย์สิน"),
            ("Prorata Date", "วันที่ใช้เริ่มคำนวณค่าเสื่อมราคา"),
        ],
    )
    build_journal(
        doc,
        [
            "ในขั้นสร้างทรัพย์สินที่ยังอยู่สถานะ Draft ระบบยังไม่สร้าง Journal Entry ค่าเสื่อมราคาให้ทันที",
            "ผู้ใช้งานควรตรวจให้ครบก่อนว่า ชื่อทรัพย์สิน วันที่รับรู้ และมูลค่าทุนถูกต้อง เพราะเมื่อเริ่มใช้งานแล้ว แผนค่าเสื่อมราคาจะอ้างอิงข้อมูลชุดนี้",
        ],
        None,
        None,
        [],
    )

    doc.add_page_break()

    build_subsection_header(
        doc,
        "6.3 การเริ่มคิดค่าเสื่อมราคา (Activation & Depreciation)",
        "หัวข้อนี้ใช้เมื่อต้องเริ่มใช้งานทรัพย์สินจริง และให้ระบบเริ่มคำนวณค่าเสื่อมราคาเป็นงวดตามต้นแบบที่กำหนดไว้",
        "เมื่อกด Confirm ให้ทรัพย์สินเริ่มใช้งาน ระบบจะเปลี่ยนสถานะจาก Draft เป็น Open หลังจากนั้นระบบจะเริ่มสร้างตารางค่าเสื่อมราคา และเมื่อถึงงวด ระบบจะบันทึกบัญชีค่าเสื่อมราคาให้อัตโนมัติ",
        [
            "เปลี่ยนทรัพย์สินจาก Draft เป็นใช้งานจริง",
            "ดูมูลค่าคงเหลือของทรัพย์สิน",
            "เปิดดู Journal Entry ค่าเสื่อมราคา",
            "ตรวจว่าเดบิตและเครดิตลงถูกบัญชี",
        ],
        samples["menus"]["assets"],
    )
    build_steps_with_images(
        doc,
        [
            (f"เปิดทรัพย์สินสถานะ Draft เช่น {draft_asset['name'].splitlines()[0]} แล้วกด Confirm เมื่อพร้อมเริ่มใช้งานจริง", image_map["ch6_asset_draft_form"], "จุดที่ใช้เริ่มใช้งานทรัพย์สินจากสถานะ Draft"),
            (f"หลังเริ่มใช้งานแล้ว ให้เปิดทรัพย์สินสถานะ Open เช่น {open_asset['name']} เพื่อดูมูลค่าคงเหลือและปุ่มจัดการค่าเสื่อมราคา", image_map["ch6_asset_open_form"], "ตัวอย่างทรัพย์สินที่เริ่มใช้งานแล้ว"),
            ("กด Open Entries เพื่อเปิดรายการบัญชีที่เกิดจากค่าเสื่อมราคา แล้วเปรียบเทียบกับแผนค่าเสื่อมราคาของทรัพย์สิน", image_map["ch6_journal_depreciation"], "ตัวอย่าง Journal Entry ของค่าเสื่อมราคา"),
        ],
    )
    build_fields(
        doc,
        [
            ("State", "ต้องเป็น Open เมื่อทรัพย์สินเริ่มใช้งานแล้ว"),
            ("Book Value", "มูลค่าคงเหลือของทรัพย์สินหลังรับรู้ค่าเสื่อมราคา"),
            ("Method", "วิธีคิดค่าเสื่อมราคา"),
            ("Method Number", "จำนวนงวดที่ใช้ตัดค่าเสื่อม"),
            ("Method Period", "ช่วงเวลาต่อหนึ่งงวด เช่น เดือน"),
            ("Open Entries", "ปุ่มที่ใช้เปิดดูรายการบัญชีของทรัพย์สิน"),
        ],
    )
    build_journal(
        doc,
        [
            f"ตัวอย่าง Journal Entry ค่าเสื่อมราคาที่ใช้จริงในระบบคือรายการอ้างอิง {dep['ref']}",
            "หลักการอ่านรายการนี้คือ ค่าใช้จ่ายค่าเสื่อมราคาจะอยู่ฝั่ง Debit และค่าเสื่อมราคาสะสมจะอยู่ฝั่ง Credit",
            "เมื่อทั้งสองบรรทัดมียอดเท่ากัน แปลว่าระบบตัดค่าเสื่อมราคาให้ทรัพย์สินในงวดนั้นเรียบร้อย",
        ],
        image_map["ch6_journal_depreciation"],
        "ตัวอย่าง Journal Entry ค่าเสื่อมราคาในระบบจริง",
        move_rows(
            dep["lines"],
            {
                "530072": "Debit บัญชีค่าใช้จ่ายค่าเสื่อมราคา เพื่อรับรู้ต้นทุนของงวดนั้น",
                "123014": "Credit บัญชีค่าเสื่อมราคาสะสม เพื่อเพิ่มยอดสะสมของค่าเสื่อมราคา",
            },
        ),
    )

    doc.add_page_break()

    build_subsection_header(
        doc,
        "6.4 การขายสินทรัพย์ (Selling Assets & Gain/Loss)",
        "หัวข้อนี้ใช้เมื่อต้องขายทรัพย์สินออกจากบริษัท โดยระบบจะมีทั้งเอกสารขายให้ลูกค้า และรายการบัญชีที่ตัดทรัพย์สินออกจากบัญชี พร้อมรับรู้ผลต่างจากการขาย",
        "การขายสินทรัพย์ในระบบจริงจะมีอย่างน้อยสองส่วนที่ต้องตรวจ คือ ใบขายที่ออกให้ลูกค้า และ Journal Entry จากการตัดทรัพย์สินออกจากระบบ เมื่อขายเสร็จสถานะของทรัพย์สินจะปิดรายการ",
        [
            "เปิดทรัพย์สินที่ต้องการขาย",
            "ตรวจใบขายที่ออกให้ลูกค้า",
            "ตรวจ Journal จากการตัดทรัพย์สินออกจากบัญชี",
            "ดูว่ามีกำไรหรือรายได้อื่นจากการขายหรือไม่",
        ],
        samples["menus"]["assets"],
    )
    build_steps_with_images(
        doc,
        [
            (f"เปิดทรัพย์สินที่ขายแล้ว เช่น {sell_asset['name']} เพื่อดูว่าสถานะถูกปิดรายการและวันที่ขายถูกต้อง", image_map["ch6_asset_sell_form"], "ตัวอย่างทรัพย์สินที่ผ่านการขายแล้ว"),
            (f"เปิดใบขายเลขที่ {sale_invoice['name']} แล้วตรวจชื่อลูกค้า วันที่เอกสาร และยอดรวมให้ตรงกับการขายจริง", image_map["ch6_asset_sale_invoice"], "ตัวอย่างใบขายสินทรัพย์ในระบบจริง"),
            (f"เปิด Journal Entry เลขที่ {sale_move['name']} เพื่อตรวจว่าระบบตัดทรัพย์สินออกจากบัญชีและลงผลต่างจากการขายเรียบร้อย", image_map["ch6_journal_sale"], "Journal จากการขายทรัพย์สิน"),
        ],
    )
    build_fields(
        doc,
        [
            ("Disposal Date", "วันที่ขายทรัพย์สิน"),
            ("State", "หลังขายแล้วทรัพย์สินจะถูกปิดรายการ"),
            ("Book Value", "มูลค่าคงเหลือก่อนปิดรายการ"),
            ("Partner", "ลูกค้าที่ซื้อทรัพย์สิน"),
            ("Amount Total", "ยอดรวมใบขาย"),
            ("Journal", "สมุดรายวันที่ใช้ลงรายการบัญชี"),
        ],
    )
    build_journal(
        doc,
        [
            f"ในตัวอย่างจริง ระบบมีใบขาย {sale_invoice['name']} ให้ลูกค้า {sale_invoice['partner']} มูลค่า {sale_invoice['amount_total']:,.2f} บาท และมีรายการบัญชีตัดทรัพย์สินเลขที่ {sale_move['name']}",
            "รายการบัญชีนี้ควรอ่านทีละบรรทัด โดยดูว่าบรรทัดใดเป็นการตัดสินทรัพย์ บรรทัดใดเป็นการล้างค่าเสื่อมสะสม และบรรทัดใดสะท้อนผลต่างจากการขาย",
            "ตัวอย่างนี้เป็นข้อมูลจริงในระบบ จึงอาจมีผลต่างจากมูลค่าขายตามใบขาย เพราะระบบอาจบันทึกผลต่างเป็นรายได้อื่นหรือผลกำไรจากการขายทรัพย์สิน",
        ],
        image_map["ch6_journal_sale"],
        "Journal Entry จากการขายทรัพย์สินในระบบจริง",
        move_rows(
            sale_move["lines"],
            {
                "124001": "Credit บัญชีทรัพย์สินถาวร เพื่อตัดมูลค่าทุนของทรัพย์สินออกจากงบ",
                "124004": "Debit บัญชีค่าเสื่อมราคาสะสม เพื่อล้างยอดสะสมที่เคยรับรู้ไว้",
                "410000": "Debit บัญชีรายได้จากการขายในประเทศ เป็นผลจากเอกสารที่ระบบเชื่อมโยงกับการขาย",
                "430000": "Credit บัญชีรายได้อื่น ใช้รับรู้ผลต่างจากการขายทรัพย์สิน",
            },
        ),
    )

    doc.add_page_break()

    build_subsection_header(
        doc,
        "6.5 การตัดจำหน่ายสินทรัพย์ (Disposal / Scrap)",
        "หัวข้อนี้ใช้เมื่อต้องตัดทรัพย์สินออกจากระบบโดยไม่มีการขาย เช่น ชำรุด ใช้งานต่อไม่ได้ หรือเลิกใช้งานถาวร",
        "การตัดจำหน่ายจะปิดทรัพย์สินออกจากระบบและสร้างรายการบัญชีเพื่อล้างสินทรัพย์และค่าเสื่อมสะสม พร้อมรับรู้มูลค่าคงเหลือเป็นค่าใช้จ่ายหรือผลขาดทุนตามโครงบัญชีที่ตั้งไว้",
        [
            "เปิดทรัพย์สินที่ต้องการตัดจำหน่าย",
            "ตรวจข้อมูลก่อนปิดรายการ",
            "เปิดดู Journal จากการตัดจำหน่าย",
            "ตรวจว่ามูลค่าคงเหลือถูกโอนไปยังบัญชีค่าใช้จ่ายแล้ว",
        ],
        samples["menus"]["assets"],
    )
    build_steps_with_images(
        doc,
        [
            (f"เปิดทรัพย์สินที่ตัดจำหน่ายแล้ว เช่น {dispose_asset['name']} เพื่อดูว่าสถานะเป็นการปิดรายการเรียบร้อย", image_map["ch6_asset_dispose_form"], "ตัวอย่างทรัพย์สินที่ถูกตัดจำหน่าย"),
            (f"เปิด Journal Entry เลขที่ {disposal_move['name']} แล้วตรวจว่าระบบตัดสินทรัพย์ ล้างค่าเสื่อมสะสม และรับรู้มูลค่าคงเหลือเป็นค่าใช้จ่ายครบทุกบรรทัด", image_map["ch6_journal_disposal"], "Journal ของการตัดจำหน่ายทรัพย์สิน"),
        ],
    )
    build_fields(
        doc,
        [
            ("Disposal Date", "วันที่ตัดจำหน่ายทรัพย์สิน"),
            ("State", "หลังตัดจำหน่ายแล้วสถานะจะปิดรายการ"),
            ("Book Value", "มูลค่าคงเหลือก่อนปิดรายการ"),
            ("Depreciation Account", "บัญชีค่าเสื่อมราคาสะสม"),
            ("Expense Account", "บัญชีค่าใช้จ่ายหรือผลขาดทุนจากการตัดทรัพย์สิน"),
        ],
    )
    build_journal(
        doc,
        [
            f"ตัวอย่างจริงใช้รายการบัญชีเลขที่ {disposal_move['name']} จากทรัพย์สิน {dispose_asset['name']}",
            "หลักการอ่านคือ บัญชีสินทรัพย์จะถูกเครดิตเพื่อตัดออก บัญชีค่าเสื่อมสะสมจะถูกเดบิตเพื่อล้างยอดสะสม และมูลค่าคงเหลือที่เหลืออยู่จะเดบิตเข้าบัญชีค่าใช้จ่ายหรือผลขาดทุน",
        ],
        image_map["ch6_journal_disposal"],
        "Journal Entry จากการตัดจำหน่ายทรัพย์สิน",
        move_rows(
            disposal_move["lines"],
            {
                "124001": "Credit บัญชีสินทรัพย์ถาวร เพื่อตัดมูลค่าทุนของทรัพย์สินออกจากบัญชี",
                "124004": "Debit บัญชีค่าเสื่อมราคาสะสม เพื่อล้างยอดสะสมที่เกี่ยวกับทรัพย์สินใบนี้",
                "510000": "Debit บัญชีต้นทุนหรือค่าใช้จ่าย เพื่อรับรู้มูลค่าคงเหลือที่ตัดจำหน่าย",
            },
        ),
    )

    doc.add_page_break()

    build_subsection_header(
        doc,
        "6.6 รายงานทรัพย์สินถาวร (Fixed Asset Report)",
        "หัวข้อนี้ใช้สำหรับดูภาพรวมทรัพย์สินทั้งหมดของบริษัท เช่น มูลค่าทุน ค่าเสื่อมราคาสะสม มูลค่าคงเหลือ และรายการที่ปิดไปแล้ว",
        "รายงานทรัพย์สินถาวรเหมาะสำหรับการตรวจสอบสิ้นเดือน สิ้นงวด หรือใช้ตอบคำถามของผู้บริหาร เพราะช่วยให้เห็นภาพรวมของทรัพย์สินทั้งระบบโดยไม่ต้องเปิดทีละใบ",
        [
            "เปิดรายงานทรัพย์สินถาวร",
            "ดูยอดรวมและกรองช่วงเวลาที่ต้องการ",
            "ใช้รายงานช่วยตรวจสอบก่อนปิดงวด",
        ],
        samples["menus"]["fixed_asset_report"],
    )
    build_steps_with_images(
        doc,
        [
            ("เข้าเมนูรายงานทรัพย์สินถาวรจากหมวด Reporting เพื่อดูภาพรวมของทรัพย์สินทั้งหมดในระบบ", image_map["ch6_nav_fixed_asset_report"], "หน้ารายงานทรัพย์สินถาวรในระบบจริง"),
            ("ตรวจยอดรวม มูลค่าทุน ค่าเสื่อมสะสม และมูลค่าคงเหลือ แล้วใช้ตัวกรองหรือช่วงเวลาเพื่อดูข้อมูลที่ต้องการ", image_map["ch6_nav_fixed_asset_report"], "พื้นที่รายงานที่ใช้วิเคราะห์ยอดทรัพย์สิน"),
        ],
    )
    build_fields(
        doc,
        [
            ("Date / Period", "ช่วงเวลาที่ใช้กรองรายงาน"),
            ("Original Value", "มูลค่าทุนรวมของทรัพย์สิน"),
            ("Accumulated Depreciation", "ค่าเสื่อมราคาสะสมรวม"),
            ("Book Value", "มูลค่าคงเหลือสุทธิ"),
            ("State", "สถานะของทรัพย์สินที่ต้องการดู เช่น ใช้งานอยู่หรือปิดรายการแล้ว"),
        ],
    )
    build_journal(
        doc,
        [
            "หน้ารายงานทรัพย์สินถาวรเป็นหน้าสำหรับดูผลรวมและวิเคราะห์ข้อมูล ไม่ได้สร้าง Journal Entry ใหม่ด้วยตัวเอง",
            "ผู้ใช้งานควรใช้รายงานนี้ร่วมกับหน้ารายการทรัพย์สินและ Open Entries เพื่อไล่ตรวจรายการบัญชีรายใบเมื่อพบยอดที่ต้องตรวจเพิ่มเติม",
        ],
        None,
        None,
        [],
    )

    add_heading(doc, "สรุปการใช้งานบทที่ 6", level=1)
    add_para(
        doc,
        "บทนี้ครอบคลุมการทำงานของทรัพย์สินถาวรตั้งแต่การสร้างทรัพย์สิน การเริ่มคิดค่าเสื่อมราคา การขาย การตัดจำหน่าย และการดูรายงานภาพรวม จุดสำคัญของการใช้งานจริงคือทุกครั้งที่มีการเปลี่ยนสถานะของทรัพย์สิน ผู้ใช้งานควรกลับไปตรวจ Journal Entry เสมอ เพื่อยืนยันว่าระบบลงบัญชีสินทรัพย์ ค่าเสื่อมสะสม ค่าใช้จ่าย และผลต่างจากการขายหรือตัดจำหน่ายถูกต้องตามรายการจริง",
    )

    DOCX_DIR.mkdir(parents=True, exist_ok=True)
    PDF_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(DOCX_PATH)
    subprocess.run(
        [str(SOFFICE), "--headless", "--convert-to", "pdf", "--outdir", str(PDF_DIR), str(DOCX_PATH)],
        check=True,
    )
    print("built fixed asset chapter 6 combined manual")


if __name__ == "__main__":
    main()

from __future__ import annotations

from copy import deepcopy
from importlib.util import module_from_spec, spec_from_file_location
from pathlib import Path
import re
import subprocess

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph


ROOT = Path(r"C:\365_project\TheCool18e\Dev")
DOCX_DIR = ROOT / "manual" / "Accouting_Manual" / "generated_20260408" / "docx"
PDF_DIR = ROOT / "manual" / "Accouting_Manual" / "generated_20260408" / "pdf_review"
SOFFICE = Path(r"C:\Program Files\LibreOffice\program\soffice.com")
PAIR_SCRIPT = ROOT / "reports" / "accounting_manual_tools" / "pair_manual_images_with_steps.py"


def load_pair_module():
    spec = spec_from_file_location("pair_manual_images_with_steps", PAIR_SCRIPT)
    module = module_from_spec(spec)
    assert spec.loader is not None
    spec.loader.exec_module(module)
    return module


PAIR = load_pair_module()
IMAGE_MAP = PAIR.IMAGE_MAP
IMAGE_DIR = PAIR.IMAGE_DIR
DETAIL_HEADING = PAIR.DETAIL_HEADING
STOP_HEADINGS = PAIR.STOP_HEADINGS


REMOVE_TOP_LEVEL = {
    "1. เริ่มจากหน้า Dashboard",
    "2. ไปที่เมนูที่ใช้",
    "3. ขั้นตอนการใช้งาน",
}


def set_run_font(run, size: int = 16, bold: bool = False, italic: bool = False) -> None:
    run.bold = bold
    run.italic = italic
    run.font.name = "TH Sarabun New"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "TH Sarabun New")
    run.font.size = Pt(size)


def make_empty_after(paragraph: Paragraph) -> Paragraph:
    new_p = deepcopy(paragraph._p)
    for child in list(new_p):
        new_p.remove(child)
    paragraph._p.addnext(new_p)
    return Paragraph(new_p, paragraph._parent)


def insert_text_after(paragraph: Paragraph, text: str, *, size: int = 16, bold: bool = False, italic: bool = False) -> Paragraph:
    new_para = make_empty_after(paragraph)
    run = new_para.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic)
    return new_para


def insert_picture_after(paragraph: Paragraph, image_path: Path, caption: str) -> Paragraph:
    pic_para = make_empty_after(paragraph)
    run = pic_para.add_run()
    run.add_picture(str(image_path), width=Inches(6.2))
    cap_para = insert_text_after(pic_para, caption, size=14, italic=True)
    return cap_para


def remove_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def doc_code_from_name(name: str) -> str | None:
    match = re.match(r"^(\d+\.\d+)_", name)
    return match.group(1) if match else None


def find_detail_heading(paragraphs: list[Paragraph]) -> int | None:
    for idx, paragraph in enumerate(paragraphs):
        if paragraph.text.strip() == DETAIL_HEADING:
            return idx
    return None


def is_numbered_step(text: str) -> bool:
    return bool(re.match(r"^\d+\.\s+", text.strip()))


def step_number(text: str) -> int | None:
    match = re.match(r"^(\d+)\.\s+", text.strip())
    return int(match.group(1)) if match else None


def find_detail_step_indices(paragraphs: list[Paragraph]) -> list[int]:
    heading_idx = find_detail_heading(paragraphs)
    if heading_idx is None:
        return []
    result: list[int] = []
    started = False
    previous = 0
    for idx, paragraph in enumerate(paragraphs[heading_idx + 1 :], start=heading_idx + 1):
        text = paragraph.text.strip()
        if not text:
            continue
        if text.startswith(STOP_HEADINGS):
            break
        if is_numbered_step(text):
            number = step_number(text) or 0
            if started and number <= previous:
                break
            started = True
            previous = number
            result.append(idx)
    return result


def strip_number(text: str) -> str:
    return re.sub(r"^\d+\.\s*", "", text.strip())


def merge_texts(texts: list[str]) -> str:
    cleaned = [strip_number(t).rstrip(" .") for t in texts if t.strip()]
    if not cleaned:
        return ""
    merged = cleaned[0]
    for extra in cleaned[1:]:
        if extra.startswith(("กด", "เลือก", "กรอก", "ตรวจ", "เปิด", "กลับ", "เมื่อ", "หาก", "แล้ว")):
            merged += f" จากนั้น{extra if extra.startswith(' ') else ' ' + extra}"
        else:
            merged += f" และ{extra if extra.startswith(' ') else ' ' + extra}"
    return merged.strip()


def build_groups(step_texts: list[str], image_specs: list[tuple[str, str]]) -> list[tuple[str, tuple[str, str]]]:
    if not step_texts:
        return []
    assigned_specs = [image_specs[min(i, len(image_specs) - 1)] for i in range(len(step_texts))]
    groups: list[tuple[list[str], tuple[str, str]]] = []
    for step_text, spec in zip(step_texts, assigned_specs):
        if not groups or groups[-1][1][0] != spec[0]:
            groups.append(([step_text], spec))
        else:
            groups[-1][0].append(step_text)
    return [(merge_texts(texts), spec) for texts, spec in groups]


def remove_existing_detail_block(doc: Document) -> tuple[Paragraph | None, list[str]]:
    paragraphs = doc.paragraphs
    for paragraph in list(paragraphs):
        if paragraph.text.strip() in REMOVE_TOP_LEVEL:
            remove_paragraph(paragraph)
    paragraphs = doc.paragraphs
    heading_idx = find_detail_heading(paragraphs)
    if heading_idx is None:
        return None, []
    step_indices = find_detail_step_indices(paragraphs)
    step_texts = [paragraphs[i].text.strip() for i in step_indices]
    start_idx = heading_idx + 1
    end_idx = None
    for idx in range(start_idx, len(paragraphs)):
        text = paragraphs[idx].text.strip()
        if text.startswith(STOP_HEADINGS):
            end_idx = idx
            break
    if end_idx is None:
        end_idx = len(paragraphs)
    for idx in range(end_idx - 1, start_idx - 1, -1):
        remove_paragraph(paragraphs[idx])
    return doc.paragraphs[heading_idx], step_texts


def enrich_doc(path: Path) -> bool:
    code = doc_code_from_name(path.name)
    if not code or code not in IMAGE_MAP:
        return False
    doc = Document(path)
    heading_para, step_texts = remove_existing_detail_block(doc)
    if heading_para is None or not step_texts:
        return False
    groups = build_groups(step_texts, IMAGE_MAP[code])
    if not groups:
        return False
    current = heading_para
    for idx, (step_text, (image_name, caption)) in enumerate(groups, start=1):
        current = insert_text_after(current, f"{idx}. {step_text}", size=16)
        image_path = IMAGE_DIR / image_name
        if image_path.exists():
            current = insert_picture_after(current, image_path, caption)
    doc.save(path)
    return True


def render_pdfs(paths: list[Path]) -> None:
    if not paths:
        return
    PDF_DIR.mkdir(parents=True, exist_ok=True)
    cmd = [str(SOFFICE), "--headless", "--convert-to", "pdf", "--outdir", str(PDF_DIR), *map(str, paths)]
    subprocess.run(cmd, check=True)


def main() -> None:
    updated: list[Path] = []
    for path in sorted(DOCX_DIR.glob("*.docx")):
        if enrich_doc(path):
            updated.append(path)
    render_pdfs(updated)
    print(f"merged flows in {len(updated)} docs")


if __name__ == "__main__":
    main()

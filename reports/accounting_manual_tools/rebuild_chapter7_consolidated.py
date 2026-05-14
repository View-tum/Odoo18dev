from __future__ import annotations

import importlib.util
import subprocess
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont
import json
import fitz


ROOT = Path(r"C:\365_project\TheCool18e\Dev")
TOOLS_DIR = ROOT / "reports" / "accounting_manual_tools"
GENERATED_DIR = ROOT / "manual" / "Accouting_Manual" / "generated_20260408"
IMAGE_DIR = GENERATED_DIR / "images"
TMP_DIR = ROOT / "tmp" / "docs" / "chapter7_rebuild"
OUTPUT_DOCX = ROOT / "Manual_Chapter_7_Full_Consolidated.docx"
OUTPUT_PDF = ROOT / "Manual_Chapter_7_Full_Consolidated.pdf"
RENDER_DIR = TMP_DIR / "rendered_pages"
SOFFICE = Path(r"C:\Program Files\LibreOffice\program\soffice.com")


def load_source_module():
    module_path = TOOLS_DIR / "build_fixed_asset_mfg_manuals.py"
    spec = importlib.util.spec_from_file_location("build_fixed_asset_mfg_manuals", module_path)
    module = importlib.util.module_from_spec(spec)
    assert spec.loader is not None
    spec.loader.exec_module(module)
    return module


def configure_doc(doc: Document):
    style = doc.styles["Normal"]
    style.font.name = "Angsana New"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Angsana New")
    style.font.size = Pt(16)
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)


def font_run(run, size=16, bold=False, color: RGBColor | None = None):
    run.bold = bold
    run.font.name = "Angsana New"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Angsana New")
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color


def add_heading(doc: Document, text: str, level=1):
    p = doc.add_paragraph()
    font_run(
        p.add_run(text),
        size=20 if level == 1 else 18 if level == 2 else 16,
        bold=True,
        color=RGBColor(31, 78, 121) if level <= 2 else None,
    )


def add_para(doc: Document, text: str, bold=False, center=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    font_run(p.add_run(text), bold=bold)


def add_bullets(doc: Document, items: Iterable[str]):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        font_run(p.add_run(str(item)))


def add_steps(doc: Document, items: Iterable[str]):
    for idx, item in enumerate(items, start=1):
        p = doc.add_paragraph()
        font_run(p.add_run(f"{idx}. {item}"))


def shade_cell(cell, color="D9EAF7"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    tc_pr.append(shd)


def set_cell(cell, text, size=14):
    cell.text = str(text)
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            font_run(run, size=size)


def add_table(doc: Document, headers: list[str], rows: list[tuple]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        set_cell(table.rows[0].cells[i], header)
        shade_cell(table.rows[0].cells[i])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell(cells[i], value)


def uniq_pairs(pairs: list[tuple[str, str]]) -> list[tuple[str, str]]:
    seen = set()
    cleaned = []
    for caption, image in pairs:
        if image in seen:
            continue
        seen.add(image)
        cleaned.append((caption, image))
    return cleaned


def crop_image(image_name: str) -> Path:
    src = IMAGE_DIR / image_name
    if not src.exists():
        return src
    TMP_DIR.mkdir(parents=True, exist_ok=True)
    render_src = src
    if src.stem.endswith("_annotated"):
        raw_name = src.name.replace("_annotated", "")
        raw_src = IMAGE_DIR / raw_name
        raw_meta = raw_src.with_suffix(raw_src.suffix + ".json")
        lite = TMP_DIR / f"{src.stem}_lite{src.suffix}"
        if raw_src.exists() and raw_meta.exists():
            if not lite.exists():
                data = json.loads(raw_meta.read_text(encoding="utf-8"))
                boxes = data.get("boxes", [])
                with Image.open(raw_src) as img:
                    draw = ImageDraw.Draw(img)
                    try:
                        font = ImageFont.truetype("arial.ttf", 28)
                    except Exception:
                        font = ImageFont.load_default()
                    for box in boxes:
                        x = box["x"]
                        y = box["y"]
                        w = box["width"]
                        h = box["height"]
                        label = str(box.get("label", ""))
                        draw.rectangle((x, y, x + w, y + h), outline=(220, 30, 30), width=5)
                        if label:
                            bx = max(0, int(x))
                            by = max(0, int(y) - 36)
                            draw.rectangle((bx, by, bx + 34, by + 30), fill=(220, 30, 30))
                            draw.text((bx + 8, by + 2), label, fill=(255, 255, 255), font=font)
                    img.save(lite)
            render_src = lite
            meta = raw_meta
        else:
            meta = src.with_suffix(src.suffix + ".json")
    else:
        meta = src.with_suffix(src.suffix + ".json")
    out = TMP_DIR / f"{src.stem}_crop{src.suffix}"
    if out.exists():
        return out
    keep_full_names = {
        "nav_manufacturing_scrap_real_annotated.png",
        "tmp_scrap_settings_real_annotated.png",
        "tmp_scrap_mo00030_real_annotated.png",
        "tmp_mo_00030_real_annotated.png",
        "tmp_scrap_done_real_annotated.png",
        "tmp_product_scarp_real_annotated.png",
        "tmp_product_scarp_accounting_real2_annotated.png",
        "tmp_scrap_journal_real_annotated.png",
    }
    keep_full = src.name.startswith("nav_dashboard_") or src.name in keep_full_names
    if keep_full:
        Image.open(render_src).save(out)
        return out
    if not meta.exists():
        Image.open(render_src).save(out)
        return out
    try:
        data = json.loads(meta.read_text(encoding="utf-8"))
        boxes = data.get("boxes", [])
        if not boxes:
            Image.open(render_src).save(out)
            return out
        xs = [float(b["x"]) for b in boxes]
        ys = [float(b["y"]) for b in boxes]
        xe = [float(b["x"]) + float(b["width"]) for b in boxes]
        ye = [float(b["y"]) + float(b["height"]) for b in boxes]
        with Image.open(render_src) as img:
            pad_x = max(24, int(img.width * 0.02))
            pad_y = max(24, int(img.height * 0.02))
            left = max(0, int(min(xs)) - pad_x)
            top = max(0, int(min(ys)) - pad_y)
            right = min(img.width, int(max(xe)) + pad_x)
            bottom = min(img.height, int(max(ye)) + pad_y)
            cropped = img.crop((left, top, right, bottom))
            cropped.save(out)
        return out
    except Exception:
        Image.open(render_src).save(out)
        return out


def add_image(doc: Document, image_name: str, caption: str):
    image_path = crop_image(image_name)
    if not image_path.exists():
        add_para(doc, f"[ไม่พบภาพ: {image_name}]")
        return
    doc.add_picture(str(image_path), width=Inches(6.2))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font_run(p.add_run(caption), size=14)


def money(value) -> str:
    try:
        return f"{float(value):,.2f}"
    except Exception:
        return str(value)


def journal_rows(lines: list[dict]) -> list[tuple[str, str, str, str]]:
    rows = []
    for line in lines:
        rows.append(
            (
                f"{line['account_code']} {line['account_name']}",
                line.get("label", ""),
                money(line.get("debit", 0.0)),
                money(line.get("credit", 0.0)),
            )
        )
    return rows


def add_title_page(doc: Document):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font_run(p.add_run("บทที่ 7 การผลิต"), size=26, bold=True, color=RGBColor(31, 78, 121))
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font_run(p2.add_run("Manufacturing - คู่มือรวม"), size=20, bold=True)
    add_para(doc, "", center=True)
    add_para(doc, "เอกสารฉบับนี้อธิบายขั้นตอนการใช้งาน การตรวจสอบผลทางบัญชี และรายงานที่เกี่ยวข้องกับงานผลิต โดยใช้ข้อมูลจริงจากระบบ", center=True)
    add_heading(doc, "สารบัญ", level=2)
    toc = [
        "7.1 ภาพรวมการไหลของข้อมูล (Overview Flow)",
        "7.2 การบันทึกบัญชีในแต่ละขั้นตอน (Journal Entries Steps)",
        "7.3 สรุปผังการตั้งค่าทางบัญชี (Configuration Guide)",
        "7.4 ตัวอย่าง",
        "7.5 ของเสีย (Scrap & Loss)",
        "7.6 Stock Valuation Report",
        "7.7 รายงานส่วนราชการ (รง.8)",
    ]
    add_steps(doc, toc)
    doc.add_section(WD_SECTION_START.NEW_PAGE)


def section_image_groups(title: str, spec: dict) -> tuple[list[tuple[str, str]], list[tuple[str, str]], list[tuple[str, str]]]:
    menu_groups = {
        "7.1": [
            ("รูป 7.1.1 หน้า Dashboard สำหรับเข้าโมดูล Manufacturing", "nav_dashboard_manufacturing_real_annotated.png"),
            ("รูป 7.1.2 เมนู Manufacturing Orders ที่ใช้เปิดใบสั่งผลิต", "nav_manufacturing_orders_real_annotated.png"),
            ("รูป 7.1.3 เมนู Bills of Materials ที่ใช้เปิดสูตรการผลิต", "nav_manufacturing_bom_real_annotated.png"),
        ],
        "7.2": [
            ("รูป 7.2.1 หน้า Dashboard สำหรับเข้าโมดูล Manufacturing", "nav_dashboard_manufacturing_real_annotated.png"),
            ("รูป 7.2.2 เมนู Inventory Valuation ที่ใช้เปิดดูรายการบัญชีที่เกิดจากสต็อก", "nav_inventory_valuation_real_annotated.png"),
        ],
        "7.3": [
            ("รูป 7.3.1 หน้า Dashboard สำหรับเข้าโมดูล Inventory", "nav_dashboard_inventory_real_annotated.png"),
            ("รูป 7.3.2 เมนู Bills of Materials ที่ใช้ตรวจสูตรการผลิต", "nav_manufacturing_bom_real_annotated.png"),
        ],
        "7.4": [
            ("รูป 7.4.1 หน้า Dashboard สำหรับเข้าโมดูล Manufacturing", "nav_dashboard_manufacturing_real_annotated.png"),
            ("รูป 7.4.2 เมนู Manufacturing Orders ที่ใช้เปิดตัวอย่างใบสั่งผลิต", "nav_manufacturing_orders_real_annotated.png"),
        ],
        "7.5": [
            ("รูป 7.5.1 หน้า Dashboard สำหรับเข้าโมดูล Manufacturing", "nav_dashboard_manufacturing_real_annotated.png"),
            ("รูป 7.5.2 เมนู Scrap ที่ใช้บันทึกของเสีย", "nav_manufacturing_scrap_real_annotated.png"),
            ("รูป 7.5.3 เมนู Valuation ที่ใช้ตรวจผลกระทบต้นทุนหลังทำ Scrap", "nav_inventory_valuation_real_annotated.png"),
        ],
        "7.6": [
            ("รูป 7.6.1 หน้า Dashboard สำหรับเข้าโมดูล Inventory", "nav_dashboard_inventory_real_annotated.png"),
            ("รูป 7.6.2 เมนู Stock Valuation Report", "nav_inventory_valuation_real_annotated.png"),
        ],
        "7.7": [
            ("รูป 7.7.1 หน้า Dashboard สำหรับเข้าโมดูล Accounting", "nav_dashboard_accounting_real_annotated.png"),
            ("รูป 7.7.2 เมนูรายงาน รง.8", "nav_accounting_rng8_real_annotated.png"),
        ],
    }
    usage_groups = {
        "7.1": [
            ("รูป 7.1.4 หน้าจอใบสั่งผลิตจริงที่ใช้เป็นตัวอย่าง", "manufacturing_order_form_real_annotated.png"),
            ("รูป 7.1.5 หน้าจอ BoM ของสินค้าตัวอย่าง", "manufacturing_bom_form_real_annotated.png"),
        ],
        "7.2": [
            ("รูป 7.2.3 หน้าจอใบสั่งผลิตที่ใช้เชื่อมกับรายการบัญชี", "manufacturing_order_form_real_annotated.png"),
            ("รูป 7.2.4 หน้าจอรายงาน Valuation ที่ใช้ตามรอย Journal", "inventory_valuation_page_real_annotated.png"),
        ],
        "7.3": [
            ("รูป 7.3.3 หมวดสินค้าสำเร็จรูปที่ใช้ในระบบจริง", "product_category_fg_form_real_annotated.png"),
            ("รูป 7.3.4 หมวดวัตถุดิบที่ใช้ในระบบจริง", "product_category_rm_form_real_annotated.png"),
            ("รูป 7.3.5 สูตรการผลิตที่ใช้ในงานตัวอย่าง", "manufacturing_bom_form_real_annotated.png"),
        ],
        "7.4": [
            ("รูป 7.4.3 ใบสั่งผลิตจริงที่ใช้เป็นตัวอย่าง", "manufacturing_order_form_real_annotated.png"),
        ],
        "7.5": [
            ("รูป 7.5.4 หน้าจอ Scrap จริงในระบบ", "manufacturing_scrap_form_real_annotated.png"),
            ("รูป 7.5.5 หน้าจอ Valuation ที่ใช้ตรวจผลของของเสีย", "inventory_valuation_page_real_annotated.png"),
        ],
        "7.6": [
        ],
        "7.7": [
        ],
    }
    je_groups = {
        "7.1": [
            ("รูป 7.1.6 Journal Entry ตอนตัดวัตถุดิบเข้าสู่งานระหว่างทำ", "journal_mfg_raw_fg02001_real_annotated.png"),
            ("รูป 7.1.7 Journal Entry ตอนรับสินค้าสำเร็จรูปเข้าคลัง", "journal_mfg_finished_real_annotated.png"),
        ],
        "7.2": [
            ("รูป 7.2.5 Journal Entry ตัดกึ่งสำเร็จรูปเข้าสู่งานระหว่างทำ", "journal_mfg_raw_fg02001_real_annotated.png"),
            ("รูป 7.2.6 Journal Entry ตัดบรรจุภัณฑ์เข้าสู่งานระหว่างทำ", "journal_mfg_raw_packaging_real_annotated.png"),
            ("รูป 7.2.7 Journal Entry รับสินค้าสำเร็จรูปเข้าคลัง", "journal_mfg_finished_real_annotated.png"),
        ],
        "7.3": [
            ("รูป 7.3.6 Journal Entry ตัวอย่างผลลัพธ์จากการตั้งค่าถูกต้อง", "journal_mfg_finished_real_annotated.png"),
        ],
        "7.4": [
            ("รูป 7.4.4 Journal Entry ตัดวัตถุดิบของตัวอย่างจริง", "journal_mfg_raw_fg02001_real_annotated.png"),
            ("รูป 7.4.5 Journal Entry รับสินค้าสำเร็จรูปของตัวอย่างจริง", "journal_mfg_finished_real_annotated.png"),
        ],
        "7.5": [
            ("รูป 7.5.6 Journal Entry ตัวอย่างที่ใช้ตรวจผลของของเสียและความสูญเสีย", "journal_mfg_raw_packaging_real_annotated.png"),
        ],
        "7.6": [
            ("รูป 7.6.4 Journal Entry ของบรรทัดวัตถุดิบในรายงาน Valuation", "journal_mfg_raw_fg02001_real_annotated.png"),
            ("รูป 7.6.5 Journal Entry ของบรรทัดรับ FG ในรายงาน Valuation", "journal_mfg_finished_real_annotated.png"),
        ],
        "7.7": [
            ("รูป 7.7.4 Journal Entry ตัวอย่างที่ใช้ตรวจเทียบกับรายงาน รง.8", "journal_mfg_finished_real_annotated.png"),
        ],
    }
    prefix = title.split(" ", 1)[0]
    return (
        uniq_pairs(menu_groups.get(prefix, [(spec["dashboard_caption"], spec["dashboard_image"]), (spec["menu_caption"], spec["menu_image"])])),
        uniq_pairs(usage_groups.get(prefix, spec.get("usage_images", []))),
        uniq_pairs(je_groups.get(prefix, spec.get("je_images", []))),
    )


def enrich_scrap_spec(specs: list[dict]) -> None:
    for spec in specs:
        if not spec["title"].startswith("7.5 "):
            continue
        spec["objectives"] = [
            "เพื่อให้ผู้ใช้งานเข้าใจว่า Scrap คือการตัดสินค้าหรือวัตถุดิบออกจากคลังเพราะไม่สามารถใช้งานต่อได้",
            "เพื่อให้เห็นว่า custom ของระบบมีทั้งการกันเลือกสินค้าผิดตอน Scrap การเติมวัตถุดิบกลับให้อัตโนมัติ และการเตรียมผลักต้นทุน Scrap กลับเข้า FG",
            "เพื่อให้ตรวจได้ทันทีว่ารายการไหนเป็นผลจริงที่เกิดแล้ว และรายการไหนเป็นความสามารถของ custom ที่ตั้งค่าไว้แต่ยังไม่เกิดรายการปิดจริงในฐานข้อมูลปัจจุบัน",
        ]
        spec["overview"] = [
            "ระบบนี้มี custom ที่เกี่ยวกับ Scrap อยู่ 3 ส่วน คือ จำกัดสินค้าที่ scrap ได้ให้เหลือเฉพาะของที่เกี่ยวกับงานผลิต, เติมของกลับให้ MO อัตโนมัติเมื่อ scrap component ผ่าน Shop Floor, และเตรียมผลักต้นทุน Scrap กลับเข้า FG ผ่าน Landed Cost",
            "ข้อมูลจริงในฐานตอนนี้มี Scrap ที่ยืนยันแล้ว 1 รายการคือ SP/00001 ของสินค้า เมนทอล และมี Scrap ร่างที่ผูกกับ MO GMP/MOPH/00030 อีก 1 รายการ",
            "บริษัทตั้ง Scrap Landed Cost Service เป็น product ชื่อ Scarp ไว้แล้ว แต่จากข้อมูลจริงตอนนี้ยังไม่พบ manufacturing landed cost ที่ validate สำเร็จจริง",
        ]
        spec["menu_paths"] = [
            "Manufacturing > Operations > Scrap",
            "Manufacturing > Operations > Manufacturing Orders",
            "Inventory > Reporting > Valuation",
            "Manufacturing > Configuration > Settings",
        ]
        spec["usage_intro"] = (
            "หัวข้อนี้ใช้ข้อมูลจริงในระบบทั้งหมด โดยแบ่งให้เห็น 2 ส่วนชัดเจน คือส่วนของการบันทึก Scrap ที่เกิดขึ้นแล้วจริง "
            "และส่วนของ custom ที่ตั้งค่าไว้เพื่อจัดการต้นทุน Scrap ต่อจากนั้น"
        )
        spec["steps"] = [
            "เริ่มจากหน้า Manufacturing Dashboard แล้วเข้าเมนู Scrap เพื่อเปิดรายการของเสียที่ต้องการตรวจสอบ",
            "เปิดรายการ Scrap ร่างที่ผูกกับ MO GMP/MOPH/00030 เพื่อดูว่าสินค้าที่ถูกตัด ปริมาณ ตำแหน่งต้นทาง ตำแหน่งของเสีย และเลข Manufacturing Order ถูกบันทึกไว้ครบหรือไม่",
            "ใช้หน้าจอ MO GMP/MOPH/00030 เพื่อตรวจต่อว่ารายการ Scrap นี้ผูกกับงานผลิตจริงใบไหน และงานผลิตใบนี้ปิดงานแล้วหรือยัง",
            "เปิดรายการ SP/00001 ซึ่งเป็น Scrap ที่ Done แล้ว เพื่อดูตัวอย่างรายการที่มีผลจริงกับสต็อก",
            "ถ้าต้องการตรวจว่าระบบเตรียมใช้ custom ด้านต้นทุนอย่างไร ให้เข้า Settings แล้วตรวจว่า Scrap Landed Cost Service ของบริษัทตั้งเป็น product ชื่อ Scarp แล้ว",
            "เปิด product Scarp เพื่อตรวจว่าเป็น service product ที่ใช้รองรับการผลักต้นทุน Scrap ผ่าน Landed Cost และดูข้อมูลบัญชีที่ผูกไว้กับสินค้านี้",
            "หลังจากนั้นเปิด Journal Entry STJ/26/03/27311 ของ SP/00001 เพื่อดูผลทางบัญชีจริงที่เกิดจากการ Scrap",
        ]
        spec["fields"] = [
            ("Product", "สินค้าที่ถูกตัดออกเป็นของเสีย"),
            ("Quantity", "จำนวนที่ตัดออกจากคลัง"),
            ("Source Location", "ตำแหน่งต้นทางก่อน Scrap"),
            ("Scrap Location", "ตำแหน่งปลายทางที่รับของเสีย"),
            ("Manufacturing Order", "เลขอ้างอิงของงานผลิตที่ผูกกับ Scrap"),
            ("Source Document", "เอกสารต้นทางที่ทำให้เกิด Scrap"),
            ("Status", "สถานะว่ายังร่างหรือยืนยันแล้ว"),
            ("Scrap Landed Cost Service", "product ที่บริษัทใช้รองรับการผลักต้นทุน Scrap กลับผ่าน Landed Cost"),
        ]
        spec["scenarios"] = [
            ("ตรวจ Scrap ระหว่างผลิต", "เปิด Scrap ที่ผูกกับ MO แล้วตรวจ Product, Quantity, Source Location, Scrap Location และเลข MO", "ผู้ใช้รู้ได้ทันทีว่าของเสียรายการนี้มาจากงานผลิตใบไหน"),
            ("ตรวจ Scrap ที่มีผลจริงกับสต็อกแล้ว", "เปิด SP/00001 ที่สถานะ Done", "เห็นตัวอย่างของเสียที่ตัดออกจากคลังจริงแล้ว"),
            ("ตรวจ custom ด้านต้นทุน", "เข้า Settings และเปิด product Scarp", "ยืนยันได้ว่าบริษัทตั้งค่าระบบรองรับการผลักต้นทุน Scrap แล้ว"),
            ("ตรวจผลทางบัญชี", "เปิด STJ/26/03/27311 แล้วดูบรรทัด Debit/Credit", "ยืนยันได้ว่าระบบย้ายมูลค่าสินค้าออกจากบัญชีสต็อกและรับรู้เป็นบัญชีสินค้าขาออก"),
        ]
        spec["je_intro"] = (
            "ผลทางบัญชีที่เกิดขึ้นจริงในฐานตอนนี้คือ Journal ของ SP/00001 ชื่อ STJ/26/03/27311 "
            "ส่วนฟังก์ชัน Scrap Landed Cost เป็น custom ที่ถูกตั้งค่าไว้แล้ว แต่ยังไม่พบรายการ manufacturing landed cost ที่ปิดจริงให้ใช้อ้างอิงเป็นภาพจริงในฐานนี้"
        )
        spec["je_tables"] = [
            (
                "รายการบัญชีจริงของ SP/00001",
                [
                    ("116902 บัญชีสินค้าขาออก", "ระบบรับรู้มูลค่าของเสียที่ถูกตัดออกจากคลัง", "120.00", "0.00"),
                    ("116002 วัตถุดิบ - สารเคมี", "ระบบลดยอดวัตถุดิบที่ถูก Scrap ออกจากคลัง", "0.00", "120.00"),
                ],
                "รายการนี้แปลว่ามูลค่าวัตถุดิบเมนทอลถูกตัดออกจากบัญชีวัตถุดิบและย้ายไปบันทึกเป็นบัญชีสินค้าขาออกทันทีเมื่อ Scrap ถูกยืนยัน",
            )
        ]
        spec["cautions"] = [
            "ถ้า Scrap ยังเป็น Draft ระบบจะยังไม่ตัดมูลค่าออกจากคลังและยังไม่เกิด Journal",
            "ถ้า Scrap ผูกกับ MO ควรตรวจให้แน่ใจว่า Product, Quantity และ Source Location ตรงกับของเสียจริงก่อนกด Validate",
            "ฟังก์ชัน Scrap Landed Cost ใช้ product Scarp เป็นตัวแทนต้นทุน แต่ในฐานข้อมูลปัจจุบันยังไม่พบเคสที่ validate สำเร็จจริง จึงไม่ควรอธิบายเหมือนมีรายการบัญชีส่วนนี้เกิดขึ้นแล้ว",
            "ถ้า Scrap เกิดจาก component ที่ทำผ่าน Shop Floor custom อาจเติมของกลับให้ MO อัตโนมัติได้ จึงควรตรวจทั้งรายการ Scrap และรายการเบิกเพิ่มประกอบกัน",
        ]
        spec["menu_images_override"] = [
            ("รูป 7.5.1 หน้า Dashboard สำหรับเข้าโมดูล Manufacturing", "nav_dashboard_manufacturing_real_annotated.png"),
            ("รูป 7.5.2 เมนู Scrap ที่ใช้บันทึกของเสีย", "nav_manufacturing_scrap_real_annotated.png"),
            ("รูป 7.5.3 หน้า Settings ที่ตั้ง Scrap Landed Cost Service", "tmp_scrap_settings_real_annotated.png"),
        ]
        spec["usage_images_override"] = [
            ("รูป 7.5.4 หน้าจอ Scrap ร่างที่ผูกกับ MO GMP/MOPH/00030", "tmp_scrap_mo00030_real_annotated.png"),
            ("รูป 7.5.5 ใบสั่งผลิต GMP/MOPH/00030 ที่อ้างอิงใน Scrap", "tmp_mo_00030_real_annotated.png"),
            ("รูป 7.5.6 หน้าจอ Scrap ที่ Done แล้วของสินค้า เมนทอล", "tmp_scrap_done_real_annotated.png"),
            ("รูป 7.5.7 product Scarp ที่ใช้ใน Scrap Landed Cost Service", "tmp_product_scarp_real_annotated.png"),
            ("รูป 7.5.8 แท็บ Accounting ของ product Scarp", "tmp_product_scarp_accounting_real2_annotated.png"),
        ]
        spec["je_images_override"] = [
            ("รูป 7.5.9 Journal Entry จริงของ Scrap SP/00001", "tmp_scrap_journal_real_annotated.png"),
        ]
        return


def build_section(doc: Document, spec: dict):
    title = spec["title"]
    add_heading(doc, title, level=1)
    add_para(doc, "Manufacturing", bold=True)
    add_heading(doc, "วัตถุประสงค์", level=2)
    add_bullets(doc, spec["objectives"])
    add_heading(doc, "ภาพรวมการทำงาน", level=2)
    add_bullets(doc, spec["overview"])
    add_heading(doc, "เมนูที่ใช้", level=2)
    add_bullets(doc, spec["menu_paths"])

    menu_images, usage_images, je_images = (
        uniq_pairs(spec["menu_images_override"]),
        uniq_pairs(spec["usage_images_override"]),
        uniq_pairs(spec["je_images_override"]),
    ) if spec.get("menu_images_override") else section_image_groups(title, spec)
    shown_images = set()
    for caption, image in menu_images:
        if image in shown_images:
            continue
        add_image(doc, image, caption)
        shown_images.add(image)

    add_heading(doc, "3. ขั้นตอนการใช้งาน", level=2)
    add_para(doc, spec["usage_intro"])
    add_steps(doc, spec["steps"])
    for caption, image in usage_images:
        if image in shown_images:
            continue
        add_image(doc, image, caption)
        shown_images.add(image)

    field_rows = [(name, meaning) for name, meaning, *_ in spec["fields"]]
    add_heading(doc, "คำอธิบาย Field สำคัญ", level=2)
    add_table(doc, ["Field Name", "Meaning"], field_rows)

    add_heading(doc, "ฟังก์ชันที่ใช้งานได้ในหัวข้อนี้", level=2)
    scenario_rows = [(a, b, c) for a, b, c in spec["scenarios"]]
    add_table(doc, ["ฟังก์ชัน", "กดหรือทำอะไร", "ผลที่ควรได้"], scenario_rows)

    add_heading(doc, "4. ตรวจสอบ Journal Entry", level=2)
    add_para(doc, spec["je_intro"])
    for caption, image in je_images:
        if image in shown_images:
            continue
        add_image(doc, image, caption)
        shown_images.add(image)
    for heading, rows, explain in spec["je_tables"]:
        add_para(doc, heading, bold=True)
        add_table(doc, ["บัญชี", "คำอธิบาย", "Debit", "Credit"], rows)
        add_para(doc, explain)
    add_heading(doc, "ข้อควรระวัง", level=2)
    add_bullets(doc, spec["cautions"])


def render_pdf(docx_path: Path, pdf_path: Path):
    if pdf_path.exists():
        pdf_path.unlink()
    subprocess.run(
        [str(SOFFICE), "--headless", "--convert-to", "pdf", "--outdir", str(docx_path.parent), str(docx_path)],
        check=True,
        cwd=str(ROOT),
    )
    generated = docx_path.with_suffix(".pdf")
    if generated != pdf_path and generated.exists():
        if pdf_path.exists():
            pdf_path.unlink()
        generated.replace(pdf_path)


def render_pages(pdf_path: Path):
    RENDER_DIR.mkdir(parents=True, exist_ok=True)
    for old in RENDER_DIR.glob("chapter7-*.png"):
        old.unlink()
    with fitz.open(pdf_path) as pdf:
        for index, page in enumerate(pdf, start=1):
            pix = page.get_pixmap(matrix=fitz.Matrix(1.6, 1.6), alpha=False)
            pix.save(RENDER_DIR / f"chapter7-{index:03d}.png")


def main():
    mod = load_source_module()
    specs = mod.manufacturing_specs()
    enrich_scrap_spec(specs)
    TMP_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_doc(doc)
    add_title_page(doc)
    for idx, spec in enumerate(specs):
        build_section(doc, spec)
        if idx != len(specs) - 1:
            doc.add_section(WD_SECTION_START.NEW_PAGE)
    doc.save(OUTPUT_DOCX)
    render_pdf(OUTPUT_DOCX, OUTPUT_PDF)
    render_pages(OUTPUT_PDF)
    print(OUTPUT_DOCX)
    print(OUTPUT_PDF)


if __name__ == "__main__":
    main()

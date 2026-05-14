from docx import Document
import os

# 1. Update 7.4 (Time Units)
file_74 = r'C:\365_project\TheCool18e\Dev\manual\Accouting_Manual\generated_20260408\docx\7.4_7.4 ตัวอย่าง.docx'
doc74 = Document(file_74)
for p in doc74.paragraphs:
    if 'เปิด MO GMP/MOPH/00030 และตรวจสินค้า วันที่เริ่มผลิต และวันที่ผลิตเสร็จ' in p.text:
        p.text = p.text.replace('และวันที่ผลิตเสร็จ', 'วันที่ผลิตเสร็จ และตรวจสอบหน้าจอ Time Summary โดย Expected Duration จะแสดงค่าเป็น "นาที (Minutes)" สำหรับ MO นี้ และ Total Expected Duration จะแสดงค่าเป็น "วัน (Days)" สำหรับงานทั้งสายการผลิต (Hierarchy) ซึ่งระบบจะคำนวณตาม critical path ของงานขนาน (Parallel) ให้อัตโนมัติ')
doc74.save(file_74)
print('Updated 7.4')

# 2. Update 7.5 (Scrap Logic)
file_75 = r'C:\365_project\TheCool18e\Dev\manual\Accouting_Manual\generated_20260408\docx\7.5_7.5 ของเสีย (Scrap & Loss).docx'
doc75 = Document(file_75)
for p in doc75.paragraphs:
    if 'เพื่อให้ผู้ใช้งานบันทึกของเสียและติดตามผลกระทบต่อสต็อกและบัญชีได้' in p.text:
        p.insert_paragraph_before('เพื่อให้เข้าใจข้อจำกัดในการ Scrap เฉพาะสินค้าสำเร็จรูป และระบบการเติมสินค้าอัตโนมัติ', style=p.style)
    
    if 'เข้าเมนู Scrap แล้วเปิดรายการของเสียที่ต้องการตรวจสอบ' in p.text:
        p.text = '1. เข้าเมนู Scrap แล้วกด Create หรือเลือกรายการที่ต้องการ\n2. ข้อกำหนดสำคัญ: ระบบจะอนุญาตให้เลือกสินค้าเฉพาะที่เป็น Finished Good หรือ By-product ของใบผลิตนั้นๆ เท่านั้น หากเลือกวัตถุดิบระบบจะแสดง Error\n3. เมื่อยืนยัน (Done) รายการ Scrap ระบบจะทำการสร้างใบเติมวัตถุดิบ (Auto-Replenish) ให้ทันทีเพื่อให้การผลิตดำเนินต่อไปได้'

doc75.save(file_75)
print('Updated 7.5')

# 3. Update 7.6 (Valuation / Landed Cost)
file_76 = r'C:\365_project\TheCool18e\Dev\manual\Accouting_Manual\generated_20260408\docx\7.6_7.6 Stock Valuation Report.docx'
doc76 = Document(file_76)
found = False
for p in doc76.paragraphs:
    if 'ข้อควรระวัง' in p.text:
        p.insert_paragraph_before('การปันส่วนต้นทุนของเสีย (Scrap Landed Cost)', style='Heading 2')
        doc76.paragraphs[-1].insert_paragraph_before('ระบบมีการตั้งค่าพิเศษเพื่อโอนมูลค่าของวัตถุดิบที่เสีย (Scrap) กลับเข้าไปเป็นต้นทุนของสินค้าสำเร็จรูป (FG) ผ่านระบบ Landed Cost อัตโนมัติ ดังนั้นมูลค่าในรายงาน Valuation ของ FG จะรวมค่าความสูญเสียเหล่านี้ไว้แล้วเพื่อให้ได้ต้นทุนขาย (COGS) ที่ถูกต้องที่สุด', style='Normal')
        break
doc76.save(file_76)
print('Updated 7.6')
